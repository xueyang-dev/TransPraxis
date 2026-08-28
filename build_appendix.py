from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path("/Users/xueyang/Dev/TransPraxis")
ST_PATH = Path("/Users/xueyang/Documents/毕业论文/2_E_原文ST.docx")
TT_PATH = Path("/Users/xueyang/Documents/毕业论文/3_C_译文TT.docx")
OUT_PATH = ROOT / "附录一_无人机感知与共同体_原文与译文_终审修订版.docx"


def table_texts(path: Path) -> list[list[str]]:
    doc = Document(path)
    table = doc.tables[0]
    return [[p.text for p in row.cells[0].paragraphs] for row in table.rows]


ST = table_texts(ST_PATH)
TT = table_texts(TT_PATH)


SOURCE_HYPHEN_FIXES = [
    ("nation- states", "nation-states"),
    ("nineteenth- century", "nineteenth-century"),
    ("Franco- Prussian", "Franco-Prussian"),
    ("first- person", "first-person"),
    ("three- dimensional", "three-dimensional"),
    ("two- dimensional", "two-dimensional"),
    ("twenty- first", "twenty-first"),
    ("twenty- six", "twenty-six"),
    ("scare- eye", "scare-eye"),
    ("cost- effective", "cost-effective"),
    ("large- scale", "large-scale"),
    ("techno- geographies", "techno-geographies"),
    ("self- preservation", "self-preservation"),
    ("community- shaping", "community-shaping"),
    ("sea- son", "season"),
    ("sea-son", "season"),
    ("hot- air", "hot-air"),
    ("remote- controlled", "remote-controlled"),
    ("data- collecting", "data-collecting"),
    ("profit- focused", "profit-focused"),
    ("game- changing", "game-changing"),
    ("media- archaeological", "media-archaeological"),
    ("life- affirming", "life-affirming"),
    ("economic- industrial", "economic-industrial"),
    ("site- specific", "site-specific"),
    ("high- tech", "high-tech"),
    ("more- than- optical", "more-than-optical"),
    ("post- carbon", "post-carbon"),
    ("drone- generated", "drone-generated"),
    ("drone- witness", "drone-witness"),
    ("community- building", "community-building"),
    ("community- led", "community-led"),
    ("low- cost", "low-cost"),
    ("post- apartheid", "post-apartheid"),
    ("inter- connectedness", "interconnectedness"),
    ("interconnect- edness", "interconnectedness"),
    ("sens- ing", "sensing"),
    ("litera- ture", "literature"),
    ("techno- optimistic", "techno-optimistic"),
    ("Bal- loons", "Balloons"),
    ("Mont- golfier", "Montgolfier"),
    ("balloon- flying", "balloon-flying"),
    ("ball- loons", "balloons"),
    ("unmanned- ness", "unmannedness"),
    ("crew- less", "crewless"),
    ("animaled", "animaled"),
    ("life- taking", "life-taking"),
    ("life- making", "life-making"),
    ("large- scale", "large-scale"),
    ("profit- maximization", "profit-maximization"),
    ("food- free", "food-free"),
]


def clean_source(text: str) -> str:
    text = text.replace("\xa0", " ").replace("\u200b", "")
    text = text.replace("'", "’")
    text = text.replace("--", "—")
    text = re.sub(r"\s+—\s+", "—", text)
    for old, new in SOURCE_HYPHEN_FIXES:
        text = text.replace(old, new)
    # Mechanical extraction errors confirmed against the supplied ST/PDF.
    for old, new in [
        ("earth.The", "earth. The"),
        ("observe,register,and", "observe, register, and"),
        ("Anthropocene.He", "Anthropocene. He"),
        ("views.planes", "views. Planes"),
        ("people.The", "people. The"),
        ("communities.Instead", "communities. Instead"),
    ]:
        text = text.replace(old, new)
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def remove_english_glosses(text: str) -> str:
    """Remove parenthetical translator glosses when the Chinese already carries the term.

    Proper names, titles, acronyms, dates, and technical abbreviations are retained.
    """

    pattern = re.compile(r"（([^（）]{1,100})）")

    def repl(match: re.Match[str]) -> str:
        inner = match.group(1).strip()
        if not re.fullmatch(r"[a-z][A-Za-z0-9 .,’'\-/:×&]+", inner):
            return match.group(0)
        if any(ch.isdigit() for ch in inner) or "×" in inner:
            return match.group(0)
        return ""

    return pattern.sub(repl, text)


TRANSLATION_REPLACEMENTS = [
    ("民用（业余）无人机", "业余无人机"),
    ("民用（业余）", "业余"),
    ("商业无人机", "商用无人机"),
    ("排斥之篱", "驱逐围栏"),
    ("后商品", "Postcommodity"),
    ("盖亚", "盖娅"),
    ("感官技术地理学", "感知性的技术地理景观"),
    ("感官集合", "感知系统"),
    ("感官特质", "感知特质"),
    ("感官属性", "感知属性"),
    ("感知域", "感知系统"),
    ("立体感知", "体积性感知"),
    ("立体感", "体积性"),
    ("立体维度", "体积维度"),
    ("立体传感器", "体积性感知传感器"),
    ("体积化感知", "体积性感知"),
    ("体积感知", "体积性感知"),
    ("远程感知技术", "遥感技术"),
    ("远程感知", "遥感"),
    ("脱碳共同体", "后碳共同体"),
    ("脱碳未来", "后碳未来"),
    ("脱碳", "后碳"),
    ("他者性", "他异性"),
    ("社区", "共同体"),
    ("社群", "共同体"),
    ("精彩绝伦", "生动"),
    ("无比珍贵", "宝贵"),
    ("心驰神往", "着迷"),
    ("视觉盛宴", "令人惊叹的景象"),
    ("惊心动魄", "壮观"),
    ("惨痛教训", "失败"),
    ("亟待保护", "需要保护"),
    ("深度剖析", "细读"),
    ("奉为最新", "视为最新"),
    ("静谧祥和地升入高空", "平静地升空"),
    ("极具美感的航空器", "美丽多彩的航空器"),
    ("互联共生", "相互联结"),
    ("大放异彩/得以升华", "因“自然之笔”而得到增强"),
    ("数据化的大众", "数据化公众"),
    ("主宰着生与死", "裁定生死"),
    ("去人类化的剥夺", "去人性化"),
    ("对世界进行去人类化的剥夺", "使世界去人性化"),
    ("戛然而止", "中断"),
    ("感情互动", "情动性互动"),
    ("情感互动", "情动性互动"),
    ("情感力量", "情动性力量"),
    ("富有情感的见证者", "情动性的见证者"),
    ("精粹媒介", "典型媒介"),
    ("赋予未来的力量", "面向未来的力量"),
    ("被奉为经典", "被视为经典"),
    ("感官感知技术", "感知技术"),
    ("技术地理", "技术地理景观"),
]


# This pass is intentionally limited to confirmed mechanical, terminology, and
# overstatement errors found in the current appendix. It is not a stylistic rewrite.
REVISION_REPLACEMENTS = [
    ("涉及的不不仅仅是", "涉及的不仅仅是"),
    ("““", "“"),
    ("””", "”"),
    ("非视觉体制的行星感知", "非视觉的行星感知"),
    ("霍特赖夫", "豪特里夫"),
    ("《分界》", "《分裂》"),
    ("这些人造边界的虚妄性", "政治边界的人为建构性"),
    ("边界的随机性与虚妄性", "边界的随机性与人为建构性"),
    ("而拒斥被固化", "而不被固定下来"),
    ("有着异曲同工之妙", "与之相似"),
    ("再次异曲同工的是", "与之类似的是"),
    ("庞然之态", "巨大规模"),
    ("具有极其重要的相关性", "具有重要意义"),
    ("产生了深远影响", "产生了影响"),
    ("《绵延的生态学》", "《绵延生态学》"),
    ("“立体性”", "“体积性”"),
    ("使用“立体”一词", "使用“体积性”一词"),
    ("在这些情境中，“立体”指向", "在这些情境中，“体积性”指向"),
    ("将视角彻底翻转", "将视角倒转"),
    ("而拒斥将阐释的过程停滞在任何一端", "而不将阐释固定于其中任何一端"),
    ("将距离感与垂直性同一更近的视角", "将距离感和垂直性与一种更近距离的视角"),
    ("奇异性", "独异性"),
    ("投影/预测", "投影"),
    ("一位敏感且情动性的见证者", "一位敏感且具有情动性的见证者"),
    ("我并无意为一种前工业时代的耕作方式去辩逐某种怀旧式的愿景。", "我无意为一种前工业时代的耕作方式辩护，或主张一种怀旧式的愿景。"),
    (
        "夏马尤指出：“尽管传统的伦理学被定义为一种‘善生’与‘善终’的学说，但死灵伦理学”却表现为一种‘善杀’的学说。死灵伦理学大肆宣扬杀人的程序，并将其转化为自满的道德评估对象。”",
        "夏马尤指出：“尽管传统的伦理学被定义为一种‘善生’与‘善终’的学说，但‘死灵伦理学’采取的是一种‘善杀’的学说。死灵伦理学讨论杀人的程序，并将其转化为自满的道德评价对象。”",
    ),
    ("““死灵伦理学”", "“死灵伦理学”"),
    ("这种“居高临下的位置”", "这种‘居高临下的位置’"),
    ("、“间隙与底层空间”中", "、‘间隙与底层空间’中"),
    (
        "他“利用无人机技术，旨在将其与监视和战争的关联，彻底翻转为某种更具“生命肯定性”的事物，将这些空中机器人视为在相互依存的行星生态系统中重塑我们自我认知的工具。”",
        "他“利用无人机技术，将其与监视和战争的关联转化为某种更具生命肯定性的事物，并将这些空中机器人视为在相互依存的行星生态系统中重塑我们自我认知的工具。”",
    ),
]


def clean_translation(text: str) -> str:
    text = text.replace("\xa0", " ").replace("\u200b", "")
    text = re.sub(r"【译者注：.*?】", "", text)
    text = re.sub(r"\[译者注：.*?\]", "", text)
    text = re.sub(r"\[\d+\]", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    for old, new in TRANSLATION_REPLACEMENTS:
        text = text.replace(old, new)
    for old, new in [
        ("技术地理景观景观", "技术地理景观"),
        ("立体化思维", "体积性思维"),
        ("立体空间", "体积空间"),
        ("“立体”一词", "“体积性”一词"),
        ("以立体的方式进行感知", "以体积性的方式进行感知"),
        ("这里的立体性", "这里的体积性"),
        ("这种立体性", "这种体积性"),
        ("立体模态", "体积性模态"),
        ("立体把握", "体积性把握"),
        ("立体界面", "体积性界面"),
        ("立体地感知", "以体积性感知的方式感知"),
        ("立体大地感知", "体积性大地感知"),
        ("感官图谱", "感知图谱"),
        ("感官体验", "感知体验"),
    ]:
        text = text.replace(old, new)
    for old, new in REVISION_REPLACEMENTS:
        text = text.replace(old, new)
    text = remove_english_glosses(text)
    text = text.replace("艺术小组“Postcommodity”（Postcommodity）", "艺术家团体 Postcommodity")
    text = text.replace("“Postcommodity”艺术小组", "Postcommodity 艺术家团体")
    text = re.sub(r"‘([^‘’'\n]{1,80})'", r"“\1”", text)
    text = re.sub(r"(?<![A-Za-z])'([^'\n]{1,80})'(?![A-Za-z])", lambda m: f"“{m.group(1)}”" if re.search(r"[\u4e00-\u9fff]", m.group(1)) else m.group(0), text)
    text = re.sub(r"“([^“”\n]{1,100})'(?=[，。；：！？、”）》])", r"“\1”", text)
    text = re.sub(r"(?<=[\u4e00-\u9fff])'(?=[\u4e00-\u9fff，。；：！？、”）》])", "", text)
    text = re.sub(r"(?<=[，。；：！？、（“])'(?=[\u4e00-\u9fff])", "", text)
    text = re.sub(r'"([^"\n]*[\u4e00-\u9fff][^"\n]*)"', r"“\1”", text)
    text = re.sub(r"[’']\s*([“《])", r"\1", text)
    text = re.sub(r"([”》])\s*[’']", r"\1", text)
    text = re.sub(r"\s+([，。；：！？、）》】])", r"\1", text)
    text = re.sub(r"([（《【])\s+", r"\1", text)
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"(?<=[\u4e00-\u9fff。；，：！？、）》】”]) +(?=[\u4e00-\u9fff“‘《（【])", "", text)
    # Quote normalization above can expose a doubled Chinese quotation mark;
    # collapse it once more at the end of the mechanical cleanup pass.
    text = text.replace("““", "“").replace("””", "”")
    # Two cells in the supplied TT contain malformed wording/quotation that the
    # earlier broad cleanup cannot match reliably after gloss removal.  These are
    # narrow repairs of the confirmed strings, not a general rewrite.
    text = text.replace(
        "但死灵伦理学”却表现为一种‘善杀’的学说",
        "但‘死灵伦理学’采取的是一种‘善杀’的学说",
    )
    text = text.replace("死灵伦理学大肆宣扬杀人的程序", "死灵伦理学讨论杀人的程序")
    text = text.replace("自满的道德评估对象", "自满的道德评价对象")
    text = text.replace(
        "利用无人机技术，旨在将其与监视和战争的关联，彻底翻转为某种更具“生命肯定性”的事物",
        "利用无人机技术，将其与监视和战争的关联转化为某种更具生命肯定性的事物",
    )
    # These two confirmed cells contain English glosses between the words being
    # repaired, so apply the final wording after gloss removal as well.
    text = text.replace(
        "而拒斥将阐释的过程停滞在任何一端",
        "而不将阐释固定于其中任何一端",
    )
    text = text.replace(
        "将距离感与垂直性同一更近的视角",
        "将距离感和垂直性与一种更近距离的视角",
    )
    # Keep the theoretical use of volumetric terminology consistent in the
    # chapter-introduction cell, where the source alternates between
    # volumetricity and three-dimensionality.
    text = text.replace("感知地球的无人机同样能够走向三维化", "感知地球的无人机同样能够转向体积性感知")
    text = text.replace("无人机感知系统的这种三维性", "无人机感知系统的这种体积性")
    text = text.replace("这种“立体性”", "这种“体积性”")
    text = text.replace("使用“立体”一词", "使用“体积性”一词")
    return text


TRANSLATION_OVERRIDES: dict[tuple[int, int], str] = {
    (0, 0): "第四章：扁平化感知与行星共同体",
    (1, 0): (
        "“从高处看世界，不仅会使其扁平化，还会使其锐化。”在奥默·法斯特（Omer Fast）的虚构短片《5000英尺是最佳高度》中，扮演无人机飞行员的人如此说道。"
        "扁平化是无人机从上方监视地球时产生的图像学特征；这些图像将世界呈现为一个没有深度的抽象表面。"
        "关于军用无人机视觉体制的研究，往往将这些图像解释为一种操作性感知形式的产物；这种感知形式会使世界去人性化，并对世界实施控制和支配。"
        "正如法斯特片中的无人机飞行员所言，扁平化的无人机视角会使事物变得更加尖锐：它仿佛一把暴力的利刃，切割并杀戮。"
        "但扁平化也可能产生另一种效果。与军用无人机不同，业余无人机所生成的扁平化视角能够唤起相互联结，而非控制与支配；无人机由此可以成为承载行星性的媒介。"
    ),
    (2, 0): (
        "近几十年来，“行星性”已成为人文学科——包括媒介研究、文化研究与文学——中的一个重要概念。尽管全球主义植根于资本主义的全球流动将世界连接起来这一观念，但行星性话语将人置于与行星的关系之中，把行星理解为生命与气候的场域。"
        "行星视角不依据民族国家、国际合作或全球网络对地球进行分类，而是将地球上的人视为一个整体，但不把他们视为同一。根据加亚特里·查克拉瓦蒂·斯皮瓦克（Gayatri Chakravorty Spivak）的观点，“行星性”一词与“世俗的”“全球的”和“大陆的”等概念相对立；这些概念依循资本主义意识形态运作，并始终带有文化与民族本质主义色彩。"
        "相比之下，对于行星主体而言，他异性并非派生属性。在确立人类主体与行星之间关系的过程中，差异化的条件被抹除：“行星是属于另一个系统的他异性物种；然而我们只是借居于此。”行星主体并不坚持对行星拥有所有权。他们将地球视为借居之地，从而反对全球金融化与资本主义同质化的进程。"
        "需要指出的是，行星性概念也因包含抹除差异与异质性的总体化、统一化倾向而受到批评；我对这些艺术作品的解读也力图呈现这一点。"
    ),
    (3, 0): (
        "本章认为，无人机的扁平化凝视具有构想行星共同体的潜力。在这类共同体愿景中，人类并不把地球据为部分人的特定领土，而是将其视为需要保护的共同行星空间。"
        "为展开这一论证，我首先对无人机进行历史化考察，寻找无人机与热气球之间的“家族相似性”。这种媒介考古学与谱系学方法反对把无人机视为最新的、能够改变局面的技术的技术乐观主义叙事，并表明，作为热气球后裔的无人机具有生成关于行星性的创造性、想象性与乌托邦式愿景的潜力。"
        "在细读无人机的美学想象及其行星愿景之前，必须先勾勒出无人机与气球之间的这一谱系关系。"
    ),
    (5, 0): (
        "乍看之下，热气球或煤气气球与无人机似乎截然不同。气球是美丽多彩、平静升空的航空器；无人机则相当丑陋。无人机主要在21世纪得到开发和使用；煤气气球则发明于18世纪。"
        "然而，两种空中技术之间仍存在一些引人注目的相似之处。我不把无人机的发明看作线性、按时间顺序展开的叙事。相反，我把空中气球视为无人机的远亲，二者通过一种感知和阐释地球的特定审美模式相联结。让我们来看一看气球飞行的开端。"
    ),
    (6, 0): (
        "1783年6月5日，也就是法国大革命前几年，蒙戈尔菲耶兄弟（Montgolfier brothers）将他们的第一个热气球升入天空。此后，工程师、爱好者、企业家、艺术家和作家对气球飞行技术进行实验、改进和发展，并撰写相关文字。"
        "例如，费利克斯·纳达尔（Félix Nadar）是19世纪著名的气球飞行家。他壮观的飞行与失败（例如其“巨人号”热气球的坠毁），以及早期利用气球进行航拍的尝试，至今仍令人着迷。除了引人注目的表演，空中气球还激发了时尚潮流，影响了服装、精制瓷器和瓷质小雕像的设计。"
        "18世纪末和19世纪的许多文学作者都对气球产生兴趣，并怀着极大热情书写这些新型飞行器。与当代无人机艺术的兴起相似，气球艺术也成为一种审美类型。18、19世纪艺术与大众文化中的“气球狂热”，可以与当今媒体围绕无人机技术制造的“无人机热潮”相比较。"
    ),
    (7, 0): (
        "除了出现在美学作品和大众娱乐活动中，气球与无人机一样，从一开始就被视为具有战争战略潜力的双重用途技术。约瑟夫·孟戈尔菲（Joseph Montgolfier）本人曾指出：“只要把气球的球囊做得足够大，就有可能装下一整支军队，乘风而行，直接从英国人的头顶越境而降。”"
        "正如无人机一样，气球很快被军方工具化，用于空中侦察、轰炸和物资运输。例如，1794年6月26日的弗勒吕斯战役中，法国军队利用“进取号”（L’Entreprenant）气球来更好地观察敌情。尽管法国赢得了这场战役——这或许部分归功于气球所产生的“震慑与敬畏”效应——但法国对气球的军事使用在1799年中断，当时拿破仑解散了航空兵团（Aerostatic Corps）。"
        "然而，在1870—1871年的普法战争期间，法国军用气球再次升空。美国南北战争期间，气球也被用作武器。"
    ),
    (8, 0): (
        "我试图论证气球与无人机之间存在“家族相似性”，但这可能会面临一种反驳：气球并非遥感技术。气球是载人飞行器，而无人机则没有机组人员；无人机的另一个名称正是无人驾驶航空器。正如我在第一章关于无人机作为赛博格的讨论中所指出的，“无人性”这一概念很复杂，因为人类与无人机紧密缠结，无法从无人机装配体中分离出来。"
        "气球飞行史也支持这一观点，因为它揭示了人类与机器之间类似的纠缠关系。最初的气球完全没有机组人员。后来，为了测试人类在大气高空生存的可行性，人们把动物（绵羊和鸡）送上天空。只有在动物毫发无损地返回之后，人类才开始升空。许多早期载人气球并不能完全自由移动，而是用绳索系留在地面上。"
        "这一过程从无机组人员的飞行，经过“载兽”、系留和非系留飞行，最终发展为载人飞行，动摇了“无人性”这一概念本身。在载人气球中，人类飞行员坐在吊篮里，通过拉动绳索和控制压舱物操纵飞行器；在无人机中，飞行员坐在地面上，通过点击鼠标控制无人机。“无人性”因此关乎相对距离和系留方式，而不是物理在场与否。"
        "不过，气球与无人机之间最重要的相似之处在于二者共同具有一种将世界扁平化的高空视角，这正是下一节要讨论的内容。无人机与气球都是使世界扁平化的感知技术。为说明这一点，我将最早的气球飞行记录之一作为分析核心，因为它为理解扁平化的美学及其情动性提供了宝贵洞见。"
    ),
    (10, 0): (
        "托马斯·鲍德温（Thomas Baldwin）的《航空百科》（1786）生动记述了气球飞行的早期历史。该书收录了他在英国切斯特上空进行一日空中旅行时记录的笔记，其中包括大气数据、气球重量及其材料等详细技术信息。某种意义上，这部作品像是为有志成为气球飞行家的人准备的手册，提供知识与指导。"
        "但《航空百科》远不止于此：它还是一篇哲学与诗性散文，讨论气球飞行作为一种新型观看技术所具有的美学与情动性。气球形成的扁平化高空视角是其中的重要部分。鲍德温这样描述这种前所未有的新视角：“物体无穷无尽、细小、清晰且彼此分离，尽管看似处在同一个平面或层次上，却无需改变视点便同时映入眼帘，令人惊异而陶醉。”"
        "这句话中的关键词是“平面”和“层次”：它们将气球视角描述为一种俯视地球而非朝向地平线的扁平化凝视。地球上原本突出的地貌特征，如丘陵、悬崖、森林和村庄，形成了缺乏空间深度的抽象图案；各个客体显得孤立而断裂。这种效果也存在于鲍德温书中那些引人注目的插图里。"
    ),
    (13, 0): (
        "我们应当如何解读鲍德温的地球图像？它们为何与无人机的感知模式相关？在我看来，鲍德温的插图是无人机图像的原型：作为遥感技术的产物，它们不仅将世界扁平化，也包含非人类视觉的要素。"
        "鲍德温的图画暗示了一种“操作性视觉”。这种视觉形式见于“使用型图像”（Gebrauchsbilder），如图表、图解和地图。西比莉·克莱默（Sybille Krämer）将这种操作性视觉的根源归于扁平性。需要指出的是，操作性视觉在功能上并不等同于军事语境中的操作性图像。正如我在第一章中所概述的，操作性图像常在无人机战争语境下被讨论；无人机的信息系统依据机器可读数据参与作出生死判断。操作性图像的扁平化效应因此与权力、侵略和支配联系在一起。"
        "但这里并非如此。鲍德温的图像与这些操作性图像有关：它们都来自空中，都具有扁平化特征，也都具有操作性视觉。但它们利用这种操作性美学来表达相互联结，而不是实施暴力。它们提供了另一种视角，使我们能够以更具创造性、审美性并诉诸感知的方式理解空中图像的操作性。正如劳拉·库尔干（Laura Kurgan）所说，不能把卫星拍摄的操作性扁平图像简化为军事暴力语境；我们也必须处理这类图像的审美面向。"
    ),
    (14, 0): (
        "我希望在鲍德温的气球图像以及当代艺术家关于无人机的美学想象中，追踪带有操作性视觉的空中图像的这一审美面向。尽管鲍德温的扁平化图像因呈现地形和地理元素而具有制图学特征，但它们并不完全属于地图。云层营造出宁静、平和的氛围；标示水道和山谷的线条具有装饰性和阿拉伯式的美感；森林所用的多种绿色让人联想到风景画。"
        "斯维特拉娜·阿尔珀斯（Svetlana Alpers）通过分析将制图元素与风景画美学结合起来的17世纪荷兰绘画，质疑了地图与风景画之间的区分。鲍德温的图像将制图性与审美性结合起来；同样，无人机图像也可以被描述为具有装饰性，并使人联想到非具象艺术。鲍德温在热情的文字描写中强调了观看的美感。他把乘气球上升描述为“奇异的”和一场“奇观”，升入天空后又“为之欢呼”。"
        "请看他对这一新视角所带来的含泪惊异的描述：“一滴纯粹喜悦的泪水在他的眼中闪过！那是纯粹而精妙的喜悦与狂喜；俯视之下，艺术与自然的作品已经发生了出人意料的变化，在新的视角下被收缩到一掌之内，缩小到难以置信的程度。”置身空中是一种“真正的崇高”体验；鲍德温在描述旅程时经常使用“崇高”这一审美范畴。他由此采取了埃德蒙·伯克（Edmund Burke）在《论崇高与美的观念之根源的哲学探讨》（Philosophical Enquiry into the Origin of Our Ideas of the Sublime and Beautiful，1757）中提出的崇高心理学路径。与伯克一样，鲍德温把崇高归因于情动、情感和身体感受，指出气球上的空中视角“超出可信性的边界”，超出“语言的力量”。他多次强调，空中视角会在身体层面触发“目眩的喜悦”，并带来“压倒性的体验”。"
    ),
    (16, 0): (
        "气球的非人类凝视高度依赖自然力量，具体而言，依赖大气条件、气候和气象状况。请注意鲍德温在《云端之上的气球景观》中画出的纤细黑色卷线，它标示了气球变化无常的路线。这一路线及其提供的感知可能性并不完全受人类飞行员控制和操纵；风、温度与气压都参与其中。因此，作为观看技术，气球并不以人为控制中心，而是由自然引导。"
        "在这一点上，鲍德温提出了一个引人注目的说法：气球传递的“风景”因“自然之笔”而得到增强。从媒介史角度看，这一说法很有意思，因为“自然之笔”这一概念主要因威廉·亨利·福克斯·塔尔博特（William Henry Fox Talbot）于1844—1846年出版的同名插图摄影书而广为人知。塔尔博特把摄影描述为一种光生绘画艺术，在这一过程中，创造美的主要是自然而非艺术家。乔安娜·齐林斯卡（Joanna Zylinska）将塔尔博特视为摄影中非人类因素的早期展示者，因为他把摄影探索为一种主要由阳光决定的技术。"
    ),
}


REVISION_OVERRIDES: dict[tuple[int, int], str] = {
    # ST row 38: the previous target cell repeated the preceding Bedoya/stewardship
    # paragraph and omitted this entire paragraph.
    (38, 0): (
        "这种关护与守护责任，与行星共同体的理念密切相连。无人机与气球表明，这片土地不属于任何一个国家或部落。"
        "Postcommodity明确表示，他们以跨原住民的方式开展工作，关注所有部落的历史，而不专指某些特定部落。"
        "气球围栏将当地居民缝合在一起，使他们彼此联结，并表明他们在保持差异的同时彼此相属。"
        "由此，这一装置中由无人机与气球促成的共同体愿景，使人联想到Rosi Braidotti关于社会联结的愿景，"
        "可以表述为：“We-Are-(All)-In-This-Together-But-We-Are-Not-All-One-And-The-Same”（“我们（所有人）都身处其中，但我们并非全然相同，也不是同一个整体”）。"
        "同样，在Postcommodity所构想的行星共同体中，没有人因国家、部落或国际边界而被排除在外：所有人都属于地球这一共同的行星家园。"
        "再次，与豪特里夫的情形一样，这一共同体愿景并不将共同体塑造成总体化、同质化的整体；相反，正如南希的哲学所强调的，它是从差异、独异性与多样性出发来思考的。"
    ),
    # ST row 97: the current cell contained a broken quote and lost the sentence
    # boundary after “future aspect.”
    (97, 0): (
        "然而，阿姆斯特朗并不只是在当地艺术展览中展示这些三维可视化图像。"
        "他的团队与当地人密切合作，对这些三维图像进行加工和改造，将其转化为展示村庄未来可能面貌的投影。"
        "这种未来维度之所以出现，是因为处理无人机图像的软件会投射出房屋和基础设施的新规划与新模型。"
        "在这里，无人机再次遵循一种体积性感知过程；它与人类行动元共同在地面上构建共同体愿景。"
        "随后，软件可以设计建筑基础设施，并模拟其可能呈现的样貌。"
        "因此，这一共同体愿景并非源自垂直测绘；相反，三维无人机图像使地面发生倾斜，并将视角倒转。"
        "无人机数据库为村庄的想象与重塑提供了一个实验性且富有创造力的矩阵。"
        "但阿姆斯特朗的工作并未止步于此：当地共同体与艺术家实际上建造了其中一些富有创造力的三维模型，"
        "用黏土、回收玻璃和轮胎等可持续、低成本材料，建造了阿姆斯特朗所谓的“后自然房屋”。"
    ),
}


HEADINGS = {
    0: "第四章：扁平化感知与行星共同体",
    4: "作为早期无人机的空中气球",
    9: "利用空中气球将世界扁平化：托马斯·鲍德温的《航空百科》（Airopaidia，1786）",
    19: "无人机的扁平化视角",
    26: "托马斯·范·豪特里夫的《分裂》（2018）",
    32: "艺术家团体 Postcommodity 的《驱逐围栏》（2015）",
    41: "伊格纳西奥·阿科斯塔的《无人机与鼓》（2018）",
    48: "行星共同体",
    56: "无人机文化、精准农业与多光谱成像",
    70: "巨型农业与作为审美媒介的无人机",
    80: "地球历史的深层感知：无人机与《绵延生态学》",
    86: "景观设计中无人机的体积性感知",
    92: "无人机艺术、体积性感知与后自然共同体",
    99: "后碳共同体",
}


CAPTIONS = {
    "4.1": (
        "William Angus after Thomas Baldwin, “A Balloon Prospect from Above the Clouds,” from Airopaidia, Containing the Narrative of a Balloon Excursion from Chester, the Eighth of September, 1785, Chester: J. Fletcher, 1786, hand-colored etching. © Yale Center for British Art, Paul Mellon Collection.",
        "图4.1 威廉·安格斯仿托马斯·鲍德温，《云端之上的气球景观》，载于《航空百科》（收录1785年9月8日从切斯特出发的气球飞行叙事），切斯特：J. Fletcher，1786年，手工着色蚀刻版画。© 耶鲁英国艺术中心，保罗·梅隆收藏。",
    ),
    "4.2": (
        "William Angus after Thomas Baldwin, “A View from the Balloon at Its Greatest Elevation,” from 1785, Chester: J. Fletcher, 1786, hand-colored etching. © Yale Center for British Art, Paul Mellon Collection.",
        "图4.2 威廉·安格斯仿托马斯·鲍德温，《气球在最高海拔处的景观》（1785），切斯特：J. Fletcher，1786年，手工着色蚀刻版画。© 耶鲁英国艺术中心，保罗·梅隆收藏。",
    ),
    "4.3": (
        "Odilon Redon, The Eye Like a Strange Balloon Mounts Toward Infinity (1882). © The Museum of Modern Art, New York Scala, Florence.",
        "图4.3 奥迪隆·雷东，《The Eye Like a Strange Balloon Mounts Toward Infinity》（1882）。© 纽约现代艺术博物馆，Scala，佛罗伦萨。",
    ),
    "4.4": (
        "Video still from Trevor Paglen, Drone Vision (2010). Archival pigment prints. 16 × 20 in. © Trevor Paglen. Courtesy of the Artist, Metro Pictures, New York and Altman Siegel, San Francisco.",
        "图4.4 特雷弗·帕格伦，《无人机视觉》（Drone Vision，2010）视频剧照。档案颜料打印，16 × 20英寸。© Trevor Paglen。经艺术家、纽约 Metro Pictures 画廊及旧金山 Altman Siegel 画廊惠允。",
    ),
    "4.5": (
        "Still from Tomas van Houtryve, Divided (2018). © Tomas van Houtryve.",
        "图4.5 托马斯·范·豪特里夫，《分裂》（Divided，2018）剧照。© Tomas van Houtryve。",
    ),
    "4.6": (
        "Through the Repellent Fence. Photo © Micheal Lundgreen, courtesy of Postcommodity and Bockley Gallery.",
        "图4.6 《穿越驱逐围栏》。摄影：Micheal Lundgreen。经 Postcommodity 和 Bockley Gallery 惠允。",
    ),
    "4.7": (
        "Still from Ignacio Acosta’s Litte ja Goabddá (Drones and Drums) (2018). © Ignacio Acosta.",
        "图4.7 伊格纳西奥·阿科斯塔《Litte ja Goabddá》（Drones and Drums，2018）剧照。© Ignacio Acosta。",
    ),
    "4.8": (
        "Drum from Lule Sami area, with hunting motifs (ca. 1673). Public domain.",
        "图4.8 吕勒萨米地区带狩猎图案的鼓（约1673年）。公有领域。",
    ),
    "4.9": (
        "Still from Ignacio Acosta’s Litte ja Goabddá (Drones and Drums) (2018). © Ignacio Acosta.",
        "图4.9 伊格纳西奥·阿科斯塔《Litte ja Goabddá》（Drones and Drums，2018）剧照。© Ignacio Acosta。",
    ),
    "5.1": (
        "Center-pivot irrigation systems on the edge of the Kubuqi Desert. Image by George Steinmetz. © George Steinmetz.",
        "图5.1 库布齐沙漠边缘的中心支轴式灌溉系统。图片：George Steinmetz。© George Steinmetz。",
    ),
    "5.2": (
        "Michele Barker and Anna Munster, Ecologies of Duration (2020). © Michele Barker and Anna Munster.",
        "图5.2 Michele Barker 与 Anna Munster，《绵延生态学》（Ecologies of Duration，2020）。© Michele Barker 与 Anna Munster。",
    ),
    "5.3": (
        "Seven Stage Futures (2017). Preparation for the Meraka: aerial view of Mokoena and Ellen’s residence, Caleb Motshabi, South Africa. © Photographer: iFlair.",
        "图5.3 《Seven Stage Futures》（2017）。Meraka 的准备工作：南非 Caleb Motshabi，Mokoena 与 Ellen 住所的航拍图。© 摄影：iFlair。",
    ),
}


# Footnote references are recovered from the PDF-derived bilingual source. The supplied
# ST DOCX omits these inline references, so they are restored here without adding footnote
# prose that is not present in the requested appendix.
FOOTNOTES: dict[tuple[str, int, int | None], list[tuple[int, str]]] = {
    ("4", 1, None): [(1, "Best."), (2, "world.")],
    ("4", 2, None): [(3, "literature."), (4, "essentialisms."), (5, "loan.”"), (6, "homogenization."), (7, "well.")],
    ("4", 5, None): [(8, "earth.")],
    ("4", 6, None): [(9, "today."), (10, "figurines."), (11, "passion."), (12, "technology.")],
    ("4", 7, None): [(13, "English.”")],
    ("4", 10, None): [(14, "ballooning."), (15, "enchanted.”")],
    ("4", 13, None): [(16, "visuality.”"), (17, "Gebrauchsbilder)"), (18, "flatness."), (19, "context."), (20, "domination."), (21, "images.")],
    ("4", 14, None): [(22, "aesthetics."), (23, "art."), (24, "joy.”"), (25, "credibility.”"), (26, "sublime,”"), (27, "(1757)."), (28, "language.”"), (29, "delight”"), (30, "experience.”")],
    ("4", 16, None): [(31, "nature.”"), (32, "sunlight.")],
    ("4", 17, None): [(33, "magnitude.”")],
    ("4", 18, None): [(34, "superorganism”"), (35, "dominated.")],
    ("4", 20, None): [(36, "discourse.")],
    ("4", 21, None): [(37, "subject."), (38, "key."), (39, "plane.”"), (40, "analysts."), (41, "data.")],
    ("4", 23, None): [(42, "planetary."), (43, "visualization”"), (44, "battlefield.”")],
    ("4", 25, None): [(45, "alterity."), (46, "critique.")],
    ("4", 27, None): [(47, "life.”"), (48, "territory.")],
    ("4", 28, None): [(49, "continuity.")],
    ("4", 29, None): [(50, "images.”")],
    ("4", 30, None): [(51, "one."), (52, "image”")],
    ("4", 31, None): [(53, "airplanes.")],
    ("4", 33, None): [(54, "Sonora."), (55, "surveilled by the US Border Patrol.")],
    ("4", 35, None): [(56, "Fence."), (57, "installation.")],
    ("4", 36, None): [(58, "performance,")],
    ("4", 37, None): [(59, "asserts.”"), (60, "ways.")],
    ("4", 38, None): [(61, "same.”")],
    ("4", 39, None): [(62, "project."), (63, "book.")],
    ("4", 40, None): [(64, "colonialism.")],
    ("4", 42, None): [(65, "world.")],
    ("4", 43, None): [(66, "signs.")],
    ("4", 44, None): [(67, "earth.")],
    ("4", 45, None): [(68, "protest."), (69, "chapter 6."), (70, "protesters.")],
    ("4", 47, None): [(71, "context.”")],
    ("5", 51, None): [(1, "contexts.”"), (2, "sensing”"), (3, "move.”")],
    ("5", 52, None): [(4, "fields.")],
    ("5", 54, None): [(5, "earth."), (6, "environment."), (7, "nature.")],
    ("5", 57, None): [(8, "population."), (9, "agriculture.")],
    ("5", 58, None): [(10, "revolution.")],
    ("5", 59, None): [(11, "yields.")],
    ("5", 60, None): [(12, "environment,”"), (13, "techno-geographies.”")],
    ("5", 61, None): [(14, "ultraviolet)."), (15, "warfare."), (16, "20 nm).”")],
    ("5", 62, None): [(17, "2020.")],
    ("5", 65, None): [(18, "monitoring."), (19, "platforms.”"), (20, "wildfires.”")],
    ("5", 66, None): [(21, "drone."), (22, "evaluation.”"), (23, "terminated.")],
    ("5", 67, None): [(24, "origins.”")],
    ("5", 71, 0): [(25, "works.")],
    ("5", 71, 1): [(26, "shipyards."), (27, "excellent example."), (28, "drawer.”")],
    ("5", 72, None): [(29, "continents."), (30, "nature.”")],
    ("5", 75, None): [(31, "atmospheres.")],
    ("5", 76, None): [(32, "oil fields.")],
    ("5", 77, None): [(33, "politics.”")],
    ("5", 81, None): [(34, "conservation."), (35, "entwinements.")],
    ("5", 82, None): [(36, "earth."), (37, "2019.")],
    ("5", 84, None): [(38, "temporalities.”"), (39, "sensing.")],
    ("5", 87, None): [(40, "gathering."), (41, "wholeness."), (42, "space."), (43, "meteorology.")],
    ("5", 88, None): [(44, "spaces.”"), (45, "satellites.")],
    ("5", 89, None): [(46, "sensing."), (47, "sites.”"), (48, "landscapes.”"), (49, "map.”")],
    ("5", 90, None): [(50, "changes."), (51, "futures."), (52, "1867.")],
    ("5", 91, None): [(53, "here."), (54, "Anthropocene.")],
    ("5", 93, None): [(55, "atmosphere.”")],
    ("5", 94, None): [(56, "ecologies.”"), (57, "practice”"), (58, "subject.")],
    ("5", 95, None): [(59, "Agent."), (60, "festivals.”")],
    ("5", 96, None): [(61, "visualizations.")],
    ("5", 98, None): [(62, "futures.”"), (63, "restoring.")],
    ("5", 100, None): [(64, "Anthropocene."), (65, "futures.”"), (66, "Neganthropocene”")],
}


# The previous appendix placed several target-side references by paragraph-length
# ratio. These anchors correct only the cases where that heuristic put a marker after
# the wrong Chinese sentence or inside a quotation.
TARGET_FOOTNOTE_ANCHORS: dict[tuple[str, int, int | None], dict[int, str]] = {
    ("4", 2, None): {
        5: "然而我们只是借居于此。”",
    },
    ("4", 13, None): {
        16: "鲍德温的图画暗示了一种“操作性视觉”。",
        17: "“使用型图像”（Gebrauchsbilder）",
        18: "操作性视觉的根源归于扁平性。",
        19: "军事语境中的操作性图像。",
        20: "参与作出生死判断。",
        21: "这类图像的审美面向。",
    },
    ("4", 14, None): {
        22: "质疑了地图与风景画之间的区分。",
        23: "并使人联想到非具象艺术。",
        24: "升入天空后又“为之欢呼”。",
        25: "缩小到难以置信的程度。”",
        26: "“真正的崇高”体验",
        27: "1757）",
        28: "超出“语言的力量”。",
        29: "“目眩的喜悦”",
        30: "“压倒性的体验”。",
    },
    ("4", 16, None): {
        31: "气球传递的“风景”因“自然之笔”而得到增强。",
        32: "一种主要由阳光决定的技术。",
    },
    ("4", 18, None): {
        34: "“超级有机体”",
        35: "以人类与非人类之间的连接性为定义特征；",
    },
    ("4", 30, None): {
        51: "且属于每一个人。",
        52: "“超理性图像”",
    },
    ("4", 38, None): {
        61: "差异、独异性与多样性出发来思考的。",
    },
    ("5", 60, None): {
        12: "“环境的可编程性”",
        13: "“技术地理景观”。",
    },
    ("5", 82, None): {
        36: "该作品将无人机感知、地质时间与地球观测相联结。",
        37: "2019年在芬兰基尔皮斯耶尔维的创作。",
    },
    ("5", 100, None): {
        65: "所称的“后碳未来”",
        66: "“负人类世（Neganthropocene）”",
    },
}


def superscript(number: int) -> str:
    chars = str(number).translate(str.maketrans("0123456789", "⁰¹²³⁴⁵⁶⁷⁸⁹"))
    # Keep multi-digit inline references together at a page/line boundary.
    return "\u2060".join(chars)


def find_anchor(text: str, anchor: str, start: int = 0) -> int:
    variants = [anchor, anchor.replace("’", "'"), anchor.replace("”", '"').replace("“", '"')]
    for item in variants:
        pos = text.find(item, start)
        if pos >= 0:
            return pos + len(item)
    # A shorter suffix keeps the marker close to the original citation if a punctuation
    # normalization changed the exact anchor.
    words = re.findall(r"[A-Za-z0-9’']+", anchor)
    for count in range(min(5, len(words)), 1, -1):
        suffix = " ".join(words[-count:])
        pos = text.find(suffix, start)
        if pos >= 0:
            return pos + len(suffix)
    return -1


def insert_at(text: str, positions: list[tuple[int, str]]) -> str:
    for pos, marker in sorted(positions, reverse=True):
        text = text[:pos] + marker + text[pos:]
    return text


def add_footnotes(source: str, target: str, key: tuple[str, int, int | None]) -> tuple[str, str]:
    items = FOOTNOTES.get(key, [])
    if not items:
        return source, target
    # A small number of ST rows already contain the PDF-extracted decimal marker;
    # remove it before restoring the authoritative superscript form.
    for number, anchor in items:
        for variant in [anchor, anchor.replace("’", "'").replace("”", '"').replace("“", '"')]:
            pattern = re.escape(variant) + r"\s*" + str(number) + r"\b"
            source = re.sub(pattern, variant, source, count=1)
    source_positions: list[tuple[int, str]] = []
    target_positions: list[tuple[int, str]] = []
    last = 0
    for number, anchor in items:
        marker = superscript(number)
        source_pos = find_anchor(source, anchor, last)
        if source_pos < 0:
            source_pos = len(source)
        else:
            last = source_pos
        source_positions.append((source_pos, marker))
        target_anchor = TARGET_FOOTNOTE_ANCHORS.get(key, {}).get(number)
        if target_anchor:
            target_pos = find_anchor(target, target_anchor)
        else:
            ratio = min(1.0, source_pos / max(1, len(source)))
            approx = int(ratio * len(target))
            # Put the marker after the nearest sentence-ending punctuation.
            candidates = [m.end() for m in re.finditer(r"[。！？；]", target) if m.end() >= approx]
            target_pos = candidates[0] if candidates else len(target)
        if target_pos < len(target) and target[target_pos] in "”》":
            target_pos += 1
        if target_pos < len(target) and target[target_pos].isdigit():
            marker += " "
        if target_positions and abs(target_pos - target_positions[-1][0]) < 2:
            marker = "\u2009" + marker
        target_positions.append((target_pos, marker))
    return insert_at(source, source_positions), insert_at(target, target_positions)


def join_paragraphs(paragraphs: list[str], keep_lines: bool = False) -> str | list[str]:
    values = [p.strip() for p in paragraphs if p.strip()]
    if keep_lines:
        return values
    return " ".join(values)


def target_for(row_index: int, paragraph_index: int, text: str) -> str:
    override = REVISION_OVERRIDES.get((row_index, paragraph_index))
    if override is None:
        override = TRANSLATION_OVERRIDES.get((row_index, paragraph_index))
    if override is not None:
        return override
    result = clean_translation(text)
    if row_index == 66:
        result = re.sub(r"\s+大地\s*$", "", result)
    if row_index in {1, 22, 28, 35, 49}:
        result = result.replace("民用无人机", "业余无人机")
    if row_index == 51:
        result = result.replace("民用无人机", "消费级无人机")
    if row_index in {42, 93, 96}:
        result = result.replace("民用无人机", "商用无人机")
    return result


def source_for(row_index: int, paragraph_index: int, text: str) -> str:
    return clean_source(text)


def build_units() -> list[dict]:
    units: list[dict] = [
        {
            "source": "Part III: The Earth",
            "target": "第三部分：地球",
            "kind": "part",
            "key": ("part", -1, None),
        }
    ]
    for i in range(len(ST)):
        if i == 51:
            units.append(
                {
                    "source": "5 Volumetric Sensing and Postcarbon Communities",
                    "target": "第五章：体积性感知与后碳共同体",
                    "kind": "chapter",
                    "key": ("chapter5", -1, None),
                }
            )
        if i == 72:
            # The source DOCX repeats the opening Steinmetz paragraph from row 71;
            # the PDF shows row 72 begins with the visual chronicle.
            st_joined = join_paragraphs(ST[i])
            tt_joined = join_paragraphs(TT[i])
            st_marker = "Steinmetz’s visual chronicle"
            tt_marker = "斯坦梅茨的视觉编年史"
            st_start = st_joined.find(st_marker)
            if st_start < 0:
                st_marker = "Steinmetz's visual chronicle"
                st_start = st_joined.find(st_marker)
            st_text = st_joined[st_start:] if st_start >= 0 else st_joined
            tt_text = tt_joined[tt_joined.find(tt_marker) :] if tt_marker in tt_joined else tt_joined
            source = source_for(i, 0, st_text)
            target = target_for(i, 0, tt_text)
            source, target = add_footnotes(source, target, ("5", i, None))
            units.append({"source": source, "target": target, "kind": "body", "key": ("5", i, None)})
            continue
        if i == 71:
            for pi, (st_p, tt_p) in enumerate(zip(ST[i], TT[i])):
                source = source_for(i, pi, st_p)
                target = target_for(i, pi, tt_p)
                source, target = add_footnotes(source, target, ("5", i, pi))
                units.append({"source": source, "target": target, "kind": "body", "key": ("5", i, pi)})
            continue
        if i == 59:
            st_lines = [source_for(i, pi, p) for pi, p in enumerate(ST[i]) if p.strip()]
            tt_lines = [target_for(i, pi, p) for pi, p in enumerate(TT[i]) if p.strip()]
            if len(st_lines) >= 4 and len(tt_lines) >= 4:
                st_lines[1:] = [f"• {x}" if not x.startswith("•") else x for x in st_lines[1:]]
                tt_lines[1:] = [f"• {x}" if not x.startswith("•") else x for x in tt_lines[1:]]
            source = "\n".join(st_lines)
            target = "\n".join(tt_lines)
            source, target = add_footnotes(source, target, ("5", i, None))
            units.append({"source": source, "target": target, "kind": "list", "key": ("5", i, None)})
            continue
        source_paras = [source_for(i, pi, p) for pi, p in enumerate(ST[i]) if p.strip()]
        # TT row 38 contains a repeated translation of the preceding Bedoya
        # paragraph plus an extra explanatory paragraph.  ST row 38 is one
        # paragraph; use the repaired paragraph override as the whole target
        # cell so the alignment remains one complete paragraph to one complete
        # paragraph.
        if i == 38:
            target_paras = [target_for(i, 0, TT[i][0])]
        else:
            target_paras = [target_for(i, pi, p) for pi, p in enumerate(TT[i]) if p.strip()]
        source = " ".join(source_paras)
        target = " ".join(target_paras)
        source, target = add_footnotes(source, target, ("4" if i < 51 else "5", i, None))
        kind = "heading" if i in HEADINGS else "body"
        if i in HEADINGS:
            target = HEADINGS[i]
        units.append({"source": source, "target": target, "kind": kind, "key": ("4" if i < 51 else "5", i, None)})
    return units


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 100, bottom: int = 90, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, width: Cm) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, east_asia: str, size: float, bold: bool = False, italic: bool = False, color: str | None = None) -> None:
    run.font.name = east_asia
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:ascii"), east_asia)
    r_fonts.set(qn("w:hAnsi"), east_asia)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), east_asia)
    r_fonts.set(qn("w:hint"), "eastAsia")
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:eastAsia"), "zh-CN")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_cell_content(cell, text: str, lang: str, kind: str, font_size_override: float | None = None) -> None:
    cell.text = ""
    paragraphs = text.split("\n") if "\n" in text else [text]
    for idx, line in enumerate(paragraphs):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(3 if kind not in {"part", "chapter", "heading"} else 1)
        p.paragraph_format.line_spacing = 1.04
        if kind in {"part", "chapter", "heading"}:
            p.paragraph_format.keep_with_next = True
        if line.startswith("• "):
            p.paragraph_format.left_indent = Cm(0.2)
            p.paragraph_format.first_line_indent = Cm(-0.2)
        run = p.add_run(line)
        if lang == "source":
            size = font_size_override if font_size_override is not None else (9.2 if kind == "body" else 9.5)
            set_run_font(run, "Times New Roman", size, bold=kind in {"part", "chapter", "heading"}, italic=kind == "caption", color="555555" if kind == "caption" else None)
        else:
            set_run_font(run, "Noto Sans CJK SC", 9.2 if kind == "body" else 9.5, bold=kind in {"part", "chapter", "heading"}, italic=kind == "caption", color="555555" if kind == "caption" else None)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_caption_units(units: list[dict], after_row: int, caption_ids: list[str]) -> None:
    # Insert in reverse order to keep the source order when multiple captions share a point.
    pos = max(i for i, unit in enumerate(units) if unit.get("row_index") == after_row) + 1
    for caption_id in caption_ids:
        source, target = CAPTIONS[caption_id]
        units.insert(pos, {"source": source, "target": target, "kind": "caption", "key": ("caption", caption_id, None)})
        pos += 1


def build_doc() -> None:
    units = build_units()
    # Record the source row for caption placement. Part/chapter inserts have no row index.
    row_cursor = -1
    for unit in units:
        if unit["key"][0] in {"4", "5"} and isinstance(unit["key"][1], int):
            unit["row_index"] = unit["key"][1]
    add_caption_units(units, 11, ["4.1"])
    add_caption_units(units, 12, ["4.2"])
    add_caption_units(units, 15, ["4.3"])
    add_caption_units(units, 21, ["4.4"])
    add_caption_units(units, 29, ["4.5"])
    add_caption_units(units, 36, ["4.6"])
    add_caption_units(units, 43, ["4.7", "4.8"])
    add_caption_units(units, 44, ["4.9"])
    add_caption_units(units, 72, ["5.1"])
    add_caption_units(units, 84, ["5.2"])
    add_caption_units(units, 98, ["5.3"])

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Noto Sans CJK SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    normal.font.size = Pt(9.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run("附录一  无人机感知与共同体：原文与译文（终审修订版）")
    set_run_font(title_run, "Noto Sans CJK SC", 14, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle_run = subtitle.add_run("Part III “The Earth”")
    set_run_font(subtitle_run, "Times New Roman", 10, italic=True, color="666666")

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.style = "Table Grid"
    width = Cm(13.35)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, label in enumerate(("原文（English Source）", "译文（中文 Translation）")):
        cell = header.cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_width(cell, width)
        set_cell_margins(cell, 100, 130, 100, 130)
        set_cell_shading(cell, "D9E2F3")
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_run_font(run, "Times New Roman" if idx == 0 else "Noto Sans CJK SC", 9.5, bold=True)

    for unit in units:
        row = table.add_row()
        kind = unit["kind"]
        for idx, (text, lang) in enumerate(((unit["source"], "source"), (unit["target"], "target"))):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_width(cell, width)
            set_cell_margins(cell, 95 if kind == "body" else 80, 120, 95 if kind == "body" else 80, 120)
            if kind in {"part", "chapter"}:
                set_cell_shading(cell, "B4C7E7")
            elif kind == "heading":
                set_cell_shading(cell, "EAF0F8")
            elif kind == "caption":
                set_cell_shading(cell, "F7F7F7")
            # The long Baldwin paragraph ends with superscript footnote 30; a slightly
            # smaller source-only size keeps that marker on the same line as its anchor.
            source_size = 8.9 if lang == "source" and unit["key"] == ("4", 14, None) else None
            add_cell_content(cell, text, lang, kind, source_size)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("附录一  ·  ")
    set_run_font(fr, "Noto Sans CJK SC", 8, color="888888")
    add_page_field(footer)

    doc.core_properties.title = "附录一：无人机感知与共同体——原文与译文（终审修订版）"
    doc.core_properties.subject = "The Sensorium of the Drone and Communities, Part III: The Earth"
    doc.core_properties.author = ""
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")
    print(f"Units: {len(units)}; table rows including header: {len(table.rows)}")


if __name__ == "__main__":
    build_doc()
