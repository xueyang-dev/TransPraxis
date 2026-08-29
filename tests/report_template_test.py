"""Template contract, routing, dependency and renderer checks."""
from io import BytesIO
from typing import get_type_hints
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

import core
from transpraxis import (academic_evidence, academic_quality, academic_validator,
                         academic_writer, final_docx, report_template,
                         thesis_constraints)


def test_academic_writer_public_type_hints_resolve():
    assert get_type_hints(academic_writer.build_report_artifact)["outline"]


def _template_bytes():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1.25)
    section.header.paragraphs[0].text = "固定页眉"
    section.footer.paragraphs[0].text = "固定页脚"
    document.add_paragraph("封面固定文字")
    document.add_heading("摘要", level=1)
    document.add_heading("1 项目概况", level=1)
    document.add_heading("1.1 项目背景", level=2)
    document.add_paragraph("模板占位内容")
    document.add_heading("2 理论框架", level=1)
    document.add_heading("2.1 理论依据", level=2)
    document.add_paragraph("模板占位内容")
    document.add_heading("3 案例分析", level=1)
    document.add_heading("3.1 案例说明", level=2)
    document.add_paragraph("模板占位内容")
    document.add_heading("4 结论与反思", level=1)
    document.add_heading("参考文献", level=1)
    result = BytesIO()
    document.save(result)
    return result.getvalue()


def _mti_template_bytes():
    document = Document()
    document.add_paragraph("专业学位硕士学位论文")
    document.add_paragraph("示例大学")
    document.add_paragraph("A Report on E-C Translation of")
    document.add_paragraph("Sample University")
    document.add_paragraph("Master of Translation and Interpreting")
    document.add_paragraph("学位论文独创性声明")
    document.add_paragraph("学位论文使用授权声明")
    document.add_heading("摘  要", level=1)
    document.add_paragraph("模板摘要占位内容")
    document.add_paragraph("关键词：模板；占位")
    document.add_heading("ABSTRACT", level=1)
    document.add_paragraph("Template abstract placeholder")
    document.add_paragraph("Keywords: template; placeholder")
    toc = document.add_paragraph()
    run = toc.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])
    headings = [
        ("第一章 引言", 1),
        ("1.1 研究背景及意义", 2),
        ("1.2 研究问题", 2),
        ("1.3 报告结构", 2),
        ("第二章 《XXX》翻译项目概述", 1),
        ("2.1 项目简介", 2),
        ("2.2 翻译流程", 2),
        ("2.2.1 译前准备", 3),
        ("2.2.2 翻译过程", 3),
        ("2.2.3 译后管理", 3),
        ("第三章 《XXX》翻译项目案例分析", 1),
        ("3.1 源语文本的类型与特征", 2),
        ("3.2 翻译难点", 2),
        ("3.3 翻译策略与解决方案", 2),
        # The real template accidentally uses Heading 2 for Chapter 4.
        ("第四章 总结与反思", 2),
        ("参考文献", 1),
        ("致 谢", 1),
        ("附录一 《XXX》原文与译文", 1),
        ("附录二 主要术语对照表", 1),
    ]
    for title, level in headings:
        document.add_heading(title, level=level)
        document.add_paragraph("模板占位内容")
    result = BytesIO()
    document.save(result)
    return result.getvalue()


def _mti_contract():
    return report_template.parse_docx_template("mti-fixture.docx", _mti_template_bytes())


def _mti_markdown(project="Drone Communities"):
    return f"""## 1 引言

### 1.1 研究背景及意义

背景正文。

### 1.2 研究问题

研究问题正文。

### 1.3 报告结构

全文共四章。

## 2 《{project}》翻译项目概述

### 2.1 项目简介

项目正文。

### 2.2 翻译流程

流程正文。

#### 2.2.1 译前准备

译前正文。

#### 2.2.2 翻译过程

翻译正文。

#### 2.2.3 译后管理

译后正文。

## 3 《{project}》翻译项目案例分析

### 3.1 源语文本的类型与特征

特征正文。

### 3.2 翻译难点

难点正文。

### 3.3 翻译策略与解决方案

策略正文。

## 4 总结与反思

总结正文。
"""


def _mti_matter(contract):
    front = []
    for item in contract["document_structure"]["front_matter"]:
        row = dict(item)
        if item["role"] in {"abstract_zh", "abstract_en"}:
            row["content"] = "真实项目摘要正文。"
        if item["role"] in {"keywords_zh", "keywords_en"}:
            row["keywords"] = ["翻译实践", "案例分析"]
        front.append(row)
    return {
        "project_title": "Drone Communities",
        "front_matter": front,
        "back_matter": [
            {**item, "content": "需要用户补充。"}
            for item in contract["document_structure"]["back_matter"]
        ],
    }


def _mti_outline(contract, cases):
    constraints = thesis_constraints.build_constraints({
        "body_language": "zh-CN", "project_name": "Drone Communities",
        "report_template_contract": contract,
    })
    return {
        "template_hash": contract["template_identity"]["sha256"],
        "sections": [
            {**chapter, "cases": list(cases) if chapter["role"] == "case_analysis"
             else [], "claims": [], "research_questions": []}
            for chapter in constraints["chapters"]
        ],
    }


def _mti_written(project="Drone Communities"):
    return [
        {"section_id": "1", "title": "引言", "content":
         "### 1.1 研究背景及意义\n\n背景正文。\n\n"
         "### 1.2 研究问题\n\n研究问题正文。\n\n"
         "### 1.3 报告结构\n\n全文共四章。"},
        {"section_id": "2", "title": f"《{project}》翻译项目概述", "content":
         "### 2.1 项目简介\n\n项目正文。\n\n"
         "### 2.2 翻译流程\n\n流程正文。\n\n"
         "#### 2.2.1 译前准备\n\n译前正文。\n\n"
         "#### 2.2.2 翻译过程\n\n翻译正文。\n\n"
         "#### 2.2.3 译后管理\n\n译后正文。"},
        {"section_id": "3", "title": f"《{project}》翻译项目案例分析", "content":
         "### 3.1 源语文本的类型与特征\n\n特征正文。\n\n"
         "### 3.2 翻译难点\n\n难点正文。\n\n"
         "### 3.3 翻译策略与解决方案\n\n策略正文。"},
        {"section_id": "4", "title": "总结与反思", "content": "总结正文。"},
    ]


def _contract():
    return report_template.parse_docx_template("fixture.docx", _template_bytes())


def _outline(contract, cases=None):
    return {
        "template_hash": contract["template_identity"]["sha256"],
        "sections": [
            {"section_id": "1", "title": "项目概况", "role": "project_overview", "cases": []},
            {"section_id": "2", "title": "理论框架", "role": "theoretical_framework", "cases": []},
            {"section_id": "3", "title": "案例分析", "role": "case_analysis", "cases": cases or []},
            {"section_id": "4", "title": "结论与反思", "role": "conclusion_reflection", "cases": []},
        ],
    }


def _report_artifact(contract):
    return {
        "template_hash": contract["template_identity"]["sha256"],
        "front_matter": [
            {**item, "content": "摘要正文。"}
            for item in contract["document_structure"]["front_matter"]
        ],
        "back_matter": [
            {**item, "content": "参考文献正文。"}
            for item in contract["document_structure"]["back_matter"]
        ],
    }


def _valid_markdown():
    return """## 1 项目概况

### 1.1 项目背景

项目背景正文。

## 2 理论框架

### 2.1 理论依据

理论依据正文。

## 3 案例分析

### 3.1 案例说明

案例正文。

## 4 结论与反思

结论正文。
"""


def test_parser_is_deterministic_and_captures_structure_and_style():
    raw = _template_bytes()
    first = report_template.parse_docx_template("fixture.docx", raw)
    second = report_template.parse_docx_template("fixture.docx", raw)
    assert first["content_hash"] == second["content_hash"]
    assert first["template_identity"]["sha256"]
    assert [x["role"] for x in first["document_structure"]["chapters"]] == [
        "project_overview", "theoretical_framework", "case_analysis",
        "conclusion_reflection",
    ]
    assert first["document_structure"]["front_matter"][0]["role"] == "abstract_zh"
    assert first["document_structure"]["back_matter"][0]["role"] == "references"
    assert first["style_contract"]["headers"] == [["固定页眉"]]
    assert first["style_contract"]["footers"] == [["固定页脚"]]
    assert first["style_contract"]["sections"][0]["geometry"]["top_margin_emu"]


def test_constraints_and_outline_keep_template_chapters_and_route_cases_by_role():
    contract = _contract()
    settings = {"body_language": "zh-CN", "report_template_contract": contract}
    constraints = thesis_constraints.build_constraints(settings)
    assert constraints["structure_source"] == "parsed_template_contract"
    assert [x["section_id"] for x in constraints["chapters"]] == ["1", "2", "3", "4"]

    research_model = academic_writer.build_research_model(
        {"project_evidence": {"document_profile": {}, "statistics": {}}},
        "自动推荐理论", settings)
    selected = {"cases": [{"case_id": "seg-case-0001", "case_type": "authentic_revision"}],
                "authentic_selection_status": "sufficient_revision_cases"}
    argument = {"claims": []}

    def fake_llm(*_args, **_kwargs):
        return '{"sections":[{"section_id":"1"},{"section_id":"2"},{"section_id":"4"}]}'

    outline = academic_writer.build_academic_outline(
        research_model, argument, selected,
        {"project_evidence": {"statistics": {}}}, fake_llm, "test", "", "model")
    assert [x["section_id"] for x in outline["sections"]] == ["1", "2", "3", "4"]
    assert outline["sections"][2]["cases"] == ["seg-case-0001"]
    assert outline["sections"][3]["cases"] == []


def test_template_validator_has_independent_compliance_gate():
    contract = _contract()
    outline = _outline(contract, ["seg-case-0001"])
    valid = academic_validator.validate_template_compliance(
        _valid_markdown(), contract, outline, _report_artifact(contract),
        {"cases": [{"case_id": "seg-case-0001"}]})
    assert valid["status"] == "pass"

    invalid = academic_validator.validate_template_compliance(
        _valid_markdown().replace("## 4 结论与反思", "## 9 新增章节"),
        contract, outline, _report_artifact(contract),
        {"cases": [{"case_id": "seg-case-0001"}]})
    issue_types = {x["type"] for x in invalid["issues"]}
    assert invalid["status"] == "fail"
    assert "template_chapter_title_mismatch" in issue_types
    assert "template_chapter_order_mismatch" in issue_types

    wrong_hash = dict(_report_artifact(contract), template_hash="different")
    hashed = academic_validator.validate_template_compliance(
        _valid_markdown(), contract, outline, wrong_hash,
        {"cases": [{"case_id": "seg-case-0001"}]})
    assert "template_hash_mismatch" in {x["type"] for x in hashed["issues"]}


def test_template_case_mapping_uses_actual_case_chapter_number():
    chapter = {
        "section_id": "4", "title": "案例分析", "role": "case_analysis",
        "required_subsections": [
            {"heading_id": "4.2", "title": "翻译难点", "level": 2},
            {"heading_id": "4.3", "title": "翻译策略", "level": 2},
        ],
    }
    contract = {
        "template_identity": {"sha256": "case-chapter-four"},
        "document_structure": {"chapters": [chapter], "top_level": 1},
    }
    report = (
        "## 4 案例分析\n\n### 4.2 翻译难点\n\n"
        "#### 4.2.1 术语难点\n\n难点。\n\n"
        "### 4.3 翻译策略\n\n#### 4.3.1 术语策略\n\n策略。")
    result = academic_validator.validate_template_compliance(
        report, contract, {"template_hash": "case-chapter-four",
                           "sections": [chapter]})
    assert "template_case_mapping_mismatch" not in {
        item["type"] for item in result["issues"]}
    assert "template_extra_subsection" not in {
        item["type"] for item in result["issues"]}


def test_template_hash_invalidates_downstream_academic_artifacts():
    state = core.new_job_state("source.docx")
    contract_a = _contract()
    contract_b = report_template.parse_docx_template("other.docx", _template_bytes() + b"x")
    academic_writer.prepare_academic_inputs(
        state, "理论", {"report_template_contract": contract_a}, [])
    academic = state["academic_state"]
    for name in ("research_model", "argument_plan", "selected_cases", "outline",
                 "sections", "validation", "review", "academic_quality", "report"):
        academic["artifacts"][name] = {"content_hash": "old"}
    academic_writer.prepare_academic_inputs(
        state, "理论", {"report_template_contract": contract_b}, [])
    assert not set(academic["artifacts"]) & {
        "research_model", "argument_plan", "selected_cases", "outline", "sections",
        "validation", "review", "academic_quality", "report",
    }
    assert state["p3_done"] is False


def test_template_bytes_and_contract_survive_state_reload(tmp_path):
    old_output = core.OUTPUT_DIR
    core.OUTPUT_DIR = tmp_path
    try:
        job_id = "template-persist"
        core.save_job_state(job_id, core.new_job_state("source.docx"))
        contract = core.save_report_template(job_id, "uploaded.docx", _template_bytes())
        loaded = core.load_report_template(job_id)
        assert loaded["contract"]["content_hash"] == contract["content_hash"]
        assert loaded["metadata"]["filename"] == "uploaded.docx"
        assert (core.job_dir(job_id) / "template-contract.json").is_file()
        assert (core.job_dir(job_id) / "report-template.docx").is_file()
        state = core.load_job_state(job_id)
        assert state["report_template_contract"]["template_identity"]["sha256"]
        core.clear_report_template(job_id)
        assert core.load_report_template(job_id) is None
    finally:
        core.OUTPUT_DIR = old_output


def test_template_renderer_preserves_word_layout_and_removes_placeholders():
    raw = _template_bytes()
    contract = _contract()
    artifact = {
        "sections": [
            {"section_id": "1", "intro_content": "项目概况正文。",
             "subsections": [{"title": "项目背景", "content": "背景正文。"}]},
            {"section_id": "2", "content": "理论正文。", "subsections": []},
            {"section_id": "3", "content": "案例正文。", "subsections": []},
            {"section_id": "4", "content": "结论正文。", "subsections": []},
        ]
    }
    rendered = Document(report_template.render_report_docx(artifact, raw, contract))
    texts = [paragraph.text for paragraph in rendered.paragraphs]
    assert "模板占位内容" not in texts
    assert "固定页眉" == rendered.sections[0].header.paragraphs[0].text
    assert "固定页脚" == rendered.sections[0].footer.paragraphs[0].text
    assert rendered.sections[0].top_margin.inches == 1.25
    assert [text for text in texts if text.startswith("固定") or "正文" in text]


def test_minimal_template_to_structured_artifact_to_rendered_docx():
    raw = _template_bytes()
    contract = _contract()
    constraints = thesis_constraints.build_constraints({
        "body_language": "zh-CN", "report_template_contract": contract})
    outline = _outline(contract, ["seg-case-0001"])
    written = [
        {"section_id": "1", "title": "项目概况",
         "content": "### 1.1 项目背景\n\n背景正文。"},
        {"section_id": "2", "title": "理论框架",
         "content": "### 2.1 理论依据\n\n理论正文。"},
        {"section_id": "3", "title": "案例分析",
         "content": "### 3.1 案例说明\n\n案例正文。"},
        {"section_id": "4", "title": "结论与反思", "content": "结论正文。"},
    ]
    markdown = "\n\n".join(
        f"## {item['section_id']} {item['title']}\n\n{item['content']}"
        for item in written) + "\n"
    artifact = academic_writer.build_report_artifact(
        markdown, written, outline, constraints, {
            "front_matter": [
                {**item, "content": "摘要正文。"}
                for item in contract["document_structure"]["front_matter"]
            ],
            "back_matter": [
                {**item, "content": "参考文献正文。"}
                for item in contract["document_structure"]["back_matter"]
            ],
        })
    assert artifact["template_hash"] == contract["template_identity"]["sha256"]
    compliance = academic_validator.validate_template_compliance(
        markdown, contract, outline, artifact,
        {"cases": [{"case_id": "seg-case-0001"}]})
    assert compliance["status"] == "pass"
    rendered = Document(report_template.render_report_docx(artifact, raw, contract))
    assert "案例正文。" in [paragraph.text for paragraph in rendered.paragraphs]


def test_legacy_report_uses_generic_renderer(tmp_path):
    old_output = core.OUTPUT_DIR
    core.OUTPUT_DIR = tmp_path
    try:
        job_id = "legacy-report"
        state = core.new_job_state("source.docx")
        state.update(p3_md="# 翻译实践报告\n\n旧任务报告。", theory="理论")
        core.save_job_state(job_id, state)
        output = core.report_docx_bytes(job_id, state)
        assert output and output[:2] == b"PK"
        assert core.load_report_template(job_id) is None
    finally:
        core.OUTPUT_DIR = old_output


def test_mti_contract_captures_full_front_body_and_back_structure():
    contract = _mti_contract()
    structure = contract["document_structure"]
    assert [item["role"] for item in structure["front_matter"]] == [
        "cover_zh", "cover_en", "originality_declaration",
        "authorization_declaration", "abstract_zh", "keywords_zh",
        "abstract_en", "keywords_en", "table_of_contents",
    ]
    assert [(item["section_id"], item["role"]) for item in structure["chapters"]] == [
        ("1", "introduction"), ("2", "project_overview"),
        ("3", "case_analysis"), ("4", "conclusion_reflection"),
    ]
    assert {item["role"] for item in structure["back_matter"]} == {
        "references", "acknowledgements", "appendix"}
    assert structure["case_requirement"]["minimum_cases"] == 6
    serialized = __import__("json").dumps(contract, ensure_ascii=False)
    assert "示例大学" not in serialized
    assert "Sample University" not in serialized


def test_mti_outline_cannot_drift_or_turn_research_questions_into_chapters():
    contract = _mti_contract()
    settings = {
        "body_language": "zh-CN", "project_name": "Drone Communities",
        "report_template_contract": contract,
        "research_questions": ["真实问题一？", "真实问题二？", "真实问题三？"],
    }
    evidence = {"project_evidence": {
        "document_profile": {"genre": "学术文本"},
        "statistics": {"total_segments": 6}, "glossary": [],
    }}
    research = academic_writer.build_research_model(evidence, "有限理论", settings)
    argument = {"claims": [
        {"claim_id": f"C{i}", "research_question": f"RQ{i}",
         "planned_sections": [str(i)]} for i in range(1, 4)
    ]}
    selected = {"cases": [
        {"case_id": f"TD-{i:04d}", "case_type": "translation_decision"}
        for i in range(1, 7)
    ]}

    def drifting_llm(*_args, **_kwargs):
        return '{"sections":[{"section_id":"A","title":"RQ1"},' \
               '{"section_id":"B","title":"RQ2"},' \
               '{"section_id":"C","title":"RQ3"}]}'

    outline = academic_writer.build_academic_outline(
        research, argument, selected, evidence, drifting_llm, "test", "", "model")
    assert [item["section_id"] for item in outline["sections"]] == ["1", "2", "3", "4"]
    assert [item["role"] for item in outline["sections"]] == [
        "introduction", "project_overview", "case_analysis", "conclusion_reflection"]
    assert outline["sections"][0]["research_questions"] == ["RQ1", "RQ2", "RQ3"]
    assert outline["sections"][2]["cases"] == [f"TD-{i:04d}" for i in range(1, 7)]
    assert outline["sections"][3]["cases"] == []
    assert outline["sections"][3]["claims"] == ["C1", "C2", "C3"]


def test_mti_required_subsection_and_case_minimum_are_hard_failures():
    contract = _mti_contract()
    cases = [{"case_id": f"TD-{i:04d}"} for i in range(1, 7)]
    outline = _mti_outline(contract, [item["case_id"] for item in cases])
    written = _mti_written()
    artifact = academic_writer.build_report_artifact(
        _mti_markdown(), written, outline,
        thesis_constraints.build_constraints({
            "project_name": "Drone Communities", "report_template_contract": contract}),
        _mti_matter(contract))
    missing = academic_validator.validate_template_compliance(
        _mti_markdown().replace("#### 2.2.2 翻译过程\n\n翻译正文。\n\n", ""),
        contract, outline, artifact, {"cases": cases})
    assert missing["status"] == "fail"
    assert "template_missing_subsection" in {item["type"] for item in missing["issues"]}
    insufficient = academic_validator.validate_template_compliance(
        _mti_markdown(), contract, outline, artifact, {
            "report_case_policy": thesis_constraints.case_policy({
                "report_stage": "proposal"}),
            "cases": cases[:5],
        })
    assert insufficient["status"] == "fail"
    assert "template_case_minimum_not_met" in {
        item["type"] for item in insufficient["issues"]}
    with_conclusion_subsections = _mti_markdown().replace(
        "## 4 总结与反思\n\n总结正文。",
        "## 4 总结与反思\n\n### 4.1 经验总结\n\n总结正文。")
    dynamic = academic_validator.validate_template_compliance(
        with_conclusion_subsections, contract, outline, artifact, {"cases": cases})
    assert "template_extra_subsection" not in {
        item["type"] for item in dynamic["issues"]}


def test_public_report_hides_ids_anonymizes_universities_and_deduplicates_headings():
    visible = report_template.public_report_markdown(
        "## 3 案例分析\n## 三、案例分析\n"
        "> [SOURCE seg-job-0001]: Sample University\n"
        "finding-abc term-secret claim-C1 AQ-001 示例大学",
        {"seg-job-0001": "例[1]"})
    assert "seg-" not in visible and "finding-" not in visible
    assert "term-secret" not in visible and "claim-C1" not in visible
    assert "AQ-001" not in visible
    assert "示例大学" not in visible
    assert "Nanjing University" not in visible
    assert "XX大学" in visible and "XX University" in visible
    assert visible.count("案例分析") == 1


def test_token_expansion_and_section_assembly_remove_known_duplications():
    evidence = {"project_evidence": {
        "statistics": {"reviewed_segments": 84, "finding_count": 42,
                       "tm_reuse_count": 0},
        "glossary": [
            {"id": "term-planetary", "preferred": "行星性"},
            {"id": "term-globalism", "preferred": "全球主义"},
        ],
    }}
    expanded = academic_validator.expand_evidence_tokens(
        "84（{{STAT:reviewed_segments}}） 42（{{STAT:finding_count}}） "
        "0（{{STAT:tm_reuse_count}}） 行星性{{TERM:term-planetary}} "
        "全球主义{{TERM:term-globalism}}", evidence)
    public = report_template.public_report_markdown(expanded)
    for duplicate in ("8484", "4242", "00", "行星性行星性", "全球主义全球主义"):
        assert duplicate not in public
    cleaned = academic_writer.finalize_report_tokens(
        "基于基于项目证据。", evidence, {"cases": []}, {"sections": []})
    assert "基于基于" not in cleaned
    section = {"title": "案例分析", "required_subsections": []}
    assembled = academic_writer._ensure_section_contract(
        "\n<!--claim:C1-->\n三、案例分析\n正文。", section)
    assert "三、案例分析" not in assembled
    conclusion = academic_writer._ensure_section_contract(
        "## 4.1 回应研究问题\n### 4.1.1 具体回应\n正文。",
        {"section_id": "4", "title": "总结与反思",
         "role": "conclusion_reflection", "research_questions": ["RQ1"],
         "required_subsections": []})
    assert not any(line.startswith("## 4.1 ") for line in conclusion.splitlines())
    assert "### 4.1 回应研究问题" in conclusion
    assert "#### 4.1.1 具体回应" in conclusion
    assert "<!--rq:RQ1-->" in conclusion
    mapped = academic_writer._ensure_section_contract(
        "### 3.2 翻译难点\n#### 3.2.1 难点一\n"
        "### 3.3 翻译策略与解决方案\n#### 3.3.1 策略一\n#### 3.3.2 多余小结",
        {"section_id": "3", "title": "案例分析", "role": "case_analysis",
         "required_subsections": [
             {"heading_id": "3.2", "title": "翻译难点", "level": 2,
              "markdown_prefix": "###", "allows_dynamic_children": True},
             {"heading_id": "3.3", "title": "翻译策略与解决方案", "level": 2,
              "markdown_prefix": "###", "allows_dynamic_children": True},
         ]})
    assert "#### 3.3.2" not in mapped and "**3.3.2 多余小结**" in mapped


def test_tm_zero_does_not_support_machine_translation_absence_claim():
    report = "## 1 项目说明\n\n机器翻译未使用，译者全程依赖人工翻译。"
    validation = academic_validator.validate_academic_report(
        report,
        {"project_evidence": {"statistics": {"tm_reuse_count": 0},
                              "segments": [], "glossary": []}},
        {"research_questions": []}, {"claims": []}, {"cases": []},
        {"sections": [{"section_id": "1", "title": "项目说明", "role": "project_overview"}]})
    assert "unsupported_claim_strength" in {
        item["type"] for item in validation["issues"]}


def test_argument_plan_cannot_treat_parenthetical_term_removal_as_logic_revision():
    segment_id = "seg-job-0001"
    evidence = {"project_evidence": {
        "segments": [{"segment_id": segment_id, "source": "A long source sentence.",
                      "initial_target": "视觉体制（scopic regime）构成论证。",
                      "final_target": "视觉体制构成论证。"}],
        "statistics": {}, "document_profile": {}, "glossary": [],
    }, "candidate_cases": [{"case_id": segment_id, "segment_id": segment_id,
                              "academic_candidate_status": "eligible"}]}
    research = academic_writer.build_research_model(
        evidence, "有限理论", {"research_questions": ["如何处理长句？"]})

    def planner(*_args, **_kwargs):
        return ('{"claims":[{"claim":"源文具有复杂句法，但修订记录表明隐含逻辑衔接得到改进。",'
                '"research_question":"RQ1","project_evidence":["seg-job-0001"],'
                '"confidence":"high","reasoning":"修订改善了信息结构"}]}')

    argument = academic_writer.build_argument_plan(
        research, evidence, planner, "test", "", "model")
    claim = argument["claims"][0]
    assert "修订记录表明" not in claim["claim"]
    assert "不能证明句法、逻辑衔接或信息结构发生修订" in claim["reasoning"]
    assert claim["confidence"] == "low"


def test_duplicate_selected_case_cannot_be_packaged_as_two_examples():
    case_id = "seg-job-0001"
    report = ("## 3 案例分析\n\n* **案例一（真实修订案例 seg-job-0001）**\n"
              "* **案例二（真实修订案例 seg-job-0001）**\n")
    validation = academic_validator.validate_academic_report(
        report,
        {"project_evidence": {"statistics": {}, "glossary": [], "segments": [
            {"segment_id": case_id, "source": "source", "initial_target": "旧译",
             "final_target": "新译"}]},
         "candidate_cases": [{"case_id": case_id,
                               "academic_candidate_status": "eligible"}]},
        {"research_questions": []}, {"claims": []},
        {"cases": [{"case_id": case_id, "case_type": "authentic_revision"}]},
        {"sections": [{"section_id": "3", "title": "案例分析",
                       "role": "case_analysis", "cases": [case_id]}]})
    assert "duplicate_selected_case_presentation" in {
        item["type"] for item in validation["issues"]}


def test_visible_examples_are_bound_to_hidden_case_provenance_by_source_text():
    cases = {"cases": [
        {"case_id": "TD-0001", "case_type": "translation_decision",
         "source_segment_id": "seg-job-0001"},
        {"case_id": "TD-0002", "case_type": "translation_decision",
         "source_segment_id": "seg-job-0002"},
    ]}
    evidence = {"project_evidence": {"segments": [
        {"segment_id": "seg-job-0001", "source": "The first exact source sentence is long enough.",
         "initial_target": "第一句。", "final_target": "第一句。"},
        {"segment_id": "seg-job-0002", "source": "The second exact source sentence is also long enough.",
         "initial_target": "第二句。", "final_target": "第二句。"},
    ], "statistics": {}, "glossary": []}}
    report = ("**例[1]（翻译决策案例）**\n* **SOURCE**：\"The first exact source "
              "sentence is long enough.\"\n* **TARGET**：第一句。\n\n"
              "**例[2]（翻译决策案例）**\n* **SOURCE**：\"The second exact source "
              "sentence is also long enough.\"\n* **TARGET**：第二句。")
    normalized = academic_writer.finalize_report_tokens(
        report, evidence, cases, {"sections": []})
    assert normalized.count("<!--case:") == 2
    assert "<!--case:TD-0001-->" in normalized
    assert "**原文**：The second exact source sentence" in normalized


def test_source_quote_list_lines_do_not_trigger_body_language_mismatch():
    source_case = (
        "* **源语（SOURCE）**：According to the source, this long English quotation "
        "contains enough words to look like prose but remains an explicitly labelled "
        "source-text example for translation analysis.\n"
        "分析：该句在中文正文中作为翻译案例处理。")
    assert academic_validator._has_english_prose_paragraph(source_case) is False
    assert academic_validator._has_english_prose_paragraph(
        "* **SOURCE**: This is another long explicitly labelled source quotation "
        "that must not be mistaken for English thesis exposition in a Chinese report.") is False
    assert academic_validator._has_english_prose_paragraph(
        "This is an unlabelled English exposition paragraph with more than twelve "
        "words and therefore violates the configured Chinese body language.") is True


def test_template_renderer_is_blocked_until_generated_and_uses_template(tmp_path):
    old_output = core.OUTPUT_DIR
    core.OUTPUT_DIR = tmp_path
    try:
        job_id = "mti-gate"
        state = core.new_job_state("source.docx")
        state.update(p3_md=_mti_markdown(), p3_done=True,
                     report_status="failed_template_validation")
        core.save_job_state(job_id, state)
        contract = core.save_report_template(
            job_id, "mti.docx", _mti_template_bytes())
        outline = _mti_outline(contract, [f"TD-{i:04d}" for i in range(1, 7)])
        artifact = academic_writer.build_report_artifact(
            _mti_markdown(), _mti_written(), outline,
            thesis_constraints.build_constraints({
                "project_name": "Drone Communities",
                "report_template_contract": contract,
                "report_stage": "proposal"}), _mti_matter(contract))
        artifact.update(report_status="failed_template_validation",
                        template_compliance="fail")
        (core.job_dir(job_id) / "academic-report.json").write_text(
            __import__("json").dumps(artifact, ensure_ascii=False), encoding="utf-8")
        assert core.report_docx_bytes(job_id, state) is None
        artifact.update(report_status="generated", template_compliance="pass")
        (core.job_dir(job_id) / "academic-report.json").write_text(
            __import__("json").dumps(artifact, ensure_ascii=False), encoding="utf-8")
        output = core.report_docx_bytes(job_id, state)
        assert output and output[:2] == b"PK"
        rendered = Document(BytesIO(output))
        all_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
        assert "XX大学" in all_text and "示例大学" not in all_text
        assert "模板占位内容" not in all_text
        assert all(title in all_text for title in (
            "引言", "翻译项目概述", "翻译项目案例分析", "总结与反思",
            "参考文献", "致 谢", "附录一"))
        with ZipFile(BytesIO(_mti_template_bytes())) as source_zip, \
                ZipFile(BytesIO(output)) as rendered_zip:
            source_parts = set(source_zip.namelist())
            rendered_parts = set(rendered_zip.namelist())
            assert "TOC" in rendered_zip.read("word/document.xml").decode("utf-8")
            assert {name for name in rendered_parts if name.startswith("word/header")} == {
                name for name in source_parts if name.startswith("word/header")}
            assert {name for name in rendered_parts if name.startswith("word/footer")} == {
                name for name in source_parts if name.startswith("word/footer")}
    finally:
        core.OUTPUT_DIR = old_output


def _case_chain_fixture(count=6):
    cases = [{
        "case_id": f"TD-{index:04d}",
        "case_type": "translation_decision",
        "source_segment_id": f"seg-case-{index:04d}",
        "academic_candidate_status": "eligible",
        "provenance": {"historical": True, "generated_for_analysis": False},
    } for index in range(1, count + 1)]
    segments = [{
        "segment_id": f"seg-case-{index:04d}",
        "source": f"Canonical source sentence number {index} is unique and sufficiently long.",
        "initial_target": f"第{index}条规范译文。",
        "final_target": f"第{index}条规范译文。",
        "process_evidence": {"findings": [{"type": "terminology"}]},
    } for index in range(1, count + 1)]
    for case, segment in zip(cases, segments):
        case.update({
            "segment_id": segment["segment_id"],
            "canonical_evidence": {
                "source": segment["source"], "initial": None,
                "target": segment["final_target"],
            },
            "focus": academic_evidence.build_case_focus(case, segment, []),
            "difficulty_group": "连贯性与质量控制",
            "strategy_group": "篇章衔接",
            "research_questions": [],
            "target_subsection": "3.3.1",
            "strategy_subsection": "3.3.1",
            "difficulty_subsection": "3.2.1",
        })
    plans = {"plans": [{
        "case_id": case["case_id"], "case_type": "translation_decision",
        "analysis_contract_type": "translation_decision",
        "problem": {"type": "cohesion", "statement": "专名连贯性决策。",
                    "grounded": True},
        "decision_rationale": "译文保持规范专名。",
        "bounded_conclusion": "结论限于本例。",
        "analysis_contract": {},
    } for case in cases]}
    selected = {
        "report_case_policy": thesis_constraints.case_policy({
            "report_stage": "proposal"}),
        "case_portfolio": {"groups": [{
            "difficulty_group": "连贯性与质量控制",
            "strategy_group": "篇章衔接",
            "difficulty_subsection": "3.2.1",
            "strategy_subsection": "3.3.1",
            "case_ids": [case["case_id"] for case in cases],
            "case_count": len(cases),
        }]},
        "cases": cases,
    }
    evidence = {"project_evidence": {
        "segments": segments, "statistics": {}, "glossary": []},
        "translation_decision_candidates": cases, "candidate_cases": []}
    outline = {"sections": [{
        "section_id": "3", "title": "案例分析", "role": "case_analysis",
        "cases": [case["case_id"] for case in cases], "claims": [],
        "research_questions": [], "required_subsections": [], "minimum_chars": 0,
    }]}
    return selected, evidence, plans, outline


def _case_chain_markdown(case_ids, *, duplicate_number=False, history=False):
    blocks = []
    for index, case_id in enumerate(case_ids, 1):
        source_index = int(case_id.split("-")[-1])
        number = 1 if duplicate_number else index
        analysis = "笔者初译错误，修改后改译为当前译文。" if history else "本例只分析可观察的译法。"
        blocks.append(
            f"**例[{number}]：案例{index}**\n"
            f"* **SOURCE**：Canonical source sentence number {source_index} is unique "
            "and sufficiently long.\n"
            f"* **TARGET**：第{source_index}条规范译文。\n"
            f"* **译法分析**：{analysis}")
    return "## 3 案例分析\n\n### 3.3 翻译策略与解决方案\n\n" + \
        "#### 3.3.4 连贯性与质量控制\n\n" + "\n\n".join(blocks) + "\n"


def _case_chain_artifact(markdown, selected, evidence, plans, outline):
    written = [{"section_id": "3", "title": "案例分析", "content": ""}]
    normalized = academic_writer.finalize_report_tokens(
        markdown, evidence, selected, outline)
    artifact = academic_writer.build_report_artifact(
        normalized, written, outline, {}, selected_cases=selected,
        evidence=evidence, case_analysis_plans=plans)
    validation = academic_validator.validate_academic_report(
        normalized, evidence, {"research_questions": []}, {"claims": []},
        selected, outline, report_artifact=artifact)
    return normalized, artifact, validation


def test_six_selected_cases_create_six_unique_nodes_and_assembly_numbers():
    selected, evidence, plans, outline = _case_chain_fixture()
    markdown = _case_chain_markdown(
        [case["case_id"] for case in selected["cases"]], duplicate_number=True)
    normalized, artifact, validation = _case_chain_artifact(
        markdown, selected, evidence, plans, outline)
    assert [node["example_number"] for node in artifact["case_nodes"]] == list(range(1, 7))
    assert len({node["case_id"] for node in artifact["case_nodes"]}) == 6
    assert artifact["case_counts"]["selected_case_count"] == 6
    assert artifact["case_counts"]["structured_case_node_count"] == 6
    assert artifact["case_counts"]["focused_case_count"] == 6
    assert artifact["case_counts"][
        "unique_provenance_bound_visible_case_count"] == 6
    assert validation["case_validation"]["status"] != "fail"
    assert [f"例[{index}]" for index in range(1, 7)] == [
        artifact["case_labels"][node["case_id"]] for node in artifact["case_nodes"]]
    public = report_template.public_report_markdown(
        artifact["sections"][0]["content"], artifact["case_labels"])
    assert "TD-" not in public and "seg-case-" not in public
    assert all(public.count(f"例[{index}]") == 1 for index in range(1, 7))
    assert normalized.count("<!--case:") == 6
    contract = _mti_contract()
    rendered_artifact = {**artifact, "template_contract": contract,
                         "template_hash": contract["template_identity"]["sha256"]}
    rendered = report_template.render_report_docx(
        rendered_artifact, _mti_template_bytes(), contract).getvalue()
    rendered_text = "\n".join(
        paragraph.text for paragraph in Document(BytesIO(rendered)).paragraphs)
    assert "TD-" not in rendered_text and "seg-case-" not in rendered_text


def test_missing_writer_case_reports_exact_id_then_targeted_repair_passes():
    selected, evidence, plans, outline = _case_chain_fixture()
    case_ids = [case["case_id"] for case in selected["cases"]]
    markdown = _case_chain_markdown(case_ids[:-1])
    _normalized, _artifact, failed = _case_chain_artifact(
        markdown, selected, evidence, plans, outline)
    mismatch = next(item for item in failed["issues"]
                    if item["type"] == "case_presentation_count_mismatch")
    assert mismatch["missing_case_ids"] == ["TD-0006"]
    packet = {
        "cases": [{**selected["cases"][-1], "evidence":
                   evidence["project_evidence"]["segments"][-1]}],
        "case_analyses": [plans["plans"][-1]],
        "case_assignments": [academic_writer._case_assignment_for_plan(
            "TD-0006", plans["plans"][-1])],
    }

    def repair_writer(*_args, **_kwargs):
        return ("**例[1]：定点补写案例**\n<!--case:TD-0006-->\n"
                "* **SOURCE**：Canonical source sentence number 6 is unique and "
                "sufficiently long.\n* **TARGET**：第6条规范译文。\n"
                "* **译法分析**：本例只分析可观察的译法。")

    section_text = markdown.split("## 3 案例分析\n\n", 1)[1]
    repaired_section = academic_writer._repair_missing_case_examples(
        section_text, packet, ["TD-0006"], repair_writer, "test", "", "model")
    repaired = "## 3 案例分析\n\n" + repaired_section + "\n"
    _normalized, artifact, passed = _case_chain_artifact(
        repaired, selected, evidence, plans, outline)
    assert artifact["case_counts"]["structured_case_node_count"] == 6
    assert passed["case_validation"]["status"] != "fail"


def test_duplicate_case_cannot_inflate_unique_provenance_count():
    selected, evidence, plans, outline = _case_chain_fixture(1)
    markdown = _case_chain_markdown(["TD-0001", "TD-0001"])
    _normalized, artifact, validation = _case_chain_artifact(
        markdown, selected, evidence, plans, outline)
    assert len(artifact["case_nodes"]) == 2
    assert artifact["case_counts"]["unique_provenance_bound_visible_case_count"] == 1
    assert "duplicate_selected_case_presentation" in {
        item["type"] for item in validation["issues"]}


def test_translation_decision_cannot_claim_historical_revision():
    selected, evidence, plans, outline = _case_chain_fixture(1)
    markdown = _case_chain_markdown(["TD-0001"], history=True)
    _normalized, _artifact, validation = _case_chain_artifact(
        markdown, selected, evidence, plans, outline)
    assert "translation_decision_presented_as_revision" in {
        item["type"] for item in validation["issues"]}


def test_marker_without_visible_node_and_visible_without_binding_do_not_count():
    selected, evidence, plans, outline = _case_chain_fixture(1)
    marker_only = ("## 3 案例分析\n\n### 3.3 翻译策略与解决方案\n\n"
                   "<!--case:TD-0001-->\n")
    _normalized, _artifact, marker_validation = _case_chain_artifact(
        marker_only, selected, evidence, plans, outline)
    assert marker_validation["case_validation"][
        "unique_provenance_bound_visible_case_count"] == 0
    assert "orphan_case_marker" in {item["type"] for item in marker_validation["issues"]}

    unbound = ("## 3 案例分析\n\n### 3.3 翻译策略与解决方案\n\n"
               "**例[1]：无绑定案例**\n* **SOURCE**：Unrelated source text that matches no "
               "selected segment at all.\n* **TARGET**：无关译文。\n")
    _normalized, _artifact, visible_validation = _case_chain_artifact(
        unbound, selected, evidence, plans, outline)
    assert visible_validation["case_validation"][
        "unique_provenance_bound_visible_case_count"] == 0
    assert "unbound_visible_case_example" in {
        item["type"] for item in visible_validation["issues"]}


def test_mti_contract_to_outline_to_report_to_validation_to_docx_e2e():
    contract = _mti_contract()
    settings = {"body_language": "zh-CN", "project_name": "Drone Communities",
                "report_template_contract": contract}
    segments = [
        {"segment_id": f"seg-{i}", "source": f"Source {i}",
         "final_target": f"译文 {i}", "coverage_zone": "middle",
         "process_evidence": {"findings": []}}
        for i in range(1, 7)
    ]
    evidence = {"project_evidence": {
        "segments": segments, "glossary": [],
        "document_profile": {"genre": "学术文本"},
        "statistics": {"total_segments": 6, "translated_segments": 6,
                       "reviewed_segments": 6, "tm_reuse_count": 0},
    }, "candidate_cases": [], "translation_decision_candidates": [
        {"case_id": f"TD-{i:04d}", "case_type": "translation_decision",
         "source_segment_id": f"seg-{i}",
         "decision_evidence": {"reasons": ["术语或句法决策"]}}
        for i in range(1, 7)
    ]}
    research = academic_writer.build_research_model(evidence, "有限理论", settings)
    argument = academic_writer.build_argument_plan(
        research, evidence, lambda *_args, **_kwargs: "{}", "test", "", "model")
    selected = academic_writer.select_academic_cases(
        research, argument, evidence, limit=6, policy="authentic_only")
    outline = academic_writer.build_academic_outline(
        research, argument, selected, evidence,
        lambda *_args, **_kwargs: '{"sections":[{"section_id":"A"}]}',
        "test", "", "model")
    written = _mti_written()
    markdown = _mti_markdown()
    matter = academic_writer.build_report_matter(
        research, evidence, selected, contract)
    artifact = academic_writer.build_report_artifact(
        markdown, written, outline, research["report_constraints"], matter)
    compliance = academic_validator.validate_template_compliance(
        markdown, contract, outline, artifact, selected)
    assert compliance["status"] == "pass"
    assert len(selected["cases"]) == 6
    assert all(item["case_type"] == "translation_decision" for item in selected["cases"])
    full_validation = academic_validator.validate_academic_report(
        markdown, evidence, research, argument, selected, outline,
        template_contract=contract, report_artifact=artifact)
    assert "outline_unknown_case" not in {
        item["type"] for item in full_validation["issues"]}
    diagnostics = academic_quality.deterministic_diagnostics(
        research, argument, selected, outline, written, evidence)
    scoped = academic_quality._scoped_inputs(
        research, argument, selected, outline, written, evidence,
        {"items": []}, diagnostics)
    assert "template_contract" not in scoped["research_model"]
    assert scoped["chapter_roles"] == {
        "1": "introduction", "2": "project_overview",
        "3": "case_analysis", "4": "conclusion_reflection"}
    artifact.update(report_status="generated", template_compliance="pass")
    output = report_template.render_report_docx(
        artifact, _mti_template_bytes(), contract).getvalue()
    rendered_text = "\n".join(
        paragraph.text for paragraph in Document(BytesIO(output)).paragraphs)
    assert "ABSTRACT" in rendered_text and "参考文献" in rendered_text
    assert "附录一" in rendered_text and "总结与反思" in rendered_text


def test_dynamic_chapter_four_toc_cache_and_surface_validation_are_complete():
    contract = _mti_contract()
    constraints = thesis_constraints.build_constraints({
        "body_language": "zh-CN", "project_name": "Drone Communities",
        "report_template_contract": contract, "report_stage": "proposal",
    })
    outline = _mti_outline(contract, [])
    written = _mti_written()
    written[-1]["content"] = (
        "### 4.1 研究问题回应\n\n"
        "#### 4.1.1 RQ1 回应\n\n第四章真实正文。\n\n"
        "### 4.2 实践经验\n\n实践经验正文。\n\n"
        "### 4.3 局限与改进\n\n局限正文。")
    markdown = "\n\n".join(
        f"## {item['section_id']} {item['title']}\n\n{item['content']}"
        for item in written) + "\n"
    matter = _mti_matter(contract)
    matter["project_title"] = "Drone Communities"
    for item in matter["front_matter"]:
        if item.get("role") == "abstract_en":
            item["content"] = "This is a complete English abstract."
        elif item.get("role") == "keywords_en":
            item["keywords"] = ["translation practice", "case analysis"]
    matter["back_matter"] = [
        {**item, "title": str(item.get("title") or "").replace(
            "《XXX》", "《Drone Communities》")}
        for item in matter["back_matter"]
    ]
    matter["report"] = {"literature_status": "literature_required"}
    artifact = academic_writer.build_report_artifact(
        markdown, written, outline, constraints, matter)
    output = report_template.render_report_docx(
        artifact, _mti_template_bytes(), contract).getvalue()
    rendered = Document(BytesIO(output))
    texts = [paragraph.text for paragraph in rendered.paragraphs]
    assert "第四章真实正文。" in texts
    assert "4.1.1 RQ1 回应" in texts
    assert "《XXX》" not in "\n".join(texts)
    assert "如果有" not in "\n".join(texts)
    xml = ZipFile(BytesIO(output)).read("word/document.xml").decode("utf-8")
    assert "TOC" in xml and "第四章 总结与反思" in xml
    assert "文化负载词处理" not in xml
    surface = final_docx.validate_final_docx(output, artifact)
    assert surface["status"] == "pass_with_warnings"
    assert surface["summary"]["nonempty_chapter_count"] == 4
    assert surface["summary"]["project_placeholder_count"] == 0


def test_english_abstract_translates_chinese_metadata_labels():
    contract = _mti_contract()
    research = {"project_metadata": {
        "project_name": "Drone Communities", "genre": "学术专著",
        "domain": "传播学/环境人文学"}, "research_questions": []}
    evidence = {"project_evidence": {
        "document_profile": {"genre": "学术专著", "domain": "传播学/环境人文学"},
        "statistics": {"total_segments": 138}, "segments": [], "glossary": []}}
    matter = academic_writer.build_report_matter(
        research, evidence, {"cases": []}, contract, {"sources": []})
    abstract = matter["report"]["abstract_en"]
    assert not any("\u3400" <= char <= "\u9fff" for char in abstract)
    assert "academic monograph" in abstract
    assert "environmental humanities" in abstract


def test_case_analysis_repetition_audit_fails_boilerplate_blocks():
    repeated = [{"case_id": f"TD-{index:04d}", "example_number": index,
                 "analysis": "同一分析句反复出现，且没有加入任何案例自己的语言形式。"}
                for index in range(1, 5)]
    audit = academic_writer.case_presentation.analysis_repetition_audit(
        repeated, {"TD-0001"})
    assert audit["status"] == "fail"
    assert audit["repeated_analysis_blocks"][0]["count"] == 4
    assert audit["repeated_sentence_count"] == 3
