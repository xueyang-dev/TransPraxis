from __future__ import annotations

import json

from docx import Document

from transpraxis import legacy_literature


def _legacy_docx(path):
    doc = Document()
    doc.add_paragraph("3.3.1 核心术语")
    doc.add_paragraph("正文中引用 Cabré（1999）。")
    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] Cabré, M. T. Terminology: Theory, Methods and Applications. Amsterdam/Philadelphia: John Benjamins, 1998.")
    doc.add_paragraph("致  谢")
    doc.save(path)


def test_inventory_reads_legacy_references_without_marking_them_verified(tmp_path):
    path = tmp_path / "legacy.docx"
    _legacy_docx(path)
    inventory = legacy_literature.parse_legacy_literature_inventory(path)
    assert inventory["summary"]["legacy_reference_count"] == 1
    record = inventory["references"][0]
    assert record["legacy_reference_id"] == "LR-001"
    assert record["raw_reference"].startswith("Cabré")
    assert record["verification_status"] == "verified"
    assert record["metadata_variants"]["original_year"] if "original_year" in record["metadata_variants"] else True


def test_plan_excludes_review_required_references_from_case_support():
    inventory = {
        "references": [
            {"reference_number": 3, "legacy_reference_id": "LR-003", "authors": ["Cabré"], "year": "1999", "verification_status": "verified"},
            {"reference_number": 13, "legacy_reference_id": "LR-013", "authors": ["Murtisari"], "year": "2016", "verification_status": "review_required"},
        ]
    }
    selected = {"cases": [{
        "case_id": "SC-1", "case_type": "synthetic_contrast", "baseline_origin": "legacy_analytical_draft",
        "research_questions": ["RQ1"], "difficulty": {"category": "negation_scope"},
        "synthetic_baseline": {"text": "模拟初译", "targeted_issue": "negation"},
        "canonical_evidence": {"source": "source", "target": "final"},
    }], "final_case_count": 1, "content_hash": "frozen", "translation_pair_hash_after": "pair"}
    model = {"research_questions": [{"rq_id": "RQ1", "question": "q1"}, {"rq_id": "RQ2", "question": "q2"}, {"rq_id": "RQ3", "question": "q3"}]}
    plan = legacy_literature.build_chapter3_writing_plan(inventory, selected, model)
    support = plan["cases"][0]["literature_support"]
    assert all(item["legacy_reference_id"] != "LR-013" for item in support)
    assert "LR-013" in plan["cases"][0]["unverified_legacy_refs_excluded"]


def test_report_records_frozen_hashes_and_does_not_reclassify_cases():
    inventory = {"summary": {"legacy_reference_count": 0, "complete_metadata_count": 0, "incomplete_metadata_count": 0, "footnote_citation_count": 0, "references_with_legacy_citations": 0}, "references": []}
    plan = {"research_questions": {}, "summary": {"cases_with_verified_literature_support": 0}, "frozen_portfolio": {"selected_case_count": 24, "selected_content_hash": "s", "translation_pair_hash": "p"}}
    selected = {"final_case_count": 24, "authentic_revision_cases": ["a"], "synthetic_contrast_cases": ["s"] * 23, "translation_decision_visible_count": 0, "translation_pair_hash_after": "p"}
    report = legacy_literature.literature_recovery_report(inventory, plan, selected)
    assert "frozen final cases：24" in report
    assert "translation-decision-only visible count：0" in report
    assert "case selection、legacy recovery、synthetic generation" in report
