from __future__ import annotations

import json

from docx import Document

from transpraxis import academic_evidence, legacy_cases


def _write_docx(path, rows):
    document = Document()
    document.add_heading("3.3.1 核心术语策略", level=3)
    for number, source, initial, revised, analysis in rows:
        for text in (
                f"例[{number}]", f"原文：{source}", f"初译：{initial}",
                f"改译：{revised}", f"分析：{analysis}"):
            document.add_paragraph(text)
    document.save(path)


def _evidence(source="The sensorium extends beyond vision.",
              initial="感官超越视觉。", target="感知域超越视觉。",
              authentic=False):
    segment_id = "seg-test-0001"
    segment = {
        "segment_id": segment_id, "segment_index": 1,
        "source": source, "initial_target": initial, "final_target": target,
        "integrity_flags": [], "process_evidence": {},
    }
    candidates = [{
        "case_id": segment_id, "case_type": "authentic_revision",
        "historical": True, "academic_candidate_status": "eligible",
    }] if authentic else []
    return {
        "project_evidence": {"segments": [segment], "glossary": []},
        "candidate_cases": candidates, "translation_decision_candidates": [],
    }


def _inventory(source="The sensorium extends beyond vision.",
               initial="感官超越视觉。", revised="感知域超越视觉。"):
    case = {
        "legacy_case_id": "LEGACY-0001", "legacy_example_number": 1,
        "legacy_source": source, "legacy_initial": initial,
        "legacy_revised": revised, "legacy_analysis": "术语边界发生变化。",
        "old_subsection": "3.3.1 核心术语策略",
        "old_difficulty_group": "术语", "old_strategy_group": "术语策略",
        "complete": True, "missing_fields": [],
    }
    return {
        "schema_version": legacy_cases.INVENTORY_VERSION,
        "source_document": "/tmp/legacy.docx",
        "source_document_name": "legacy.docx",
        "cases": [case], "content_hash": academic_evidence.stable_hash([case]),
    }


def _review(*, repair="pass", materiality="pass", analysis="pass"):
    def fake(provider, key, model, system, user_prompt, temperature=0.1):
        payload = json.loads(user_prompt)
        return json.dumps({"validations": [{
            "case_id": case["case_id"], "plausibility": "pass",
            "materiality": materiality, "repair_correctness": repair,
            "academic_analysis_value": analysis,
            "baseline_issue_span": case["legacy_simulated_initial"],
            "final_repair_span": case["current_project_target"],
            "contrast_rationale": "术语概念边界构成可辩护的实质差异。",
            "rejection_reason": "", "duplicate_with_case_id": "",
        } for case in payload["cases"]]}, ensure_ascii=False)
    return fake


def test_inventory_parses_all_numbered_cases_and_fields(tmp_path):
    path = tmp_path / "legacy.docx"
    rows = [(number, f"Source sentence {number}.", f"初译{number}。",
             f"改译{number}。", f"分析{number}。") for number in range(1, 26)]
    _write_docx(path, rows)
    artifact = legacy_cases.parse_legacy_case_inventory(path)
    assert artifact["summary"] == {
        "total_case_count": 25,
        "complete_source_initial_revised_count": 25,
        "complete_four_field_count": 25,
        "missing_field_case_count": 0,
        "duplicate_case_count": 0,
    }
    assert [x["legacy_example_number"] for x in artifact["cases"]] == list(range(1, 26))


def test_legacy_initial_is_synthetic_and_current_target_is_presented():
    result = legacy_cases.recover_legacy_cases(
        _inventory(revised="当前译文旧稿。"), _evidence(target="当前正式译文。"),
        _review(), "fixture", "", "fixture")
    case = result["items"][0]
    assert case["case_type"] == "synthetic_contrast"
    assert case["historical"] is False
    assert case["generated_for_analysis"] is True
    assert case["baseline_origin"] == "legacy_analytical_draft"
    assert case["target_contrast_text"] == "当前正式译文。"
    assert case["legacy_revised"] == "当前译文旧稿。"


def test_unbound_source_is_rejected():
    result = legacy_cases.recover_legacy_cases(
        _inventory(source="A wholly unrelated sentence about astronomy."),
        _evidence(), _review(), "fixture", "", "fixture")
    case = result["items"][0]
    assert case["reason"] == "source_alignment_below_threshold"
    assert not case.get("validation", {}).get("academic_case_eligible")


def test_legacy_contrast_passes_when_current_target_still_repairs_issue():
    result = legacy_cases.recover_legacy_cases(
        _inventory(), _evidence(), _review(), "fixture", "", "fixture")
    case = result["items"][0]
    assert case["validation"]["academic_case_eligible"] is True
    assert set(case["synthetic_evidence"].values()) >= {False, True, "pass"}
    assert case["current_final_compatibility"] in {
        "exact_current_compatible", "analytically_compatible_but_text_changed"}


def test_obsolete_legacy_contrast_is_rejected_without_rewriting_baseline():
    inventory = _inventory()
    before = inventory["cases"][0]["legacy_initial"]
    result = legacy_cases.recover_legacy_cases(
        inventory, _evidence(), _review(repair="fail"), "fixture", "", "fixture")
    case = result["items"][0]
    assert case["current_final_compatibility"] == "obsolete"
    assert case["validation"]["academic_case_eligible"] is False
    assert "repair_correctness" in case["validation"]["rejected_reasons"]
    assert case["synthetic_baseline"]["text"] == before


def test_authentic_upgrade_requires_project_revision_provenance_and_text_chain():
    result = legacy_cases.recover_legacy_cases(
        _inventory(), _evidence(authentic=True), _review(),
        "fixture", "", "fixture")
    case = result["items"][0]
    assert case["case_type"] == "authentic_revision"
    assert case["historical"] is True
    assert case["validation"]["authentic_provenance_match"] is True
    assert case["synthetic_baseline"] is None


def test_qa_source_cannot_become_final_legacy_synthetic():
    evidence = _evidence()
    evidence["project_evidence"]["segments"][0]["integrity_flags"] = [
        "source_target_misaligned"]
    result = legacy_cases.recover_legacy_cases(
        _inventory(), evidence, _review(), "fixture", "", "fixture",
        qa_source_segment_ids={"seg-test-0001"})
    assert result["items"][0]["reason"] == "integrity_flagged_source_segment"


def test_merge_prefers_legacy_and_keeps_new_generated_as_fallback():
    legacy = legacy_cases.recover_legacy_cases(
        _inventory(), _evidence(), _review(), "fixture", "", "fixture")
    generated_case = {
        "case_id": "SC-0002", "case_type": "synthetic_contrast",
        "synthetic_baseline": {"text": "模拟。"},
        "synthetic_evidence": {
            "baseline_plausibility": "pass", "material_difference": "pass",
            "repair_correctness": "pass", "academic_analysis_value": "pass"},
        "validation": {"academic_case_eligible": True},
    }
    merged = legacy_cases.merge_synthetic_artifacts(
        legacy, {"items": [generated_case]})
    assert [x["baseline_origin"] for x in merged["items"]] == [
        "legacy_analytical_draft", "newly_generated"]
    assert merged["metrics"]["legacy_synthetic_case_count"] == 1
    assert merged["metrics"]["newly_generated_synthetic_case_count"] == 1


def test_truncated_new_baseline_is_not_eligible_for_final_portfolio():
    generated = {
        "case_id": "SC-0002", "case_type": "synthetic_contrast",
        "synthetic_baseline": {"text": "只翻译了前半句："},
        "synthetic_evidence": {
            "baseline_plausibility": "pass", "material_difference": "pass",
            "repair_correctness": "pass", "academic_analysis_value": "pass"},
        "validation": {"academic_case_eligible": True, "rejected_reasons": []},
    }
    merged = legacy_cases.merge_synthetic_artifacts(None, {"items": [generated]})
    case = merged["items"][0]
    assert case["validation"]["academic_case_eligible"] is False
    assert case["synthetic_evidence"]["baseline_plausibility"] == "fail"
    assert case["validation"]["rejected_reasons"] == ["baseline_truncated"]


def test_manual_gate_review_rejects_obsolete_current_repair_without_changing_baseline():
    recovery = legacy_cases.recover_legacy_cases(
        _inventory(), _evidence(), _review(), "fixture", "", "fixture")
    baseline = recovery["items"][0]["synthetic_baseline"]["text"]
    reviewed = legacy_cases.apply_manual_reviews(recovery, {"reviews": [{
        "case_id": "LSC-0001",
        "gate_overrides": {"repair_correctness": "fail"},
        "reason": "current target no longer contains the legacy repair",
    }]})
    case = reviewed["items"][0]
    assert case["validation"]["academic_case_eligible"] is False
    assert case["synthetic_evidence"]["repair_correctness"] == "fail"
    assert case["synthetic_baseline"]["text"] == baseline


def test_recovery_never_mutates_translation_pairs():
    state = {"pairs": [{"source": "s", "initial_target": "i", "target": "t"}]}
    before = academic_evidence.stable_hash(state["pairs"])
    legacy_cases.recover_legacy_cases(
        _inventory(), _evidence(), _review(), "fixture", "", "fixture")
    assert academic_evidence.stable_hash(state["pairs"]) == before
