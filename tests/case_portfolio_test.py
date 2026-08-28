"""Final/proposal case policy, focus provenance and structured-node regressions."""
from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from transpraxis import (academic_evidence, academic_quality, academic_validator,
                         academic_writer, case_presentation, final_docx, report_template,
                         thesis_constraints)


def _span(text: str, excerpt: str) -> dict:
    start = text.index(excerpt)
    return {"start": start, "end": start + len(excerpt), "text": excerpt,
            "word_count": len(excerpt.split()), "selection_reason": "fixture"}


def _case(index: int, *, case_type: str = "translation_decision",
          group: int | None = None) -> dict:
    source_words = [f"source{index}_{word}" for word in range(300)]
    source = " ".join(source_words)
    source_focus = " ".join(source_words[80:150])
    target_words = [f"译文{index}_{word}" for word in range(180)]
    target = " ".join(target_words)
    target_focus = " ".join(target_words[45:95])
    initial = None
    initial_span = None
    if case_type == "authentic_revision":
        initial_words = [f"初译{index}_{word}" for word in range(180)]
        initial = " ".join(initial_words)
        initial_span = _span(initial, " ".join(initial_words[45:95]))
    elif case_type == "synthetic_contrast":
        initial_words = [f"模拟初译{index}_{word}" for word in range(180)]
        initial = " ".join(initial_words)
        initial_span = _span(initial, " ".join(initial_words[45:95]))
    group = group if group is not None else index % 4
    difficulty = f"项目难点组{group + 1}"
    strategy = f"项目策略组{group + 1}"
    case_id = (f"AR-{index:04d}" if case_type == "authentic_revision" else
               f"SC-{index:04d}" if case_type == "synthetic_contrast" else
               f"TD-{index:04d}")
    row = {
        "case_id": case_id,
        "case_type": case_type,
        "segment_id": f"seg-focus-{index:04d}",
        "source_segment_id": f"seg-focus-{index:04d}",
        "canonical_analytical_case_identity": f"seg-focus-{index:04d}|issue-{index}",
        "canonical_evidence": {"source": source, "initial": initial, "target": target},
        "focus": {
            "issue": f"issue-{index}",
            "source_span": _span(source, source_focus),
            "initial_span": initial_span,
            "target_span": _span(target, target_focus),
            "soft_max_words": 180,
        },
        "difficulty_group": difficulty,
        "strategy_group": strategy,
        "difficulty_subsection": f"3.2.{group + 1}",
        "strategy_subsection": f"3.3.{group + 1}",
        "target_subsection": f"3.3.{group + 1}",
        "research_questions": [f"RQ{index % 3 + 1}"],
        "provenance": {"historical": case_type == "authentic_revision",
                        "generated_for_analysis": case_type == "synthetic_contrast"},
    }
    if case_type == "synthetic_contrast":
        row.update({
            "historical": False, "generated_for_analysis": True,
            "baseline_origin": "newly_generated",
            "synthetic_baseline": {"text": initial, "provenance": "analytical_simulation",
                                     "baseline_origin": "newly_generated"},
            "synthetic_evidence": {
                "historical": False, "generated_for_analysis": True,
                "baseline_plausibility": "pass", "material_difference": "pass",
                "repair_correctness": "pass", "academic_analysis_value": "pass",
            },
            "validation": {"academic_case_eligible": True},
        })
    if case_type in {"authentic_revision", "synthetic_contrast"}:
        row.update({
            "final_case_eligible": True, "contrast_ready": True,
            "contrast_type": "authentic" if case_type == "authentic_revision"
            else "synthetic",
        })
    return row


def _selected(count: int, *, stage: str = "final_report") -> dict:
    if stage == "final_report":
        cases = [_case(1, case_type="authentic_revision")]
        cases.extend(_case(index, case_type="synthetic_contrast")
                     for index in range(2, count + 1))
    else:
        cases = [_case(index) for index in range(1, count + 1)]
    return {
        "report_case_policy": thesis_constraints.case_policy({"report_stage": stage}),
        "cases": cases,
    }


def _node(case: dict, number: int, subsection: str | None = None) -> dict:
    focus = case["focus"]
    node = {
        "type": "case_example", "case_id": case["case_id"],
        "case_type": case["case_type"], "example_number": number,
        "subsection_id": subsection or case["target_subsection"],
        "focus": {"source": focus["source_span"],
                  "initial": focus["initial_span"], "target": focus["target_span"]},
        "source": focus["source_span"]["text"],
        "initial_target": (focus["initial_span"] or {}).get("text"),
        "target": focus["target_span"]["text"],
        "difficulty": {"statement": case["focus"]["issue"]},
        "strategy": case["strategy_group"],
        "effect": {"dimension": "terminological_precision",
                   "demonstrated_by": "译法在当前文本中保持一致。"},
        "bounded_claim": "结论仅限于本例。",
        "evidence": {"level": "fixture", "can_support": [], "cannot_support": []},
        "analysis_fields": {
            "difficulty": {"statement": case["focus"]["issue"]},
            "strategy": case["strategy_group"],
            "effect": {"dimension": "terminological_precision",
                       "demonstrated_by": "译法在当前文本中保持一致。"},
            "bounded_claim": "结论仅限于本例。",
            "visible_analysis": ["该译法在当前文本中保持一致。"],
        },
        "provenance": case["provenance"],
        "final_case_eligible": case.get("final_case_eligible", False),
        "contrast_ready": case.get("contrast_ready", False),
        "contrast_type": case.get("contrast_type"),
        "visible": True, "provenance_bound": True,
    }
    if case["case_type"] == "synthetic_contrast":
        node["synthetic_baseline"] = {
            **case["synthetic_baseline"],
            "text": focus["initial_span"]["text"],
        }
        node["synthetic_evidence"] = case["synthetic_evidence"]
    presentation = case_presentation.build_case_presentation(node)
    node.update(presentation=presentation,
                content=case_presentation.render_case_presentation_markdown(presentation),
                analysis=presentation["analysis"])
    return node


def _artifact(selected: dict) -> dict:
    return {"case_nodes": [_node(case, index) for index, case in enumerate(
        selected["cases"], 1)]}


def test_final_report_19_cases_fails_and_20_unique_provenance_safe_passes():
    nineteen = _selected(19)
    failed = academic_validator.validate_case_portfolio(
        nineteen, _artifact(nineteen), "final_report")
    assert failed["status"] == "fail"
    assert failed["provenance_safe_case_count"] == 19
    assert "case_minimum_not_met" in {item["type"] for item in failed["issues"]}

    twenty = _selected(20)
    passed = academic_validator.validate_case_portfolio(
        twenty, _artifact(twenty), "final_report")
    assert passed["status"] == "pass_with_warnings"
    assert passed["unique_provenance_safe_case_count"] == 20
    assert "case_minimum_not_met" not in {item["type"] for item in passed["issues"]}


def test_final_report_decision_only_cases_fail_the_contrast_contract():
    selected = {
        "report_case_policy": thesis_constraints.case_policy({"report_stage": "final_report"}),
        "cases": [_case(index) for index in range(1, 21)],
    }
    result = academic_validator.validate_case_portfolio(
        selected, _artifact(selected), "final_report")
    issue_types = {item["type"] for item in result["issues"]}
    assert result["final_case_count"] == 0
    assert result["translation_decision_visible_count"] == 40
    assert "final_case_type_contract_violation" in issue_types
    assert "translation_decision_visible_in_final_report" in issue_types


def test_final_report_20_contrast_cases_pass_and_decision_stays_backend_only():
    selected = _selected(20)
    result = academic_validator.validate_case_portfolio(
        selected, _artifact(selected), "final_report")
    issue_types = {item["type"] for item in result["issues"]}
    assert result["final_case_count"] == 20
    assert result["contrast_case_count"] == 20
    assert result["contrast_ready_case_count"] == 20
    assert result["translation_decision_visible_count"] == 0
    assert not issue_types & {
        "final_case_count_below_minimum", "final_case_contrast_not_ready",
        "final_case_type_contract_violation", "translation_decision_visible_in_final_report",
    }


def test_final_report_can_be_one_authentic_plus_nineteen_synthetic_or_all_synthetic():
    mixed = _selected(20)
    all_synthetic = {**mixed, "cases": [
        _case(index, case_type="synthetic_contrast") for index in range(1, 21)]}
    result = academic_validator.validate_case_portfolio(
        all_synthetic, _artifact(all_synthetic), "final_report")
    assert result["status"] != "fail"
    assert result["final_case_count"] == result["contrast_ready_case_count"] == 20


def test_final_report_rejects_missing_synthetic_label_or_rewrite_label():
    selected = _selected(20)
    artifact = _artifact(selected)
    artifact["case_nodes"][1]["content"] = artifact["case_nodes"][1][
        "content"].replace("模拟初译", "初译", 1)
    result = academic_validator.validate_case_portfolio(
        selected, artifact, "final_report")
    assert "synthetic_label_count_mismatch" in {item["type"] for item in result["issues"]}

    artifact = _artifact(selected)
    artifact["case_nodes"][2]["content"] = artifact["case_nodes"][2][
        "content"].replace("改译", "译文", 1)
    result = academic_validator.validate_case_portfolio(
        selected, artifact, "final_report")
    assert "final_case_rewrite_label_count_mismatch" in {
        item["type"] for item in result["issues"]}


def test_final_report_rejects_qa_source_and_academic_value_failures():
    selected = _selected(20)
    selected["qa_excluded_source_segment_ids"] = [
        selected["cases"][1]["source_segment_id"]]
    selected["cases"][2]["synthetic_evidence"]["academic_analysis_value"] = "fail"
    result = academic_validator.validate_case_portfolio(
        selected, _artifact(selected), "final_report")
    issue_types = {item["type"] for item in result["issues"]}
    assert "qa_case_selected_as_final_synthetic" in issue_types
    assert "synthetic_gate_failed_in_final_case" in issue_types


def test_final_selector_keeps_translation_decision_as_backend_candidate_only():
    state = {
        "paras": ["A recorded source.", "A second source.",
                  "Although the complex term planetarity appears in the source, the current "
                  "sentence retains its focus; therefore the decision is observable."],
        "pairs": [
            {"source": "A recorded source.", "target": "终译一", "initial_target": "初译一"},
            {"source": "A second source.", "target": "终译二", "initial_target": "初译二"},
            {"source": "Although the complex term planetarity appears in the source, the current "
             "sentence retains its focus; therefore the decision is observable.",
             "target": "终译三", "initial_target": "终译三"},
        ], "findings": [{"segment_index": 2, "severity": "actionable",
                          "type": "review", "reason": "术语决策需要保留语义边界。"}],
        "human_actions": [], "glossary": [],
    }
    evidence = academic_evidence.build_academic_evidence(state, "final-policy")
    policy = thesis_constraints.case_policy({"report_stage": "final_report"})
    selected = academic_writer.select_academic_cases(
        {}, {"claims": []}, evidence, limit=20, synthetic_artifact={"items": []},
        policy="mixed", report_case_policy=policy)
    assert all(item["case_type"] != "translation_decision"
               for item in selected["cases"])
    assert selected["translation_decision_candidate_pool_count"] >= 1
    assert selected["selection_status"] == "insufficient_contrast_cases"


def _surface_docx(*, bad_first_case: bool = False) -> bytes:
    document = Document()
    for index in range(1, 21):
        document.add_paragraph(f"例[{index}]")
        document.add_paragraph("原文：source")
        document.add_paragraph("模拟初译：baseline")
        document.add_paragraph("改译：final")
        document.add_paragraph("分析：该对比显示具体翻译策略。")
    if bad_first_case:
        document.paragraphs[2].text = "译文：final"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_final_docx_requires_a_contrast_near_every_formal_case():
    artifact = {
        "report_stage": "final_report",
        "case_policy": thesis_constraints.case_policy({"report_stage": "final_report"}),
    }
    good = final_docx.validate_final_docx(_surface_docx(), artifact)
    assert good["summary"]["case_count"] == 20
    assert good["summary"]["synthetic_label_count"] == 20
    assert good["summary"]["rewrite_label_count"] == 20
    assert "docx_translation_decision_only_case" not in {
        item["type"] for item in good["issues"]}
    bad = final_docx.validate_final_docx(_surface_docx(bad_first_case=True), artifact)
    assert "docx_translation_decision_only_case" in {
        item["type"] for item in bad["issues"]}


def test_proposal_six_cases_passes():
    selected = _selected(6, stage="proposal")
    result = academic_validator.validate_case_portfolio(
        selected, _artifact(selected), "proposal")
    assert result["status"] != "fail"
    assert result["minimum_cases"] == 6
    assert result["provenance_safe_case_count"] == 6


def test_duplicate_canonical_cases_reduce_unique_count_to_15_and_fail():
    selected = _selected(15)
    duplicates = []
    for index, original in enumerate(selected["cases"][:5], 16):
        duplicate = {**original, "case_id": f"TD-{index:04d}"}
        duplicates.append(duplicate)
    selected["cases"].extend(duplicates)
    result = academic_validator.validate_case_portfolio(selected, report_stage="final_report")
    assert result["selected_case_count"] == 20
    assert result["unique_case_count"] == 15
    assert result["provenance_safe_case_count"] == 15
    assert result["status"] == "fail"


def test_focus_offsets_inside_canonical_pass_and_outside_fail():
    selected = _selected(20)
    assert academic_validator.validate_case_portfolio(
        selected, report_stage="final_report")["status"] != "fail"
    selected["cases"][0]["focus"]["source_span"]["text"] = "not canonical"
    failed = academic_validator.validate_case_portfolio(
        selected, report_stage="final_report")
    assert failed["status"] == "fail"
    assert "focus_span_outside_canonical" in {
        item["type"] for item in failed["issues"]}


def _render_template() -> tuple[bytes, dict]:
    document = Document()
    document.add_heading("第三章 案例分析", level=1)
    document.add_heading("3.3 翻译策略与解决方案", level=2)
    buffer = BytesIO()
    document.save(buffer)
    data = buffer.getvalue()
    return data, report_template.parse_docx_template("focus.docx", data)


def _case_markdown(cases: list[dict]) -> str:
    groups: dict[str, list[dict]] = {}
    for case in cases:
        groups.setdefault(case["target_subsection"], []).append(case)
    chunks = ["## 3 案例分析", "", "### 3.3 翻译策略与解决方案", ""]
    for subsection, members in groups.items():
        chunks.extend([f"#### {subsection} {members[0]['strategy_group']}", ""])
        for case in members:
            focus = case["focus"]
            chunks.extend([
                "**例[1]：聚焦案例**",
                f"<!--case:{case['case_id']}-->",
                f"* **SOURCE**：{focus['source_span']['text']}",
            ])
            if case["case_type"] == "authentic_revision":
                chunks.append(f"* **INITIAL**：{focus['initial_span']['text']}")
            chunks.extend([
                f"* **TARGET**：{focus['target_span']['text']}",
                "* **分析**：本例只分析该聚焦翻译点。", "",
            ])
    return "\n".join(chunks)


def _build_artifact(selected: dict) -> tuple[str, dict]:
    markdown = _case_markdown(selected["cases"])
    segments = [{
        "segment_id": case["segment_id"],
        "source": case["canonical_evidence"]["source"],
        "initial_target": case["canonical_evidence"]["initial"],
        "final_target": case["canonical_evidence"]["target"],
        "process_evidence": {},
    } for case in selected["cases"]]
    evidence = {"project_evidence": {"segments": segments, "glossary": [],
                                      "statistics": {}}}
    outline = {"sections": [{"section_id": "3", "title": "案例分析",
                              "role": "case_analysis",
                              "cases": [x["case_id"] for x in selected["cases"]]}]}
    normalized = academic_writer.finalize_report_tokens(
        markdown, evidence, selected, outline)
    artifact = academic_writer.build_report_artifact(
        normalized, [{"section_id": "3", "title": "案例分析", "content": ""}],
        outline, thesis_constraints.build_constraints({"report_stage": "proposal"}),
        selected_cases=selected, evidence=evidence)
    return normalized, artifact


def test_renderer_uses_70_word_focus_not_300_word_segment():
    selected = _selected(1, stage="proposal")
    _markdown, artifact = _build_artifact(selected)
    template, contract = _render_template()
    artifact.update(template_contract=contract,
                    template_hash=contract["template_identity"]["sha256"])
    rendered = report_template.render_report_docx(artifact, template, contract).getvalue()
    text = "\n".join(paragraph.text for paragraph in Document(BytesIO(rendered)).paragraphs)
    assert selected["cases"][0]["focus"]["source_span"]["text"] in text
    assert "source1_299" not in text


def test_decision_omits_initial_but_revision_uses_all_focused_spans():
    decision = _selected(1, stage="proposal")
    _markdown, decision_artifact = _build_artifact(decision)
    public = report_template.public_report_markdown(
        decision_artifact["sections"][0]["content"],
        decision_artifact["case_labels"], decision_artifact["case_types"])
    assert "**译文**：" in public and "**初译**：" not in public \
        and "**改译**：" not in public
    assert decision_artifact["case_nodes"][0]["focus"]["initial"] is None

    revision = _selected(1, stage="proposal")
    revision["cases"] = [_case(1, case_type="authentic_revision")]
    _markdown, revision_artifact = _build_artifact(revision)
    node = revision_artifact["case_nodes"][0]
    assert node["focus"]["source"] and node["focus"]["initial"] and node["focus"]["target"]


def test_twenty_cases_number_continuously_across_multiple_subsections():
    selected = _selected(20)
    _markdown, artifact = _build_artifact(selected)
    assert [node["example_number"] for node in artifact["case_nodes"]] == list(range(1, 21))
    assert len({node["subsection_id"] for node in artifact["case_nodes"]}) == 4


def test_last_decision_case_node_stops_before_next_chapter():
    selected = _selected(1, stage="proposal")
    markdown = _case_markdown(selected["cases"]) + \
        "\n## 4 总结与反思\n\n本章讨论真实修订案例的初译记录。\n"
    case = selected["cases"][0]
    evidence = {"project_evidence": {"segments": [{
        "segment_id": case["segment_id"],
        "source": case["canonical_evidence"]["source"],
        "initial_target": None, "final_target": case["canonical_evidence"]["target"],
        "process_evidence": {},
    }], "glossary": [], "statistics": {}}}
    outline = {"sections": [
        {"section_id": "3", "title": "案例分析", "role": "case_analysis",
         "cases": [case["case_id"]]},
        {"section_id": "4", "title": "总结与反思", "role": "conclusion_reflection",
         "cases": []},
    ]}
    normalized = academic_writer.finalize_report_tokens(
        markdown, evidence, selected, outline)
    artifact = academic_writer.build_report_artifact(
        normalized, [], outline, thesis_constraints.build_constraints({
            "report_stage": "proposal"}), selected_cases=selected, evidence=evidence)
    node = artifact["case_nodes"][0]
    assert "## 4" not in node["content"] and "初译记录" not in node["content"]


def test_writer_packet_excludes_full_canonical_segment():
    selected = _selected(1, stage="proposal")
    case = selected["cases"][0]
    evidence = {"project_evidence": {"segments": [{
        "segment_id": case["segment_id"],
        "source": case["canonical_evidence"]["source"],
        "initial_target": None, "final_target": case["canonical_evidence"]["target"],
        "process_evidence": {},
    }], "glossary": [], "statistics": {}}}
    section = {"section_id": "3", "role": "case_analysis",
               "cases": [case["case_id"]], "claims": [],
               "required_statistics": []}
    packet = academic_writer._section_packet(
        section, {}, {"claims": []}, selected, evidence,
        {"sections": [section]}, [])
    serialized = json.dumps(packet, ensure_ascii=False)
    assert case["focus"]["source_span"]["text"] in serialized
    assert "source1_299" not in serialized


def test_quality_packet_excludes_full_canonical_segment():
    selected = _selected(1, stage="proposal")
    case = selected["cases"][0]
    payload = academic_quality._scoped_inputs(
        {}, {"claims": []}, selected,
        {"sections": [{"section_id": "3", "role": "case_analysis"}]},
        [{"section_id": "3", "content": ""}],
        {"project_evidence": {"segments": [], "glossary": [], "statistics": {}}},
        {"items": []}, {})
    serialized = json.dumps(payload, ensure_ascii=False)
    assert case["focus"]["source_span"]["text"] in serialized
    assert "source1_299" not in serialized


def test_quality_evidence_utilization_recognizes_structured_case_markers():
    selected = _selected(2, stage="proposal")
    sections = [{
        "section_id": "3",
        "content": "\n".join(
            f"<!--case:{case['case_id']}-->" for case in selected["cases"]),
    }]
    result = academic_quality.evidence_utilization(
        sections, selected,
        {"project_evidence": {"segments": [], "glossary": [], "statistics": {}}})
    assert result["cases_used"] == 2
    assert all(row["used_in_report"] for row in result["rows"])


def _visible_field_labels(content: str) -> list[str]:
    return __import__("re").findall(
        r"^\s*\*\*([^*：:\n]+)\*\*\s*[：:]", content,
        __import__("re").MULTILINE)


def test_case_presentation_contract_is_case_type_specific_and_hides_internal_fields():
    decision = _node(_case(1), 1)
    assert _visible_field_labels(decision["content"]) == ["原文", "译文", "分析"]
    assert not any(label in decision["content"] for label in (
        "翻译难点：", "译法分析：", "翻译效果：", "有界结论：", "证据边界："))
    assert all(key in decision for key in (
        "difficulty", "strategy", "effect", "bounded_claim", "evidence", "provenance"))
    assert not any(key in decision["presentation"] for key in (
        "difficulty", "strategy", "effect", "bounded_claim", "evidence", "provenance"))

    revision = _node(_case(2, case_type="authentic_revision"), 2)
    assert _visible_field_labels(revision["content"]) == ["原文", "初译", "改译", "分析"]
    assert revision["presentation"]["initial"]


def test_case_presentation_focus_emphasis_is_bounded_to_the_language_point():
    node = {
        "case_id": "TD-FOCUS", "case_type": "translation_decision",
        "example_number": 1,
        "focus": {
            "source": {"text": "The multiperspectival sensorium shapes perception."},
            "initial": None,
            "target": {"text": "多视角感知域塑造了感知方式。"},
            "issue": "multiperspectival sensorium → 多视角感知域",
        },
        "analysis_fields": {"visible_analysis": ["该译法保持了术语的一致性。"]},
    }
    presentation = case_presentation.build_case_presentation(node)
    assert "**multiperspectival sensorium**" in presentation["source"]
    assert "**多视角感知域**" in presentation["target"]
    assert not presentation["source"].startswith("**The")


def test_preview_markdown_and_docx_share_the_same_case_presentation():
    selected = _selected(1, stage="proposal")
    normalized, artifact = _build_artifact(selected)
    node = artifact["case_nodes"][0]
    rendered_case = case_presentation.render_case_presentation_markdown(
        node["presentation"])
    assert rendered_case in normalized
    preview = report_template.public_report_markdown(
        artifact["sections"][0]["content"], artifact["case_labels"],
        artifact["case_types"])
    assert "**原文**：" in preview and "**译文**：" in preview \
        and "**分析**：" in preview
    template, contract = _render_template()
    artifact.update(template_contract=contract,
                    template_hash=contract["template_identity"]["sha256"])
    docx = report_template.render_report_docx(artifact, template, contract).getvalue()
    text = "\n".join(p.text for p in Document(BytesIO(docx)).paragraphs)
    assert all(label in text for label in ("原文：", "译文：", "分析："))
    assert not any(label in text for label in ("翻译难点：", "译法分析：", "有界结论："))


def test_twenty_four_nodes_all_meet_the_presentation_contract():
    selected = _selected(24)
    selected["cases"][0] = _case(1, case_type="authentic_revision", group=1)
    _markdown, artifact = _build_artifact(selected)
    assert len(artifact["case_nodes"]) == 24
    result = academic_validator.validate_case_portfolio(
        selected, artifact, "final_report")
    assert "case_presentation_contract_violation" not in {
        item["type"] for item in result["issues"]}
    for node in artifact["case_nodes"]:
        labels = _visible_field_labels(node["content"])
        expected = ["原文", "初译", "改译", "分析"] \
            if node["case_type"] == "authentic_revision" else ["原文", "模拟初译", "改译", "分析"]
        assert labels == expected


def test_current_real_job_has_20_reliable_cases_or_explicit_insufficiency():
    root = Path(__file__).resolve().parent.parent / "outputs" / "c7fc0af0d6626931"
    if not (root / "academic-evidence.json").is_file():
        pytest.skip("local real-job artifacts are intentionally not versioned")
    load = lambda name: json.loads((root / name).read_text())
    policy = thesis_constraints.case_policy({"report_stage": "final_report"})
    selected = academic_writer.select_academic_cases(
        load("research-model.json"), load("argument-plan.json"),
        load("academic-evidence.json"), limit=policy["target_cases"],
        synthetic_artifact={"items": []}, policy="mixed",
        report_case_policy=policy)
    validation = academic_validator.validate_case_portfolio(
        selected, report_stage="final_report")
    assert validation["provenance_safe_case_count"] >= 20 \
        or selected["case_coverage_status"] == "insufficient"
