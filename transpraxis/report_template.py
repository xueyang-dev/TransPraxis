"""Canonical DOCX report-template contracts and template-aware rendering.

The parser is deliberately deterministic.  A template may be semantically
ambiguous, but its heading hierarchy, styles, sections, and fixed text are
still preserved instead of being re-invented by a model.
"""
from __future__ import annotations

import hashlib
import io
import json
import re
from copy import deepcopy
from typing import Any, Dict, List, Mapping, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.text.paragraph import Paragraph
from . import case_provenance


SCHEMA_VERSION = "report-template-contract-v3"
RENDERER_VERSION = "report-template-renderer-v4"


class TemplateParseError(ValueError):
    """Raised when an uploaded DOCX cannot be converted to a contract."""


_ROLE_RULES = (
    ("references", ("参考文献", "references", "bibliography")),
    ("acknowledgements", ("致谢", "acknowledg")),
    ("appendix", ("附录", "appendix", "annex")),
    ("abstract", ("摘要", "abstract")),
    ("table_of_contents", ("目录", "contents", "table of contents", "toc")),
    ("introduction", ("引言", "绪论", "introduction", "background")),
    ("project_overview", ("翻译项目概述", "翻译项目概况", "项目描述", "项目概述",
                           "项目概况", "研究方法", "project overview",
                           "project description", "methodology", "research method")),
    ("theoretical_framework", ("理论框架", "文献综述", "theoretical framework",
                                "literature review", "theory")),
    ("case_analysis", ("案例分析", "案例研究", "case analysis", "case study")),
    ("conclusion_reflection", ("结论", "反思", "conclusion", "reflection", "discussion")),
)


def _norm(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _title_key(value: Any) -> str:
    value = _norm(value).casefold()
    value = re.sub(r"^第\s*[一二三四五六七八九十百千万]+\s*章\s*", "", value)
    value = re.sub(r"^[一二三四五六七八九十百千万]+[、.．]\s*", "", value)
    value = re.sub(r"^\d+(?:\.\d+)*[.)、．]?\s*", "", value)
    return re.sub(r"[\s:：.。、()（）\[\]【】_-]+", "", value)


def _display_title(value: Any) -> str:
    text = _norm(value)
    text = re.sub(r"^第\s*[一二三四五六七八九十百千万]+\s*章\s*", "", text)
    text = re.sub(r"^\d+(?:\.\d+)*[.)、．]?\s+", "", text)
    return text or _norm(value)


def _role_for_title(value: Any) -> str:
    key = _title_key(value)
    for role, candidates in _ROLE_RULES:
        if any(_title_key(candidate) in key or key in _title_key(candidate)
               for candidate in candidates):
            return role
    return "generic_section"


def _section_id(value: Any, fallback: str) -> str:
    text = _norm(value)
    numeric = re.match(r"^\s*(\d+(?:\.\d+)*)[.)、．]?\s+", text)
    if numeric:
        return numeric.group(1)
    chinese = re.match(r"^\s*第\s*([一二三四五六七八九十百千万]+)\s*章", text)
    if chinese:
        digits = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
                  "六": 6, "七": 7, "八": 8, "九": 9, "十": 10}
        value = chinese.group(1)
        if value in digits:
            return str(digits[value])
    return str(fallback)


def _paragraph_outline_level(paragraph) -> Optional[int]:
    style_name = _norm(getattr(getattr(paragraph, "style", None), "name", ""))
    match = re.search(r"(?:heading|标题)\s*([1-9])", style_name, re.IGNORECASE)
    if match:
        return int(match.group(1))
    ppr = paragraph._p.pPr
    outline = ppr.find(qn("w:outlineLvl")) if ppr is not None else None
    if outline is not None:
        try:
            return int(outline.get(qn("w:val"))) + 1
        except (TypeError, ValueError):
            return None
    return None


def _paragraph_style_name(paragraph) -> str:
    return _norm(getattr(getattr(paragraph, "style", None), "name", ""))


def _heading_records(document) -> List[Dict[str, Any]]:
    records = []
    for index, paragraph in enumerate(document.paragraphs):
        text = _norm(paragraph.text)
        level = _paragraph_outline_level(paragraph)
        # Reachable template defect: the supplied MTI template styles
        # “第四章” as Heading 2. The explicit chapter number is stronger
        # structural evidence than that accidental style assignment.
        if re.match(r"^第\s*[一二三四五六七八九十百千万]+\s*章(?:\s|$)", text):
            level = 1
        if not text or level is None:
            continue
        records.append({
            "paragraph_index": index,
            "title": text,
            "level": level,
            "style": _paragraph_style_name(paragraph),
        })
    return records


def _front_matter_records(document) -> List[Dict[str, Any]]:
    """Identify required front matter even when it is not a Heading style."""
    texts = [_norm(paragraph.text) for paragraph in document.paragraphs]
    compact = [re.sub(r"\s+", "", text) for text in texts]
    xml = document.element.body.xml

    def has(pattern: str, *, flags: int = 0) -> bool:
        return any(re.search(pattern, text, flags) for text in texts if text)

    def record(role: str, title: str, source_title: str = "") -> Dict[str, Any]:
        return {
            "section_id": role,
            "title": title,
            "source_title": source_title or title,
            "role": role,
            "level": 0,
            "required": True,
        }

    values: List[Dict[str, Any]] = []
    if any("专业学位硕士学位论文" in text for text in compact):
        values.append(record("cover_zh", "中文封面"))
    if has(r"\bA\s+(?:Report|Thesis)\b", flags=re.IGNORECASE) and has(
            r"\b(?:Master|Translation and Interpreting)\b", flags=re.IGNORECASE):
        values.append(record("cover_en", "英文封面"))
    if any("独创性声明" in text for text in compact):
        values.append(record("originality_declaration", "独创性声明"))
    if any("使用授权声明" in text for text in compact):
        values.append(record("authorization_declaration", "使用授权声明"))
    zh_abstract = next((text for text in texts
                        if "摘要" in re.sub(r"\s+", "", text)
                        and "ABSTRACT" not in text.upper()), "")
    if zh_abstract:
        values.append(record("abstract_zh", "中文摘要", zh_abstract))
    zh_keywords = next((text for text in texts
                        if re.match(r"^\s*关键词\s*[：:]", text)), "")
    if zh_keywords:
        values.append(record("keywords_zh", "关键词", zh_keywords))
    en_abstract = next((text for text in texts if re.search(
        r"\bABSTRACT\b", text, re.IGNORECASE)), "")
    if en_abstract:
        values.append(record("abstract_en", "ABSTRACT", en_abstract))
    en_keywords = next((text for text in texts if re.match(
        r"^\s*Keywords\s*[：:]", text, re.IGNORECASE)), "")
    if en_keywords:
        values.append(record("keywords_en", "Keywords", en_keywords))
    if " TOC " in xml or any(text == "目录" for text in compact):
        values.append(record("table_of_contents", "目录"))
    return values


def _case_requirement(chapters: List[Mapping[str, Any]]) -> Dict[str, Any]:
    """Return the case floor encoded by this MTI report structure."""
    roles = [str(item.get("role") or "") for item in chapters]
    required = {
        str(sub.get("heading_id") or "")
        for chapter in chapters
        for sub in chapter.get("required_subsections") or []
    }
    is_mti_practice = roles == [
        "introduction", "project_overview", "case_analysis",
        "conclusion_reflection",
    ] and {"1.1", "1.2", "1.3", "2.1", "2.2", "2.2.1", "2.2.2",
           "2.2.3", "3.1", "3.2", "3.3"}.issubset(required)
    return {
        "minimum_cases": 6 if is_mti_practice else 0,
        "applies_to_report_stage": "proposal" if is_mti_practice else None,
        "status_when_insufficient": "failed_template_validation",
        "source": "proposal_requirement_in_canonical_mti_structure"
        if is_mti_practice else "template_unspecified",
    }


def _style_info(document, name: str) -> Dict[str, Any]:
    try:
        style = document.styles[name]
    except KeyError:
        return {"present": False}
    font = style.font
    size = font.size.pt if font.size else None
    color = str(font.color.rgb) if font.color and font.color.rgb else None
    return {
        "present": True,
        "font_name": font.name,
        "font_size_pt": size,
        "bold": font.bold,
        "italic": font.italic,
        "color": color,
        "line_spacing": style.paragraph_format.line_spacing,
        "space_after_pt": (style.paragraph_format.space_after.pt
                            if style.paragraph_format.space_after else None),
        "first_line_indent_twips": (
            round(style.paragraph_format.first_line_indent.twips)
            if style.paragraph_format.first_line_indent else None),
    }


def _section_geometry(section) -> Dict[str, Any]:
    return {
        "page_width_emu": int(section.page_width or 0),
        "page_height_emu": int(section.page_height or 0),
        "top_margin_emu": int(section.top_margin or 0),
        "bottom_margin_emu": int(section.bottom_margin or 0),
        "left_margin_emu": int(section.left_margin or 0),
        "right_margin_emu": int(section.right_margin or 0),
        "header_distance_emu": int(section.header_distance or 0),
        "footer_distance_emu": int(section.footer_distance or 0),
    }


def _fixed_text(document) -> List[str]:
    values = []
    for paragraph in document.paragraphs:
        text = _norm(paragraph.text)
        if text and _paragraph_outline_level(paragraph) is None:
            values.append(anonymize_sensitive_institutions(text))
    return values[:100]


def _numbering_profile(document) -> List[Dict[str, Any]]:
    values = []
    for index, paragraph in enumerate(document.paragraphs):
        ppr = paragraph._p.pPr
        num_pr = ppr.find(qn("w:numPr")) if ppr is not None else None
        if num_pr is None:
            continue
        num_id = num_pr.find(qn("w:numId"))
        ilvl = num_pr.find(qn("w:ilvl"))
        values.append({
            "paragraph_index": index,
            "num_id": num_id.get(qn("w:val")) if num_id is not None else None,
            "level": ilvl.get(qn("w:val")) if ilvl is not None else None,
            "style": _paragraph_style_name(paragraph),
        })
    return values


def _matter_record(record: Mapping[str, Any], order: int) -> Dict[str, Any]:
    return {
        "section_id": _section_id(record.get("title"), str(order)),
        "title": _display_title(record.get("title")),
        "source_title": record.get("title"),
        "role": _role_for_title(record.get("title")),
        "level": record.get("level", 1),
        "required": True,
        "style": record.get("style", ""),
    }


def _chapter_record(record: Mapping[str, Any], order: int,
                    all_records: List[Mapping[str, Any]], top_level: int = 1) -> Dict[str, Any]:
    source_title = record.get("title") or f"章节 {order}"
    title = _display_title(source_title)
    section_id = _section_id(source_title, str(order))
    start = next(index for index, item in enumerate(all_records)
                 if item is record)
    children = []
    next_order = 1
    for child in all_records[start + 1:]:
        if child.get("level", 1) <= record.get("level", 1):
            break
        source_child_title = child.get("title") or f"小节 {next_order}"
        child_title = _display_title(source_child_title)
        child_id = _section_id(source_child_title, f"{section_id}.{next_order}")
        child_item = {
            "heading_id": child_id,
            "title": child_title,
            "source_title": source_child_title,
            "level": int(child.get("level") or 2),
            "required": True,
            "style": child.get("style", ""),
            "markdown_prefix": "#" * max(
                3, int(child.get("level") or 2) - top_level + 2),
        }
        if child_id in {"3.2", "3.3"}:
            child_item.update({
                "allows_dynamic_children": True,
                "mapping_group": "difficulty_strategy",
                "mapping_side": "problem" if child_id == "3.2" else "solution",
            })
        children.append(child_item)
        next_order += 1
    role = _role_for_title(title)
    purposes = {
        "introduction": "提出项目背景、2—3 个真实研究问题并说明四章结构。",
        "project_overview": "仅依据项目记录说明项目简介及译前、译中、译后流程。",
        "case_analysis": "按源语特征、翻译难点、对应策略与真实例证分析组织正文。",
        "conclusion_reflection": "回应研究问题，总结策略、局限与改进方向，不引入新案例。",
    }
    return {
        "section_id": section_id,
        "title": title,
        "source_title": source_title,
        "role": role,
        "level": int(record.get("level") or 1),
        "purpose": purposes.get(role, "按模板章节标题与顺序组织翻译实践证据。"),
        "allows_dynamic_subsections": role == "conclusion_reflection",
        "required_subsections": children,
        "style": record.get("style", ""),
    }


def parse_docx_template(filename: str, template_bytes: bytes) -> Dict[str, Any]:
    """Parse a DOCX into a stable, serializable Template Contract."""
    if not template_bytes:
        raise TemplateParseError("模板文件为空。")
    try:
        document = Document(io.BytesIO(template_bytes))
    except Exception as exc:  # python-docx raises several parser-specific errors
        raise TemplateParseError(f"DOCX 模板无法读取：{exc}") from exc
    records = _heading_records(document)
    if not records:
        raise TemplateParseError("未识别到 Heading 样式或 outline level；请上传包含章节标题的 DOCX 模板。")
    top_level = min(record["level"] for record in records)
    body_records = [record for record in records if record["level"] == top_level]
    front_roles = {"abstract", "abstract_zh", "abstract_en", "table_of_contents"}
    back_roles = {"references", "acknowledgements", "appendix"}
    front = _front_matter_records(document)
    body = []
    back = []
    body_started = False
    for record in body_records:
        role = _role_for_title(record["title"])
        item = _matter_record(record, len(front) + len(back) + len(body) + 1)
        if not body_started and role in front_roles:
            if not any(_title_key(x.get("source_title")) == _title_key(
                    item.get("source_title")) for x in front):
                front.append(item)
            continue
        if role in back_roles:
            back.append(item)
            continue
        body_started = True
        body.append(record)
    chapters = [_chapter_record(record, index, records, top_level)
                for index, record in enumerate(body, start=1)]
    if not chapters:
        raise TemplateParseError("模板没有可识别的正文一级章节。")

    sections = [{"geometry": _section_geometry(section)}
                for section in document.sections]
    headers = [
        [anonymize_sensitive_institutions(_norm(p.text))
         for p in section.header.paragraphs if _norm(p.text)]
        for section in document.sections
    ]
    footers = [
        [anonymize_sensitive_institutions(_norm(p.text))
         for p in section.footer.paragraphs if _norm(p.text)]
        for section in document.sections
    ]
    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows[:3]:
            rows.append([anonymize_sensitive_institutions(_norm(cell.text))
                         for cell in row.cells])
        tables.append({"rows": rows, "columns": len(table.columns)})
    template_hash = hashlib.sha256(template_bytes).hexdigest()
    warnings = []
    for record in records:
        source_level = _paragraph_outline_level(document.paragraphs[
            int(record["paragraph_index"])])
        if source_level and source_level != record["level"] and re.match(
                r"^第\s*[一二三四五六七八九十百千万]+\s*章", record["title"]):
            warnings.append(
                f"promoted numbered chapter from Heading {source_level}: {record['title']}")
    contract = {
        "schema_version": SCHEMA_VERSION,
        "template_identity": {
            "filename": str(filename or "template.docx"),
            "sha256": template_hash,
            "template_id": template_hash[:16],
            "source_type": "docx",
        },
        "document_structure": {
            "front_matter": front,
            "chapters": chapters,
            "back_matter": back,
            "case_requirement": _case_requirement(chapters),
            "heading_count": len(records),
            "top_level": top_level,
        },
        "style_contract": {
            "sections": sections,
            "styles": {name: _style_info(document, name) for name in (
                "Normal", "Title", "Heading 1", "Heading 2", "Heading 3",
                "Intense Quote", "List Bullet", "List Number")},
            "headers": headers,
            "footers": footers,
            "tables": tables,
            "fixed_text": _fixed_text(document),
            "numbering": _numbering_profile(document),
        },
        "source_provenance": {
            "parser": "deterministic-docx-heading-v1",
            "source_filename": str(filename or "template.docx"),
            "warnings": warnings,
        },
        "strictness": "strict_structure",
    }
    contract["content_hash"] = hashlib.sha256(
        json.dumps(contract, ensure_ascii=False, sort_keys=True,
                   separators=(",", ":")).encode("utf-8")).hexdigest()
    return contract


def contract_summary(contract: Optional[Mapping[str, Any]]) -> Dict[str, Any]:
    if not contract:
        return {"configured": False, "status": "template_not_configured"}
    structure = contract.get("document_structure") or {}
    chapters = structure.get("chapters") or []
    subsection_count = sum(len(x.get("required_subsections") or []) for x in chapters)
    return {
        "configured": True,
        "status": "parsed",
        "filename": (contract.get("template_identity") or {}).get("filename"),
        "template_id": (contract.get("template_identity") or {}).get("template_id"),
        "template_hash": (contract.get("template_identity") or {}).get("sha256"),
        "chapter_count": len(chapters),
        "subsection_count": subsection_count,
        "front_matter": [x.get("title") for x in structure.get("front_matter") or []],
        "back_matter": [x.get("title") for x in structure.get("back_matter") or []],
        "minimum_cases": int((structure.get("case_requirement") or {}).get(
            "minimum_cases") or 0),
    }


def _delete_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def _heading_level(paragraph: Paragraph) -> Optional[int]:
    return _paragraph_outline_level(paragraph)


def _find_heading(document, title: str, level: Optional[int] = None) -> Optional[Paragraph]:
    key = _title_key(title)
    for paragraph in document.paragraphs:
        if level is not None and _heading_level(paragraph) != level:
            continue
        if _title_key(paragraph.text) == key:
            return paragraph
    return None


def _insert_paragraph_after(anchor: Paragraph, text: str, style: Optional[str] = None) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addnext(element)
    paragraph = Paragraph(element, anchor._parent)
    if style:
        try:
            paragraph.style = style
        except KeyError:
            pass
    paragraph.add_run(text)
    return paragraph


def _add_markdown_runs(paragraph: Paragraph, text: str) -> None:
    pieces = re.split(r"(\*\*[^*]+\*\*)", str(text or ""))
    for piece in pieces:
        if not piece:
            continue
        bold = piece.startswith("**") and piece.endswith("**")
        value = piece[2:-2] if bold else piece
        run = paragraph.add_run(value)
        if bold:
            run.bold = True


def _insert_content_after(anchor: Paragraph, content: str) -> Paragraph:
    current = anchor
    for raw_line in str(content or "").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("<!--"):
            continue
        style = None
        heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if heading:
            style = f"Heading {max(1, min(3, len(heading.group(1)) - 1))}"
            line = heading.group(2).strip()
        elif line.startswith("> "):
            style = "Intense Quote"
            line = line[2:].strip()
        elif re.match(r"^[-*]\s+", line):
            style = "List Bullet"
            line = re.sub(r"^[-*]\s+", "", line)
        paragraph = _insert_paragraph_after(current, "", style)
        paragraph._p.remove(paragraph.runs[0]._r)
        _add_markdown_runs(paragraph, line)
        current = paragraph
    return current


_PUBLIC_QUOTE = re.compile(
    r"^>\s*\[(SOURCE|INITIAL|TARGET|SYNTHETIC_SOURCE|SIMULATED|OPTIMIZED)\s+"
    r"([^\]]+)\]:\s*(.*)$")
_INTERNAL_ID = re.compile(
    r"\b(?:seg-[A-Za-z0-9_.:-]+|TD-\d+|SC-\d+|finding-[A-Za-z0-9_.:-]+|"
    r"term-[A-Za-z0-9_.:-]+|claim-[A-Za-z0-9_.:-]+|"
    r"(?:AQ|AV|AR|LR)-\d+)\b", re.IGNORECASE)


def anonymize_sensitive_institutions(text: str) -> str:
    """Apply the report's user-facing university anonymisation policy."""
    value = str(text or "")
    value = re.sub(r"[\u3400-\u9fff]{2,20}大学", "XX大学", value)
    value = re.sub(
        r"\b(?:[A-Z][A-Za-z&.'’-]*(?:\s+|$)){1,8}University\b(?!\s+Press)",
        "XX University", value)
    return value


def public_report_markdown(
    text: str, case_labels: Optional[Mapping[str, str]] = None,
    case_types: Optional[Mapping[str, str]] = None,
) -> str:
    """Hide provenance while retaining academic case labels and quotations."""
    labels = dict(case_labels or {})
    case_types = dict(case_types or {})
    case_ids = []
    for line in str(text or "").splitlines():
        match = _PUBLIC_QUOTE.match(line.strip())
        if match and match.group(2) not in case_ids:
            case_ids.append(match.group(2))
    for value in _INTERNAL_ID.findall(str(text or "")):
        if value.lower().startswith(("seg-", "td-", "sc-")) and value not in case_ids:
            case_ids.append(value)
    for case_id in case_ids:
        labels.setdefault(case_id, f"例[{len(labels) + 1}]")
    visible_case_labels = {
        case_id for case_id, label in labels.items()
        if label and label in str(text or "")}

    quote_labels = {
        "SOURCE": "原文", "INITIAL": "初译", "TARGET": "改译",
        "SYNTHETIC_SOURCE": "原文", "SIMULATED": "模拟初译",
        "OPTIMIZED": "改译",
    }
    seen = set()
    lines = []
    raw = re.sub(r"<!--.*?-->", "", str(text or ""), flags=re.DOTALL)
    raw = re.sub(r"\{\{(?:STAT|TERM):[^}]+\}\}", "", raw)
    for line in raw.splitlines():
        match = _PUBLIC_QUOTE.match(line.strip())
        if match:
            kind, case_id, value = match.groups()
            if case_id not in seen:
                if case_id not in visible_case_labels:
                    if lines and lines[-1].strip():
                        lines.append("")
                    lines.extend([f"**{labels.get(case_id, '例')}**", ""])
                seen.add(case_id)
            case_type = case_types.get(case_id)
            if not case_type and case_id.upper().startswith("TD-"):
                case_type = "translation_decision"
            case = {"case_type": case_type}
            display = case_provenance.display_contract(case)
            label = display["target_label"] if kind in {"TARGET", "OPTIMIZED"} else (
                "模拟初译" if kind == "SIMULATED" else
                display["initial_label"] if kind == "INITIAL" and
                display["initial_label"] else quote_labels[kind])
            lines.append(f"> {label}：{value}")
            continue
        for case_id, label in labels.items():
            line = line.replace(case_id, label)
        line = _INTERNAL_ID.sub("", line)
        lines.append(line.rstrip())
    cleaned = []
    for line in lines:
        heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", line.strip())
        previous = next((x for x in reversed(cleaned) if x.strip()), "")
        previous_heading = re.match(r"^#{1,6}\s+(.+?)\s*$", previous.strip())
        if heading and previous_heading and _title_key(heading.group(2)) == _title_key(
                previous_heading.group(1)):
            continue
        cleaned.append(line)
    return anonymize_sensitive_institutions(
        re.sub(r"\n{3,}", "\n\n", "\n".join(cleaned)).strip())


def _remove_body_examples(document, preserve: List[Paragraph], start: Paragraph) -> None:
    """Remove sample paragraphs/tables while retaining headings and section breaks."""
    body = document.element.body
    children = list(body)
    start_index = children.index(start._p)
    keep = {paragraph._p for paragraph in preserve}
    for child in children[start_index + 1:]:
        if child.tag == qn("w:sectPr") or child in keep \
                or child.find(".//" + qn("w:sectPr")) is not None:
            continue
        body.remove(child)


def _clear_between(anchor: Paragraph, next_anchor: Optional[Paragraph]) -> None:
    parent = anchor._p.getparent()
    children = list(parent)
    start = children.index(anchor._p) + 1
    end = children.index(next_anchor._p) if next_anchor is not None else len(children)
    for child in children[start:end]:
        if child.tag == qn("w:sectPr") or child.find(".//" + qn("w:sectPr")) is not None:
            continue
        parent.remove(child)


def _replace_visible_text(document, replacements: Mapping[str, str]) -> None:
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for relationship in document.part.rels.values():
        if relationship.reltype not in {RT.HEADER, RT.FOOTER}:
            continue
        part = relationship.target_part
        paragraphs.extend(Paragraph(element, part.element)
                          for element in part.element.xpath(".//w:p"))
    for paragraph in paragraphs:
        for run in paragraph.runs:
            value = run.text
            for old, new in replacements.items():
                value = value.replace(old, new)
            value = anonymize_sensitive_institutions(value)
            if value != run.text:
                run.text = value

    # python-docx intentionally omits paragraphs nested in structured document
    # tags, including the cached result of a Word TOC.  Replace text nodes in
    # the main part as well so project-title placeholders cannot survive there.
    for node in document.element.xpath(".//w:t"):
        value = str(node.text or "")
        for old, new in replacements.items():
            value = value.replace(old, new)
        node.text = anonymize_sensitive_institutions(value)


def _fill_cover(document, project_title: str) -> None:
    if not project_title:
        return
    paragraphs = document.paragraphs
    for index, paragraph in enumerate(paragraphs[:45]):
        text = _norm(paragraph.text)
        if re.match(r"^《.+》$", text):
            paragraph.text = f"《{project_title}》"
        elif "汉英翻译实践报告" in text or "英汉翻译实践报告" in text:
            paragraph.text = "英汉翻译实践报告"
        elif re.match(r"^A Report on [CE]-[CE] Translation of", text,
                      re.IGNORECASE):
            paragraph.text = "A Report on E-C Translation of"
            if index + 1 < len(paragraphs):
                paragraphs[index + 1].text = project_title


def _set_update_fields(document) -> None:
    settings = document.settings.element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


_CHAPTER_NUMBERS = {
    "1": "一", "2": "二", "3": "三", "4": "四", "5": "五",
    "6": "六", "7": "七", "8": "八", "9": "九", "10": "十",
}


def canonical_chapter_title(section_id: Any, title: Any) -> str:
    """Return the visible chapter title used by body headings and the TOC."""
    text = _norm(title)
    if re.match(r"^第\s*[一二三四五六七八九十百千万]+\s*章", text):
        return text
    number = _CHAPTER_NUMBERS.get(str(section_id))
    return f"第{number}章 {text}" if number else text


def _surface_matter_title(title: Any, project_title: Any = "") -> str:
    text = re.sub(r"\s*[（(]如果有\s*[，,、]?\s*另起一页[）)]\s*$", "", _norm(title))
    project = _norm(project_title)
    if "《XXX》" in text:
        text = text.replace("《XXX》", f"《{project}》" if project else "《当前翻译项目》")
    return text


def canonical_toc_entries(report_artifact: Mapping[str, Any]) -> List[Dict[str, Any]]:
    """Build the deterministic TOC cache from the final structured artifact."""
    entries: List[Dict[str, Any]] = [
        {"title": "摘  要", "level": 1},
        {"title": "ABSTRACT", "level": 1},
    ]
    for section in report_artifact.get("sections") or []:
        section_id = str(section.get("section_id") or "")
        entries.append({
            "title": canonical_chapter_title(section_id, section.get("title")),
            "level": 1,
        })
        for subsection in section.get("subsections") or []:
            heading_id = str(subsection.get("heading_id") or "").strip()
            title = _norm(subsection.get("title"))
            if not heading_id or not title:
                continue
            entries.append({
                "title": f"{heading_id} {title}",
                "level": min(3, heading_id.count(".") + 1),
            })
    for item in report_artifact.get("back_matter") or []:
        title = _surface_matter_title(item.get("title"), report_artifact.get("project_title"))
        if title:
            entries.append({"title": title, "level": 1})
    return entries


def _field_run(kind: str, text: str = "") -> Any:
    run = OxmlElement("w:r")
    if kind == "instruction":
        node = OxmlElement("w:instrText")
        node.set(qn("xml:space"), "preserve")
        node.text = text
    else:
        node = OxmlElement("w:fldChar")
        node.set(qn("w:fldCharType"), kind)
    run.append(node)
    return run


def _text_run(text: str) -> Any:
    run = OxmlElement("w:r")
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    return run


def _refresh_toc_cache(document, report_artifact: Mapping[str, Any]) -> None:
    """Replace a template TOC's stale cache while retaining its Word field.

    Word or LibreOffice may later refresh page numbers.  Until that happens,
    the visible cache still contains the exact canonical headings instead of
    template samples.
    """
    toc = next((node for node in document.element.xpath(".//w:sdt")
                if "TOC" in " ".join(str(item.text or "")
                                      for item in node.iter(qn("w:instrText")))), None)
    if toc is None:
        field_paragraph = next((paragraph for paragraph in
                                document.element.xpath(".//w:p")
                                if "TOC" in " ".join(str(item.text or "")
                                                     for item in paragraph.iter(
                                                         qn("w:instrText")))), None)
        if field_paragraph is None:
            return
        parent = field_paragraph.getparent()
        position = parent.index(field_paragraph)
        parent.remove(field_paragraph)
        title_paragraph = OxmlElement("w:p")
        title_props = OxmlElement("w:pPr")
        title_style = OxmlElement("w:pStyle")
        title_style.set(qn("w:val"), "TOCHeading")
        title_props.append(title_style)
        title_paragraph.append(title_props)
        title_paragraph.append(_text_run("目录"))
        parent.insert(position, title_paragraph)
        position += 1
        entries = canonical_toc_entries(report_artifact)
        for index, entry in enumerate(entries):
            paragraph = OxmlElement("w:p")
            props = OxmlElement("w:pPr")
            style = OxmlElement("w:pStyle")
            style.set(qn("w:val"), f"TOC{int(entry.get('level') or 1)}")
            props.append(style)
            paragraph.append(props)
            if index == 0:
                paragraph.extend([
                    _field_run("begin"),
                    _field_run("instruction", ' TOC \\o "1-3" \\h \\z \\u '),
                    _field_run("separate"),
                ])
            paragraph.append(_text_run(str(entry.get("title") or "")))
            if index == len(entries) - 1:
                paragraph.append(_field_run("end"))
            parent.insert(position, paragraph)
            position += 1
        return
    content = toc.find(qn("w:sdtContent"))
    if content is None:
        return
    old_paragraphs = list(content.findall(qn("w:p")))
    title_props = deepcopy(old_paragraphs[0].find(qn("w:pPr"))) \
        if old_paragraphs and old_paragraphs[0].find(qn("w:pPr")) is not None else None
    style_props: Dict[int, Any] = {}
    for index, paragraph in enumerate(old_paragraphs[1:4], start=1):
        props = paragraph.find(qn("w:pPr"))
        if props is not None:
            style_props[index] = deepcopy(props)
    for child in list(content):
        content.remove(child)

    title_paragraph = OxmlElement("w:p")
    if title_props is not None:
        title_paragraph.append(title_props)
    title_paragraph.append(_text_run("目录"))
    content.append(title_paragraph)

    entries = canonical_toc_entries(report_artifact)
    for index, entry in enumerate(entries):
        paragraph = OxmlElement("w:p")
        level = int(entry.get("level") or 1)
        source_props = style_props.get(level)
        if source_props is None:
            source_props = style_props.get(1)
        props = deepcopy(source_props) if source_props is not None else None
        if props is not None:
            paragraph.append(props)
        if index == 0:
            paragraph.extend([
                _field_run("begin"),
                _field_run("instruction", ' TOC \\o "1-3" \\h \\z \\u '),
                _field_run("separate"),
            ])
        paragraph.append(_text_run(str(entry.get("title") or "")))
        if index == len(entries) - 1:
            paragraph.append(_field_run("end"))
        content.append(paragraph)


def _update_anchor_lineage(
    anchors: Dict[str, Paragraph], heading_id: str, paragraph: Paragraph,
) -> None:
    parts = [part for part in str(heading_id or "").split(".") if part]
    for end in range(1, len(parts) + 1):
        anchors[".".join(parts[:end])] = paragraph


def _section_content(report_artifact: Mapping[str, Any], section_id: str) -> Dict[str, Any]:
    for section in report_artifact.get("sections") or []:
        if str(section.get("section_id")) == str(section_id):
            return dict(section)
    return {}


def render_report_docx(
    report_artifact: Mapping[str, Any],
    template_bytes: bytes,
    template_contract: Mapping[str, Any],
) -> io.BytesIO:
    """Render a structured report into a copy of the user's DOCX template.

    The template remains the document authority: its sections, styles,
    headers/footers, and existing heading paragraphs are retained.  Generated
    prose is inserted beneath matching canonical headings.
    """
    if not template_bytes:
        raise TemplateParseError("模板文件为空。")
    document = Document(io.BytesIO(template_bytes))
    chapters = ((template_contract.get("document_structure") or {}).get("chapters")
                or [])
    top_level = int(((template_contract.get("document_structure") or {}).get(
        "top_level") or 1))
    body_titles = {_title_key(x.get("title")) for x in chapters}
    chapter_paragraphs = [paragraph for paragraph in document.paragraphs
                          if _title_key(paragraph.text) in body_titles]
    if not chapter_paragraphs:
        # Do not silently rebuild a structurally incompatible template.
        raise TemplateParseError("模板正文标题与 Template Contract 不一致，无法安全填充。")
    back_items = ((template_contract.get("document_structure") or {}).get("back_matter")
                  or [])
    back_titles = {
        _title_key(x.get("title"))
        for x in back_items
    }
    preserve = list(chapter_paragraphs)
    for chapter in chapters:
        for required in chapter.get("required_subsections") or []:
            found = _find_heading(document, required.get("title"))
            if found is not None and found not in preserve:
                preserve.append(found)
    for item in back_items:
        found = _find_heading(document, item.get("title"))
        if found is not None and found not in preserve:
            preserve.append(found)
    _remove_body_examples(document, preserve, chapter_paragraphs[0])

    case_labels = report_artifact.get("case_labels") or {}
    case_types = report_artifact.get("case_types") or {}

    for chapter in chapters:
        section_id = str(chapter.get("section_id"))
        heading = _find_heading(document, chapter.get("title"))
        if heading is None:
            continue
        section = _section_content(report_artifact, section_id)
        if section.get("title"):
            heading.text = canonical_chapter_title(section_id, section["title"])
        # A real MTI template in the supported corpus marks Chapter 4 as
        # Heading 2.  The parser promotes it logically; the final surface must
        # promote its Word style as well so a refreshed TOC keeps four peers.
        try:
            heading.style = "Heading 1"
        except KeyError:
            pass
        section_subsections = list(section.get("subsections") or [])
        subsections = {_title_key(x.get("title")): x for x in section_subsections}
        subsections_by_id = {str(x.get("heading_id") or ""): x
                             for x in section_subsections if x.get("heading_id")}
        if section_subsections:
            chapter_content = section.get("intro_content") or ""
            chapter_anchor = heading
            if chapter_content:
                chapter_anchor = _insert_content_after(
                    heading, public_report_markdown(
                        chapter_content, case_labels, case_types))
            anchor_by_id: Dict[str, Paragraph] = {section_id: chapter_anchor}
            required_ids = set()
            for required in chapter.get("required_subsections") or []:
                sub_heading = _find_heading(document, required.get("title"))
                if sub_heading is None:
                    continue
                heading_id = str(required.get("heading_id") or "")
                required_ids.add(heading_id)
                item = subsections_by_id.get(heading_id) or subsections.get(
                    _title_key(required.get("title"))) or {}
                last = _insert_content_after(
                    sub_heading, public_report_markdown(
                        item.get("content") or "", case_labels, case_types))
                _update_anchor_lineage(anchor_by_id, heading_id, last)
            for item in section_subsections:
                heading_id = str(item.get("heading_id") or "")
                if not heading_id or heading_id in required_ids:
                    continue
                parent_id = heading_id.rsplit(".", 1)[0] if "." in heading_id else ""
                anchor = anchor_by_id.get(parent_id)
                if anchor is None:
                    continue
                level = max(1, min(3, int(item.get("level") or 4) - 1))
                dynamic_heading = _insert_paragraph_after(
                    anchor, f"{heading_id} {item.get('title')}", f"Heading {level}")
                last = _insert_content_after(
                    dynamic_heading, public_report_markdown(
                        item.get("content") or "", case_labels, case_types))
                _update_anchor_lineage(anchor_by_id, heading_id, last)
        else:
            _insert_content_after(heading, public_report_markdown(
                section.get("content") or "", case_labels, case_types))

    matter = {str(x.get("role")): x for x in report_artifact.get("front_matter") or []}
    front_contract = (template_contract.get("document_structure") or {}).get(
        "front_matter") or []
    abstract_contracts = [x for x in front_contract
                          if x.get("role") in {"abstract_zh", "abstract_en"}]
    abstract_anchors = [_find_heading(document, x.get("source_title") or x.get("title"))
                        for x in abstract_contracts]
    toc_anchor = next((paragraph for paragraph in document.paragraphs
                       if _title_key(paragraph.text) == _title_key("目录")
                       or "toc heading" in _paragraph_style_name(
                           paragraph).casefold()
                       or ("TOC" in paragraph._p.xml and
                           ("instrText" in paragraph._p.xml or
                            "fldChar" in paragraph._p.xml))), None)
    for index, contract_item in enumerate(abstract_contracts):
        anchor = abstract_anchors[index]
        if anchor is None:
            continue
        next_anchor = abstract_anchors[index + 1] if index + 1 < len(
            abstract_anchors) else (toc_anchor or chapter_paragraphs[0])
        _clear_between(anchor, next_anchor)
        role = str(contract_item.get("role"))
        item = matter.get(role) or {}
        anchor.text = "摘  要" if role == "abstract_zh" else "ABSTRACT"
        content = str(item.get("content") or "需要用户补充。")
        if role == "abstract_en":
            report = report_artifact.get("report") or {}
            content = str(report.get("abstract_en") or content).strip()
            if not content or re.search(r"[\u3400-\u9fff]", content):
                content = (
                    "This report presents the translation project and its case "
                    "analysis based on the documented source and target texts.")
        keywords_role = "keywords_zh" if role == "abstract_zh" else "keywords_en"
        keywords = (matter.get(keywords_role) or {}).get("keywords") or []
        if role == "abstract_en" and (
                not keywords or any(re.search(r"[\u3400-\u9fff]", str(x))
                                    for x in keywords)):
            keywords = ["translation practice", "case analysis"]
        label = "关键词" if role == "abstract_zh" else "Keywords"
        _insert_content_after(anchor, content + "\n\n" + label + "：" +
                              "，".join(str(x) for x in keywords))

    back_matter = {str(x.get("role")): x for x in report_artifact.get("back_matter") or []}
    for contract_item in back_items:
        heading = _find_heading(document, contract_item.get("title"))
        if heading is None:
            continue
        role = str(contract_item.get("role") or "")
        candidates = [x for x in report_artifact.get("back_matter") or []
                      if x.get("role") == role]
        item = next((x for x in candidates if str(x.get("section_id") or "") ==
                     str(contract_item.get("section_id") or "")), None)
        item = item or next((x for x in candidates if _title_key(x.get("title")) ==
                             _title_key(contract_item.get("title"))), None)
        item = item or back_matter.get(role) or {}
        if item.get("title"):
            heading.text = _surface_matter_title(
                item.get("title"), report_artifact.get("project_title"))
        content = item.get("content") or "需要用户补充。"
        _insert_content_after(heading, public_report_markdown(
            str(content), case_labels, case_types))

    project_title = str(report_artifact.get("project_title") or "")
    _replace_visible_text(document, {
        "《XXX》": f"《{project_title}》" if project_title else "《当前翻译项目》",
    })
    _fill_cover(document, project_title)
    _refresh_toc_cache(document, report_artifact)
    _set_update_fields(document)
    out = io.BytesIO()
    document.save(out)
    out.seek(0)
    return out
