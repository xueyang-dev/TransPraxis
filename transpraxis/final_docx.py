"""Post-render integrity checks for the user-facing academic DOCX."""
from __future__ import annotations

import hashlib
import io
import json
import re
from typing import Any, Dict, List, Mapping

from docx import Document
from docx.oxml.ns import qn

from . import case_presentation, case_provenance, report_template


SCHEMA_VERSION = "final-docx-validation-v2"
_CJK = re.compile(r"[\u3400-\u9fff]")
_CASE_NUMBER = re.compile(r"例\s*[\[（(](\d+)[\]）)]")
_INTERNAL_ID = re.compile(
    r"\b(?:TD-\d+|SC-\d+|seg-[A-Za-z0-9_.:-]+|(?:AQ|AV|AR|LR)-\d+)\b",
    re.IGNORECASE,
)
_INTERNAL_LANGUAGE = re.compile(
    r"(?:segment\s*id|validation\s*gate|candidate\s*pool|provenance|"
    r"修订资格门禁|统计门禁|候选池|证据边界|有界结论)", re.IGNORECASE)
_SAMPLE_TITLES = (
    "文化负载词处理", "多模态文本协调", "机器翻译误译问题处理",
    "文化负载词翻译策略", "多模态文本协调方案", "机器翻译误译率较高的解决方案",
)
_TEMPLATE_INSTRUCTION = re.compile(r"[（(]如果有\s*[，,、]?\s*另起一页[）)]")


def _norm(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _all_xml_text(document) -> str:
    return "\n".join(str(node.text or "")
                     for node in document.element.xpath(".//w:t"))


def _toc_lines(document) -> List[str]:
    toc = next((node for node in document.element.xpath(".//w:sdt")
                if "TOC" in " ".join(str(item.text or "")
                                      for item in node.iter(
                                          qn("w:instrText")))), None)
    if toc is None:
        field_paragraphs = [paragraph for paragraph in document.element.xpath(".//w:p")
                            if "TOC" in " ".join(str(item.text or "")
                                                 for item in paragraph.iter(
                                                     qn("w:instrText")))]
        if not field_paragraphs:
            return []
        field_paragraph = field_paragraphs[0]
        parent = field_paragraph.getparent()
        siblings = list(parent)
        start = siblings.index(field_paragraph)
        lines = []
        for paragraph in siblings[max(0, start - 1):]:
            if paragraph.tag != qn("w:p"):
                continue
            text = _norm("".join(str(node.text or "") for node in paragraph.iter(
                qn("w:t"))))
            if text:
                lines.append(text)
            if any(str(node.get(
                    qn("w:fldCharType")) or "") == "end"
                   for node in paragraph.iter(
                       qn("w:fldChar"))):
                break
        return lines
    content = toc.find(qn("w:sdtContent"))
    if content is None:
        return []
    lines = []
    for paragraph in content.findall(qn("w:p")):
        text = _norm("".join(str(node.text or "")
                             for node in paragraph.iter(qn("w:t"))))
        if text:
            lines.append(text)
    return lines


def _direct_toc_elements(document) -> set[Any]:
    """Return direct-body TOC cache paragraphs so body checks can ignore them."""
    paragraphs = list(document.paragraphs)
    field_index = next((index for index, paragraph in enumerate(paragraphs)
                        if "TOC" in " ".join(str(item.text or "")
                                             for item in paragraph._p.iter(
                                                 qn("w:instrText")))), None)
    if field_index is None:
        return set()
    start = field_index - 1 if field_index and _norm(
        paragraphs[field_index - 1].text) == "目录" else field_index
    elements = set()
    for paragraph in paragraphs[start:]:
        elements.add(paragraph._p)
        if any(str(node.get(
                qn("w:fldCharType")) or "") == "end"
               for node in paragraph._p.iter(
                   qn("w:fldChar"))):
            break
    return elements


def _expected_chapter_titles(report_artifact: Mapping[str, Any]) -> List[str]:
    return [report_template.canonical_chapter_title(
        section.get("section_id"), section.get("title"))
        for section in report_artifact.get("sections") or []]


def _expected_toc_titles(report_artifact: Mapping[str, Any]) -> List[str]:
    return [_norm(item.get("title"))
            for item in report_template.canonical_toc_entries(report_artifact)]


def _issue(kind: str, reason: str, severity: str = "error", **extra: Any) -> Dict[str, Any]:
    return {"type": kind, "severity": severity, "reason": reason, **extra}


def validate_final_docx(
    docx_bytes: bytes, report_artifact: Mapping[str, Any],
) -> Dict[str, Any]:
    """Validate the rendered DOCX, not merely its pre-render artifact."""
    document = Document(io.BytesIO(docx_bytes))
    excluded_toc = _direct_toc_elements(document)
    paragraphs = [paragraph for paragraph in document.paragraphs
                  if paragraph._p not in excluded_toc]
    texts = [_norm(paragraph.text) for paragraph in paragraphs]
    visible_text = _all_xml_text(document)
    issues: List[Dict[str, Any]] = []

    chapter_titles = _expected_chapter_titles(report_artifact)
    back_titles = [report_template._surface_matter_title(
        item.get("title"), report_artifact.get("project_title"))
                   for item in report_artifact.get("back_matter") or []]
    top_titles = [*chapter_titles, *back_titles]
    positions: Dict[str, int] = {}
    for title in top_titles:
        try:
            positions[title] = texts.index(_norm(title))
        except ValueError:
            positions[title] = -1

    for index, title in enumerate(chapter_titles):
        start = positions.get(title, -1)
        if start < 0:
            issues.append(_issue(
                "docx_missing_chapter", f"最终 DOCX 缺少正文标题：{title}",
                chapter=index + 1))
            continue
        next_positions = [value for current, value in positions.items()
                          if current != title and value > start]
        end = min(next_positions) if next_positions else len(paragraphs)
        body = [text for text in texts[start + 1:end]
                if text and text not in top_titles]
        if not body:
            issues.append(_issue(
                "docx_empty_chapter", f"最终 DOCX 的 {title} 只有标题、没有正文。",
                chapter=index + 1))
        if paragraphs[start].style.name != "Heading 1":
            issues.append(_issue(
                "docx_malformed_chapter_heading",
                f"{title} 使用 {paragraphs[start].style.name}，不是 Heading 1。",
                chapter=index + 1))

    toc_lines = _toc_lines(document)
    toc_text = "\n".join(toc_lines)
    expected_toc = _expected_toc_titles(report_artifact)
    if not toc_lines:
        issues.append(_issue("docx_toc_missing", "最终 DOCX 没有可见目录缓存。"))
    else:
        cursor = -1
        for title in expected_toc:
            location = toc_text.find(title, cursor + 1)
            if location < 0:
                issues.append(_issue(
                    "docx_toc_title_mismatch", f"目录缺少 canonical 标题：{title}"))
            else:
                cursor = location
        stale = [title for title in _SAMPLE_TITLES if title in toc_text]
        if stale:
            issues.append(_issue(
                "docx_toc_template_cache_stale",
                "目录仍含模板示例标题：" + "、".join(stale)))

    placeholders = sorted(set(re.findall(r"《XXX》|\bXXX\b", visible_text)))
    if placeholders:
        issues.append(_issue(
            "docx_project_placeholder_visible",
            "最终 DOCX 仍含项目占位符：" + "、".join(placeholders)))
    if _TEMPLATE_INSTRUCTION.search(visible_text):
        issues.append(_issue(
            "docx_template_instruction_visible",
            "最终 DOCX 标题中仍含“如果有，另起一页”等模板编写说明。"))

    try:
        abstract_start = texts.index("ABSTRACT")
    except ValueError:
        abstract_start = -1
    first_chapter = min((positions.get(title, -1) for title in chapter_titles
                         if positions.get(title, -1) >= 0), default=len(texts))
    abstract_text = " ".join(text for text in texts[abstract_start + 1:first_chapter]
                             if text) if abstract_start >= 0 else ""
    if abstract_start < 0 or not abstract_text:
        issues.append(_issue("docx_english_abstract_missing", "ABSTRACT 缺少英文正文。"))
    elif _CJK.search(abstract_text):
        issues.append(_issue(
            "docx_english_abstract_contains_chinese",
            "ABSTRACT 或 English Keywords 中仍含中文字段值。"))

    reference_title = next((title for title in back_titles if "参考文献" in title), "")
    reference_start = positions.get(reference_title, -1)
    if reference_start < 0:
        issues.append(_issue("docx_references_missing", "最终 DOCX 缺少参考文献标题。"))
    else:
        next_positions = [value for value in positions.values() if value > reference_start]
        reference_end = min(next_positions) if next_positions else len(texts)
        if not any(texts[reference_start + 1:reference_end]):
            issues.append(_issue("docx_references_empty", "参考文献标题后没有任何条目。"))

    numbers = [int(match.group(1)) for text in texts
               for match in [_CASE_NUMBER.fullmatch(text)] if match]
    expected_numbers = list(range(1, len(numbers) + 1))
    if numbers != expected_numbers:
        issues.append(_issue(
            "docx_case_numbering_not_continuous",
            f"案例编号为 {numbers}，应为 {expected_numbers}。"))

    final_policy = report_artifact.get("case_policy") or {}
    final_contract = report_artifact.get("report_stage") == "final_report" \
        and bool(final_policy.get("contrast_required"))
    final_case_nodes = [node for node in report_artifact.get("case_nodes") or []
                        if node.get("type") == "case_example"]
    case_blocks = []
    for index, text in enumerate(texts):
        if not _CASE_NUMBER.fullmatch(text):
            continue
        end = next((position for position in range(index + 1, len(texts))
                    if _CASE_NUMBER.fullmatch(texts[position])), len(texts))
        case_blocks.append((text, "\n".join(texts[index:end])))
    synthetic_label_count = sum(bool(re.search(
        r"^\s*(?:模拟初译|对比译法（模拟）)\s*[：:]", block, re.MULTILINE))
        for _heading, block in case_blocks)
    authentic_initial_label_count = sum(bool(re.search(
        r"^\s*初译\s*[：:]", block, re.MULTILINE))
        for _heading, block in case_blocks)
    rewrite_label_count = sum(bool(re.search(
        r"^\s*改译\s*[：:]", block, re.MULTILINE))
        for _heading, block in case_blocks)
    decision_only_count = sum(bool(re.search(
        r"^\s*译文\s*[：:]", block, re.MULTILINE)) and not re.search(
            r"^\s*(?:初译|模拟初译|对比译法（模拟）)\s*[：:]", block,
            re.MULTILINE) for _heading, block in case_blocks)
    countable_case_count = len(numbers)
    if final_contract:
        minimum = int(final_policy.get("minimum_cases") or 20)
        if final_case_nodes:
            countable_case_count = sum(
                case_provenance.counts_toward_minimum(node, final_policy)
                for node in final_case_nodes)
        else:
            type_map = report_artifact.get("case_types") or {}
            countable_case_count = sum(
                case_provenance.counts_toward_minimum({"case_type": case_type}, final_policy)
                for case_type in type_map.values()) if type_map else len(numbers)
        if countable_case_count < minimum:
            issues.append(_issue(
                "docx_final_case_count_below_minimum",
                f"最终 DOCX 只有 {countable_case_count} 个可计数正式案例，"
                f"总案例 {len(numbers)} 个，至少需要 {minimum} 个。"))
        if final_case_nodes and len(numbers) != len(final_case_nodes):
            issues.append(_issue(
                "docx_final_case_node_count_mismatch",
                f"DOCX 可见案例数 {len(numbers)} 与 structured case node 数"
                f" {len(final_case_nodes)} 不一致。"))
        if decision_only_count:
            issues.append(_issue(
                "docx_translation_decision_only_case",
                f"最终 DOCX 有 {decision_only_count} 个案例只有‘译文’而没有翻译对比。"))
        if rewrite_label_count != len(numbers):
            issues.append(_issue(
                "docx_case_rewrite_label_mismatch",
                f"最终 DOCX 只有 {rewrite_label_count}/{len(numbers)} 个案例含‘改译’。"))
        if synthetic_label_count + authentic_initial_label_count != len(numbers):
            issues.append(_issue(
                "docx_case_contrast_label_mismatch",
                "最终 DOCX 每个案例必须包含‘初译’或明确的‘模拟初译’字段。"))

    internal_ids = sorted(set(_INTERNAL_ID.findall(visible_text)))
    internal_terms = sorted(set(match.group(0) for match in
                                _INTERNAL_LANGUAGE.finditer(visible_text)))
    if internal_ids or internal_terms:
        issues.append(_issue(
            "docx_internal_language_visible",
            "最终 DOCX 泄漏内部标记：" + "、".join([*internal_ids, *internal_terms])))

    analysis_paragraphs = [value for value in texts if value.startswith("分析：")]
    if len(analysis_paragraphs) != len(numbers):
        issues.append(_issue(
            "docx_case_analysis_count_mismatch",
            f"最终 DOCX 含 {len(numbers)} 个案例，但只有 "
            f"{len(analysis_paragraphs)} 个可见‘分析’段。"))
    analyses = []
    for number, text in zip(numbers, analysis_paragraphs):
        analyses.append({
            "case_id": f"surface-case-{number}",
            "example_number": number,
            "analysis": text.split("：", 1)[1].strip(),
        })
    repetition = case_presentation.analysis_repetition_audit(analyses)
    if repetition.get("status") == "fail":
        issues.append(_issue(
            "docx_case_analysis_boilerplate",
            "最终 DOCX 的案例分析存在明显重复块或高相似模板。"))
    elif repetition.get("status") == "pass_with_warnings":
        issues.append(_issue(
            "docx_case_analysis_repetition_warning",
            "最终 DOCX 的案例分析仍有少量相似句，需要人工确认。",
            severity="warning"))

    literature_status = str((report_artifact.get("report") or {}).get(
        "literature_status") or report_artifact.get("literature_status") or "")
    if literature_status != "complete":
        issues.append(_issue(
            "docx_literature_support_required",
            "正文结构与案例已形成，但学术文献支持尚未建立。",
            severity="warning"))

    errors = sum(item["severity"] == "error" for item in issues)
    warnings = sum(item["severity"] == "warning" for item in issues)
    status = "fail" if errors else ("pass_with_warnings" if warnings else "pass")
    result = {
        "schema_version": SCHEMA_VERSION,
        "status": status,
        "summary": {
            "errors": errors,
            "warnings": warnings,
            "chapter_count": len(chapter_titles),
            "nonempty_chapter_count": len(chapter_titles) - sum(
                item["type"] in {"docx_missing_chapter", "docx_empty_chapter"}
                for item in issues),
            "toc_entry_count": max(0, len(toc_lines) - 1),
            "expected_toc_entry_count": len(expected_toc),
            "project_placeholder_count": len(placeholders),
            "abstract_chinese_character_count": len(_CJK.findall(abstract_text)),
            "case_count": len(numbers),
            "countable_case_count": countable_case_count,
            "synthetic_label_count": synthetic_label_count,
            "authentic_initial_label_count": authentic_initial_label_count,
            "rewrite_label_count": rewrite_label_count,
            "translation_decision_visible_count": decision_only_count,
            "internal_marker_count": len(internal_ids) + len(internal_terms),
            "literature_status": literature_status or "unknown",
        },
        "issues": issues,
        "analysis_repetition": repetition,
        "toc": {"expected": expected_toc, "actual": toc_lines},
    }
    result["content_hash"] = hashlib.sha256(json.dumps(
        result, ensure_ascii=False, sort_keys=True,
        separators=(",", ":")).encode("utf-8")).hexdigest()
    return result
