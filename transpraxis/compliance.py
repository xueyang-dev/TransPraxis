"""Source-backed compliance and deterministic language constraints.

This is deliberately not a rules DSL.  The default profile is an anonymous
reference architecture distilled from real MTI practice, not an institution's
published policy.  Institution-specific requirements belong in future custom
profiles supplied by the user.
"""
from __future__ import annotations

import io
import json
import re
from copy import deepcopy
from pathlib import Path
from typing import Any, Dict, Iterable, List, Mapping

from docx import Document
from docx.oxml.ns import qn

from . import case_provenance

VERSION = "compliance-result-v1"
DEFAULT_PROFILE_ID = "MTI_PRACTICE_REPORT_DEFAULT"
REFERENCE_SOURCE_ID = "mti_practice_report_reference_v1"
SOURCE_ROOT = Path(__file__).resolve().parents[1]
PROJECT_TRACE_SOURCE = "docs/mti-practice-driven-roadmap.md"
CUSTOM_PROFILE_SOURCE_UNAVAILABLE = "院校自定义规范（未配置）"
PROJECT_SOURCE_TRACE_ONLY = "项目实现追踪（非规范来源）"
SOURCE_RECORDS = {
    REFERENCE_SOURCE_ID: {
        "source_id": REFERENCE_SOURCE_ID,
        "title": "默认 MTI 实践报告参考结构",
        "source_type": "reference_template",
        "authority_level": "reference_template",
        "scope": "default_profile",
        "source_date": "2026",
        "source_version": "v1",
        "source_url": "",
        "document_url": "",
        "page_or_clause": "default_profile",
        "source_summary": (
            "基于匿名真实 MTI 实践样本抽象出的默认报告结构和确定性检查；"
            "用于产品默认 profile，不代表全国统一或院校强制要求。"
        ),
        "source_file_present": False,
        "source_verified": True,
        "original_retained_outside_repo": False,
    },
}
AUTHORITY_RANK = {
    "custom_profile": 30,
    "reference_template": 10,
    "project": 0,
}
ENFORCEMENT = {"enforced", "manual_review", "recommended", "project_constraint"}


def _rule(rule_id, category, description, *, authority_level="project",
          source_document=None, source_date="", page_or_clause="", excerpt="",
          source_id=None, source_verified=False, source_url="",
          implementation_source=PROJECT_TRACE_SOURCE, implementation_clause="",
          scope="report", check_type="deterministic",
          enforcement="project_constraint", expected="", conflicts_with=(),
          supersedes=()):
    source_record = SOURCE_RECORDS.get(str(source_id)) if source_id else None
    if source_record:
        source_document = source_document or source_record["title"]
        source_date = source_date or source_record.get("source_date", "")
        page_or_clause = page_or_clause or source_record.get("page_or_clause", "")
        excerpt = excerpt or source_record.get("source_summary", "")
        source_url = source_url or source_record.get("source_url", "")
        source_verified = bool(source_verified or source_record.get("source_verified"))

    # The roadmap explains implementation only.  Reference-template and future
    # custom-profile rules need a reliable structured source mapping before
    # they may be enforced.
    if source_document == PROJECT_TRACE_SOURCE:
        source_document = (PROJECT_SOURCE_TRACE_ONLY if authority_level == "project"
                           else CUSTOM_PROFILE_SOURCE_UNAVAILABLE)
        if authority_level != "project":
            page_or_clause = "待提供"
    if (source_record and authority_level != "project" and
            not source_verified and not source_record.get("source_file_present")):
        source_document = CUSTOM_PROFILE_SOURCE_UNAVAILABLE
        page_or_clause = "待提供"
        source_url = ""
    if not source_document:
        source_document = (PROJECT_SOURCE_TRACE_ONLY if authority_level == "project"
                           else CUSTOM_PROFILE_SOURCE_UNAVAILABLE)
    if not page_or_clause:
        page_or_clause = ("待提供" if authority_level != "project" else
                          implementation_clause or "implementation trace")
    path = SOURCE_ROOT / str(source_document) if source_document else None
    source_file_present = bool(
        source_record.get("source_file_present") if source_record else
        path and path.is_file()
    )
    available = bool(source_file_present and page_or_clause)
    structured_mapping = bool(
        source_record and source_verified and page_or_clause and excerpt
    )
    reliable_mapping = bool(available or structured_mapping)
    if authority_level != "project" and enforcement == "enforced" \
            and not reliable_mapping:
        enforcement = "manual_review"
    source_kind = (str((source_record or {}).get("source_type") or "custom_profile")
                   if authority_level != "project" and reliable_mapping
                   else "unmapped_authority" if authority_level != "project"
                   else "implementation_trace")
    return {
        "rule_id": rule_id, "category": category, "description": description,
        "authority_level": authority_level,
        "source_id": source_id,
        "source_type": (source_record or {}).get("source_type"),
        "source_scope": (source_record or {}).get("scope"),
        "source_document": source_document, "source_date": source_date,
        "source_url": source_url,
        "page_or_clause": page_or_clause,
        "source_excerpt_or_summary": excerpt,
        "source_available": available,
        "source_file_present": source_file_present,
        "source_recorded": bool(source_record),
        "reliable_source_mapping": reliable_mapping,
        "source_kind": source_kind,
        "implementation_source": implementation_source,
        "implementation_clause": implementation_clause,
        "scope": scope, "check_type": check_type,
        "severity": "error" if enforcement == "enforced" else "warning",
        "enforcement": enforcement,
        "expected": expected, "actual": None, "status": "not_checked",
        "location": "—", "message": "",
        "conflicts_with": list(conflicts_with),
        "supersedes": list(supersedes),
    }


def compliance_profile(profile_id: str = DEFAULT_PROFILE_ID) -> Dict[str, Any]:
    """Return the anonymous default MTI practice-report profile."""
    profile_id = DEFAULT_PROFILE_ID
    roadmap = PROJECT_TRACE_SOURCE
    rules = [
        _rule("abstract_zh_length", "abstract",
              "中文摘要为 400—600 个汉字。",
              authority_level="reference_template",
              source_id=REFERENCE_SOURCE_ID,
              page_or_clause="default_profile.abstract_zh",
              excerpt="默认参考范围：中文摘要400—600个汉字。",
              implementation_clause="Stage 4 / abstract",
              expected="400-600 CJK chars", enforcement="enforced"),
        _rule("keywords_count", "keywords",
              "中英文关键词各有对应项，数量为 5—8 个并使用逗号分隔。",
              authority_level="reference_template",
              source_id=REFERENCE_SOURCE_ID,
              page_or_clause="default_profile.keywords",
              excerpt="默认参考范围：中英文关键词5—8个，使用一致分隔符。",
              implementation_clause="Stage 4 / keywords",
              expected="5-8 each; comma-separated", enforcement="enforced"),
        _rule("toc_depth", "toc", "目录可见层级最多三级。",
              authority_level="reference_template",
              source_id=REFERENCE_SOURCE_ID,
              page_or_clause="default_profile.toc",
              excerpt="默认目录展示不超过三级标题。",
              implementation_clause="Stage 4 / toc", expected="<=3 levels",
              enforcement="enforced"),
        _rule("citation_reference_bidirectional", "citation",
              "正文引用与参考文献双向对应；无重复或缺失 ID。",
              authority_level="reference_template",
              source_id=REFERENCE_SOURCE_ID,
              page_or_clause="default_profile.references",
              excerpt="默认 profile 要求正文引用与参考文献双向对应。",
              implementation_clause="Stage 4 / citation",
              expected="references are cited and citations have references",
              enforcement="enforced"),
        _rule("citation_style", "citation",
              "引用格式由用户模板或所在院校要求决定。",
              authority_level="reference_template",
              source_id=REFERENCE_SOURCE_ID,
              implementation_clause="Stage 4 / configurable citation style",
              page_or_clause="default_profile.citation_style",
              excerpt="不同培养单位可能采用脚注、顺序编码或其他格式，默认 profile 不裁决。",
              check_type="manual", enforcement="manual_review",
              conflicts_with=[]),
        _rule("figure_table_numbering", "figures",
              "图表使用章内编号，captions 存在，无重复或明显跳号。",
              authority_level="reference_template",
              source_id=REFERENCE_SOURCE_ID,
              page_or_clause="default_profile.figures_tables",
              excerpt="默认 profile 使用章号.序号并要求图题、表题。",
              implementation_clause="Stage 4 / figures",
              expected="图/表 chapter.item with caption", enforcement="enforced"),
        _rule("bilingual_appendix", "appendix",
              "双语附录存在且 source/translation 角色明确。",
              authority_level="reference_template",
              source_id=REFERENCE_SOURCE_ID,
              page_or_clause="default_profile.appendix",
              excerpt="默认 profile 使用原文与译文角色明确的双语附录。",
              implementation_clause="Stage 4 / appendix",
              expected="source + translation appendix", enforcement="enforced"),
        _rule("source_length", "source", "原文长度按语源规则核对。",
              authority_level="reference_template",
              source_id=REFERENCE_SOURCE_ID,
              page_or_clause="default_profile.source_length",
              excerpt="中文源文可按汉字数检查；英文源文折算需根据所在院校要求确认。",
              implementation_clause="Stage 4 / source length",
              expected="Chinese source >= 10000 Han characters; English conversion manual_review",
              enforcement="manual_review"),
        _rule("case_conclusion_structure", "structure",
              "案例分析报告包含任务描述、任务过程、案例分析和实践总结；充分性仍需人工复核。",
              authority_level="reference_template",
              source_id=REFERENCE_SOURCE_ID,
              page_or_clause="default_profile.report_structure",
              excerpt="默认结构包括引言、项目概述、案例分析和总结反思。",
              implementation_clause="Stage 4 / structure",
              expected="task description + task process + case analysis + practice summary",
              enforcement="enforced"),
        _rule("docx_layout", "layout",
              "DOCX 使用 A4、规定页边距/页眉页脚距离和正文固定 20 磅行距。",
              authority_level="reference_template",
              source_id=REFERENCE_SOURCE_ID,
              page_or_clause="default_profile.docx_layout",
              excerpt=("默认参考版式使用A4、固定20磅行距及一组可确定性检查的页面距离；"
                       "院校特殊封面、页眉和目录要求需另行确认。"),
              implementation_clause="Stage 4 / DOCX layout",
              expected={
                  "paper_cm": {"width": 21.0, "height": 29.7},
                  "margins_cm": {"top": 3.3, "bottom": 3.3,
                                 "left": 3.0, "right": 2.8,
                                 "header": 2.6, "footer": 2.6},
                  "line_spacing": {"rule": "exact", "points": 20},
              }, enforcement="enforced"),
        _rule("synthetic_case_policy", "case",
              "synthetic 案例的项目严格策略与学校计数要求分开显示。",
              authority_level="reference_template",
              source_id=REFERENCE_SOURCE_ID,
              page_or_clause="default_profile.synthetic_policy",
              implementation_clause="Stage 4 / synthetic policy",
              excerpt="synthetic case 是否计入正式案例数量需根据所在院校要求确认。",
              enforcement="manual_review"),
        _rule("author_placeholders", "manual",
              "作者、导师、日期等占位符必须解决。",
              implementation_clause="Stage 4 / placeholders",
              excerpt="未完成信息进入 manual review。"),
    ]
    return {
        "schema_version": VERSION,
        "profile_id": DEFAULT_PROFILE_ID,
        "display_name": "默认 MTI 实践报告规范",
        "profile_type": "default_mti_practice_report",
        "program": "MTI 翻译实践报告",
        "effective_date": "2026",
        "authority_order": ["custom_profile", "reference_template", "project"],
        "authority_hierarchy": {
            "priority": ["custom_profile", "reference_template", "project"],
            "resolution_basis": ["specificity", "authority", "date", "manual_review"],
        },
        "authority_mapping_status": "reference_template_mapped",
        "synthetic_policy": {
            "allowed_in_formal_report": "manual_review",
            "counts_toward_minimum": "manual_review",
            "supplementary_only": "manual_review",
            "requires_special_disclosure": "manual_review",
            "source_status": "not_confirmed",
        },
        "sources": [
            {
                **deepcopy(record),
                "document": record["title"],
                "source_kind": record.get("source_type") or "reference_template",
                "available": bool(record.get("source_verified") or
                                  record.get("source_file_present")),
                "note": ("匿名参考结构已登记；不代表全国统一或院校强制要求。"
                          if record.get("source_verified") else
                          record.get("availability_reason") or "来源待核对。"),
            }
            for record in SOURCE_RECORDS.values()
        ],
        "implementation_sources": [{
            "document": roadmap,
            "authority_level": "project",
            "source_kind": "implementation_trace",
            "note": "仅说明项目如何实现检查；不是院校学术规范。",
        }],
        "conflicts": [{
            "conflict_id": "citation_style_requires_profile_choice",
            "rule": "citation_style",
            "options": ["footnote", "numeric_sequence", "author_date",
                        "template_defined"],
            "resolved_by": "manual_review",
            "status": "manual_review",
            "message": "不同培养单位可能采用不同引文格式，请根据实际模板确认。",
        }],
        "rules": rules,
    }


def _cjk_count(value: Any) -> int:
    return len(re.findall(r"[\u3400-\u9fff]", str(value or "")))


def _english_words(value: Any) -> int:
    return len(re.findall(r"\b[A-Za-z][A-Za-z'-]*\b", str(value or "")))


def _keyword_values(value: Any) -> tuple[List[str], bool]:
    if isinstance(value, str):
        parts = [item.strip() for item in re.split(r"[,，;；、\n]", value)
                 if item.strip()]
        return parts, True
    values = list(value or [])
    malformed = any(re.search(r"[,，;；、\n]", str(item))
                    for item in values)
    return [str(item).strip() for item in values if str(item).strip()], not malformed


def _sections(text: str) -> List[Dict[str, Any]]:
    current = {"section_id": "front", "title": "front", "start": 0, "lines": []}
    sections = [current]
    for line in str(text or "").splitlines():
        match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if match:
            current = {"section_id": match.group(2), "title": match.group(2),
                       "start": line, "lines": []}
            sections.append(current)
        current["lines"].append(line)
    return sections


def _toc_titles(value: Any) -> List[str]:
    if isinstance(value, Mapping):
        value = value.get("actual") or value.get("entries") or value.get("titles") or []
    titles = []
    for item in value if isinstance(value, (list, tuple)) else []:
        if isinstance(item, Mapping):
            item = item.get("title") or item.get("text") or ""
        item = str(item or "").strip()
        if item:
            titles.append(item)
    return titles


def _result(rule: Mapping[str, Any], status: str, message: str,
            location: str = "", actual: Any = None) -> Dict[str, Any]:
    enforcement = str(rule.get("enforcement") or "project_constraint")
    if enforcement == "enforced" and not rule.get("reliable_source_mapping",
                                                     rule.get("source_available")):
        enforcement = "manual_review"
    if enforcement == "enforced" and rule.get("authority_level") == "project":
        enforcement = "project_constraint"
    return {
        "rule_id": rule["rule_id"], "category": rule["category"],
        "description": rule["description"], "authority_level": rule["authority_level"],
        "source_id": rule.get("source_id"),
        "source_type": rule.get("source_type"),
        "source_scope": rule.get("source_scope"),
        "source_document": rule["source_document"], "source_date": rule["source_date"],
        "source_url": rule.get("source_url", ""),
        "page_or_clause": rule["page_or_clause"],
        "source_excerpt_or_summary": rule["source_excerpt_or_summary"],
        "source_available": rule["source_available"], "scope": rule["scope"],
        "source_file_present": rule.get("source_file_present", rule.get("source_available")),
        "source_recorded": rule.get("source_recorded", bool(rule.get("source_id"))),
        "source_kind": rule.get("source_kind"),
        "implementation_source": rule.get("implementation_source"),
        "implementation_clause": rule.get("implementation_clause"),
        "reliable_source_mapping": rule.get("reliable_source_mapping",
                                             rule.get("source_available")),
        "check_type": rule["check_type"], "severity": rule["severity"],
        "enforcement": enforcement, "expected": rule["expected"],
        "actual": actual, "status": status,
        "conflicts_with": rule["conflicts_with"], "supersedes": rule["supersedes"],
        "location": location or "—", "message": message,
    }


def inspect_docx_layout(docx_bytes: bytes) -> Dict[str, Any]:
    """Read reliable OOXML layout facts; do not infer school compliance."""
    document = Document(io.BytesIO(docx_bytes))
    section = document.sections[0] if document.sections else None
    emu_per_inch = 914400
    styles = []
    for paragraph in document.paragraphs[:200]:
        if paragraph.text.strip() and paragraph.style.name not in styles:
            styles.append(paragraph.style.name)
    spacings = []
    spacing_details = []
    fonts = []
    for paragraph in document.paragraphs[:300]:
        xml = paragraph._p.xpath(".//w:spacing/@w:line")
        if xml:
            spacings.extend(str(item) for item in xml)
        for spacing in paragraph._p.xpath(".//w:spacing"):
            line = spacing.get(qn("w:line"))
            if line is not None:
                spacing_details.append({
                    "line": str(line),
                    "line_rule": str(spacing.get(qn("w:lineRule")) or ""),
                })
        for run in paragraph.runs[:4]:
            names = run._r.xpath(".//w:rFonts/@w:eastAsia|.//w:rFonts/@w:ascii")
            sizes = run._r.xpath(".//w:sz/@w:val")
            fonts.extend(str(item) for item in names)
            fonts.extend(str(float(item) / 2) for item in sizes)
    page_number = bool(document.element.xpath(
        ".//w:instrText[contains(., 'PAGE')]"))
    return {
        "sections": len(document.sections),
        "page_width_emu": section.page_width if section else None,
        "page_height_emu": section.page_height if section else None,
        "page_width_inches": round(section.page_width / emu_per_inch, 3)
        if section else None,
        "page_height_inches": round(section.page_height / emu_per_inch, 3)
        if section else None,
        "a4_detected": bool(section and
                            abs(section.page_width / emu_per_inch - 8.27) < .03 and
                            abs(section.page_height / emu_per_inch - 11.69) < .03),
        "margins_emu": {
            "top": section.top_margin, "bottom": section.bottom_margin,
            "left": section.left_margin, "right": section.right_margin,
            "header": section.header_distance, "footer": section.footer_distance,
        } if section else {},
        "margins_inches": {
            "top": round(section.top_margin / emu_per_inch, 3),
            "bottom": round(section.bottom_margin / emu_per_inch, 3),
            "left": round(section.left_margin / emu_per_inch, 3),
            "right": round(section.right_margin / emu_per_inch, 3),
            "header": round(section.header_distance / emu_per_inch, 3),
            "footer": round(section.footer_distance / emu_per_inch, 3),
        } if section else {},
        "paragraph_styles": styles[:20], "line_spacings": sorted(set(spacings))[:20],
        "line_spacing_details": spacing_details[:100],
        "fonts": sorted(set(fonts))[:30], "page_number_field": page_number,
    }


def _evaluate_docx_layout(layout: Mapping[str, Any], expected: Mapping[str, Any]):
    """Compare only the source-backed layout facts; absent facts stay manual."""
    if not layout:
        return "not_applicable", "当前没有 DOCX bytes。", layout
    missing = []
    mismatches = []

    paper = expected.get("paper_cm") or {}
    width = layout.get("page_width_inches")
    height = layout.get("page_height_inches")
    if width is None or height is None:
        missing.append("paper size")
    else:
        actual_cm = {"width": width * 2.54, "height": height * 2.54}
        for key, wanted in paper.items():
            if abs(actual_cm[key] - float(wanted)) > 0.08:
                mismatches.append(f"paper {key}={actual_cm[key]:.2f}cm")

    margins = layout.get("margins_inches") or {}
    wanted_margins = expected.get("margins_cm") or {}
    for key, wanted in wanted_margins.items():
        if key not in margins or margins[key] is None:
            missing.append(f"{key} margin")
        elif abs(float(margins[key]) * 2.54 - float(wanted)) > 0.08:
            mismatches.append(f"{key}={float(margins[key]) * 2.54:.2f}cm")

    line_expected = expected.get("line_spacing") or {}
    details = list(layout.get("line_spacing_details") or [])
    if not details:
        missing.append("fixed line spacing")
    elif not any(
            item.get("line") == str(int(float(line_expected.get("points", 20)) * 20))
            and item.get("line_rule") == line_expected.get("rule", "exact")
            for item in details):
        mismatches.append("fixed line spacing is not exact 20pt")

    actual = dict(layout)
    actual["source_backed_checks"] = {
        "paper": "pass" if "paper size" not in missing and not any(
            item.startswith("paper ") for item in mismatches) else "fail",
        "margins": "pass" if not any("margin" in item for item in missing + mismatches)
        else "fail",
        "line_spacing": "pass" if "fixed line spacing" not in missing and not any(
            "line spacing" in item for item in mismatches) else "fail",
    }
    if mismatches:
        status = "fail"
        message = "DOCX layout source-backed check failed：" + "；".join(mismatches)
    elif missing:
        status = "manual_review"
        message = "DOCX layout facts incomplete，需人工确认：" + "；".join(missing)
    else:
        status = "pass"
        message = "A4、页面距离与固定20磅行距均与来源记录一致。"
    return status, message, actual


def evaluate_compliance(
    state: Mapping[str, Any], artifacts: Mapping[str, Any],
    profile: Mapping[str, Any], report_text: str = "",
    docx_bytes: bytes = b"",
) -> Dict[str, Any]:
    """Evaluate deterministic facts and explicit manual/conflict records."""
    raw_report = artifacts.get("report") or {}
    report = (raw_report.get("report") if isinstance(raw_report, Mapping)
              and "report" in raw_report else raw_report) or {}
    evidence = (artifacts.get("evidence") or {}).get("project_evidence") or {}
    selected = artifacts.get("selected_cases") or {}
    outline = artifacts.get("outline") or {}
    sources = (artifacts.get("literature_sources") or {}).get("sources") or []
    text = str(report_text or state.get("p3_md") or "")
    report_available = bool(report) or bool(text.strip())
    results = {}
    rule_map = {x["rule_id"]: x for x in profile.get("rules") or []}

    abstract = _cjk_count(report.get("abstract_zh"))
    results["abstract_zh_length"] = _result(
        rule_map["abstract_zh_length"],
        "not_applicable" if not report_available else
        "pass" if 400 <= abstract <= 600 else "fail",
        f"统计规则：Unicode CJK unified/extension A 汉字数，不含英文、数字、标点。",
        "中文摘要", abstract)

    zh, zh_format_ok = _keyword_values(report.get("keywords_zh"))
    en, en_format_ok = _keyword_values(report.get("keywords_en"))
    separator_ok = zh_format_ok and en_format_ok
    kw_ok = 5 <= len(zh) <= 8 and 5 <= len(en) <= 8 and separator_ok
    results["keywords_count"] = _result(
        rule_map["keywords_count"],
        "not_applicable" if not report_available else "pass" if kw_ok else "fail",
        f"中文 {len(zh)}、英文 {len(en)}；分隔符{'一致' if separator_ok else '异常'}。",
        "关键词", {"zh": len(zh), "en": len(en), "separator_ok": separator_ok})

    heading_depths = [len(x.group(1)) for x in
                      re.finditer(r"^(#{1,6})\s+", text, re.MULTILINE)]
    toc_depth = max(heading_depths, default=0)
    final_docx = (artifacts.get("final_docx_validation") or {})
    toc_record = final_docx.get("toc") or report.get("toc") or {}
    toc_expected = _toc_titles(toc_record.get("expected")
                               if isinstance(toc_record, Mapping) else [])
    toc_actual = _toc_titles(toc_record.get("actual")
                             if isinstance(toc_record, Mapping) else toc_record)
    toc_missing = [title for title in toc_expected if title not in toc_actual]
    if toc_missing:
        toc_status = "fail"
    elif toc_expected or toc_actual:
        toc_status = "pass"
    elif report_available:
        toc_status = "manual_review"
    else:
        toc_status = "not_applicable"
    results["toc_depth"] = _result(
        rule_map["toc_depth"], "fail" if toc_depth > 3 else toc_status,
        f"正文标题最大 {toc_depth} 级；目录必要组成部分缺失 {len(toc_missing)} 个。",
        "目录 / 标题", {"maximum_heading_level": toc_depth,
                       "expected": toc_expected, "actual": toc_actual,
                       "missing": toc_missing})

    citations = set(re.findall(r"<!--cite:([A-Za-z0-9_.:-]+)-->", text))
    source_ids = [str(x.get("source_id")) for x in sources if x.get("source_id")]
    duplicate_ids = sorted({x for x in source_ids if source_ids.count(x) > 1})
    missing = sorted(citations - set(source_ids))
    unused = sorted(set(source_ids) - citations)
    status = "not_applicable" if not citations and not source_ids else \
        "pass" if not missing and not unused and not duplicate_ids else "fail"
    results["citation_reference_bidirectional"] = _result(
        rule_map["citation_reference_bidirectional"], status,
        f"引用 {len(citations)}，references {len(source_ids)}，缺失 {len(missing)}，"
        f"未用 {len(unused)}，重复 {len(duplicate_ids)}。",
        "正文引用 / references",
        {"missing": missing, "unused": unused, "duplicate_ids": duplicate_ids})

    results["citation_style"] = _result(
        rule_map["citation_style"], "manual_review",
        "不同培养单位可能采用不同引文格式，请根据实际模板确认。",
        "citation policy", "configurable")

    labels = []
    for line_number, line in enumerate(text.splitlines(), 1):
        match = re.search(r"(图|表)\s*(\d+(?:\.\d+)?)", line)
        if match:
            labels.append((match.group(1), match.group(2),
                           line[match.end():].strip(" ：:|\t"), line_number))
    invalid = [f"{kind}{number}" for kind, number, _, _ in labels
               if "." not in number]
    numbers = [(kind, number) for kind, number, _, _ in labels]
    duplicate_labels = sorted({f"{kind}{number}" for kind, number in numbers
                               if numbers.count((kind, number)) > 1})
    missing_captions = [f"{kind}{number}" for kind, number, caption, _ in labels
                        if not caption.strip()]
    gaps = []
    for kind in ("图", "表"):
        by_chapter = {}
        for current_kind, number, _, _ in labels:
            if current_kind != kind or "." not in number:
                continue
            chapter, item = (int(part) for part in number.split(".", 1))
            by_chapter.setdefault(chapter, []).append(item)
        for chapter, items in by_chapter.items():
            for previous, current in zip(sorted(set(items)), sorted(set(items))[1:]):
                if current - previous > 1:
                    gaps.append(f"{kind} {chapter}.{previous} → {chapter}.{current}")
    status = "not_applicable" if not labels else \
        "fail" if invalid or duplicate_labels or missing_captions or gaps else "pass"
    results["figure_table_numbering"] = _result(
        rule_map["figure_table_numbering"], status,
        f"检测到 {len(labels)} 个图表 caption；格式异常 {len(invalid)}，"
        f"重复 {len(duplicate_labels)}，缺 caption {len(missing_captions)}，"
        f"明显跳号 {len(gaps)}。",
        "图 / 表", {"invalid": invalid, "duplicates": duplicate_labels,
                    "missing_captions": missing_captions, "gaps": gaps})

    appendices = list(report.get("appendices") or [])
    def _appendix_roles(value: Any) -> tuple[bool, bool]:
        if isinstance(value, Mapping):
            source = value.get("source") or value.get("original")
            target = value.get("translation") or value.get("target")
            if source or target:
                return bool(source), bool(target)
            value = " ".join(str(value.get(key) or "")
                              for key in ("title", "content", "text"))
        text_value = str(value or "")
        return ("原文" in text_value or "source" in text_value.casefold(),
                "译文" in text_value or "translation" in text_value.casefold())

    has_bilingual = any(all(_appendix_roles(x)) for x in appendices)
    results["bilingual_appendix"] = _result(
        rule_map["bilingual_appendix"],
        "not_applicable" if not report_available else
        "pass" if has_bilingual else "fail",
        "双语附录 source/translation 角色已登记。" if has_bilingual else
        "附录中没有同时可识别的原文与译文角色。", "附录",
        appendices)

    raw_source = "\n".join(str(x.get("source") or "")
                           for x in evidence.get("segments") or [])
    raw_source = raw_source or "\n".join(str(x or "") for x in state.get("paras") or [])
    cjk = _cjk_count(raw_source)
    english = _english_words(raw_source)
    if not raw_source.strip():
        source_status, detail = "not_applicable", "无源文可检查。"
    elif cjk >= 10000:
        source_status, detail = "pass", f"中文版面汉字 {cjk}，满足 10,000。"
    elif english and cjk < 100:
        source_status = "manual_review"
        detail = f"英文源文 {english} words；最低长度折算需要根据所在院校要求确认。"
    else:
        source_status, detail = "fail", f"中文版面汉字 {cjk}，低于 10,000。"
    results["source_length"] = _result(
        rule_map["source_length"], source_status, detail, "source",
        {"cjk_layout_chars": cjk, "english_words_observed": english})

    roles = {str(x.get("role")) for x in outline.get("sections") or []}
    rq_count = len(re.findall(r"\bRQ\d+\b", text))
    source_structure_ok = all([
        bool({"task_description", "project_overview", "introduction"} & roles),
        bool({"task_process", "project_overview"} & roles),
        "case_analysis" in roles,
        bool({"practice_summary", "conclusion_reflection", "conclusion"} & roles),
    ])
    # Existing project outlines use the established RQ/conclusion contract;
    # retain that valid shape while accepting the source's named report parts.
    legacy_structure_ok = "case_analysis" in roles and rq_count and bool(
        {"conclusion_reflection", "conclusion"} & roles)
    structure_ok = source_structure_ok or legacy_structure_ok
    results["case_conclusion_structure"] = _result(
        rule_map["case_conclusion_structure"],
        "not_applicable" if not roles else "pass" if structure_ok else "fail",
        f"案例分析={'有' if 'case_analysis' in roles else '无'}，RQ={rq_count}，"
        f"结论={'有' if bool({'conclusion_reflection','conclusion'} & roles) else '无'}；"
        f"报告基本结构={'有' if source_structure_ok else '需补充'}；"
        "充分性需人工复核。",
        "案例分析 / 结论", {"roles": sorted(roles), "rq_markers": rq_count})

    layout = inspect_docx_layout(docx_bytes) if docx_bytes else \
        dict((artifacts.get("final_docx_validation") or {}).get(
            "layout_facts") or {})
    layout_status, layout_message, layout_actual = _evaluate_docx_layout(
        layout, rule_map["docx_layout"].get("expected") or {})
    results["docx_layout"] = _result(
        rule_map["docx_layout"], layout_status, layout_message, "DOCX",
        layout_actual)

    cases = list(selected.get("cases") or [])
    synthetic = [case_provenance.with_provenance(x) for x in cases
                 if case_provenance.is_synthetic(x)]
    results["synthetic_case_policy"] = _result(
        rule_map["synthetic_case_policy"], "manual_review" if synthetic else "not_applicable",
        f"项目 strict policy 是 project constraint；是否计入最低正式案例数量需要根据所在院校要求确认。"
        f"当前 synthetic {len(synthetic)} 个。", "case policy",
        {"project_constraint": selected.get("synthetic_count_policy") or
         selected.get("synthetic_case_count_policy") or
         (selected.get("report_case_policy") or {}).get(
             "synthetic_count_policy"),
         "profile_policy": dict(profile.get("synthetic_policy") or {}),
         "synthetic_count": len(synthetic)})

    placeholder_patterns = (
        "【待作者填写】", "unknown author", "unknown supervisor", "unknown date")
    found = []
    for section in _sections(text):
        for index, line in enumerate(section["lines"], 1):
            if any(value in line.casefold() for value in placeholder_patterns):
                found.append({"section": section["section_id"], "line": index,
                              "excerpt": line.strip()[:200]})
    report_blob = json.dumps(report, ensure_ascii=False)
    if any(value.casefold() in report_blob.casefold()
           for value in placeholder_patterns) and not any(
               item.get("section") == "report artifact" for item in found):
        found.append({"section": "report artifact", "line": None,
                      "excerpt": "报告结构化字段含未解决作者信息占位符。"})
    results["author_placeholders"] = _result(
        rule_map["author_placeholders"],
        "not_applicable" if not report_available else
        "manual_review" if found else "pass",
        f"发现 {len(found)} 个未解决作者信息占位符。" if found else
        "未发现已知作者信息占位符。", "全文", found)

    # A profile extension without a registered deterministic checker is still
    # visible as an unresolved rule; it must not silently disappear or turn
    # into an institution-enforced result.
    for rule_id, rule in rule_map.items():
        if rule_id not in results:
            fallback_status = "manual_review" if not rule.get(
                "reliable_source_mapping", rule.get("source_available")) else "not_applicable"
            results[rule_id] = _result(
                rule, fallback_status,
                "当前没有对应的确定性检查器，需人工确认或补充来源映射。",
                "—", None)

    counts = {key: sum(1 for item in results.values() if item["status"] == key)
              for key in ("pass", "fail", "manual_review", "not_applicable")}
    enforced = [x for x in results.values()
                if x["enforcement"] == "enforced" and x["authority_level"] != "project"]
    profile_manual = [x for x in results.values()
                      if x["authority_level"] != "project" and
                      x["status"] == "manual_review"]
    source_missing = [x["rule_id"] for x in results.values()
                      if x.get("authority_level") != "project" and
                      not x.get("reliable_source_mapping")]
    profile_status = "fail" if any(
        x["status"] == "fail" and x["enforcement"] == "enforced" for x in enforced) \
        else "manual_review" if profile_manual \
        else "pass"
    project_fail = any(x["status"] == "fail" and
                       x["enforcement"] in {"project_constraint", "enforced"}
                       for x in results.values())
    return {
        "schema_version": VERSION, "profile_id": profile.get("profile_id"),
        "display_name": profile.get("display_name"),
        "profile_type": profile.get("profile_type"), "program": profile.get("program"),
        "profile_compliance": {
            "status": profile_status,
            "enforced_rule_count": len(enforced),
            "blocking_failures": [x["rule_id"] for x in enforced
                                  if x["status"] == "fail"],
        },
        "project_constraints": {
            "status": "fail" if project_fail else "pass",
            "failures": [x["rule_id"] for x in results.values()
                         if x["status"] == "fail" and x["enforcement"] ==
                         "project_constraint"],
        },
        "manual_reviews": [x["rule_id"] for x in results.values()
                           if x["status"] == "manual_review"],
        "conflicts": deepcopy(profile.get("conflicts") or [{
            "conflict_id": "citation_style_requires_profile_choice",
            "rule": "citation_style", "resolved_by": "manual_review",
            "status": "manual_review",
        }]),
        "source_audit": {
            "enforced_rule_count": len(enforced),
            "manual_review_rule_count": sum(
                1 for x in results.values()
                if x["authority_level"] != "project" and
                x["enforcement"] == "manual_review"),
            "evaluated_manual_review_count": sum(
                1 for x in results.values() if x["status"] == "manual_review"),
            "conflict_count": len(profile.get("conflicts") or []),
            "rules_without_source_mapping": source_missing,
        },
        "unresolved_items": [x for x in results.values()
                             if x["status"] in {"manual_review", "fail"}],
        "counts": {**counts, "manual_review": counts["manual_review"],
                   "not_checked": 0},
        "status": "fail" if project_fail or profile_status == "fail"
        else "manual_review" if counts["manual_review"] else "pass",
        "rules": list(results.values()),
    }


def _constraint_result(kind, value, allowed, sections_text):
    found = []
    for section in _sections(sections_text):
        for line_number, line in enumerate(section["lines"], 1):
            if value.casefold() in line.casefold():
                found.append({"section": section["section_id"], "line": line_number,
                              "excerpt": line.strip()[:200]})
    return {"kind": kind, "value": value, "allowed": allowed,
            "status": "fail" if not allowed and found else "manual_review"
            if found else "pass", "occurrences": found}


def evaluate_language_constraints(state: Mapping[str, Any], report_text: str) -> Dict[str, Any]:
    """Evaluate only explicit, localizable project constraints."""
    settings = state.get("research_settings") or {}
    text = str(report_text or state.get("p3_md") or "")
    required = state.get("required_terminology") or settings.get(
        "required_terminology") or []
    results = []
    configs = [("forbidden_report_phrase", x, False)
               for x in settings.get("forbidden_report_phrases") or []]
    configs += [("theory_label", x, True)
                for x in settings.get("allowed_theory_labels") or []]
    configs += [("protected_name", x, True)
                for x in settings.get("protected_names") or []]
    configs += [("protected_work_title", x, True)
                for x in settings.get("protected_work_titles") or []]
    detected_theory = set(re.findall(r"(目的论|功能对等|关联理论|诠释学|变译理论)", text))
    allowed_theory = set(settings.get("allowed_theory_labels") or [])
    for value in configs:
        value, allowed = value[1], value[2]
        found = []
        for section in _sections(text):
            for line_number, line in enumerate(section["lines"], 1):
                if value.casefold() in line.casefold():
                    found.append({"section": section["section_id"],
                                  "line": line_number,
                                  "excerpt": line.strip()[:200]})
        results.append({"kind": "constraint", "value": value, "allowed": allowed,
                        "status": "fail" if not allowed and found else
                        "pass" if found or allowed else "not_applicable",
                        "occurrences": found})
    unexpected_theory = sorted(detected_theory - allowed_theory) if allowed_theory else []
    if unexpected_theory:
        results.append({"kind": "theory_label", "value": unexpected_theory,
                        "allowed": sorted(allowed_theory), "status": "manual_review",
                        "occurrences": [{"section": "全文", "line": None,
                                         "excerpt": "理论标签需要定位复核"}]})
    missing_terms = []
    for item in required:
        if isinstance(item, Mapping):
            source = str(item.get("source") or "").strip()
            target = str(item.get("target") or "").strip()
            if source and source.casefold() not in text.casefold() and \
                    (not target or target not in text):
                missing_terms.append(source)
        else:
            source = str(item or "").strip()
            if source and source.casefold() not in text.casefold():
                missing_terms.append(source)
    if missing_terms:
        results.append({"kind": "required_terminology", "value": missing_terms,
                        "allowed": True, "status": "manual_review",
                        "occurrences": []})
    failures = [x for x in results if x["status"] == "fail"]
    return {
        "schema_version": VERSION, "status": "fail" if failures else
        "manual_review" if any(x["status"] == "manual_review" for x in results)
        else "pass", "constraints": results, "failures": failures,
    }
