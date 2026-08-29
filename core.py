"""TransPraxis / 译践 —— 核心逻辑层（与 Streamlit UI 解耦，便于测试）。

职责：大模型路由、文档清洗、术语抽取、双语翻译、报告生成、任务进度持久化。
"""
import hashlib
import io
import json
import os
import re
import shutil
import subprocess
import threading
import tempfile
import time
import traceback
import uuid
import zipfile
from datetime import datetime, timedelta, timezone
from html import escape as html_escape
from pathlib import Path

import fitz  # PyMuPDF
import httpx
import pandas as pd
from defusedxml import ElementTree as ET
from docx import Document
from google import genai
from openai import OpenAI

from transpraxis import models as _models
from transpraxis import state_migration as _state_migration
from transpraxis import context as _context
from transpraxis import knowledge as _knowledge
from transpraxis import repair as _repair
from transpraxis import translation_evidence as _translation_evidence
from transpraxis import checkpoint as _checkpoint
from transpraxis import snapshots as _snapshots
from transpraxis import entity_registry as _entity_registry
from transpraxis import model_roles as _model_roles
from transpraxis import pdf_ingestion as _pdf_ingestion
from transpraxis import translation_protocol as _translation_protocol
from transpraxis import translation_target as _translation_target
from transpraxis import finalization as _finalization
from transpraxis import rendered_qa as _rendered_qa

# ================= 常量 =================
# 任务进度与过程文件的本地存储目录（已加入 .gitignore）
OUTPUT_DIR = Path("outputs")

DELIVERY_CONFIG_DEFAULTS = {
    "enable_annotate": False,
    "enable_report": False,
    "deliver_plain_docx": True,
    "deliver_bilingual_docx": True,
    "deliver_pdf": False,
    "deliver_terms_xlsx": True,
    "deliver_tbx": False,
    "deliver_tmx": False,
    "deliver_jsonl": False,
    "deliver_evidence": True,
    "deliver_cases": False,
    "deliver_academic_workspace": False,
    "deliver_review_report": False,
}


def default_delivery_config():
    return dict(DELIVERY_CONFIG_DEFAULTS)


def normalize_delivery_config(config, *, enable_report=None, enable_annotate=None):
    normalized = default_delivery_config()
    if isinstance(config, dict):
        normalized.update({key: bool(config[key]) for key in normalized if key in config})
    if enable_report is not None:
        normalized["enable_report"] = bool(enable_report)
    if enable_annotate is not None:
        normalized["enable_annotate"] = bool(enable_annotate)
    return normalized

# 提供商注册表：kind=openai_compat 走 OpenAI SDK 兼容接口（官方与中转站通用）。
# base_url 为 None 表示使用官方 SDK 默认路由；custom_base_url 表示地址由用户在
# UI 中填写（通用中转站）。OpenCode Go 只列出官方标记为
# /chat/completions 的模型；/messages 与 /responses 模型不走此路由。
PROVIDERS = {
    "OpenCode Go": {
        "kind": "openai_compat",
        "base_url": "https://opencode.ai/zen/go/v1",
        "proxy_bypass": True,
        "default_model": "glm-5.2",
        "models": [
            "glm-5.2", "glm-5.1",
            "deepseek-v4-pro", "deepseek-v4-flash",
            "kimi-k3", "kimi-k2.7-code", "kimi-k2.6",
            "mimo-v2.5-pro", "mimo-v2.5", "hy3", "grok-4.5",
        ],
    },
    "DeepSeek": {
        "kind": "openai_compat",
        "base_url": "https://api.deepseek.com",
        "default_model": "deepseek-v4-flash",
        "models": ["deepseek-v4-flash", "deepseek-v4-pro"],
    },
    "OpenAI": {
        "kind": "openai",
        "models": ["gpt-4o-mini", "gpt-4o", "gpt-4.1-mini", "gpt-4.1"],
        "capabilities": {
            "supports_json_schema": True,
            "supports_json_object": True,
            "supports_response_format": True,
        },
    },
    "Gemini": {
        "kind": "gemini",
        "default_model": "gemini-3.6-flash",
        "models": ["gemini-3.6-flash", "gemini-3.5-flash",
                   "gemini-3.5-flash-lite", "gemini-2.5-flash", "gemini-2.5-pro"],
    },
    "OpenRouter": {
        "kind": "openai_compat",
        "base_url": "https://openrouter.ai/api/v1",
        "models": [],
        "model_hint": "如 anthropic/claude-sonnet-4、openai/gpt-5、google/gemini-3-flash",
    },
    "SiliconFlow": {
        "kind": "openai_compat",
        "base_url": "https://api.siliconflow.cn/v1",
        "models": [],
        "model_hint": "如 Qwen/Qwen3-235B-A22B、deepseek-ai/DeepSeek-V3.2",
    },
    "Moonshot (Kimi)": {
        "kind": "openai_compat",
        "base_url": "https://api.moonshot.cn/v1",
        "models": [],
        "model_hint": "如 kimi-k2.5、moonshot-v1-8k",
    },
    "Zhipu (GLM)": {
        "kind": "openai_compat",
        "base_url": "https://open.bigmodel.cn/api/paas/v4",
        "models": [],
        "model_hint": "如 glm-4.5、glm-5",
    },
    "Qwen (DashScope)": {
        "kind": "openai_compat",
        "base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1",
        "models": [],
        "model_hint": "如 qwen-max、qwen3-235b-a22b",
    },
    "自定义中转站": {
        "kind": "openai_compat",
        "base_url": None,
        "custom_base_url": True,
        "models": [],
        "model_hint": "填写中转站提供的模型名（OpenAI /chat/completions 兼容）",
    },
}

# 兼容旧引用：仅保留 模型名 -> 列表 的视图
MODELS = {name: cfg["models"] for name, cfg in PROVIDERS.items()}

# 会话线程级中转地址：自定义中转站的 base_url 由 UI 写入，call_llm 自动优先使用。
# 每个 Streamlit 会话线程独立，多设备同时使用不会互相串扰。
_LLM_CTX = threading.local()

# Runtime status is deliberately file-backed: the UI can poll it while the
# worker is inside a blocking provider request, and a browser refresh does not
# erase the last known activity.
_RUNTIME_CTX = threading.local()
_RUNTIME_LOCK = threading.RLock()
_RUNTIME_WORKERS = {}
_RUNTIME_WORKERS_LOCK = threading.RLock()


def normalize_openai_base_url(base_url):
    """接受基址或误填的 OpenAI 兼容具体端点，并统一回基址。"""
    base = (base_url or "").strip().rstrip("/")
    for suffix in ("/chat/completions", "/completions", "/responses", "/models"):
        if base.endswith(suffix):
            return base[:-len(suffix)]
    return base


def set_llm_base_url(base_url):
    """为当前线程设置 OpenAI 兼容中转站地址（空值清除）。"""
    _LLM_CTX.base_url = normalize_openai_base_url(base_url) or None


def _runtime_job_id():
    return getattr(_RUNTIME_CTX, "job_id", None)


def _runtime_cancel_requested(job_id):
    return bool((load_runtime_state(job_id) or {}).get("cancel_requested"))


# ================= 基础工具函数 =================
def clean_xml_chars(text):
    if not isinstance(text, str):
        return str(text)
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)


def parse_json_array(text):
    """从 LLM 输出中稳健地解析 JSON 数组；解析失败返回 None。

    依次尝试：整体解析 -> 去掉 Markdown 代码块后解析 -> 从每个 '[' 位置做
    raw_decode（可容忍输出前后夹带解释文字）。
    """
    if not isinstance(text, str) or not text.strip():
        return None
    candidate = text.strip()
    candidate = re.sub(r'^```(?:json)?\s*', '', candidate, flags=re.DOTALL)
    candidate = re.sub(r'\s*```$', '', candidate, flags=re.DOTALL).strip()

    try:
        obj = json.loads(candidate)
        if isinstance(obj, list):
            return obj
    except Exception:
        pass

    decoder = json.JSONDecoder()
    for m in re.finditer(r'\[', candidate):
        try:
            obj, _ = decoder.raw_decode(candidate[m.start():])
        except Exception:
            continue
        if isinstance(obj, list):
            return obj
    return None


def parse_translation_array(res, expected):
    """Compatibility API backed by the canonical translation parser."""
    return _translation_protocol.parse_translation_array(res, expected)

def is_rate_limited(err):
    s = str(err)
    return '429' in s or 'RESOURCE_EXHAUSTED' in s or 'rate limit' in s.lower()


# ================= PDF 确定性段落提取 =================
# 经验（来自 localize-anything 与全书实测）：分段/清洗必须确定性，不能交给 LLM。
# 旧流程把 ~2500 字符的任意文本块交给模型"清洗"，导致两类系统性缺陷：
#   1. 块边界落在句中 -> 句子被拦腰截断（"…pecking at" / "crumbs and bones…"）；
#   2. 模型自由裁量 -> 对白、引语被随意拆分或合并，分段结果不可复现。
# 新流程直接读 PDF 版面：块(block)->行(line)->首行缩进判定段落，连字符修复、
# 跨页段落合并、页眉页脚/页码剔除全部确定完成。

# 句末终结符（用于判断段落是否未完结、需要与下一段合并）
_SENTENCE_TERMINAL = set('.!?"”’…:;)')

# 纯装饰符号行（章节分隔花饰等），无字母/数字，不是正文
_ORNAMENT_RE = re.compile(r"^[\s*•·▪◦‣❦❧—–\-]{1,12}$")

# 常见缩写（句点不计入句界）
_ABBREV_RE = re.compile(
    r"\b(?:Lt|Col|Gen|Maj|Capt|Sgt|Brig|Mr|Mrs|Ms|Dr|St|No|Vol|pp|"
    r"e\.g|i\.e|vs|etc|a\.m|p\.m|U\.S|A\.F|B\.C|A\.D)\.", re.IGNORECASE)


def extract_pdf_paragraphs(file_bytes):
    """Compatibility API backed by the layout-aware PDF ingestion module."""
    return _pdf_ingestion.extract_pdf_paragraphs(file_bytes)


def _ocr_pdf_text(file_bytes, max_pages=3):
    """扫描件 OCR：只渲染代表性页（前/中/后），交给 tesseract 识别。

    任何环节失败都返回空字符串（由调用方降级），绝不让 OCR 阻塞流程。
    """
    import shutil
    import subprocess
    import tempfile

    if shutil.which("tesseract") is None:
        return ""
    try:
        import fitz
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            count = doc.page_count
            if count == 0:
                return ""
            indices = sorted({0, count // 2, count - 1})[:max_pages]
            chunks = []
            for idx in indices:
                pix = doc[idx].get_pixmap(dpi=200)
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                    pix.save(tmp.name)
                    tmp_path = tmp.name
                try:
                    proc = subprocess.run(
                        ["tesseract", tmp_path, "-", "-l", "chi_sim+eng"],
                        capture_output=True, timeout=120)
                    text = proc.stdout.decode("utf-8", errors="ignore")
                    if text.strip():
                        chunks.append(text.strip())
                except Exception:  # noqa: BLE001 - OCR 失败不阻断
                    pass
                finally:
                    Path(tmp_path).unlink(missing_ok=True)
            return "\n".join(chunks)
    except Exception:  # noqa: BLE001
        return ""


def extract_document_paragraphs(filename, file_bytes):
    """上传后的轻量预提取：PDF 走确定性解析，扫描件尝试 OCR 代表页；
    DOCX 走 python-docx。返回 (paragraphs, warnings)。"""
    warnings = []
    paragraphs = []
    name = (filename or "").lower()
    try:
        if name.endswith(".pdf"):
            paragraphs = [clean_xml_chars(p) for p in extract_pdf_paragraphs(file_bytes)]
            if not paragraphs:
                warnings.append("PDF 无文本层，正在尝试 OCR 前/中/后代表页…")
                ocr_text = _ocr_pdf_text(file_bytes)
                if ocr_text:
                    paragraphs = [clean_xml_chars(p.strip())
                                  for p in re.split(r"\n+", ocr_text)
                                  if p.strip() and len(p.strip()) > 1]
                else:
                    warnings.append("OCR 不可用或失败，无法自动画像；可手动选择风格")
        elif name.endswith(".docx"):
            doc_word = Document(io.BytesIO(file_bytes))
            for p in doc_word.paragraphs:
                for sub_p in re.split(r"\n+", clean_xml_chars(p.text)):
                    t = sub_p.strip()
                    if len(t) > 1 and not _ORNAMENT_RE.match(t):
                        paragraphs.append(t)
        else:
            warnings.append("不支持的文档格式，无法自动画像")
    except Exception as exc:  # noqa: BLE001
        warnings.append(f"预提取失败：{exc}")
    return paragraphs, warnings


def call_llm(provider, api_key, model, system_prompt, user_prompt,
             temperature=0.1, base_url=None, response_format=None):
    """底层大模型统一路由（超时 150 秒，模型可配置）。

    支持官方接口与 OpenAI /chat/completions 兼容中转站：
    base_url 显式传入 > 提供商默认 base_url > 会话线程级自定义中转地址。
    """
    cfg = PROVIDERS.get(provider)
    if not cfg:
        return ""
    runtime_job_id = _runtime_job_id()
    if runtime_job_id:
        if _runtime_cancel_requested(runtime_job_id):
            raise RuntimeError("任务已请求取消")
        update_runtime_state(
            runtime_job_id, status="waiting_external", phase="waiting_llm",
            phase_label="等待模型响应", event="已向模型发送请求",
            event_name="llm_request_started")
    if cfg["kind"] == "gemini":
        try:
            client = genai.Client(api_key=api_key,
                                  http_options=genai.types.HttpOptions(timeout=150_000))
        except (AttributeError, TypeError):
            client = genai.Client(api_key=api_key)
        config_kwargs = {"temperature": temperature}
        if response_format and cfg.get("capabilities", {}).get("supports_response_format"):
            config_kwargs["response_mime_type"] = "application/json"
        res = client.models.generate_content(
            model=model,
            contents=user_prompt,
            system_instruction=system_prompt,
            config=genai.types.GenerateContentConfig(**config_kwargs),
        )
        result = (res.text or "").strip()
        if runtime_job_id:
            update_runtime_state(runtime_job_id, status="running", phase="running",
                                 phase_label="处理模型结果", event="已收到模型响应",
                                 event_name="llm_response_received")
            if _runtime_cancel_requested(runtime_job_id):
                raise RuntimeError("任务已请求取消")
        return result

    # OpenAI 官方与所有 OpenAI 兼容接口共用 SDK 路由
    kwargs = {"api_key": api_key, "timeout": (15.0, 180.0), "max_retries": 1}
    resolved_base = normalize_openai_base_url(base_url or cfg.get("base_url"))
    if cfg.get("custom_base_url") and not resolved_base:
        resolved_base = getattr(_LLM_CTX, "base_url", None)
    if resolved_base:
        kwargs["base_url"] = resolved_base
    http_client = None
    if cfg.get("proxy_bypass"):
        # 该接口经用户本地代理时 TLS 失败：显式关闭环境代理。
        http_client = httpx.Client(trust_env=False, timeout=(15.0, 180.0))
        kwargs["http_client"] = http_client
    try:
        client = OpenAI(**kwargs)
        request = {
            "model": model,
            "messages": [{"role": "system", "content": system_prompt},
                         {"role": "user", "content": user_prompt}],
            "temperature": temperature,
        }
        if response_format and cfg.get("capabilities", {}).get(
                "supports_response_format"):
            request["response_format"] = response_format
        res = client.chat.completions.create(**request)
        result = (res.choices[0].message.content or "").strip()
        if runtime_job_id:
            update_runtime_state(runtime_job_id, status="running", phase="running",
                                 phase_label="处理模型结果", event="已收到模型响应",
                                 event_name="llm_response_received")
            if _runtime_cancel_requested(runtime_job_id):
                raise RuntimeError("任务已请求取消")
        return result
    finally:
        if http_client:
            http_client.close()


def test_provider(provider, api_key, model, base_url=None):
    """连通性测试：发送一个最小请求，返回 (ok, message)。"""
    t0 = time.time()
    try:
        out = call_llm(provider, api_key, model,
                       "你是连接测试助手。", "请只回复两个字：OK",
                       temperature=0.0, base_url=base_url)
    except Exception as exc:
        return False, f"请求失败：{exc}"[:320]
    elapsed = time.time() - t0
    if not (out or "").strip():
        return False, "返回内容为空（请检查 API Key / 模型名 / 余额）"
    return True, f"响应「{(out or '').strip()[:24]}」· 耗时 {elapsed:.1f}s"


def fetch_provider_models(provider, api_key, base_url=None):
    """获取 OpenAI 兼容服务商的模型目录，返回 (ok, models, message)。"""
    cfg = PROVIDERS.get(provider) or {}
    if cfg.get("kind") not in ("openai", "openai_compat"):
        return False, [], "当前服务商不支持 OpenAI 兼容模型目录"
    if not (api_key or "").strip():
        return False, [], "请先填写 API Key"
    resolved_base = normalize_openai_base_url(base_url or cfg.get("base_url"))
    if not resolved_base:
        return False, [], "请先填写 API 地址"
    try:
        response = httpx.get(
            f"{resolved_base}/models",
            headers={"Authorization": f"Bearer {api_key.strip()}",
                     "Accept": "application/json"},
            timeout=20.0,
        )
        response.raise_for_status()
        payload = response.json()
    except httpx.HTTPStatusError as exc:
        return False, [], f"获取模型失败：HTTP {exc.response.status_code}"
    except (httpx.HTTPError, ValueError) as exc:
        return False, [], f"获取模型失败：{str(exc)[:240]}"

    rows = payload.get("data") if isinstance(payload, dict) else payload
    if not rows and isinstance(payload, dict):
        rows = payload.get("models")
    models = []
    for row in rows or []:
        model_id = row if isinstance(row, str) else (
            row.get("id") or row.get("model") or row.get("name")
            if isinstance(row, dict) else None)
        if model_id and str(model_id).strip():
            models.append(str(model_id).strip())
    models = sorted(set(models), key=str.casefold)
    if not models:
        return False, [], "模型目录为空或返回格式不受支持"
    return True, models, f"已获取 {len(models)} 个可用模型"


def parse_termbase(file_stream):
    """解析用户上传的术语库 Excel，返回概念化术语条目列表。

    必选列：Source / Target；可选列：Behavior（translate/preserve）、
    Status（locked/provisional）、Preferred（首选译名）、Forbidden（禁止译名，
    可用 ; 或 , 分隔）、Scope、Note。解析失败抛出 ValueError（不再静默吞错）。
    """
    try:
        df = pd.read_excel(file_stream)
    except Exception as e:
        raise ValueError(f"无法读取 Excel 文件：{e}") from e
    df.columns = [str(c).strip() for c in df.columns]
    return _termbase_df_to_entries(df)


def _termbase_df_to_entries(df):
    if "Source" not in df.columns or "Target" not in df.columns:
        raise ValueError("术语库缺少 Source / Target 列，请检查表头")
    df = df.dropna(subset=["Source", "Target"])
    entries = []
    for _, row in df.iterrows():
        entry = {"source": str(row["Source"]).strip(),
                 "target": str(row["Target"]).strip()}
        for col, key in (("Behavior", "behavior"), ("Status", "status"),
                         ("Preferred", "preferred"), ("Scope", "scope"),
                         ("Note", "note")):
            if col in df.columns and pd.notna(row[col]):
                entry[key] = str(row[col]).strip()
        if "Forbidden" in df.columns and pd.notna(row["Forbidden"]):
            entry["forbidden"] = [x.strip() for x in
                                  re.split(r"[;；,，]", str(row["Forbidden"])) if x.strip()]
        entries.append(entry)
    return entries


def parse_termbase_csv(file_stream):
    """CSV 术语库：列约定与 Excel 一致（Source/Target 必选）。"""
    try:
        df = pd.read_csv(file_stream)
    except Exception as e:
        raise ValueError(f"无法读取 CSV 文件：{e}") from e
    df.columns = [str(c).strip() for c in df.columns]
    return _termbase_df_to_entries(df)


def _local_name(tag):
    return str(tag).rsplit("}", 1)[-1]


def parse_termbase_tbx(file_stream):
    """解析 TBX 术语库（Trados MultiTerm 等标准格式）。

    每个 termEntry 取前两种语言的第一个 term，作为 source/target；
    导入条目默认 status=locked（导入即视为已固定译名）。
    """
    try:
        root = ET.fromstring(file_stream.read())
    except Exception as e:
        raise ValueError(f"无法解析 TBX 文件：{e}") from e
    entries = []
    for te in root.iter():
        if _local_name(te.tag) != "termEntry":
            continue
        langs = []
        for ls in te.iter():
            if _local_name(ls.tag) != "langSet":
                continue
            # ElementTree 会把 xml: 前缀自动展开为命名空间键
            lang = ls.get("{http://www.w3.org/XML/1998/namespace}lang") \
                or ls.get("xml:lang") or ls.get("lang") or ""
            term = None
            for node in ls.iter():
                if _local_name(node.tag) == "term" and (node.text or "").strip():
                    term = node.text.strip()
                    break
            if term and lang and lang not in [x[0] for x in langs]:
                langs.append((lang, term))
            if len(langs) >= 2:
                break
        if len(langs) >= 2:
            entries.append({"source": langs[0][1], "target": langs[1][1],
                            "behavior": "translate", "status": "locked"})
    if not entries:
        raise ValueError("TBX 中未找到可导入的术语（需要至少两种语言的 langSet）")
    return entries


def import_tmx(file_stream):
    """导入 TMX 翻译记忆（Trados / memoQ 等导出的标准格式）。

    按 <tu> 的 <tuv><seg> 文本对入库：仅接受源文含字母/数字且译文非空的
    单元；与现有翻译记忆冲突的源文跳过（不覆盖项目内已审校条目）。
    返回 {"added": n, "skipped": m}。
    """
    try:
        root = ET.fromstring(file_stream.read())
    except Exception as e:
        raise ValueError(f"无法解析 TMX 文件：{e}") from e
    existing = load_tm()
    added = skipped = 0
    for tu in root.iter():
        if _local_name(tu.tag) != "tu":
            continue
        texts = []
        for tuv in tu:
            if _local_name(tuv.tag) != "tuv":
                continue
            seg = next((c for c in tuv if _local_name(c.tag) == "seg"), None)
            if seg is not None and (seg.text or "").strip():
                texts.append(seg.text.strip())
            if len(texts) >= 2:
                break
        if len(texts) >= 2:
            src, tgt = texts[0], texts[-1]
            if _tm_eligible(src, tgt):
                if src not in existing:
                    existing[src] = {"target": tgt, "reviewed": True,
                                     "source": "tmx_import"}
                    added += 1
                else:
                    skipped += 1
    if not added:
        raise ValueError("TMX 中未找到可导入的新翻译单元（源文需含字母/数字，且不与现有记忆冲突）")
    save_tm(existing)
    return {"added": added, "skipped": skipped}


def extract_auto_terms(paragraphs, target_lang, provider, api_key, model):
    """自动抽取术语库（兼容旧接口：返回 {source: target}）。

    新实现（transpraxis.terminology.extract_auto_terms_v2）：分布式采样、
    全量 occurrences、candidate 状态与 model_knowledge 证据。
    """
    from transpraxis.terminology import extract_auto_terms_v2
    entries, _warnings = extract_auto_terms_v2(
        paragraphs, target_lang, provider, api_key, model)
    return {e["source"]: e["target"] for e in entries}


# ================= 概念化术语表（对齐 localize-anything 的 Glossary 模型）=================
def normalize_glossary(entries):
    """标准化术语条目（委托 transpraxis.models，兼容旧字段并新增 id/occurrences/evidence）。"""
    return _models.normalize_glossary(entries)


def glossary_block(glossary):
    """把术语表渲染成注入翻译/审校 prompt 的文本块。"""
    locked_translate = [e for e in glossary if e["behavior"] == "translate" and e["status"] == "locked"]
    preserve = [e for e in glossary if e["behavior"] == "preserve"]
    provisional = [e for e in glossary if e["behavior"] == "translate" and e["status"] != "locked"]
    lines = []
    if locked_translate:
        lines.append("【锁定术语（必须使用首选译名，不得使用禁止译名）】：")
        for e in locked_translate:
            seg = f"- {e['source']} -> {e['preferred']}"
            if e["forbidden"]:
                seg += f"（禁止：{'、'.join(e['forbidden'])}）"
            lines.append(seg)
    if preserve:
        lines.append("【必须保留原文的术语/名称】：" + "、".join(e["source"] for e in preserve))
    if provisional:
        lines.append("【建议术语（仅供参考，请优先采用）】：")
        for e in provisional:
            lines.append(f"- {e['source']} -> {e['target']}")
    return "\n".join(lines)


def glossary_to_terms(glossary):
    """翻译行为术语 -> 扁平 dict（供报告生成等场景使用）。"""
    return {e["source"]: (e["preferred"] or e["target"])
            for e in glossary if e["behavior"] == "translate" and e["target"]}


def check_glossary_compliance(src, tgt, glossary, segment_id=None,
                              section_profile=None):
    """锁定术语的确定性合规检查（委托 transpraxis.terminology：entry_id/segment_id 级）。"""
    from transpraxis.terminology import check_glossary_compliance as _qa
    return _qa(src, tgt, glossary, segment_id=segment_id,
               section_profile=section_profile)


# ================= 确定性检查（对齐 localize-anything 的机械检查）=================
PRESERVE_RE = re.compile(
    r'(?P<placeholder>%[sd]|%1\$[sd]|\{[A-Za-z_][A-Za-z0-9_]*\}|\{\{[A-Za-z_][A-Za-z0-9_]*\}\})'
    r'|(?P<url>https?://\S+|www\.\S+)'
    r'|(?P<email>[\w.+-]+@[\w-]+(?:\.[\w-]+)+)'
    r'|(?P<doi>10\.\d{4,9}/[^\s]+)'
    r'|(?P<citation>\[\d+(?:[-,]\s*\d+)*\])',
    re.IGNORECASE,
)

PRESERVE_SEVERITY = {
    "placeholder": "blocking",   # 占位符损坏 = 结构破坏，绝不可自动放行
    "url": "actionable",
    "email": "actionable",
    "doi": "actionable",
    "citation": "actionable",
}


def extract_preserved_tokens(text):
    """提取源文本中必须原样保留的 token（占位符/URL/邮箱/DOI/引用标注）。"""
    return {m.group(0): m.lastgroup for m in PRESERVE_RE.finditer(text or "")}


def find_residuals(src, tgt, target_lang):
    """检测目标语言中残留的源语言片段。

    返回 [(片段, severity)]：连续 ≥2 个源语单词/较长汉字串 -> actionable；
    单个词（可能是专有名词）-> informational。启发式，不替代审校。
    """
    tgt_clean = PRESERVE_RE.sub(" ", tgt or "")
    if target_lang == "English":
        source_runs = set(re.findall(r'[\u4e00-\u9fff]{2,}', src or ""))
        return [(c, "actionable" if len(c) >= 4 else "informational")
                for c in re.findall(r'[\u4e00-\u9fff]{2,}', tgt_clean)
                if any(c in run for run in source_runs)]
    src_words = set(w.lower() for w in re.findall(r'[A-Za-z]{5,}', src or ""))
    allowed = {"mti"}  # 产品名等明确保留词白名单（审校负责语义判断）
    words = re.findall(r'[A-Za-z]{5,}', tgt_clean)
    hits = [w for w in words if w.lower() in src_words and w.lower() not in allowed]
    result, run = [], []
    for w in words:
        if w in hits:
            run.append(w)
        else:
            if run:
                result.append((" ".join(run),
                               "actionable" if len(run) >= 2 else "informational"))
                run = []
    if run:
        result.append((" ".join(run), "actionable" if len(run) >= 2 else "informational"))
    return result


def _count_sentences(text):
    """粗粒度句数统计：按终结符切分（引号/括号闭合归并到前一句）。"""
    text = _ABBREV_RE.sub(" ", text)
    parts = re.split(r"[.!?…。！？]+[”\"'’)\]]*", text)
    return sum(1 for p in parts if p.strip())


def is_incomplete_translation(src, tgt):
    """疑似漏译/截断判定（双重规则，实测调优）：
    1. 字符级：长原文（≥120 字符）配极短译文（<15%）——只拦灾难性截断，
       英译中正常比例可低至 0.2-0.3，不能用高阈值；
    2. 句子级：原文 ≥2 句而译文不足一半句数，且字符占比 <35%——
       截断译文必然句数对不上，完整译文即使语言再凝练也很少掉一半句。
    """
    tgt = (tgt or "").strip()
    if not tgt:
        return True
    if len(src) >= 120 and len(tgt) < 0.15 * len(src):
        return True
    src_sents = _count_sentences(src)
    if src_sents >= 2:
        tgt_sents = _count_sentences(tgt)
        if tgt_sents < src_sents * 0.5 and len(tgt) < 0.35 * len(src):
            return True
    return False


def _deterministic_finding(segment_index, severity, category, summary,
                           explanation, recommendation, *, source_span=None,
                           target_span=None, reason=None, kind=None,
                           detected_text=None, **extra):
    """Build an evidence-free but actionable deterministic QA finding."""
    finding = {
        "segment_index": segment_index, "segment_id": segment_index,
        "type": "check", "severity": severity,
        "category": category, "summary": summary,
        "source_span": source_span, "target_span": target_span,
        "explanation": explanation, "recommendation": recommendation,
        "confidence": None, "detector": "Deterministic QA",
        "diagnostic_version": 1,
        "reason": reason or summary,
    }
    if kind:
        finding["kind"] = kind
    if detected_text:
        finding["detected_text"] = detected_text
    finding.update({key: value for key, value in extra.items() if value is not None})
    return finding


def check_translation_batch(sources, targets, glossary, target_lang,
                            section_profile=None):
    """确定性检查一批译文：空译、保留项丢失、源语残留、锁定术语合规（scope 感知）。"""
    findings = []
    for i, (src, tgt) in enumerate(zip(sources, targets)):
        raw_src, raw_tgt = src, tgt
        src = "" if src is None else str(src)
        tgt = "" if tgt is None else str(tgt)
        if not tgt.strip():
            findings.append(_deterministic_finding(
                i, "blocking", "completeness", "译文为空",
                "原文存在，但当前段落没有任何译文内容，无法确认信息是否被完整传达。",
                "补译本段后，检查是否覆盖原文的全部句子、限制条件和专有名词。",
                source_span=src, target_span=""))
            continue
        target_report = _translation_target.validate_translation_target(
            raw_src, raw_tgt, segment_index=i)
        for issue in target_report["issues"]:
            findings.append(_deterministic_finding(
                i, "blocking", "format_integrity",
                issue["message"],
                "目标文本仍带有模型 transport 或解释包装，不能作为普通正文交付。",
                "重新解析或重新翻译本段，只保留最终目标文本后再检查。",
                source_span=src, target_span=tgt,
                reason=issue["message"], kind="translation_target_invariant",
                invariant_code=issue["code"]))
        # 完整性检查：拦截截断译文。
        # 实测根因：审校/修复环节的整段替换把长段译文换成了一句修正。
        if is_incomplete_translation(src, tgt):
            reason = (f"疑似漏译/截断：原文 {len(src)} 字符/{_count_sentences(src)} 句，"
                      f"译文仅 {len(tgt.strip())} 字符/{_count_sentences(tgt)} 句")
            findings.append(_deterministic_finding(
                i, "blocking", "completeness", "译文疑似遗漏或被截断",
                "原文长度和句子数量与当前译文不匹配，译文可能没有覆盖完整内容。",
                "对照原文逐句补齐缺失内容，并确认修复后的译文通过完整性复验。",
                source_span=src, target_span=tgt, reason=reason))
        source_words = re.findall(r"[A-Za-z]+", src or "")
        whole_source_preserved = any(
            str(entry.get("behavior") or "") == "preserve"
            and str(entry.get("status") or "") == "locked"
            and re.sub(r"\s+", " ", str(entry.get("source") or "")).strip().casefold()
            == re.sub(r"\s+", " ", src or "").strip().casefold()
            for entry in glossary or [])
        if target_lang != "English" and len(source_words) >= 2 \
                and re.sub(r"\s+", " ", src or "").strip().casefold() \
                == re.sub(r"\s+", " ", tgt or "").strip().casefold() \
                and not whole_source_preserved:
            findings.append(_deterministic_finding(
                i, "actionable", "translation_completion", "译文仍保留整段源文",
                "当前译文与英文原文完全相同，除非这是明确的保留项，否则可能尚未完成翻译。",
                "确认该段是否属于项目规定的保留内容；若不是，请重新翻译并保留专有名词的必要形式。",
                source_span=src, target_span=tgt,
                reason="译文与英文源段完全相同，疑似整段未翻译"))
        for token, kind in extract_preserved_tokens(src).items():
            if token not in tgt:
                severity = PRESERVE_SEVERITY.get(kind, "actionable")
                findings.append(_deterministic_finding(
                    i, severity, "format_integrity",
                    f"译文遗漏必须保留的{kind}「{token}」",
                    f"原文包含必须保留的 {kind}「{token}」，但当前译文中找不到该内容。",
                    "补回该保留项后重新检查占位符、链接或引用格式，确认其余译文未被破坏。",
                    source_span=token, target_span="",
                    reason=f"保留项 {kind}「{token}」在译文中丢失", kind=kind))
        for residual, sev in find_residuals(src, tgt, target_lang):
            findings.append(_deterministic_finding(
                i, sev, "source_language_residue", f"译文残留源语片段「{residual}」",
                "当前译文仍包含原文语言片段，可能意味着该部分未完成翻译，或未经确认地保留了源语。",
                "确认该片段是否为有意保留的专名或术语；若不是，请翻译后检查术语和上下文一致性。",
                source_span=residual if residual in src else None,
                target_span=residual, reason=f"疑似残留源语片段「{residual}」",
                kind="source_residue", detected_text=residual))
        findings.extend(check_glossary_compliance(
            src, tgt, glossary, segment_id=i, section_profile=section_profile))
        for f in findings:
            if "segment_index" not in f:
                f["segment_index"] = i
            if "segment_id" not in f or f.get("segment_id") is None:
                f["segment_id"] = i
    return findings


def _globalize_batch_findings(findings, offset):
    """Persist repair findings in document-global, never batch-local, space."""
    return [
        {**finding,
         "segment_id": offset + finding["segment_id"],
         "segment_index": offset + finding["segment_index"]}
        for finding in findings
    ]


# ================= 语义批次（对齐 localize-anything 的上下文批次）=================
BATCH_SIZE = 4
MAX_BATCH_CHARS = 1600
TRANSLATION_MAX_BATCH_CHARS = 2400


def make_batches(paragraphs, batch_size=BATCH_SIZE, max_chars=MAX_BATCH_CHARS,
                 semantic_units=None):
    """把段落聚成批次；提供 semantic_units 时不跨单元边界。"""
    if semantic_units:
        batches = []
        ranges = []
        for unit in semantic_units:
            if not isinstance(unit, dict):
                continue
            try:
                start = int(unit.get("start_segment"))
                end = int(unit.get("end_segment"))
            except (TypeError, ValueError):
                continue
            if 0 <= start <= end < len(paragraphs):
                ranges.append((start, end))
        for start, end in sorted(ranges):
            batches.extend(make_batches(paragraphs[start:end + 1], batch_size, max_chars))
        covered = {index for start, end in ranges for index in range(start, end + 1)}
        range_size = sum(end - start + 1 for start, end in ranges)
        if batches and len(covered) == len(paragraphs) and range_size == len(paragraphs):
            return batches
    batches, cur, n = [], [], 0
    for p in paragraphs:
        if cur and (len(cur) >= batch_size or n + len(p) > max_chars):
            batches.append(cur)
            cur, n = [], 0
        cur.append(p)
        n += len(p)
    if cur:
        batches.append(cur)
    return batches


def _translation_evidence_index(
    paras, pairs, batch_pairs, glossary, document_profile,
    document_synopsis, section_digests, findings_all, blind=False,
    candidate_targets=None,
):
    return _translation_evidence.TranslationEvidenceIndex(
        paras, pairs + batch_pairs, glossary, document_profile,
        document_synopsis, section_digests, findings_all,
        blind=blind, candidate_targets=candidate_targets)


def _batch_section_profile(document_profile, offset, batch_len):
    """按全局段区间匹配 section profile（用于相关术语的 section:<id> scope）。"""
    if not document_profile:
        return None
    for sec in document_profile.get("sections") or []:
        if not isinstance(sec, dict):
            continue
        try:
            start = int(sec.get("start_segment"))
            end = int(sec.get("end_segment"))
        except (TypeError, ValueError):
            continue
        if start <= offset and offset + batch_len - 1 <= end:
            return sec
    return None


# ================= 翻译记忆（对齐 localize-anything 的 TM：仅收录审校通过段落）=================
def tm_path():
    return OUTPUT_DIR / "translation_memory.json"


def _tm_eligible(source, target):
    """翻译记忆资格：源文必须有字母/数字（纯符号装饰行不入库），译文非空。"""
    return bool(re.search(r"[A-Za-z0-9\u4e00-\u9fff]", source or "")) \
        and bool((target or "").strip()) \
        and not _translation_target.is_translation_transport_wrapper(target)


def load_tm():
    """加载翻译记忆并自清洗：非法条目（无字母源文/空译文/未过审校）直接丢弃。

    翻译记忆是错误放大器（一次错译会复制到全书），因此加载即消毒，
    防止旧版本或异常写入留下的污染条目继续命中。
    """
    p = tm_path()
    if p.is_file():
        try:
            raw = json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            return {}
        return {k: v for k, v in raw.items()
                if isinstance(v, dict) and v.get("reviewed")
                and _tm_eligible(k, v.get("target"))}
    return {}


def save_tm(tm):
    p = tm_path()
    p.parent.mkdir(parents=True, exist_ok=True)
    tmp = p.with_suffix(".json.tmp")
    tmp.write_text(json.dumps(tm, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(p)


# ================= 翻译 / 修复 / 审校（对齐 localize-anything 的三通道）=================
def _translator_system(glossary_text, style_rules, target_lang):
    return (f"你是一位学术翻译专家，请将用户提供的段落翻译成{target_lang}。\n"
            f"规则：只翻译可翻译的正文；作者姓名、机构名、品牌名、URL、邮箱、DOI、"
            f"引用标注（如 [12]）等保留原文；译文须与原文一一对应并保持顺序。\n"
            f"{glossary_text}\n"
            f"{style_rules}\n"
            "翻译前请在内部检查：指代与照应关系、术语和专名、理论/非字面用法、"
            "临时造词、长句逻辑关系与修辞功能；优先复用已提供的术语、实体和连续性选择，"
            "人工锁定实体/术语优先于审校、TM 和生成式建议；生成式实体提示不得覆盖锁定术语。"
            "避免明显的逐词直译和词典首义机械替换。\n"
            "不要输出分析过程，只输出最终译文。\n"
            "请严格输出合法的 JSON 字符串数组，不要包含任何解释文字。")


def _invoke_llm(call_fn, provider, api_key, model, system_prompt, user_prompt,
                temperature, response_format=None):
    """Call old and new provider/test doubles through one compatibility gate."""
    kwargs = {"temperature": temperature}
    if response_format is not None:
        kwargs["response_format"] = response_format
    try:
        return call_fn(provider, api_key, model, system_prompt, user_prompt, **kwargs)
    except TypeError:
        try:
            return call_fn(provider, api_key, model, system_prompt, user_prompt,
                           temperature=temperature)
        except TypeError:
            return call_fn(provider, api_key, model, system_prompt, user_prompt)


def _native_translation_response_format(provider, model, expected):
    capabilities = _model_roles.provider_capabilities(PROVIDERS, provider, model)
    if capabilities.get("supports_json_schema") and capabilities.get(
            "supports_response_format"):
        return _translation_protocol.json_schema_for_translations(expected)
    return None


def translate_batch(segments, ctx_prev, ctx_next, glossary_text, style_rules, target_lang,
                    provider, api_key, model, context_packet=None, call_llm_fn=None):
    """翻译一个语义批次，返回与 segments 等长的译文列表；失败抛出 RuntimeError。"""
    numbered = "\n".join(f"{i + 1}. {s}" for i, s in enumerate(segments))
    if context_packet is not None:
        # Keep the system prefix invariant across batches.  Glossary/style are
        # carried once in the ordered context packet below.
        sys_prompt = _translator_system("", "", target_lang)
        user_prompt = _context.render_context_packet(context_packet)
    else:
        sys_prompt = _translator_system(glossary_text, style_rules, target_lang)
        context = ""
        if ctx_prev:
            context += "【前文上下文】：\n" + "\n".join(f"- {s}" for s in ctx_prev) + "\n\n"
        if ctx_next:
            context += "【后文上下文】：\n" + "\n".join(f"- {s}" for s in ctx_next) + "\n\n"
        user_prompt = f"{context}待翻译段落（按序号返回等长译文数组）：\n{numbered}"
    call_fn = call_llm_fn or call_llm
    response_format = _native_translation_response_format(
        provider, model, len(segments))
    last_err = None
    for _attempt in range(3):
        try:
            res = _invoke_llm(
                call_fn, provider, api_key, model, sys_prompt, user_prompt,
                temperature=0.3, response_format=response_format)
            return _translation_protocol.parse_translation_response(
                res, len(segments))
        except Exception as e:
            last_err = e
            if is_rate_limited(e):
                time.sleep(15)
    raise RuntimeError(
        f"批次翻译失败（{len(segments)} 段）："
        f"{last_err or '模型返回格式异常或数量不匹配'}"
    )


def repair_batch(sources, targets, findings, glossary_text, style_rules, target_lang,
                 provider, api_key, model, call_llm_fn=None):
    """根据确定性检查发现的问题自动修复一批译文；返回与 sources 等长的译文列表。"""
    numbered = "\n".join(
        f"{i + 1}. 原文：{s}\n   译文：{t}" for i, (s, t) in enumerate(zip(sources, targets)))
    issues = "\n".join(f"- 段落 {f['segment_index'] + 1}: [{f['severity']}] {f['reason']}"
                       for f in findings)
    sys_prompt = _translator_system(glossary_text, style_rules, target_lang)
    user_prompt = ("以下译文未通过检查，请仅修正有问题的段落，其余段落保持原样，"
                   f"返回与段落数相同的 JSON 字符串数组：\n\n{numbered}\n\n问题清单：\n{issues}")
    call_fn = call_llm_fn or call_llm
    response_format = _native_translation_response_format(
        provider, model, len(sources))
    for _attempt in range(3):
        try:
            res = _invoke_llm(
                call_fn, provider, api_key, model, sys_prompt, user_prompt,
                temperature=0.2, response_format=response_format)
            arr = _translation_protocol.parse_translation_response(
                res, len(sources))
            return [clean_xml_chars(item).strip() for item in arr]
        except Exception as e:
            if is_rate_limited(e):
                time.sleep(15)
    raise RuntimeError("自动修复失败：模型返回格式异常")


def review_translation_batch(sources, targets, glossary_text, style_rules, target_lang,
                             provider, api_key, model):
    """独立审校一个批次（与翻译分离的 prompt/上下文），返回 (findings, failed)。"""
    numbered = "\n".join(
        f"{i + 1}. 原文：{s}\n   译文：{t}" for i, (s, t) in enumerate(zip(sources, targets)))
    sys_prompt = (f"你是一位独立的翻译审校专家，负责审查机器译文。请检查：语义准确性、术语一致性、"
                  f"漏译/增译、目标语言自然度与风格。只报告真实存在的问题，"
                  f"不要为低风险或主观偏好制造 finding。\n"
                  f"severity 只允许以下三种：blocking（结构/占位符/语义严重错误）、"
                  f"actionable（应修正的问题）、informational（建议）。\n"
                  "如果整批译文没有问题，请严格返回空数组 []，不要输出任何 informational 备注。\n"
                  f"{glossary_text}\n"
                  f"{style_rules}\n"
                  '请严格输出 JSON 数组，每项格式：{"segment_index": 0, "severity": "actionable", '
                  '"reason": "问题说明", "suggested_target": "可选：修正后的译文"}')
    user_prompt = f"待审校段落（目标语言：{target_lang}）：\n{numbered}"
    for _attempt in range(3):
        try:
            res = call_llm(provider, api_key, model, sys_prompt, user_prompt, temperature=0.2)
            arr = parse_json_array(res)
            if arr is None:
                return [], True
            return arr, False
        except Exception as e:
            if is_rate_limited(e):
                time.sleep(15)
            else:
                break
    return [], True


# ================= 自动标注（三色学习重点）=================
# 红=生僻词/难词；黄=专业名词（特殊译法）；青绿=翻译难点句（特别译法）。
ANNOT_BATCH_SIZE = 10
ANNOT_MAX_PER_SEG = {"rare": 3, "domain": 3, "hard": 2}

# 常用英语词表（en_50k 字幕语料前 14000 词），用于把 LLM 滥标的"生僻词"挡回去
_DATA_DIR = Path(_models.__file__).resolve().parent / "resources"
_COMMON_WORDS = None


def _common_words():
    global _COMMON_WORDS
    if _COMMON_WORDS is None:
        try:
            _COMMON_WORDS = set(
                _DATA_DIR.joinpath("en_common.txt").read_text(encoding="utf-8").splitlines())
        except OSError:
            _COMMON_WORDS = set()
    return _COMMON_WORDS


_INFLECTION_SUFFIXES = (("ily", "y"), ("ness", ""), ("ment", ""), ("tion", ""),
                        ("sion", ""), ("ing", ""), ("ed", ""), ("er", ""),
                        ("est", ""), ("es", ""), ("ly", ""), ("s", ""))


def _base_form(word):
    """词形还原（一次）：grinning->grin、speedily->speedy、boxes->box。"""
    w = word
    for suffix, repl in _INFLECTION_SUFFIXES:
        if w.endswith(suffix) and len(w) > len(suffix) + 2:
            w = w[:-len(suffix)] + repl
            break
    if len(w) > 3 and w[-1] == w[-2] and w[-1] not in "eio":
        w = w[:-1]  # grinning -> grinn -> grin
    return w


def _is_common_word(token):
    """token 或其词形还原是否在常用词表内。"""
    common = _common_words()
    if not common:
        return False
    w = token.lower().strip('"\'(),;:!?')
    if not w:
        return True
    if w in common:
        return True
    base = _base_form(w)
    return base != w and base in common


_KINSHIP_TITLE_RE = re.compile(
    r"^(?:grandma|grandpa|grandmother|grandfather|aunt|uncle|mr|mrs|ms|dr|sir|lady|lord|"
    r"mother|father|mom|dad|brother|sister|captain|colonel|major|general|rabbi|"
    r"professor|doctor)\b", re.IGNORECASE)


def _rare_ok(span_text, token_freq):
    """生僻词门槛：单 token、非常用词（含词形还原）、全书出现次数少。"""
    span = (span_text or "").strip()
    if not span or " " in span or "\u00a0" in span:
        return False  # 只接受单词（可带连字符），短语/名句一律不要
    w = span.lower().strip('"\'(),;:!?')
    if len(w) < 4:
        return False
    if _is_common_word(w):
        return False
    if token_freq.get(w, 0) >= 8:
        return False  # 全书反复出现的词不算生僻
    return True


def _domain_ok(span_text):
    """专业名词门槛：称谓+人名、全常用词短语不算专业名词。"""
    span = (span_text or "").strip()
    if not span:
        return False
    if _KINSHIP_TITLE_RE.match(span):
        return False  # Grandma Sarah / Mr. Smith 之类
    tokens = re.findall(r"[A-Za-z]+(?:['’\-][A-Za-z]+)*", span)
    if not tokens:
        return False
    if all(_is_common_word(t) for t in tokens):
        return False  # Translation from Hebrew 之类全是常用词
    return True


def _normalize_annotations(annotations):
    """标注字典键归一化：JSON 落盘后键变字符串，统一回 int；越界/非列表值丢弃。

    这是本项目反复踩过的坑（标注渲染、过滤、续跑三处各自处理过一次），
    统一入口避免再次各修各的。
    """
    out = {}
    for k, v in (annotations or {}).items():
        try:
            gi = int(k)
        except (TypeError, ValueError):
            continue
        if isinstance(v, list):
            out[gi] = v
    return out


def _clean_annotations(annotations, pairs):
    """过滤 + 去重 + 数量上限；术语表覆盖的 domain（note 以"术语："开头）不参与过滤。"""
    token_freq = {}
    for pr in pairs:
        for tok in re.findall(r"[A-Za-z]+", pr["source"].lower()):
            token_freq[tok] = token_freq.get(tok, 0) + 1
    cleaned = {}
    for gi, items in _normalize_annotations(annotations).items():
        if not (0 <= gi < len(pairs)):
            continue  # 段已被删除等场景：越界键丢弃
        src_text = pairs[gi]["source"]
        seen, kept, counts = set(), [], {"rare": 0, "domain": 0, "hard": 0}
        for it in items:
            atype = it["type"]
            s_span = it.get("src_span")
            span_text = src_text[s_span[0]:s_span[1]] if s_span else ""
            overlay = atype == "domain" and str(it.get("note", "")).startswith("术语：")
            if atype == "rare" and not _rare_ok(span_text, token_freq):
                continue
            if atype == "domain" and not overlay and not _domain_ok(span_text):
                continue
            key = (atype, tuple(s_span) if s_span else None,
                   tuple(it["tgt_span"]) if it.get("tgt_span") else None)
            if key in seen:
                continue
            seen.add(key)
            if counts[atype] >= ANNOT_MAX_PER_SEG[atype]:
                continue
            counts[atype] += 1
            kept.append(it)
        if kept:
            cleaned[gi] = kept
    return cleaned


def _annotator_system(target_lang):
    return (f"你是一位翻译教学专家。请从下列{target_lang}双语对照中标注三类学习重点：\n"
            "1. rare：真正的生僻词/难词——英语母语者也未必认识的低频书面词，"
            "如 chicory、muezzin、cacophony。只标注单个单词（最多一个带连字符的复合词）。\n"
            "   严禁标注日常常用词（如 production、grin、rooster、speedily、elementary），"
            "严禁标注短语、引用句或名句（如 'Elementary, my dear Watson'）。\n"
            "2. domain：专业领域名词，译文采用了专门/约定俗成的译法（如术语表、行话、专名译法）；\n"
            "   严禁标注亲属称谓（如 Grandma Sarah）、全常用词短语（如 Translation from Hebrew）、"
            "普通日常表达。\n"
            "3. hard：翻译难度高的句子，译文使用了特别翻译技巧（语序调整、词性转换、拆合句、"
            "文化负载词处理、比喻/双关处理等）；普通直译句不要标注。\n"
            "规则：\n"
            "- 每个段落最多 2 个 rare、2 个 domain、1 个 hard；标注真正有价值的，宁缺毋滥；\n"
            "- src 必须是原文中的原文字符串，tgt 必须是对应译文中的字符串（hard 可以是整句/整段）；\n"
            "- note 用一句话说明标注理由或所用译法；\n"
            '严格输出 JSON 数组，每项格式：{"seg": 1, "type": "rare", "src": "...", '
            '"tgt": "...", "note": "..."}。seg 为段落序号（从 1 开始）。不要输出任何解释文字。')


def annotate_batch(pairs_slice, target_lang, provider, api_key, model):
    """标注一个批次，返回 [{seg, type, src, tgt, note}]（seg 为 0 基）；解析失败返回 []。"""
    numbered = "\n\n".join(
        f"--- 段落 {i + 1} ---\n原文：{p['source']}\n译文：{p['target']}"
        for i, p in enumerate(pairs_slice))
    sys_prompt = _annotator_system(target_lang)
    user_prompt = f"待标注双语段落：\n{numbered}"
    for _attempt in range(3):
        try:
            res = call_llm(provider, api_key, model, sys_prompt, user_prompt, temperature=0.2)
            arr = parse_json_array(res)
            if arr is None:
                return []
            out = []
            for item in arr:
                if not isinstance(item, dict):
                    continue
                seg = item.get("seg")
                if not isinstance(seg, int) or not (1 <= seg <= len(pairs_slice)):
                    continue
                atype = item.get("type")
                if atype not in ANNOTATION_COLORS:
                    continue
                src = str(item.get("src") or "").strip()
                tgt = str(item.get("tgt") or "").strip()
                note = str(item.get("note") or "").strip()
                if not src:
                    continue
                out.append({"seg": seg - 1, "type": atype, "src": src,
                            "tgt": tgt, "note": note})
            return out
        except Exception as e:
            if is_rate_limited(e):
                time.sleep(10)  # 限流退避后重试，避免整批标注静默丢失
                continue
            return []
    return []


def _compose_spans(spans, text_len):
    """按边界切分 + 优先级（rare>domain>hard）覆盖，返回互不重叠的 (start,end,type) 列表。

    难点句常覆盖整段，词级标注嵌在其中：切到所有起点/终点后，每段取覆盖它的最高优先级。
    """
    priority = {"rare": 0, "domain": 1, "hard": 2}
    clipped = []
    for s, e, t in spans:
        s = max(0, min(int(s), text_len))
        e = max(s, min(int(e), text_len))
        if s < e:
            clipped.append((s, e, t))
    if not clipped:
        return []
    bounds = sorted({0, text_len} | {x for s, e, _ in clipped for x in (s, e)})
    composed = []
    for a, b in zip(bounds, bounds[1:]):
        if a >= b:
            continue
        mid = (a + b) / 2
        covering = [(priority[t], t) for s, e, t in clipped if s <= mid < e]
        if covering:
            composed.append((a, b, min(covering)[1]))
    return composed


def annotate_stage(state, job_id, glossary, provider, api_key, model, target_lang,
                   on_caption=None):
    """三色自动标注：LLM 识别 + 术语表确定性覆盖（专业名词=黄色必标）。"""
    if state.get("annotations_done"):
        return state
    pairs = state["pairs"]
    annotations = _normalize_annotations(state.get("annotations"))
    batches = make_batches([p["source"] for p in pairs], batch_size=ANNOT_BATCH_SIZE,
                           max_chars=2600)
    # 断点按“已标注段数”记录（而非批次号），避免批大小调整后续跑错位
    done_offset = state.get("annotations_done_offset") or 0
    start_bi = next((bi for bi, b in enumerate(batches)
                     if sum(len(x) for x in batches[:bi]) >= done_offset), len(batches))
    if done_offset:
        if on_caption:
            on_caption(f"↩️ 从第 {done_offset + 1} 段（批次 {start_bi + 1}/{len(batches)}）继续标注...")

    # 1) LLM 标注
    failed_batches = 0
    for bi in range(start_bi, len(batches)):
        batch_srcs = batches[bi]
        offset = sum(len(b) for b in batches[:bi])
        slice_pairs = pairs[offset:offset + len(batch_srcs)]
        if on_caption and bi % 10 == 0:
            on_caption(f"🎨 自动标注第 {offset + 1}-{offset + len(slice_pairs)} 段"
                       f"（共 {len(pairs)} 段，批次 {bi + 1}/{len(batches)}）...")
        items = annotate_batch(slice_pairs, target_lang, provider, api_key, model)
        if not items:
            failed_batches += 1
        for item in items:
            gi = offset + item["seg"]
            src, tgt, atype, note = item["src"], item["tgt"], item["type"], item["note"]
            src_span = _find_span(pairs[gi]["source"], src)
            tgt_span = _find_span(pairs[gi]["target"], tgt) if tgt else None
            if atype == "hard":
                # 难句兜底：找不到精确片段时标整段
                if src_span is None:
                    src_span = (0, len(pairs[gi]["source"]))
                if tgt_span is None:
                    tgt_span = (0, len(pairs[gi]["target"])) if pairs[gi]["target"] else None
            elif src_span is None:
                continue  # 词级标注必须在原文中定位
            annotations.setdefault(gi, []).append(
                {"type": atype, "src_span": list(src_span) if src_span else None,
                 "tgt_span": list(tgt_span) if tgt_span else None, "note": note})
        # 每批落盘：断点粒度 = 一个批次
        state["annotations"] = annotations
        state["annotations_done_offset"] = offset + len(slice_pairs)
        save_job_state(job_id, state)

    # 2) 术语表确定性覆盖：专业名词（特殊译法）-> 黄色
    for gi, pr in enumerate(pairs):
        for entry in glossary:
            term = (entry.get("source") or "").strip()
            if len(term) < 2 or term not in pr["source"]:
                continue
            span = _find_span(pr["source"], term)
            tgt_span = None
            tgt_term = (entry.get("target") or "").strip()
            if tgt_term:
                tgt_span = _find_span(pr["target"], tgt_term)
            annotations.setdefault(gi, []).append(
                {"type": "domain", "src_span": list(span) if span else None,
                 "tgt_span": list(tgt_span) if tgt_span else None,
                 "note": f"术语：{term} -> {tgt_term or '保留原文'}"})

    # 3) 确定性过滤（常用词/称谓/全常用词短语）+ 去重 + 数量上限
    cleaned = _clean_annotations(annotations, pairs)
    state["annotations"] = cleaned
    state["annotations_done"] = True
    state["annotations_failed_batches"] = failed_batches
    save_job_state(job_id, state)
    if on_caption:
        total = sum(len(v) for v in cleaned.values())
        on_caption(f"✅ 自动标注完成：{total} 处（失败批次 {failed_batches}/{len(batches)}）")
    return state


def translate_stage(state, job_id, glossary, provider, api_key, model, target_lang,
                    style_rules, enable_review, use_tm=True, document_profile=None,
                    on_status=None, on_caption=None, translator_config=None,
                    reviewer_config=None):
    """阶段二：语义批次翻译 + 确定性检查/修复 + 独立审校 + 翻译记忆。

    对齐 localize-anything 经验：
    - 语义批次：≤4 段一组，携带前后文，保留每批落盘的断点粒度；
    - 概念化术语表：锁定术语强制首选译名/禁止译名，保留项强制原样；
    - 确定性检查：占位符/URL/引用等保留项、残留原文、锁定术语合规，问题自动修复一轮；
    - 独立审校：actionable 建议经确定性复验后应用，blocking 记录给用户确认；
    - 翻译记忆：仅审校通过的段落入库，精确命中直接复用。
    """
    # 严格术语治理门禁：存在待审核候选术语且未冻结（且未显式跳过）时，
    # 任何入口都禁止开始翻译。导入的锁定术语视为已固定，不构成阻塞。
    pending = [e for e in (glossary or [])
               if (e.get("status") or "").lower() == "candidate"]
    if state.get("quality_mode") and pending \
            and not state.get("glossary_frozen") \
            and not state.get("quality_bypass"):
        raise RuntimeError(
            "严格术语治理：仍有候选术语未审核（术语表尚未冻结），"
            "禁止开始翻译（请在术语审核面板冻结后继续）")
    translator_config = _model_roles.normalize_role_config(
        translator_config, fallback_provider=provider, fallback_model=model,
        fallback_api_key=api_key)
    reviewer_config = _model_roles.normalize_role_config(
        reviewer_config, fallback_provider=provider, fallback_model=model,
        fallback_api_key=api_key)
    translator_call = _model_roles.make_role_call(call_llm, translator_config)
    reviewer_call = _model_roles.make_role_call(call_llm, reviewer_config)
    tm = load_tm() if use_tm else {}
    if use_tm:
        recovered, pending_events = _checkpoint.reconcile_translation_memory(
            tm, state, job_dir(job_id))
        if recovered:
            save_tm(tm)
            state["tm_recovered_count"] = state.get("tm_recovered_count", 0) + pending_events
            save_job_state(job_id, state)
    paras = state["paras"]
    pairs = state["pairs"]
    truncated_indexes = []

    def _commit_translation_batch(batch_pairs, offset):
        nonlocal truncated_indexes
        pairs.extend(batch_pairs)
        changed_indexes = list(range(offset, offset + len(batch_pairs)))
        if changed_indexes:
            _mark_translation_truth_changed(
                job_id, state, changed_indexes,
                "翻译流水线写入 CURRENT_TRANSLATION；记录该批次影响范围",
                actor="pipeline", action="translation_batch")
        truncated_indexes = []
        save_job_state(job_id, state)

    def _save_translation_failure():
        nonlocal truncated_indexes
        if truncated_indexes:
            _mark_translation_truth_changed(
                job_id, state, truncated_indexes,
                "断点恢复截断未完成批次；CURRENT_TRANSLATION 已改变",
                actor="pipeline", action="translation_checkpoint_truncate")
            truncated_indexes = []
        save_job_state(job_id, state)

    batches = make_batches(
        paras, max_chars=TRANSLATION_MAX_BATCH_CHARS,
        semantic_units=state.get("semantic_units")
        or state.get("section_digests") or None)
    registry = _entity_registry.EntityRegistry(state.get("entity_registry") or [])

    # 断点：从第一个未完成批次继续；若中间批次不完整则截断重译
    cum_end, start_batch = 0, 0
    for bi, b in enumerate(batches):
        prev_end = cum_end
        cum_end += len(b)
        if cum_end > len(pairs):
            if len(pairs) > prev_end:
                truncated_indexes = list(range(prev_end, len(pairs)))
                del pairs[prev_end:]
            start_batch = bi
            break
    else:
        start_batch = len(batches)

    stats = state.setdefault("review_stats", {
        "reviewed_segments": 0, "batches_reviewed": 0,
        "blocking": 0, "actionable": 0, "informational": 0, "review_failed": 0,
    })
    findings_all = state.setdefault("findings", [])
    from transpraxis.terminology import (
        detect_glossary_conflicts as _detect_conflicts,
        glossary_block as _glossary_block,
        select_glossary_for_segments as _select_glossary,
    )
    frozen_hash = (state.get("glossary_frozen") or {}).get("glossary_hash")
    sections = (document_profile or {}).get("sections") or []
    section_digests = state.get("section_digests") or []
    document_synopsis = state.get("document_synopsis") or {}

    for bi in range(start_batch, len(batches)):
        batch = batches[bi]
        offset = sum(len(b) for b in batches[:bi])
        _checkpoint.append_event(job_dir(job_id), {
            "batch": bi, "offset": offset, "phase": "generation_started",
            "segment_count": len(batch),
        })
        if on_status:
            on_status(f"【阶段二】双语翻译与术语严格注入...（批次 {bi + 1}/{len(batches)}）")
        if on_caption:
            on_caption(f"🌍 正在翻译第 {offset + 1}-{offset + len(batch)} 段（共 {len(paras)} 段）...")

        ctx_prev = paras[max(0, offset - 2):offset]
        ctx_next = paras[min(len(paras), offset + len(batch)):min(len(paras), offset + len(batch) + 2)]
        previous_target = _context.select_target_context(pairs, offset, limit=2)
        section_digest = _context.digest_for_segment(section_digests, offset)

        # 1) 翻译记忆精确命中直接复用
        batch_pairs = [None] * len(batch)
        to_translate = []  # (index, clean_source)
        for i, para in enumerate(batch):
            clean_src = para.replace('\n', ' ')
            hit = tm.get(clean_src)
            if hit and hit.get("reviewed") and hit.get("target"):
                batch_pairs[i] = {"source": clean_src, "target": hit["target"],
                                  "initial_target": hit["target"],
                                  "accepted_target": hit["target"],
                                  "target_provenance": "tm_approved",
                                  "reviewed": True, "review_status": "tm_approved",
                                  "from_tm": True}
                state["tm_used_count"] = state.get("tm_used_count", 0) + 1
            elif not re.search(r"[A-Za-z0-9\u4e00-\u9fff]", clean_src):
                # 纯符号段落（章节分隔装饰等）：不是正文，原样保留，不调模型
                batch_pairs[i] = {"source": clean_src, "target": clean_src,
                                  "initial_target": clean_src,
                                  "accepted_target": clean_src,
                                  "target_provenance": "reviewed",
                                  "reviewed": True, "review_status": "reviewed_clean",
                                  "from_tm": False}
            else:
                to_translate.append((i, clean_src))

        # 相关术语选择：只注入本批实际出现的 locked translate / preserve 条目，
        # provisional 仅作受限建议；记录实际注入的 entry IDs（审计）。
        texts = [t for _, t in to_translate]
        section_profile = _batch_section_profile(document_profile, offset, len(batch))
        selected, injected_ids = _select_glossary(
            texts, glossary + _knowledge.provisional_hints(
                state.get("knowledge_candidates") or [],
                authoritative_entries=glossary),
            document_profile, section_profile)
        glossary_text = _glossary_block(selected)
        entity_hints = registry.hints_for(
            texts, glossary_entries=glossary, limit=12)
        context_packet = _context.compile_context_packet(
            document_profile=document_profile,
            document_synopsis=document_synopsis,
            section_digest=section_digest,
            glossary_text=glossary_text,
            previous_source=ctx_prev,
            previous_target=previous_target,
            next_source=ctx_next,
            current_batch=texts,
            style_rules=style_rules,
            entity_hints=entity_hints,
        )
        state.setdefault("context_packet_log", []).append({
            "batch": bi,
            "offset": offset,
            **_context.context_metadata(context_packet),
        })
        state.setdefault("glossary_injection_log", []).append({
            "batch": bi,
            "offset": offset,
            "entry_ids": injected_ids,
            "glossary_version": (state.get("glossary_frozen") or {}).get("version"),
            "glossary_hash": (state.get("glossary_frozen") or {}).get("glossary_hash"),
            "entity_hint_count": len(entity_hints),
            "document_synopsis_summary": (document_synopsis or {}).get("summary", ""),
            "section_digest_summary": (section_digest or {}).get("summary", ""),
        })

        # 2) 未命中段落批次翻译
        if to_translate:
            try:
                targets = translate_batch(texts, ctx_prev, ctx_next, glossary_text, style_rules,
                                          target_lang, translator_config["provider"],
                                          translator_config["api_key"],
                                          translator_config["model"],
                                          context_packet=context_packet,
                                          call_llm_fn=translator_call)
            except RuntimeError as batch_error:
                if "任务已请求取消" in str(batch_error):
                    raise
                failure = {
                    "batch": bi, "offset": offset, "segment_count": len(texts),
                    "reason": str(batch_error)[:500], "recovered": False,
                }
                state.setdefault("translation_failures", []).append(failure)
                _checkpoint.append_event(job_dir(job_id), {
                    "batch": bi, "offset": offset,
                    "phase": "translation_protocol_failed",
                    "reason": str(batch_error)[:240],
                })
                # 批次解析失败时降级为逐段翻译，保证进度不中断
                if len(texts) == 1:
                    _save_translation_failure()
                    raise
                if on_caption:
                    on_caption("⚠️ 批次翻译返回格式异常，降级为逐段翻译...")
                targets = []
                try:
                    for t in texts:
                        single_packet = dict(context_packet, current_batch=[t])
                        targets.append(translate_batch(
                            [t], ctx_prev, ctx_next, glossary_text, style_rules, target_lang,
                            translator_config["provider"], translator_config["api_key"],
                            translator_config["model"], context_packet=single_packet,
                            call_llm_fn=translator_call)[0])
                except Exception as single_error:
                    failure["reason"] = str(single_error)[:500]
                    _save_translation_failure()
                    raise
                failure["recovered"] = True
            for (i, src), tgt in zip(to_translate, targets):
                cleaned_tgt = clean_xml_chars(tgt).replace('\n', ' ')
                batch_pairs[i] = {"source": src, "target": cleaned_tgt,
                                  "initial_target": cleaned_tgt,
                                  "target_provenance": "generated",
                                  "reviewed": False, "review_status": "not_reviewed",
                                  "from_tm": False}

        batch_sources = [p["source"] for p in batch_pairs]
        batch_targets = [p["target"] for p in batch_pairs]
        for p in batch_pairs:
            p["glossary_entry_ids"] = list(injected_ids)
            p["glossary_hash_used"] = frozen_hash
        _checkpoint.append_event(job_dir(job_id), {
            "batch": bi, "offset": offset, "phase": "generation_done",
        })
        findings = check_translation_batch(
            batch_sources, batch_targets, glossary, target_lang,
            section_profile=section_profile)

        # 3) 确定性问题自动修复（一轮）
        fixable = [f for f in findings if f["severity"] in ("blocking", "actionable")]
        if fixable and len(fixable) <= 8:
            if on_caption:
                on_caption(f"🔧 发现 {len(fixable)} 个确定性问题，正在自动修复...")
            try:
                repaired = repair_batch(batch_sources, batch_targets, fixable, glossary_text,
                                        style_rules, target_lang,
                                        translator_config["provider"],
                                        translator_config["api_key"],
                                        translator_config["model"],
                                        call_llm_fn=translator_call)
                if repaired and len(repaired) == len(batch_pairs):
                    formal_targets = list(batch_targets)
                    shadow_targets = list(formal_targets)
                    for j, p in enumerate(batch_pairs):
                        if not p["from_tm"] and repaired[j] and repaired[j].strip():
                            candidate = clean_xml_chars(repaired[j]).replace('\n', ' ')
                            # 修复结果本身截断时，不接受更差的译文
                            if is_incomplete_translation(batch_sources[j], candidate) \
                                    and not is_incomplete_translation(batch_sources[j], formal_targets[j]):
                                continue
                            shadow_targets[j] = candidate
                    overlay = _repair.create_overlay(
                        formal_targets, shadow_targets, fixable, "deterministic",
                        sources=batch_sources,
                        finding_segment_ids=[offset + f["segment_index"] for f in fixable])
                    shadow_findings = check_translation_batch(
                        batch_sources, shadow_targets, glossary, target_lang,
                        section_profile=section_profile)
                    shadow_findings = _globalize_batch_findings(shadow_findings, offset)
                    blind_findings, blind_failed, blind_trace = [], False, None
                    if enable_review and not any(
                            f["severity"] in ("blocking", "actionable")
                            for f in shadow_findings):
                        shadow_index = _translation_evidence_index(
                            paras, pairs, batch_pairs, glossary, document_profile,
                            document_synopsis, section_digests, findings_all,
                            blind=True,
                            candidate_targets={offset + j: shadow_targets[j]
                                               for j in range(len(shadow_targets))})
                        blind_findings, blind_failed, blind_trace = \
                            _translation_evidence.review_translation_batch_with_evidence(
                                batch_sources, shadow_targets, glossary_text, style_rules,
                                target_lang, reviewer_config["provider"],
                                reviewer_config["api_key"], reviewer_config["model"],
                                shadow_index, call_llm=reviewer_call, blind=True,
                                segment_ids=list(range(offset, offset + len(batch_pairs))),
                                review_identity={
                                    "input_hash": overlay["input_hash"],
                                    "candidate_hash": overlay["candidate_hash"],
                                })
                        if blind_trace:
                            state.setdefault("review_evidence", []).append({
                                "batch": bi, "phase": "shadow_repair", **blind_trace})
                    overlay = _repair.evaluate_overlay(
                        overlay, shadow_findings, blind_findings, blind_failed,
                        review_identity=(blind_trace or {}).get("review_identity"))
                    state.setdefault("repair_overlays", []).append({
                        "batch": bi, "offset": offset, **overlay,
                        "blind_trace": blind_trace,
                    })
                    promoted = _repair.promoted_targets(overlay)
                    for j, p in enumerate(batch_pairs):
                        if not p["from_tm"]:
                            p["target"] = promoted[j]
                batch_targets = [p["target"] for p in batch_pairs]
                findings = check_translation_batch(
                    batch_sources, batch_targets, glossary, target_lang,
                    section_profile=section_profile)
            except Exception as exc:
                state.setdefault("repair_overlays", []).append({
                    "batch": bi, "offset": offset, "source": "deterministic",
                    "status": "rejected", "rejection": "repair_error",
                    "error": str(exc)[:240],
                })
                # 修复失败则保留原译文与 finding。
        _checkpoint.append_event(job_dir(job_id), {
            "batch": bi, "offset": offset, "phase": "deterministic_qa_done",
        })

        # 4) 独立审校：actionable 建议复验后应用；blocking 记录给用户
        review_succeeded = False
        if enable_review:
            stats["batches_reviewed"] += 1
            evidence_index = _translation_evidence_index(
                paras, pairs, batch_pairs, glossary, document_profile,
                document_synopsis, section_digests, findings_all)
            rfindings, failed, review_trace = \
                _translation_evidence.review_translation_batch_with_evidence(
                    batch_sources, batch_targets, glossary_text, style_rules, target_lang,
                    reviewer_config["provider"], reviewer_config["api_key"],
                    reviewer_config["model"], evidence_index, call_llm=reviewer_call,
                    segment_ids=list(range(offset, offset + len(batch_pairs))))
            review_event_id = (
                f"translation-review-{job_id}-{bi}-formal-"
                f"{len(state.get('review_evidence') or [])}")
            review_trace["review_event_id"] = review_event_id
            state.setdefault("review_evidence", []).append({
                "batch": bi, "phase": "formal_review", **review_trace,
            })
            if failed:
                stats["review_failed"] += 1
            if failed:
                # Provider/JSON/protocol failure is not a clean review.  Do
                # not interpret an empty findings list as acceptance.
                for pair in batch_pairs:
                    if not pair.get("from_tm"):
                        pair["review_status"] = "review_failed"
            else:
                review_succeeded = True
                global_to_local = {
                    offset + index: index for index in range(len(batch_pairs))
                }
                for rf in rfindings:
                    sev = rf.get("severity")
                    if sev not in ("blocking", "actionable", "informational"):
                        continue
                    segment_id = rf.get("segment_id")
                    idx = global_to_local.get(segment_id)
                    if idx is None:
                        continue
                    record = {
                        "segment_id": segment_id, "segment_index": segment_id,
                        "severity": sev, "type": "review",
                        "category": rf.get("category") or "semantic_accuracy",
                        "summary": str(rf.get("summary") or rf.get("reason") or "审校发现问题"),
                        "source_span": rf.get("source_span"),
                        "target_span": rf.get("target_span"),
                        "explanation": rf.get("explanation"),
                        "recommendation": rf.get("recommendation"),
                        "confidence": rf.get("confidence"),
                        "detector": rf.get("detector") or "Semantic QA",
                        "diagnostic_version": rf.get("diagnostic_version"),
                        "reason": str(rf.get("reason") or rf.get("summary") or "审校发现问题"),
                        "evidence_refs": list(rf.get("evidence_refs") or []),
                        "review_event_id": review_event_id,
                    }
                    if sev == "actionable" and rf.get("suggested_target") \
                            and not batch_pairs[idx]["from_tm"]:
                        suggested = clean_xml_chars(rf["suggested_target"]).replace('\n', ' ').strip()
                        if suggested:
                            old_target = batch_pairs[idx]["target"]
                            overlay = _repair.create_overlay(
                                [old_target], [suggested], [rf], "review_suggested",
                                sources=[batch_sources[idx]],
                                finding_segment_ids=[segment_id])
                            recheck = _globalize_batch_findings(check_translation_batch(
                                [batch_sources[idx]], [suggested], glossary, target_lang,
                                section_profile=section_profile), segment_id)
                            blind_trace = None
                            blind_findings = []
                            blind_failed = False
                            if not any(f["severity"] in ("blocking", "actionable")
                                       for f in recheck):
                                blind_index = _translation_evidence_index(
                                    paras, pairs, batch_pairs, glossary, document_profile,
                                    document_synopsis, section_digests, findings_all,
                                    blind=True, candidate_targets={segment_id: suggested})
                                blind_findings, blind_failed, blind_trace = \
                                    _translation_evidence.review_translation_batch_with_evidence(
                                        [batch_sources[idx]], [suggested], glossary_text,
                                        style_rules, target_lang,
                                        reviewer_config["provider"],
                                        reviewer_config["api_key"],
                                        reviewer_config["model"], blind_index,
                                        call_llm=reviewer_call, blind=True,
                                        segment_ids=[segment_id], review_identity={
                                            "input_hash": overlay["input_hash"],
                                            "candidate_hash": overlay["candidate_hash"],
                                        })
                            overlay = _repair.evaluate_overlay(
                                overlay, recheck, blind_findings, blind_failed,
                                review_identity=(blind_trace or {}).get("review_identity"))
                            state.setdefault("repair_overlays", []).append({
                                "batch": bi, "offset": offset,
                                "batch_local_ordinal": idx,
                                "segment_id": segment_id, **overlay,
                                "blind_trace": blind_trace,
                            })
                            if blind_trace:
                                state.setdefault("review_evidence", []).append({
                                    "batch": bi, "phase": "suggested_shadow_review",
                                    **blind_trace,
                                })
                            if overlay["status"] == "accepted":
                                batch_pairs[idx]["target"] = suggested
                            else:
                                record["suggested_target"] = suggested
                                findings_all.append(record)
                            continue
                    findings_all.append(record)
            _checkpoint.append_event(job_dir(job_id), {
                "batch": bi, "offset": offset, "phase": "semantic_review_done",
            })
        else:
            _checkpoint.append_event(job_dir(job_id), {
                "batch": bi, "offset": offset, "phase": "semantic_review_skipped",
            })

        # 审校可能修改过译文：对最终译文整体复验一次确定性检查
        findings = check_translation_batch(
            batch_sources, [p["target"] for p in batch_pairs], glossary, target_lang,
            section_profile=section_profile)

        # 批内冲突检测（跨段同术语多译法）——在 TM 入库前执行
        for cf in _detect_conflicts(batch_pairs, glossary, sections=sections):
            segment_id = offset + cf["segment_id"]
            cf["segment_index"] = segment_id
            cf["segment_id"] = segment_id
            findings_all.append(cf)

        # 记录仍未解决的确定性问题
        for f in findings:
            if f["severity"] in ("blocking", "actionable", "informational"):
                segment_id = offset + f["segment_index"]
                record = dict(f)
                record.update({"segment_id": segment_id,
                               "segment_index": segment_id,
                               "severity": f["severity"], "type": "check",
                               "detector": f.get("detector") or "Deterministic QA"})
                findings_all.append(record)

        # 5) 批后知识反馈：只进入 candidate queue，不改变 frozen glossary。
        accepted_for_knowledge = []
        if enable_review and review_succeeded:
            for j, _pair in enumerate(batch_pairs):
                segment_id = offset + j
                segment_findings = [
                    finding for finding in findings_all
                    if finding.get("segment_id", finding.get("segment_index")) == segment_id
                ]
                if not segment_findings:
                    accepted_for_knowledge.append(j)
        # Standard mode also records low-authority continuity observations.
        accepted_for_knowledge = []
        review_bad_segments = {
            finding.get("segment_id")
            for finding in (rfindings if enable_review and review_succeeded else [])
            if finding.get("severity") in ("blocking", "actionable")
        }
        for j, pair in enumerate(batch_pairs):
            local_findings = [
                finding for finding in findings
                if finding.get("segment_index") == j
                and finding.get("severity") in ("blocking", "actionable")
            ]
            if (not pair.get("from_tm") and not local_findings
                    and offset + j not in review_bad_segments):
                accepted_for_knowledge.append(j)
        if accepted_for_knowledge:
            knowledge_segment_ids = [offset + j for j in accepted_for_knowledge]
            knowledge_candidates, knowledge_events, knowledge_warning = \
                _knowledge.observe_batch(
                    [batch_sources[j] for j in accepted_for_knowledge],
                    [batch_pairs[j]["target"] for j in accepted_for_knowledge],
                    paras, pairs,
                    glossary, offset, provider, api_key, model,
                    existing_candidates=state.get("knowledge_candidates") or [],
                    call_llm=translator_call, segment_ids=knowledge_segment_ids,
                    observation_provenance=(
                        "reviewed" if review_succeeded else "generated_continuity"
                    ))
        else:
            # Review failure or any finding leaves no trustworthy observation.
            knowledge_candidates, knowledge_events, knowledge_warning = \
                state.get("knowledge_candidates") or [], [], None
        for event in knowledge_events:
            event["batch"] = bi
            state.setdefault("knowledge_events", []).append(event)
            if event.get("type") == "entity_observation":
                registry.observe(
                    event.get("source"), event.get("observed_target"),
                    entity_type=_entity_registry.entity_type_from_kind(
                        event.get("kind")),
                    segment_id=event.get("segment_id"),
                    provenance=event.get("provenance") or "generated_observation",
                    confidence=0.7 if event.get("provenance") == "reviewed" else 0.35,
                )
            if event.get("type") == "target_conflict":
                segment_id = event.get("segment_id", offset)
                findings_all.append({
                    "segment_id": segment_id, "segment_index": segment_id,
                    "severity": "actionable",
                    "type": "knowledge_conflict",
                    "category": "terminology_consistency",
                    "summary": "观察到的译法与锁定术语不一致",
                    "source_span": event.get("source"),
                    "target_span": event.get("observed_target"),
                    "explanation": "翻译流中观察到的译法与项目锁定术语的首选译名不同，可能造成术语漂移。",
                    "recommendation": "核对当前语境是否构成合理例外；若不是，请统一为锁定术语的首选译名。",
                    "confidence": None, "detector": "Knowledge QA",
                    "diagnostic_version": 1,
                    "reason": event.get("reason") or "翻译流观察译法与锁定术语不一致",
                    "source": event.get("source"),
                    "observed_target": event.get("observed_target"),
                    "preferred_target": event.get("preferred_target"),
                })
        state["knowledge_candidates"] = knowledge_candidates
        state["translation_continuity"] = list(knowledge_candidates)
        state["entity_registry"] = registry.to_list()
        if knowledge_warning:
            state["knowledge_feedback_failures"] = \
                state.get("knowledge_feedback_failures", 0) + 1
            state.setdefault("knowledge_events", []).append({
                "type": "extract_failed", "batch": bi, "offset": offset,
                "reason": knowledge_warning,
            })
        _checkpoint.append_event(job_dir(job_id), {
            "batch": bi, "offset": offset, "phase": "knowledge_feedback_done",
            "candidate_count": len(knowledge_candidates),
        })

        # 6) 审校通过的段落 -> 翻译记忆
        if enable_review:
            for j, p in enumerate(batch_pairs):
                segment_id = offset + j
                seg_findings = [
                    f for f in findings_all
                    if f.get("segment_id", f.get("segment_index")) == segment_id
                ]
                if not p.get("from_tm"):
                    if not review_succeeded:
                        p["review_status"] = "review_failed"
                        p["reviewed"] = False
                    elif seg_findings:
                        p["review_status"] = "reviewed_with_findings"
                        p["reviewed"] = False
                    else:
                        p["review_status"] = "reviewed_clean"
                if review_succeeded and not seg_findings and not p["from_tm"] \
                        and _tm_eligible(p["source"], p["target"]):
                    p["reviewed"] = True
                    p["accepted_target"] = p["target"]
                    p["target_provenance"] = "tm_approved" if p.get("from_tm") else "reviewed"
                    if use_tm:
                        tm[p["source"]] = {"target": p["target"], "reviewed": True}
                    stats["reviewed_segments"] += 1

        _commit_translation_batch(batch_pairs, offset)  # 正式状态先提交，TM 只随后晋升
        _checkpoint.append_event(job_dir(job_id), {
            "batch": bi, "offset": offset, "phase": "state_commit_done",
            "pairs_count": len(pairs),
        })
        tm_entries = _checkpoint.batch_entries([
            pair for pair in batch_pairs
            if not pair.get("from_tm")
            and pair.get("review_status") == "reviewed_clean"
        ]) if enable_review and use_tm and review_succeeded else []
        if tm_entries:
            _checkpoint.append_event(job_dir(job_id), {
                "batch": bi, "offset": offset, "phase": "tm_promotion_pending",
                "entries": tm_entries,
            })
            save_tm(tm)
            _checkpoint.append_event(job_dir(job_id), {
                "batch": bi, "offset": offset, "phase": "tm_promotion_done",
                "entries": tm_entries,
            })

    # 全局冲突检测（跨批次），与批内结果去重
    batch_conflict_keys = {(f.get("type"), f.get("entry_id"),
                            f.get("segment_index"), True)
                           for f in findings_all if f.get("conflict")}
    for cf in _detect_conflicts(pairs, glossary, sections=sections):
        cf["segment_index"] = cf["segment_id"]
        key = (cf.get("type"), cf.get("entry_id"), cf.get("segment_index"), True)
        if key not in batch_conflict_keys:
            findings_all.append(cf)
    state["entity_registry"] = registry.to_list()
    existing_entity_conflicts = {
        (item.get("type"), item.get("segment_index"), item.get("source"))
        for item in findings_all
    }
    for conflict in registry.consistency_findings():
        key = (conflict.get("type"), conflict.get("segment_index"),
               conflict.get("source"))
        if key not in existing_entity_conflicts:
            findings_all.append(conflict)
            existing_entity_conflicts.add(key)

    stats["blocking"] = sum(1 for f in findings_all if f["severity"] == "blocking")
    stats["actionable"] = sum(1 for f in findings_all if f["severity"] == "actionable")
    stats["informational"] = sum(1 for f in findings_all if f["severity"] == "informational")
    state["has_blocking"] = stats["blocking"] > 0
    return state


def findings_report_md(state):
    """把审查结果渲染成 Markdown 报告（下载/展示用）。"""
    stats = state.get("review_stats") or {}
    lines = [
        "# 翻译审查报告", "",
        "## 概览",
        f"- 已审校段落：{stats.get('reviewed_segments', 0)}",
        f"- 审校批次：{stats.get('batches_reviewed', 0)}",
        f"- 审校失败批次：{stats.get('review_failed', 0)}",
        f"- 翻译记忆复用：{state.get('tm_used_count', 0)} 段",
        f"- blocking：{stats.get('blocking', 0)}",
        f"- actionable：{stats.get('actionable', 0)}",
        f"- informational：{stats.get('informational', 0)}",
        "", "## 待处理问题",
    ]
    findings = state.get("findings") or []
    if not findings:
        lines.append("无。")
    else:
        for f in findings:
            line = f"- 第 {f.get('segment_index', -1) + 1} 段 [{f.get('severity')}] {f.get('reason')}"
            if f.get("suggested_target"):
                line += f"（建议译文：{f['suggested_target']}）"
            lines.append(line)
    return "\n".join(lines)


# ================= 文档/表格生成 =================
EN_FONT = "Times New Roman"
CN_FONT = "宋体"

# 自动标注三色：生僻词=红、专业名词（特殊译法）=黄、翻译难点句=青绿
ANNOTATION_COLORS = {"rare": "C00000", "domain": "BF8F00", "hard": "008080"}
ANNOTATION_LABELS = {"rare": "生僻词/难词", "domain": "专业名词（特殊译法）",
                     "hard": "翻译难点句（特别译法）"}


def _apply_doc_fonts(doc):
    """默认字体与可编辑学术文档的基础版式。"""
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    style_sizes = {
        "Normal": 12, "Heading 1": 15, "Heading 2": 13.5,
        "Heading 3": 12.5, "Heading 4": 12, "Title": 18,
    }
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3",
                       "Heading 4", "Title", "Intense Quote", "List Bullet",
                       "List Number"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = EN_FONT
        if style_name in style_sizes:
            style.font.size = Pt(style_sizes[style_name])
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), EN_FONT)
        rfonts.set(qn("w:hAnsi"), EN_FONT)
        rfonts.set(qn("w:eastAsia"), CN_FONT)
        if style_name == "Normal":
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.first_line_indent = Cm(0.74)
        elif style_name not in {"Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"}:
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.first_line_indent = None


def _apply_run_fonts(run):
    """单个 run 的字体（表格单元格里的 run 不受 Normal 样式继承影响时兜底）。"""
    from docx.oxml.ns import qn
    run.font.name = EN_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), EN_FONT)
    rfonts.set(qn("w:hAnsi"), EN_FONT)
    rfonts.set(qn("w:eastAsia"), CN_FONT)


def dict_to_excel(term_dict):
    df = pd.DataFrame(list(term_dict.items()), columns=["Source", "Target"])
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    output.seek(0)
    return output


def paragraphs_to_word(paragraphs):
    doc = Document()
    _apply_doc_fonts(doc)
    doc.add_heading('阶段一：清洗后原文提取', 0)
    for p in paragraphs:
        doc.add_paragraph(p)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def translations_to_word(pairs):
    doc = Document()
    _apply_doc_fonts(doc)
    doc.add_heading("译文", 0)
    for pair in pairs or []:
        doc.add_paragraph(str(pair.get("target") or ""))
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def translations_to_pdf(pairs):
    """Render a paginated, Unicode-aware translation PDF with PyMuPDF Story."""
    paragraphs = "".join(
        '<table class="entry"><tr><td dir="auto">'
        f'{html_escape(str(pair.get("target") or ""))}</td></tr></table>'
        for pair in pairs or [])
    html = f'<article><h1>译文</h1>{paragraphs}</article>'
    css = (
        "@page { size: A4; } "
        "body { font-family: sans-serif; font-size: 11pt; line-height: 1.65; "
        "color: #1f2937; } h1 { font-size: 20pt; margin: 0 0 20pt; "
        "color: #111827; } table.entry { width: 100%; border-collapse: collapse; "
        "margin: 0 0 10pt; break-inside: avoid; page-break-inside: avoid; } "
        "td { padding: 0; text-align: start; unicode-bidi: plaintext; }"
    )
    story = fitz.Story(html=html, user_css=css, em=11)
    page_rect = fitz.paper_rect("a4")
    content_rect = fitz.Rect(54, 54, page_rect.width - 54, page_rect.height - 54)

    def rectfn(_rect_num, _filled):
        return page_rect, content_rect, None

    pdf = story.write_with_links(rectfn)
    for number, page in enumerate(pdf, start=1):
        page.insert_text(
            (page_rect.width / 2 - 4, page_rect.height - 25), str(number),
            fontsize=9, fontname="helv", color=(0.45, 0.48, 0.52))
    data = pdf.tobytes(garbage=4, deflate=True)
    pdf.close()
    return data


def glossary_to_excel(entries, fallback=None):
    normalized = normalize_glossary(entries or [])
    if normalized:
        rows = [{
            "Source": entry.get("source") or "",
            "Target": entry.get("preferred") or entry.get("target") or "",
            "Status": entry.get("status") or "",
            "Domain": entry.get("domain") or "",
            "Note": entry.get("note") or "",
        } for entry in normalized]
        frame = pd.DataFrame(rows)
    else:
        frame = pd.DataFrame(list((fallback or {}).items()),
                             columns=["Source", "Target"])
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        frame.to_excel(writer, index=False)
    output.seek(0)
    return output


def _find_span(text, needle):
    """宽容定位子串：统一引号/破折号/省略号、折叠空白后查找，返回 (start, end) 或 None。"""
    if not needle or not text:
        return None
    if needle in text:
        pos = text.find(needle)
        return pos, pos + len(needle)
    mapping = []
    norm_chars = []
    for orig_idx, ch in enumerate(text):
        ch2 = ch.replace("’", "'").replace("‘", "'").replace("“", '"').replace("”", '"')
        ch2 = ch2.replace("–", "-").replace("—", "-").replace("…", "...")
        if ch2.isspace():
            if norm_chars and norm_chars[-1] != " ":
                norm_chars.append(" ")
                mapping.append(None)
            continue
        norm_chars.append(ch2)
        mapping.append(orig_idx)
    norm_text = "".join(norm_chars)
    needle2 = needle.replace("’", "'").replace("‘", "'").replace("“", '"').replace("”", '"')
    needle2 = needle2.replace("–", "-").replace("—", "-").replace("…", "...")
    needle2 = re.sub(r"\s+", " ", needle2).strip()
    pos = norm_text.find(needle2)
    if pos < 0:
        return None
    start = None
    for i in range(pos, len(mapping)):
        if mapping[i] is not None:
            start = mapping[i]
            break
    end = None
    for i in range(min(pos + len(needle2), len(mapping)) - 1, -1, -1):
        if mapping[i] is not None:
            end = mapping[i] + 1
            break
    if start is None or end is None:
        return None
    return start, end


def _colored_cell(cell, text, spans, colors=None):
    """把一个单元格按 spans（(start,end,type) 已排序不重叠）拆成带色 run。"""
    from docx.shared import RGBColor
    palette = colors or ANNOTATION_COLORS
    cursor = 0
    first = True
    for start, end, atype in spans:
        if start > cursor:
            run = cell.paragraphs[0].add_run(text[cursor:start])
            _apply_run_fonts(run)
        run = cell.paragraphs[0].add_run(text[start:end])
        _apply_run_fonts(run)
        run.font.color.rgb = RGBColor.from_string(
            str(palette.get(atype, ANNOTATION_COLORS[atype])).lstrip("#"))
        if atype in ("rare", "domain"):
            run.bold = True
        cursor = end
        first = False
    if cursor < len(text):
        run = cell.paragraphs[0].add_run(text[cursor:])
        _apply_run_fonts(run)
    if first and not text:
        cell.paragraphs[0].add_run("")


def pairs_to_word(pairs, annotations=None, colors=None):
    """双语对照表 -> Word 表格。

    annotations: {seg: [{"type": "rare|domain|hard", "src_span": [s,e]|None,
                         "tgt_span": [s,e]|None, "note": str}]}
    colors: {"rare"|"domain"|"hard": "RRGGBB"}，可自定义三类标注颜色。
    """
    palette = dict(ANNOTATION_COLORS)
    if colors:
        palette.update({k: str(v).lstrip("#") for k, v in colors.items()})
    doc = Document()
    _apply_doc_fonts(doc)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "原文"
    table.rows[0].cells[1].text = "译文"
    annot = _normalize_annotations(annotations)
    for i, pair in enumerate(pairs):
        row = table.add_row().cells
        seg_annot = annot.get(i) or []
        src_spans = _compose_spans(
            [(it["src_span"][0], it["src_span"][1], it["type"])
             for it in seg_annot if it.get("src_span")], len(pair['source']))
        tgt_spans = _compose_spans(
            [(it["tgt_span"][0], it["tgt_span"][1], it["type"])
             for it in seg_annot if it.get("tgt_span")], len(pair['target']))
        _colored_cell(row[0], pair['source'], src_spans, palette)
        _colored_cell(row[1], pair['target'], tgt_spans, palette)
    # 图例（放表格后，避免挤占首行）
    p_legend = doc.add_paragraph()
    legend_parts = [
        f"{ANNOTATION_LABELS[k]}（#{palette.get(k, ANNOTATION_COLORS[k])}）"
        for k in ("rare", "domain", "hard")]
    run = p_legend.add_run("图例：" + "；".join(legend_parts) + "。")
    _apply_run_fonts(run)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _add_formatted_runs(paragraph, text):
    parts = text.split('**')
    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        _apply_run_fonts(run)
        if i % 2 != 0:
            run.bold = True


def _markdown_table_row(line):
    value = str(line or "").strip()
    if value.startswith("|"):
        value = value[1:]
    if value.endswith("|"):
        value = value[:-1]
    return [cell.strip() for cell in value.split("|")]


def _is_markdown_table_separator(line):
    cells = _markdown_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def _add_markdown_table(doc, rows):
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=1, cols=columns)
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        cells = table.rows[0].cells if row_index == 0 else table.add_row().cells
        for col_index in range(columns):
            cell = cells[col_index]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.line_spacing = 1.35
            paragraph.paragraph_format.space_after = 0
            paragraph.paragraph_format.first_line_indent = None
            run = paragraph.add_run(values[col_index] if col_index < len(values) else "")
            _apply_run_fonts(run)
            if row_index == 0:
                run.bold = True
    return table


def markdown_to_word(md_text, theory):
    doc = Document()
    _apply_doc_fonts(doc)
    md_text = re.sub(r'```markdown|```', '', md_text).strip()
    md_text = re.sub(r'<!--.*?-->', '', md_text, flags=re.DOTALL)
    quote_labels = {
        "SYNTHETIC_SOURCE": "真实源文",
        "SIMULATED": "模拟初译",
        "OPTIMIZED": "优化译文",
    }
    quote_labels.update({"SOURCE": "原文", "INITIAL": "初译", "TARGET": "终译"})
    md_text = re.sub(
        r'(?m)^>\s*\[(SYNTHETIC_SOURCE|SIMULATED|OPTIMIZED|SOURCE|INITIAL|TARGET)\s+'
        r'(?:SC-\d{4,}|seg-[A-Za-z0-9_-]+)\]:\s*',
        lambda match: f"> {quote_labels[match.group(1)]}：", md_text)
    lines = md_text.splitlines()
    while lines and not lines[0].strip():
        lines.pop(0)
    if lines and re.match(r"^#\s+翻译实践报告", lines[0], re.IGNORECASE):
        lines.pop(0)
    deduped = []
    for line in lines:
        heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", line.strip())
        if heading:
            title = re.sub(r"\s+#+\s*$", "", heading.group(2)).strip()
            previous = next((value for value in reversed(deduped) if value.strip()), "")
            previous_heading = re.match(r"^#{1,6}\s+(.+?)\s*$", previous.strip())
            if previous_heading:
                previous_title = re.sub(r"\s+#+\s*$", "", previous_heading.group(1)).strip()
                if re.sub(r"^\d+(?:\.\d+)*[.、．]?\s*", "", title).casefold() == \
                        re.sub(r"^\d+(?:\.\d+)*[.、．]?\s*", "", previous_title).casefold():
                    continue
        deduped.append(line)
    title = doc.add_heading(f'翻译实践报告：基于{theory}', 0)
    title.alignment = 1
    index = 0
    while index < len(deduped):
        line = deduped[index].strip()
        if ("|" in line and index + 1 < len(deduped)
                and _is_markdown_table_separator(deduped[index + 1])):
            rows = [_markdown_table_row(line)]
            index += 2
            while index < len(deduped) and "|" in deduped[index] and deduped[index].strip():
                rows.append(_markdown_table_row(deduped[index]))
                index += 1
            _add_markdown_table(doc, rows)
            continue
        line = line.strip()
        if not line:
            index += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if heading:
            heading_level = max(1, min(4, len(heading.group(1)) - 1))
            doc.add_heading(heading.group(2).strip(), level=heading_level)
        elif line.startswith(('- ', '* ')):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_runs(p, line[2:])
        elif line.startswith('> '):
            p = doc.add_paragraph(style='Intense Quote')
            _add_formatted_runs(p, line[2:])
        else:
            p = doc.add_paragraph()
            _add_formatted_runs(p, line)
        index += 1
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


# ================= 任务持久化（真正的断点续传）=================
def _ensure_output_dir():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def is_onboarded():
    """首次使用引导标记：成功配置并测试通过 AI 引擎后置位。"""
    return (OUTPUT_DIR / ".onboarded").exists()


def mark_onboarded():
    _ensure_output_dir()
    (OUTPUT_DIR / ".onboarded").touch()


def provider_config_path():
    return OUTPUT_DIR / "provider_config.json"


def save_provider_config(provider, model, api_key, base_url=None, reviewer=None):
    """把 AI 引擎配置落盘（本地单机工具，0600 权限），重启后自动加载。"""
    _ensure_output_dir()
    cfg = {
        "provider": provider or "",
        "model": model or "",
        "api_key": api_key or "",
        "base_url": base_url or "",
    }
    if isinstance(reviewer, dict) and reviewer.get("provider") and reviewer.get("model"):
        cfg["reviewer"] = {
            "provider": reviewer.get("provider", ""),
            "model": reviewer.get("model", ""),
            "api_key": reviewer.get("api_key", ""),
            "base_url": reviewer.get("base_url", ""),
        }
    path = provider_config_path()
    path.write_text(json.dumps(cfg, ensure_ascii=False, indent=2),
                    encoding="utf-8")
    try:
        path.chmod(0o600)
    except OSError:
        pass
    return path


def load_provider_config():
    """读取已保存的 AI 引擎配置；文件缺失或损坏时返回 None。"""
    try:
        path = provider_config_path()
        if not path.exists():
            return None
        data = json.loads(path.read_text(encoding="utf-8"))
        if not isinstance(data, dict) or not data.get("provider"):
            return None
        result = {
            "provider": str(data.get("provider", "")),
            "model": str(data.get("model", "") or ""),
            "api_key": str(data.get("api_key", "") or ""),
            "base_url": str(data.get("base_url", "") or ""),
        }
        reviewer = data.get("reviewer")
        if isinstance(reviewer, dict) and reviewer.get("provider"):
            result["reviewer"] = {
                "provider": str(reviewer.get("provider", "")),
                "model": str(reviewer.get("model", "") or ""),
                "api_key": str(reviewer.get("api_key", "") or ""),
                "base_url": str(reviewer.get("base_url", "") or ""),
            }
        return result
    except Exception:  # noqa: BLE001 - 配置文件损坏时按未保存处理
        return None


def job_dir(job_id):
    return OUTPUT_DIR / job_id


def job_state_path(job_id):
    return job_dir(job_id) / "state.json"


RUNTIME_SCHEMA_VERSION = "2"
RUNTIME_STATE_FILE = "runtime_state.json"
RUNTIME_EVENTS_FILE = "runtime_events.jsonl"
RUNTIME_TECHNICAL_LOG = "runtime_technical.log"
RUNTIME_HEARTBEAT_SECONDS = 5
RUNTIME_LEASE_SECONDS = 60
RUNTIME_STALL_SECONDS = 45
RUNTIME_ACTIVE_STATUSES = {
    "resume_requested", "queued", "starting", "running",
    "waiting_external", "cancelling",
}

_RUNTIME_EVENT_META = {
    "resume_requested": ("user", "lifecycle"),
    "pipeline_resumed": ("user", "progress"),
    "llm_request_started": ("user", "external_wait"),
    "llm_response_received": ("user", "progress"),
    "job_completed": ("user", "lifecycle"),
    "job_failed": ("user", "error"),
    "job_cancelled": ("user", "lifecycle"),
    "cancel_requested": ("user", "lifecycle"),
    "interrupted": ("user", "error"),
    "stalled": ("user", "error"),
    "retry_requested": ("user", "lifecycle"),
    "job_queued": ("technical", "orchestration"),
    "worker_started": ("technical", "orchestration"),
    "worker_released": ("technical", "orchestration"),
    "job_checkpointed": ("technical", "checkpoint"),
    "checkpoint_saved": ("technical", "checkpoint"),
    "state_flushed": ("technical", "checkpoint"),
    "lease_acquired": ("technical", "orchestration"),
    "lease_renewed": ("technical", "orchestration"),
    "evidence": ("user", "progress"),
    "literature_evidence": ("user", "progress"),
    "research_model": ("user", "progress"),
    "literature_claims": ("user", "progress"),
    "argument_plan": ("user", "progress"),
    "selected_cases": ("user", "progress"),
    "outline": ("user", "progress"),
    "sections": ("user", "progress"),
    "validation": ("user", "progress"),
    "review": ("user", "progress"),
    "quality_repair": ("user", "progress"),
    "academic_quality": ("user", "progress"),
}


def runtime_state_path(job_id):
    return job_dir(job_id) / RUNTIME_STATE_FILE


def runtime_events_path(job_id):
    return job_dir(job_id) / RUNTIME_EVENTS_FILE


def runtime_technical_log_path(job_id):
    return job_dir(job_id) / RUNTIME_TECHNICAL_LOG


def _utc_now_iso():
    return datetime.now(timezone.utc).isoformat(timespec="seconds")


def _runtime_defaults():
    return {
        "schema_version": RUNTIME_SCHEMA_VERSION,
        "status": "idle",
        "pipeline": "document_pipeline",
        "stage": "",
        "stage_id": "",
        "stage_index": None,
        "stage_total": None,
        "operation": "",
        "operation_id": "",
        "operation_label": "",
        "section_id": None,
        "phase": "",
        "phase_label": "",
        "started_at": None,
        "operation_started_at": None,
        "last_heartbeat_at": None,
        "last_progress_at": None,
        "completed_units": 0,
        "total_units": 0,
        "overall_progress": None,
        "last_event": "",
        "events": [],
        "last_technical_event": "",
        "cancel_requested": False,
        "attempt": 0,
        "resume_request_id": None,
        "error": None,
        "worker": {
            "owner_pid": None,
            "worker_id": None,
            "lease_expires_at": None,
        },
    }


def load_runtime_state(job_id):
    """Read the durable worker status without changing the workflow state."""
    if not job_id:
        return _runtime_defaults()
    try:
        raw = json.loads(runtime_state_path(job_id).read_text(encoding="utf-8"))
    except (OSError, TypeError, ValueError):
        return _runtime_defaults()
    if not isinstance(raw, dict):
        return _runtime_defaults()
    result = _runtime_defaults()
    result.update(raw)
    result["schema_version"] = RUNTIME_SCHEMA_VERSION
    if result.get("status") == "cancel_requested":
        result["status"] = "cancelling"
    worker = result.get("worker")
    result["worker"] = {**_runtime_defaults()["worker"], **worker} \
        if isinstance(worker, dict) else dict(_runtime_defaults()["worker"])
    result["events"] = [item for item in result.get("events") or []
                         if isinstance(item, dict)
                         and item.get("visibility") == "user"][-24:]
    return result


def _runtime_event_code(message):
    value = re.sub(r"[^a-z0-9]+", "_", str(message or "").lower()).strip("_")
    return value[:80] or "runtime_event"


def _runtime_event_meta(event_name, visibility=None, category=None):
    default_visibility, default_category = _RUNTIME_EVENT_META.get(
        str(event_name or ""), ("technical", "debug"))
    return visibility or default_visibility, category or default_category


def _append_runtime_event(job_id, event, *, stage=None, operation=None, metadata=None,
                          visibility=None, category=None):
    event_name = event.get("event") if isinstance(event, dict) else _runtime_event_code(event)
    visibility, category = _runtime_event_meta(event_name, visibility, category)
    record = {
        "timestamp": _utc_now_iso(),
        "job_id": job_id,
        "pipeline": event.get("pipeline") if isinstance(event, dict) else None,
        "stage": stage or "",
        "operation": operation or "",
        "event": event_name,
        "message": event.get("message") if isinstance(event, dict) else str(event),
        "visibility": visibility,
        "category": category,
    }
    if record["pipeline"] is None:
        record["pipeline"] = "document_pipeline"
    if metadata:
        record["metadata"] = metadata
    path = runtime_events_path(job_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as stream:
        stream.write(json.dumps(record, ensure_ascii=False, sort_keys=True) + "\n")
        stream.flush()
        os.fsync(stream.fileno())
    return record


def read_runtime_events(job_id, limit=None, *, visibility=None, category=None):
    path = runtime_events_path(job_id)
    if not path.is_file():
        return []
    events = []
    try:
        for line in path.read_text(encoding="utf-8").splitlines():
            try:
                value = json.loads(line)
            except (TypeError, ValueError):
                continue
            if isinstance(value, dict):
                event_visibility, event_category = _runtime_event_meta(
                    value.get("event"), value.get("visibility"), value.get("category"))
                normalized = {**value, "visibility": event_visibility,
                              "category": event_category}
                if visibility and event_visibility != visibility:
                    continue
                if category and event_category != category:
                    continue
                events.append(normalized)
    except OSError:
        return []
    return events[-limit:] if limit else events


def update_runtime_state(job_id, *, event=None, progress=False, heartbeat=False,
                         event_name=None, event_metadata=None,
                         event_visibility=None, event_category=None, **changes):
    """Atomically merge a worker status update into runtime_state.json."""
    if not job_id:
        return _runtime_defaults()
    now = _utc_now_iso()
    with _RUNTIME_LOCK:
        current = load_runtime_state(job_id)
        current.update({key: value for key, value in changes.items()
                        if value is not None})
        if heartbeat:
            current["last_heartbeat_at"] = now
            worker = dict(current.get("worker") or {})
            if worker.get("worker_id"):
                worker["lease_expires_at"] = (
                    datetime.now(timezone.utc) + timedelta(seconds=RUNTIME_LEASE_SECONDS)
                ).isoformat(timespec="seconds")
                current["worker"] = worker
        if progress:
            current["last_progress_at"] = now
        if event:
            message = str(event).strip()
            if message:
                event_name = event_name or _runtime_event_code(message)
                visibility, category = _runtime_event_meta(
                    event_name, event_visibility, event_category)
                inline = {
                    "at": now, "timestamp": now, "event": event_name,
                    "message": message, "visibility": visibility,
                    "category": category,
                }
                if event_metadata:
                    inline["metadata"] = event_metadata
                duplicate = visibility == "user" and bool(current.get("events")) \
                    and current["events"][-1].get("event") == event_name \
                    and current["events"][-1].get("message") == message \
                    and current["events"][-1].get("metadata") == event_metadata
                if visibility == "user":
                    current["last_event"] = message
                    if not duplicate:
                        current["events"] = (current.get("events") or []) + [inline]
                        current["events"] = current["events"][-24:]
                else:
                    current["last_technical_event"] = message
                if not duplicate:
                    _append_runtime_event(
                        job_id,
                        {"event": event_name, "message": message,
                         "pipeline": current.get("pipeline")},
                        stage=current.get("stage_id") or current.get("stage"),
                        operation=current.get("operation_id") or current.get("operation"),
                        metadata=event_metadata, visibility=visibility, category=category)
        directory = job_dir(job_id)
        directory.mkdir(parents=True, exist_ok=True)
        tmp = directory / f"{RUNTIME_STATE_FILE}.tmp"
        tmp.write_text(json.dumps(current, ensure_ascii=False, indent=2),
                       encoding="utf-8")
        tmp.replace(runtime_state_path(job_id))
        return current


def _runtime_pid_alive(pid):
    if not isinstance(pid, int) or pid <= 0:
        return False
    try:
        os.kill(pid, 0)
    except OSError:
        return False
    return True


def _runtime_parse_time(value):
    if not value:
        return None
    try:
        parsed = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
        return parsed if parsed.tzinfo else parsed.replace(tzinfo=timezone.utc)
    except (TypeError, ValueError):
        return None


def _runtime_worker_registered(job_id, worker_id=None):
    with _RUNTIME_WORKERS_LOCK:
        worker = _RUNTIME_WORKERS.get(job_id)
        return bool(worker and worker.is_alive() and (
            worker_id is None or getattr(worker, "worker_id", None) == worker_id))


def _runtime_business_complete(state):
    return bool(state and state.get("p1_done") and state.get("p2_done") and (
        state.get("p3_done") or not state.get("report_enabled", True)) and (
        state.get("annotations_done") or not state.get("enable_annotate", True)))


def _runtime_mark_lost(job_id, status, message):
    return update_runtime_state(
        job_id, status=status, phase=status, phase_label=message,
        worker={"owner_pid": None, "worker_id": None, "lease_expires_at": None},
        event=message, event_name=status)


def get_job_runtime_status(job_id, state=None):
    """Return the durable runtime status and reconcile dead local workers."""
    state = state if isinstance(state, dict) else load_job_state(job_id)
    runtime = load_runtime_state(job_id)
    status = runtime.get("status") or "idle"
    if status in RUNTIME_ACTIVE_STATUSES:
        worker = runtime.get("worker") or {}
        pid = worker.get("owner_pid")
        worker_id = worker.get("worker_id")
        lease = _runtime_parse_time(worker.get("lease_expires_at"))
        heartbeat = _runtime_parse_time(runtime.get("last_heartbeat_at"))
        last_progress = _runtime_parse_time(runtime.get("last_progress_at"))
        now = datetime.now(timezone.utc)
        registered = _runtime_worker_registered(job_id, worker_id)
        transition_age = (now - last_progress).total_seconds() if last_progress else None
        if status == "resume_requested" and not pid \
                and transition_age is not None and transition_age <= RUNTIME_STALL_SECONDS:
            return runtime
        if not _runtime_pid_alive(pid):
            return _runtime_mark_lost(job_id, "interrupted", "上次运行已中断")
        if lease and lease < now or heartbeat and (
                now - heartbeat).total_seconds() > RUNTIME_STALL_SECONDS:
            if registered:
                return _runtime_mark_lost(job_id, "stalled", "暂无新的运行信号")
            return _runtime_mark_lost(job_id, "interrupted", "上次运行已中断")
        if not registered and pid == os.getpid() \
                and status in {"queued", "starting"} \
                and transition_age is not None and transition_age <= RUNTIME_STALL_SECONDS:
            return runtime
        if not registered and pid == os.getpid():
            return _runtime_mark_lost(job_id, "interrupted", "上次运行已中断")
        return runtime
    if status == "idle":
        inferred = "completed" if _runtime_business_complete(state) else "idle_incomplete"
        if inferred != status:
            return update_runtime_state(job_id, status=inferred,
                                        phase=inferred,
                                        phase_label="已完成" if inferred == "completed"
                                        else "尚未完成")
    return runtime


def build_job_runtime_view(job_id, state=None):
    """Canonical user-facing runtime view shared by every product surface."""
    state = state if isinstance(state, dict) else load_job_state(job_id)
    runtime = get_job_runtime_status(job_id, state)
    status = runtime.get("status") or "idle_incomplete"
    if status == "idle" and not _runtime_business_complete(state):
        status = "idle_incomplete"
    if status == "idle_incomplete" and state and \
            state.get("stage") == "TERMS_PREPARED" and state.get("quality_mode") \
            and state.get("glossary") is not None and not state.get("glossary_frozen") \
            and not state.get("quality_bypass"):
        status = "waiting_manual"
    labels = {
        "resume_requested": "正在恢复任务", "queued": "正在恢复任务",
        "starting": "正在恢复任务", "running": "正在运行",
        "waiting_external": "正在运行",
        "stalled": "暂无运行信号", "interrupted": "上次运行已中断",
        "failed": "当前步骤失败", "cancelling": "正在取消", "cancelled": "任务已取消",
        "completed": "已完成", "idle_incomplete": "未完成",
        "waiting_manual": "待术语确认",
    }
    actions = {
        "resume_requested": ["details"], "queued": ["cancel", "details"],
        "starting": ["cancel", "details"],
        "running": ["cancel", "details"],
        "waiting_external": ["cancel", "details"],
        "stalled": ["retry", "cancel", "details"],
        "interrupted": ["resume", "retry", "details"],
        "failed": ["retry", "resume", "details"],
        "cancelling": ["details"], "cancelled": ["resume", "details"],
        "idle_incomplete": ["resume"], "waiting_manual": ["view"],
        "completed": ["view"],
    }
    events = read_runtime_events(job_id, 5, visibility="user") \
        or runtime.get("events") or []
    completed, total = int(runtime.get("completed_units") or 0), \
        int(runtime.get("total_units") or 0)
    academic_present = bool(state and ((state.get("academic_state") or {}).get("artifacts")
                                       or (state.get("academic_state") or {}).get(
                                           "current_stage") not in {None, "", "not_started"}))
    if state and state.get("p2_done") and (
            state.get("report_enabled", True) or academic_present):
        completed, total = _academic_runtime_progress(job_id, state)
    if status == "completed" and total:
        completed = total
    progress = (completed, total) if total else None
    operation = runtime.get("operation_label") or ""
    is_report = bool(state and state.get("p2_done") and (
        state.get("report_enabled", True) or academic_present))
    surface_label = "实践报告" if is_report else "任务处理"
    if is_report and (not operation or operation in {"准备工作流", "准备任务", "继续处理"}):
        operation = _academic_resume_context(state, runtime)["operation_label"]
    if status in {"resume_requested", "queued", "starting"}:
        detail = "正在读取最近检查点…"
        if not operation or operation in {"准备工作流", "准备任务"}:
            operation = "正在从最近检查点继续学术写作" if is_report \
                else "正在从最近进度继续处理"
    elif status == "waiting_external":
        detail = "正在等待模型响应"
    elif status == "running":
        detail = runtime.get("phase_label") or "已恢复 · 正在执行"
    elif status == "interrupted":
        detail = "当前进度已安全保存，可以继续处理。"
    elif status == "stalled":
        detail = "暂时没有新的运行信号，可以重试当前步骤。"
    elif status == "idle_incomplete":
        detail = "当前进度已保存，可以继续处理。"
    elif status == "failed":
        error = runtime.get("error") or {}
        detail = error.get("message") if isinstance(error, dict) else str(error)
    else:
        detail = runtime.get("phase_label") or ""
    primary_action = {
        "idle_incomplete": "resume", "interrupted": "resume",
        "cancelled": "resume", "failed": "retry", "stalled": "retry",
    }.get(status)
    return {
        "status": status,
        "status_label": labels.get(status, status),
        "runtime_status": status,
        "headline_status": labels.get(status, status),
        "badge": status,
        "surface_label": surface_label,
        "headline": operation or surface_label,
        "detail": detail,
        "current_operation": operation,
        "operation_id": runtime.get("operation_id") or runtime.get("operation") or "",
        "progress": progress,
        "progress_completed": completed,
        "progress_total": total,
        "available_actions": actions.get(status, []),
        "primary_action": primary_action,
        "show_no_worker_warning": status in {
            "idle_incomplete", "interrupted", "cancelled"},
        "last_activity": events[-1].get("message") if events else "",
        "last_activity_at": events[-1].get("timestamp") if events else None,
        "user_events": events,
        "events": events,
        "runtime": runtime,
    }


def _write_runtime_technical_log(job_id, exc):
    path = runtime_technical_log_path(job_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as stream:
        stream.write(f"[{_utc_now_iso()}] {type(exc).__name__}: {exc}\n")
        stream.write(traceback.format_exc())
        stream.write("\n")
        stream.flush()


def _runtime_stage_info(label):
    """Extract the human-facing stage contract from existing status labels."""
    text = str(label or "").strip()
    match = re.search(r"学术写作\s+(\d+)\s*/\s*(\d+)", text)
    if match:
        index, total = int(match.group(1)), int(match.group(2))
        operation_label = re.sub(r"^.*?】\s*", "", text).strip(" .…") or text
        section_match = re.search(r"第\s*([\w.-]+)\s*节", text)
        stage_id = "academic_writing"
        for needle, candidate in (
                ("证据库", "evidence"), ("研究问题", "research_model"),
                ("论点", "argument_plan"), ("案例", "selected_cases"),
                ("提纲", "outline"), ("撰写正文", "sections"),
                ("确定性", "validation"), ("审稿", "review"),
                ("修订", "quality_repair"), ("质量", "academic_quality")):
            if needle in text:
                stage_id = candidate
                break
        return {
            "stage": "academic_writing",
            "stage_id": stage_id,
            "stage_index": index,
            "stage_total": total,
            "operation": stage_id,
            "operation_id": "section_rewrite" if section_match else stage_id,
            "operation_label": operation_label,
            "section_id": section_match.group(1) if section_match else None,
        }
    operation_label = re.sub(r"^.*?】\s*", "", text).strip(" .…") or text
    if "文档" in text or "排版" in text:
        stage = "document_processing"
    elif "术语" in text:
        stage = "terminology"
    elif "翻译" in text:
        stage = "translation"
    elif "标注" in text:
        stage = "annotation"
    elif "报告" in text or "学术" in text:
        stage = "academic_writing"
    else:
        stage = "pipeline"
    return {
        "stage": stage,
        "stage_id": stage,
        "operation": stage,
        "operation_id": stage,
        "operation_label": operation_label,
    }


def _runtime_overall_progress(stage_info, state):
    """Return a conservative progress estimate; only completion may be 1.0."""
    if not isinstance(state, dict):
        return 0.0
    if state.get("p1_done") and state.get("p2_done") \
            and (state.get("p3_done") or not state.get("report_enabled", True)) \
            and (state.get("annotations_done") or not state.get("enable_annotate", True)):
        return 1.0
    return None


def _academic_runtime_progress(job_id, state):
    """Count only durable academic artifact checkpoints; no estimated percent."""
    if not isinstance(state, dict):
        return 0, 0
    academic = state.get("academic_state") or {}
    artifacts = academic.get("artifacts") or {}
    names = ("evidence", "research_model", "argument_plan", "selected_cases",
             "outline", "sections", "validation", "review",
             "literature_support_review", "academic_quality", "report")
    completed = sum(1 for name in names if (artifacts.get(name) or {}).get("file"))
    return completed, len(names)


_ACADEMIC_RESUME_META = {
    "evidence": (1, "evidence", "继续构建学术证据"),
    "literature_evidence": (2, "literature_evidence", "继续整理文献证据"),
    "research_model": (3, "research_model", "继续建立研究模型"),
    "literature_claims": (4, "literature_claims", "继续整理文献主张"),
    "argument_plan": (5, "argument_plan", "继续规划论点"),
    "case_analysis": (6, "outline", "继续规划案例分析"),
    "outline": (6, "outline", "继续生成学术提纲"),
    "writing": (7, "sections", "继续撰写报告章节"),
    "validation": (8, "validation", "继续验证报告"),
    "review": (9, "review", "继续执行学术复核"),
    "repair": (10, "quality_repair", "继续修订受影响章节"),
    "academic_quality": (11, "academic_quality", "继续评估学术质量"),
}


def _academic_resume_context(state, runtime=None):
    academic = (state or {}).get("academic_state") or {}
    current_stage = str(academic.get("current_stage") or "")
    index, operation, label = _ACADEMIC_RESUME_META.get(
        current_stage, (1, "academic_writing", "继续学术写作"))
    runtime = runtime or {}
    section_id = runtime.get("section_id")
    if current_stage == "repair" and section_id:
        operation = "section_rewrite"
        label = f"继续重新生成第 {section_id} 节"
    return {
        "pipeline": "academic_writing",
        "stage": "academic_writing",
        "stage_id": operation,
        "stage_index": index,
        "stage_total": 11,
        "operation": operation,
        "operation_id": operation,
        "operation_label": label,
        "section_id": section_id or "",
    }


def _runtime_heartbeat_loop(job_id, stop_event):
    while not stop_event.wait(RUNTIME_HEARTBEAT_SECONDS):
        current = load_runtime_state(job_id)
        if current.get("status") not in RUNTIME_ACTIVE_STATUSES:
            return
        update_runtime_state(job_id, heartbeat=True)


def _runtime_status_callback(job_id, state, label):
    stage_info = _runtime_stage_info(label)
    current = load_runtime_state(job_id)
    operation = stage_info.get("operation") or current.get("operation")
    operation_started_at = current.get("operation_started_at")
    if operation != current.get("operation") or not operation_started_at:
        operation_started_at = _utc_now_iso()
    saved_state = load_job_state(job_id)
    state = saved_state or (state if isinstance(state, dict) else {})
    changes = dict(stage_info)
    completed_units, total_units = _academic_runtime_progress(job_id, state)
    changes.update({
        "operation": operation,
        "operation_id": stage_info.get("operation_id") or operation,
        "operation_label": stage_info.get("operation_label"),
        "section_id": stage_info.get("section_id") or "",
        "operation_started_at": operation_started_at,
        "overall_progress": _runtime_overall_progress(stage_info, state),
        "status": "running",
        "phase": "running",
        "phase_label": "正在处理",
        "completed_units": completed_units,
        "total_units": total_units,
    })
    update_runtime_state(job_id, progress=True, event=label,
                         event_name=stage_info.get("stage_id") or operation,
                         event_visibility="user", event_category="progress",
                         **changes)
    if _runtime_cancel_requested(job_id):
        raise RuntimeError("任务已请求取消")


def _runtime_caption_callback(job_id, text):
    message = str(text or "").strip()
    if not message:
        return
    update_runtime_state(job_id, progress=True, event=message, status="running",
                         event_visibility="user", event_category="progress")
    if _runtime_cancel_requested(job_id):
        raise RuntimeError("任务已请求取消")


def _run_job_worker(job_id, filename, file_bytes, pipeline_kwargs, base_url=None):
    stop_event = threading.Event()
    heartbeat = threading.Thread(target=_runtime_heartbeat_loop,
                                 args=(job_id, stop_event), daemon=True)
    _RUNTIME_CTX.job_id = job_id
    set_llm_base_url(base_url)
    heartbeat.start()
    try:
        kwargs = dict(pipeline_kwargs or {})
        kwargs.pop("on_status", None)
        kwargs.pop("on_caption", None)
        state = load_job_state(job_id) or new_job_state(filename)
        queued = load_runtime_state(job_id)
        update_runtime_state(
            job_id, status="starting", phase="starting",
            phase_label="正在读取最近检查点", heartbeat=True, progress=True,
            event="后台 worker 已接管任务", event_name="worker_started",
            event_visibility="technical", event_category="orchestration",
            pipeline="academic_writing" if kwargs.get("enable_report")
            else "document_pipeline")
        update_runtime_state(
            job_id, status="running", phase="running",
            phase_label="已恢复 · 正在执行" if queued.get("resume_request_id")
            else "正在执行", heartbeat=True)
        result = run_job_pipeline(
            job_id, filename, file_bytes, **kwargs,
            on_status=lambda label: _runtime_status_callback(job_id, state, label),
            on_caption=lambda text: _runtime_caption_callback(job_id, text),
        )
        result_state = load_job_state(job_id) or state
        if _runtime_cancel_requested(job_id):
            update_runtime_state(job_id, status="cancelled", phase="cancelled",
                                 phase_label="已取消", event="任务已取消",
                                 event_name="job_cancelled", progress=True,
                                 event_visibility="user", event_category="lifecycle",
                                 worker={"owner_pid": None, "worker_id": None,
                                         "lease_expires_at": None})
        elif _runtime_business_complete(result_state):
            update_runtime_state(job_id, status="completed", phase="completed",
                                 phase_label="已完成", event="任务已完成",
                                 event_name="job_completed", progress=True,
                                 event_visibility="user", event_category="lifecycle",
                                 heartbeat=True, overall_progress=1.0,
                                 worker={"owner_pid": None, "worker_id": None,
                                         "lease_expires_at": None})
        else:
            update_runtime_state(job_id, status="idle_incomplete", phase="idle_incomplete",
                                 phase_label="等待继续", event="阶段已保存，等待继续",
                                 event_name="job_checkpointed", progress=True,
                                 event_visibility="technical", event_category="checkpoint",
                                 heartbeat=True,
                                 worker={"owner_pid": None, "worker_id": None,
                                         "lease_expires_at": None})
        return result
    except Exception as exc:  # noqa: BLE001 - worker must publish failure to UI
        cancelled = _runtime_cancel_requested(job_id) or "请求取消" in str(exc)
        if not cancelled:
            _write_runtime_technical_log(job_id, exc)
        current = load_runtime_state(job_id)
        error = None if cancelled else {
            "type": type(exc).__name__,
            "message": str(exc)[:500] or "任务执行失败",
            "stage": current.get("stage_id") or current.get("stage"),
            "operation": current.get("operation_id") or current.get("operation"),
            "timestamp": _utc_now_iso(),
            "technical_log": RUNTIME_TECHNICAL_LOG,
        }
        update_runtime_state(
            job_id, status="cancelled" if cancelled else "failed",
            phase="cancelled" if cancelled else "failed",
            phase_label="已取消" if cancelled else "步骤失败",
            error=error,
            event="任务已取消" if cancelled else f"步骤失败：{str(exc)[:180]}",
            event_name="job_cancelled" if cancelled else "job_failed",
            event_visibility="user",
            event_category="lifecycle" if cancelled else "error",
            progress=True, heartbeat=True,
            worker={"owner_pid": None, "worker_id": None,
                    "lease_expires_at": None})
        return None
    finally:
        stop_event.set()
        set_llm_base_url(None)
        _RUNTIME_CTX.__dict__.clear()
        update_runtime_state(
            job_id, event="后台运行已释放", event_name="worker_released",
            event_visibility="technical", event_category="orchestration")
        with _RUNTIME_WORKERS_LOCK:
            _RUNTIME_WORKERS.pop(job_id, None)


def start_job_worker(job_id, filename, file_bytes, pipeline_kwargs, base_url=None,
                     resume_request_id=None):
    """Start one resumable pipeline worker; repeated UI reruns are idempotent."""
    with _RUNTIME_WORKERS_LOCK:
        worker = _RUNTIME_WORKERS.get(job_id)
        if worker and worker.is_alive():
            return False
        runtime = load_runtime_state(job_id)
        status = runtime.get("status") or "idle"
        matching_resume = status == "resume_requested" and resume_request_id \
            and runtime.get("resume_request_id") == resume_request_id
        if status in RUNTIME_ACTIVE_STATUSES and not matching_resume:
            return False
        if not matching_resume:
            runtime = get_job_runtime_status(job_id)
            if runtime.get("status") in RUNTIME_ACTIVE_STATUSES:
                return False
        state = load_job_state(job_id) or new_job_state(filename)
        worker_id = uuid.uuid4().hex
        attempt = int(runtime.get("attempt") or 0) + 1
        started_at = _utc_now_iso()
        lease_expires_at = (datetime.now(timezone.utc) +
                            timedelta(seconds=RUNTIME_LEASE_SECONDS)).isoformat(
                                timespec="seconds")
        completed_units, total_units = (0, 0)
        if state.get("p2_done") and pipeline_kwargs.get("enable_report"):
            completed_units, total_units = _academic_runtime_progress(job_id, state)
        context = _academic_resume_context(state, runtime) \
            if matching_resume and pipeline_kwargs.get("enable_report") else {
                "pipeline": "document_pipeline", "stage": "pipeline",
                "stage_id": "pipeline", "operation": "pipeline",
                "operation_id": "pipeline", "operation_label": "准备任务",
                "section_id": "", "stage_index": None, "stage_total": None,
            }
        update_runtime_state(
            job_id, status="queued", phase="starting", phase_label="准备中",
            **context, started_at=runtime.get("started_at") if matching_resume
            else started_at,
            operation_started_at=started_at, last_progress_at=started_at,
            last_heartbeat_at=started_at, cancel_requested=False, error=None,
            attempt=attempt, resume_request_id=resume_request_id or "",
            completed_units=completed_units, total_units=total_units,
            overall_progress=None,
            worker={"owner_pid": os.getpid(), "worker_id": worker_id,
                    "lease_expires_at": lease_expires_at},
            event="已排入后台 worker", event_name="job_queued",
            event_visibility="technical", event_category="orchestration")
        worker = threading.Thread(
            target=_run_job_worker,
            args=(job_id, filename, file_bytes, dict(pipeline_kwargs or {}), base_url),
            name=f"transpraxis-{job_id}", daemon=True)
        worker.worker_id = worker_id
        _RUNTIME_WORKERS[job_id] = worker
        worker.start()
        return True


def resume_job(job_id, filename, pipeline_kwargs, base_url=None,
               resume_request_id=None):
    """Idempotently request resume and publish the transition before worker start."""
    with _RUNTIME_WORKERS_LOCK:
        runtime = get_job_runtime_status(job_id)
        if runtime.get("status") in RUNTIME_ACTIVE_STATUSES:
            return False
        state = load_job_state(job_id) or new_job_state(filename)
        request_id = resume_request_id or uuid.uuid4().hex
        if runtime.get("resume_request_id") == request_id:
            return False
        pipeline_kwargs = dict(pipeline_kwargs or {})
        academic = state.get("academic_state") or {}
        if state.get("p2_done") and (academic.get("artifacts") or academic.get(
                "current_stage") not in {None, "", "not_started"}):
            pipeline_kwargs["enable_report"] = True
        completed_units, total_units = (0, 0)
        context = {
            "pipeline": "document_pipeline", "stage": "pipeline",
            "stage_id": "pipeline", "operation": "pipeline",
            "operation_id": "pipeline", "operation_label": "继续处理",
            "section_id": "", "stage_index": None, "stage_total": None,
        }
        if state.get("p2_done") and pipeline_kwargs.get("enable_report"):
            completed_units, total_units = _academic_runtime_progress(job_id, state)
            context = _academic_resume_context(state, runtime)
        now = _utc_now_iso()
        update_runtime_state(
            job_id, status="resume_requested", phase="resume_requested",
            phase_label="正在恢复任务", resume_request_id=request_id,
            started_at=now, operation_started_at=now,
            last_progress_at=now, completed_units=completed_units,
            total_units=total_units, overall_progress=None,
            worker={"owner_pid": None, "worker_id": None,
                    "lease_expires_at": None},
            event="已从断点恢复任务", event_name="resume_requested",
            event_visibility="user", event_category="lifecycle",
            event_metadata={"resume_request_id": request_id}, **context)
        return start_job_worker(
            job_id, filename, None, pipeline_kwargs, base_url=base_url,
            resume_request_id=request_id)


def is_job_worker_alive(job_id):
    with _RUNTIME_WORKERS_LOCK:
        worker = _RUNTIME_WORKERS.get(job_id)
        return bool(worker and worker.is_alive())


def request_job_cancel(job_id):
    """Request cancellation between provider calls; an active HTTP call finishes first."""
    if not is_job_worker_alive(job_id):
        return False
    update_runtime_state(job_id, status="cancelling", cancel_requested=True,
                         phase="cancelling", phase_label="正在取消",
                         event="已请求取消任务", event_name="cancel_requested",
                         progress=True)
    return True


def retry_job_step(job_id):
    """Invalidate only the failed academic operation and its downstream work."""
    runtime = get_job_runtime_status(job_id)
    if runtime.get("status") not in {"failed", "stalled", "interrupted", "cancelled"}:
        return False
    operation = runtime.get("operation_id") or runtime.get("operation") or ""
    scopes = {
        "validation": "validation", "review": "review",
        "literature_support_review": "literature_review",
        "academic_quality": "quality", "quality_repair": "quality",
        "section_rewrite": "section", "sections": "writer", "repair": "writer",
        "outline": "planning",
        "argument_plan": "planning", "selected_cases": "planning",
        "research_model": "planning", "evidence": "all",
    }
    state = load_job_state(job_id)
    if state and (runtime.get("stage") == "academic_writing" or
                  operation in scopes):
        error_message = str((runtime.get("error") or {}).get("message") or "")
        scope = "case_analysis" if operation == "section_rewrite" and \
            "missing case target subsection" in error_message else \
            scopes.get(operation, "writer")
        invalidate_academic_report(job_id, scope, runtime.get("section_id"))
    update_runtime_state(
        job_id, status="idle_incomplete", phase="retry_ready", phase_label="等待重试",
        cancel_requested=False, error=None, event="已准备重试当前步骤",
        event_name="retry_requested", worker={"owner_pid": None, "worker_id": None,
                                               "lease_expires_at": None})
    return True


def _invalidate_final_delivery_state(state):
    """Invalidate only the mutable working approval; snapshot history stays on disk."""
    state["delivery_status"] = "draft"
    state["delivery_approved_by_human"] = False
    state["delivery_approval"] = None
    if state.get("stage") == "FINAL":
        state["stage"] = _state_migration.derive_stage(state)
    return state


def _finalization_artifacts(job_id):
    """Load only the small set of artifacts needed for impact explanations."""
    names = ("selected_cases", "outline", "argument_plan", "sections")
    return {name: load_academic_artifact(job_id, name) for name in names}


def _reset_final_qa(state, reason=""):
    qa = _finalization.normalize_final_qa(state.get("final_qa"))
    qa.update({
        "structural_qa": "NOT_RUN",
        "libreoffice_render": "NOT_RUN",
        "author_visual_review": "NOT_CONFIRMED",
        "word_final_review": "NOT_CONFIRMED",
        "rendered_at": None,
        "page_count": None,
        "updated_at": _finalization.now_iso(),
    })
    if reason:
        notes = dict(qa.get("notes") or {})
        notes["stale_reason"] = reason
        qa["notes"] = notes
    state["final_qa"] = qa
    return qa


def _mark_translation_truth_changed(
    job_id, state, indexes, reason, *, actor="user", action="translation_changed",
):
    """Record one canonical CURRENT_TRANSLATION mutation and its impact slice."""
    indexes = sorted({int(index) for index in indexes
                      if isinstance(index, int) or str(index).lstrip("-").isdigit()})
    truth = dict(state.get("translation_truth") or {})
    truth["authority"] = _finalization.CURRENT_TRANSLATION
    truth["version"] = int(truth.get("version") or 0) + 1
    truth["last_changed_at"] = _finalization.now_iso()
    truth["last_change"] = {
        "action": action,
        "actor": actor,
        "reason": reason,
        "segment_indexes": indexes,
        "segment_ids": [
            _finalization.segment_id(
                job_id, index, (state.get("pairs") or [])[index]
                if 0 <= index < len(state.get("pairs") or []) else {})
            for index in indexes
        ],
    }
    state["translation_truth"] = truth
    changed_segment_ids = list(truth["last_change"]["segment_ids"])
    from transpraxis import academic_writer
    propagated = academic_writer.propagate_artifact_staleness(
        state, input_segment_ids=changed_segment_ids)
    # Make the read-only artifact inputs available to the pure impact helper,
    # then remove them before state is persisted.
    enriched = dict(state)
    enriched["_finalization_artifacts"] = _finalization_artifacts(job_id)
    impact = _finalization.build_dependency_impact(
        enriched, job_id, indexes, reason)
    state["dependency_impact"] = impact
    _finalization.mark_case_reviews_stale(
        state, impact.get("affected_case_ids") or [], reason)
    stale_names = [item.get("id") for item in impact.get("affected") or []
                   if item.get("kind") == "artifact"]
    academic = state.setdefault("academic_state", {})
    # Legacy records have no direct edges and retain their historical
    # invalidation behavior. Canonical records are changed only by the graph
    # propagation above, preserving their own direct inputs.
    if not propagated and stale_names:
        academic_writer._invalidate_names(
            state, [name for name in stale_names
                    if name not in {"delivery_assets", "libreoffice_render"}], reason)
    if propagated:
        if any(name in propagated for name in {"sections", "report", "validation", "review"}):
            state["p3_done"] = False
            academic["status"] = "stale"
        if any(name in propagated for name in {"report", "sections"}):
            state["p3_md"] = ""
            state["p3_sections"] = []
    _reset_final_qa(state, reason)
    _invalidate_final_delivery_state(state)
    state.setdefault("human_actions", []).append({
        "finding_id": f"segment-mutation:{','.join(map(str, indexes)) or 'none'}",
        "action": action,
        "note": reason,
        "timestamp": _finalization.now_iso(),
        "actor": actor,
    })
    return state


def translation_truth_view(job_id, state=None):
    """Return the user-facing authority/version summary for current targets."""
    state = state if state is not None else load_job_state(job_id) or {}
    truth = dict(state.get("translation_truth") or {})
    pairs = state.get("pairs") or []
    return {
        "authority": truth.get("authority") or _finalization.CURRENT_TRANSLATION,
        "version": int(truth.get("version") or 0),
        "segment_count": len(pairs),
        "last_changed_at": truth.get("last_changed_at"),
        "last_change": dict(truth.get("last_change") or {}),
        "label": "CURRENT_TRANSLATION · 当前工作译文",
    }


def dependency_impact_view(job_id, state=None):
    state = state if state is not None else load_job_state(job_id) or {}
    return _finalization.normalize_dependency_impact(state.get("dependency_impact"))


def compliance_profile_view(job_id, state=None):
    """Evaluate source-backed compliance and explicit project constraints."""
    state = state if state is not None else load_job_state(job_id) or {}
    from transpraxis import compliance
    artifacts = {
        name: load_academic_artifact(job_id, name)
        for name in ("evidence", "report", "validation", "outline",
                     "selected_cases", "literature_sources",
                     "final_docx_validation")
    }
    profile_id = str(state.get("compliance_profile_id") or
                     compliance.DEFAULT_PROFILE_ID)
    profile = compliance.compliance_profile(profile_id)
    result = compliance.evaluate_compliance(
        state, artifacts, profile, state.get("p3_md") or "")
    language = compliance.evaluate_language_constraints(
        state, state.get("p3_md") or "")
    result["language_constraints"] = language
    language_constraints = language.get("constraints") or []
    result["counts"]["pass"] += sum(
        item.get("status") == "pass" for item in language_constraints)
    result["counts"]["fail"] += len(language.get("failures") or [])
    result["counts"]["manual_review"] += sum(
        item.get("status") == "manual_review" for item in language_constraints)
    project = result.setdefault("project_constraints", {})
    project.setdefault("failures", []).extend(
        f"language:{item.get('kind')}:{item.get('value')}"
        for item in language.get("failures") or [])
    if language.get("failures"):
        project["status"] = "fail"
        result["status"] = "fail"
    if language.get("status") == "manual_review" and \
            result.get("status") != "fail":
        result["status"] = "manual_review"
    return result


def current_translation_hash(state=None):
    from transpraxis import academic_evidence
    state = state or {}
    return academic_evidence.stable_hash({
        "pairs": [
            {key: pair.get(key) for key in ("source", "initial_target", "target")}
            for pair in state.get("pairs") or []
        ],
    })


def generate_report_qa(job_id, state=None, *, save_file=False):
    """Build the concise QA report and bind it to current artifact hashes."""
    from transpraxis import academic_evidence, academic_writer
    state = state or load_job_state(job_id) or {}
    report_record = load_academic_artifact(job_id, "report") or {}
    render_record = load_academic_artifact(job_id, "libreoffice_render") or {}
    final_docx = load_academic_artifact(job_id, "final_docx_validation") or {}
    compliance = compliance_profile_view(job_id, state)
    case_review = _finalization.case_review_gate(
        state, load_academic_artifact(job_id, "selected_cases"))
    final_qa = _finalization.normalize_final_qa(state.get("final_qa"))
    final_qa["structural_qa"] = "PASS" if final_docx.get("status") in {
        "pass", "pass_with_warnings"} else "FAIL" if final_docx.get(
        "status") == "fail" else "NOT_RUN"
    translation_hash = current_translation_hash(state)
    report_hash = str(report_record.get("content_hash") or "")
    docx_hash = str(render_record.get("source_docx_hash") or
                    final_docx.get("source_docx_hash") or "")
    markdown = _rendered_qa.render_qa_markdown(
        translation_hash=translation_hash, report_hash=report_hash,
        docx_hash=docx_hash, render_record=render_record,
        pdf_qa=render_record.get("analysis") or {}, compliance=compliance,
        case_review=case_review, final_qa=final_qa,
        placeholders=next((x.get("actual") or [] for x in compliance.get(
            "rules") or [] if x.get("rule_id") == "author_placeholders"), []))
    value = {
        "schema_version": _rendered_qa.VERSION,
        "generated_at": _finalization.now_iso(),
        "translation_truth_hash": translation_hash,
        "report_content_hash": report_hash,
        "source_docx_hash": docx_hash,
        "rendered_pdf_hash": render_record.get("rendered_pdf_hash"),
        "final_qa": final_qa,
        "compliance_status": compliance.get("status"),
        "case_review_status": case_review.get("status"),
        "content_hash": academic_evidence.stable_hash({
            "translation": translation_hash, "report": report_hash,
            "docx": docx_hash, "pdf": render_record.get("rendered_pdf_hash"),
            "compliance": compliance.get("status"),
            "case_review": case_review.get("status"), "final_qa": final_qa,
        }),
        "markdown": markdown,
    }
    if save_file:
        academic_writer._save_artifact(
            state, job_dir(job_id), "report_qa", value, str(value["content_hash"]),
            _rendered_qa.VERSION,
            input_artifact_ids=["report", "final_docx_validation",
                                "libreoffice_render"],
            input_segment_ids=[])
        (job_dir(job_id) / "report-qa.md").write_text(markdown, encoding="utf-8")
        save_job_state(job_id, state)
    return value


def save_compliance_record(job_id, state=None):
    """Persist the current compliance and language-constraint artifacts."""
    from transpraxis import academic_evidence, academic_writer, compliance
    state = state or load_job_state(job_id) or {}
    result = compliance_profile_view(job_id, state)
    academic = state.get("academic_state") or {}
    records = academic.get("artifacts") or {}
    compliance_dependency = academic_evidence.stable_hash({
        "profile_id": result.get("profile_id"),
        "report": (records.get("report") or {}).get("content_hash"),
        "evidence": (records.get("evidence") or {}).get("content_hash"),
        "selected_cases": (records.get("selected_cases") or {}).get("content_hash"),
        "literature_sources": (records.get("literature_sources") or {}).get(
            "content_hash"),
        "outline": (records.get("outline") or {}).get("content_hash"),
    })
    academic_writer._save_artifact(
        state, job_dir(job_id), "compliance", result,
        compliance_dependency, compliance.compliance_profile(
            str(state.get("compliance_profile_id") or compliance.DEFAULT_PROFILE_ID)
        ).get("schema_version"),
        input_artifact_ids=["report", "evidence", "selected_cases",
                            "literature_sources"])
    language = result.get("language_constraints") or {}
    settings = state.get("research_settings") or {}
    language_dependency = academic_evidence.stable_hash({
        "report": (records.get("report") or {}).get("content_hash"),
        "language_constraints": {
            key: settings.get(key) for key in (
                "forbidden_report_phrases", "allowed_theory_labels",
                "required_terminology", "protected_names",
                "protected_work_titles")
        },
    })
    academic_writer._save_artifact(
        state, job_dir(job_id), "language_constraints", language,
        language_dependency, compliance.VERSION,
        input_artifact_ids=["report"])
    state["compliance_record"] = result
    state["language_constraint_record"] = language
    save_job_state(job_id, state)
    return result


def _case_artifact_case(job_id, case_id):
    selected = load_academic_artifact(job_id, "selected_cases") or {}
    case = next((item for item in selected.get("cases") or []
                 if str(item.get("case_id")) == str(case_id)), None)
    return selected, case


def _mark_case_downstream_stale(job_id, state, case_id, reason, *, actor="user",
                                action="case_changed", root_stale=True):
    from transpraxis import academic_writer
    root_id = f"case:{case_id}"
    before_root = dict((state.get("academic_state") or {}).get(
        "artifacts", {}).get(root_id) or {})
    propagated = academic_writer.propagate_artifact_staleness(
        state, input_artifact_ids=[f"case:{case_id}"])
    if not root_stale and before_root:
        academic = state.setdefault("academic_state", {})
        academic.setdefault("artifacts", {})[root_id] = before_root
        academic_writer._write_status_mirror(academic, root_id, before_root)
    enriched = dict(state)
    enriched["_finalization_artifacts"] = _finalization_artifacts(job_id)
    impact = _finalization.build_dependency_impact(
        enriched, job_id, [], reason, changed_case_ids=[str(case_id)])
    state["dependency_impact"] = impact
    stale_names = [item.get("id") for item in impact.get("affected") or []
                   if item.get("kind") == "artifact"]
    academic = state.setdefault("academic_state", {})
    # The old fallback is retained only for pre-Stage-2 records. Canonical
    # records are invalidated by their case direct edge and reverse traversal.
    if not propagated and stale_names:
        academic_writer._invalidate_names(state, stale_names, reason)
    if propagated and any(name in propagated for name in {"sections", "report"}):
        state["p3_done"] = False
        academic["status"] = "stale"
    _reset_final_qa(state, reason)
    _invalidate_final_delivery_state(state)
    state.setdefault("human_actions", []).append({
        "finding_id": f"case:{case_id}", "action": action, "note": reason,
        "timestamp": _finalization.now_iso(), "actor": actor,
    })
    return state


def review_academic_case(job_id, case_id, status, note="", actor="user"):
    """Record one author decision without changing provenance or content."""
    state = load_job_state(job_id)
    if state is None:
        return None, False, "任务不存在"
    selected, case = _case_artifact_case(job_id, case_id)
    if case is None:
        return state, False, "找不到案例"
    from transpraxis import case_provenance
    normalized = case_provenance.with_provenance(case)
    review_status = str(status or "").strip().lower()
    if review_status not in case_provenance.REVIEW_STATUSES:
        return state, False, "案例审校状态无效"
    state.setdefault("case_reviews", {})[str(case_id)] = {
        "review_status": review_status,
        "case_origin": normalized.get("case_origin"),
        "text_role": dict(normalized.get("text_role") or {}),
        "review_reason": str(note or "")[:700],
        "note": str(note or "")[:700],
        "reviewed_at": _finalization.now_iso(),
        "updated_at": _finalization.now_iso(),
        "actor": actor,
        "translation_truth_version": int(
            (state.get("translation_truth") or {}).get("version") or 0),
        "content_stale": False,
        "stale_reason": None,
        "stale_at": None,
    }
    state.setdefault("human_actions", []).append({
        "finding_id": f"case:{case_id}",
        "action": "case_approved" if review_status == "approved"
        else "case_excluded",
        "note": "批准案例纳入学术分析" if review_status == "approved"
        else f"排除案例：{str(note or '人工排除')[:180]}",
        "timestamp": _finalization.now_iso(),
        "actor": actor,
    })
    _mark_case_downstream_stale(
        job_id, state, str(case_id),
        "作者更新案例审核状态，需重组案例相关写作下游",
        actor=actor, action="case_review_changed", root_stale=False)
    save_job_state(job_id, state)
    return state, True, "已保存案例审校状态"


def replace_rejected_case(job_id, case_id, *, actor="user"):
    """Replace one author-rejected case from the already validated pool."""
    from transpraxis import academic_writer, academic_quality, case_provenance
    state = load_job_state(job_id)
    if state is None:
        return None, False, ["任务不存在"]
    selected, old_case = _case_artifact_case(job_id, case_id)
    if selected is None or old_case is None:
        return state, False, ["找不到案例"]
    review = (state.get("case_reviews") or {}).get(str(case_id)) or {}
    if str(review.get("review_status")) != "rejected":
        return state, False, ["只能替换已被作者排除的案例"]
    selected_ids = {str(x.get("case_id")) for x in selected.get("cases") or []}
    reviews = state.get("case_reviews") or {}
    overrides = state.get("case_review_overrides") or {}
    candidates = []
    if case_provenance.is_synthetic(old_case):
        synthetic = load_academic_artifact(job_id, "synthetic_validation") or {}
        for candidate in synthetic.get("items") or []:
            cid = str(candidate.get("case_id") or "")
            validation = candidate.get("validation") or {}
            review_record = reviews.get(cid) if isinstance(reviews, dict) else None
            override = overrides.get(cid) if isinstance(overrides, dict) else None
            if not cid or cid in selected_ids or candidate is old_case:
                continue
            if review_record and review_record.get("review_status") == "rejected":
                continue
            if override and override.get("baseline_status") == "rejected":
                continue
            if not validation.get("academic_case_eligible"):
                continue
            difficulty = candidate.get("difficulty") or {}
            evidence = candidate.get("synthetic_evidence") or {}
            old_group = str(old_case.get("difficulty_group") or "")
            target_match = bool(old_group and str(difficulty.get("group") or "") == old_group)
            candidates.append((not target_match,
                               difficulty.get("academic_value") != "high",
                               difficulty.get("confidence") != "high",
                               evidence.get("material_difference") != "pass",
                               candidate.get("segment_index", 0), candidate))
    else:
        evidence_artifact = load_academic_artifact(job_id, "evidence") or {}
        argument_plan = load_academic_artifact(job_id, "argument_plan") or {}
        candidate = academic_quality.select_replacement_case(
            str(case_id), list(old_case.get("supports_claims") or []),
            selected, argument_plan, evidence_artifact)
        if candidate:
            candidates.append((False, False, False, False,
                               candidate.get("segment_index", 0), candidate))
    candidates.sort(key=lambda item: item[:-1])
    if not candidates:
        return state, False, ["没有可用的已验证替换候选；请保持该案例排除并调整案例数量"]
    candidate = case_provenance.with_provenance(dict(candidates[0][-1]))
    candidate.update({
        "supports_claims": sorted(set(old_case.get("supports_claims") or [])),
        "research_questions": sorted(set(old_case.get("research_questions") or [])),
        "argument_role": old_case.get("argument_role", "supporting"),
        "difficulty_group": old_case.get("difficulty_group") or
        candidate.get("difficulty_group"),
        "difficulty_subsection": old_case.get("difficulty_subsection"),
        "strategy_subsection": old_case.get("strategy_subsection"),
        "target_subsection": old_case.get("target_subsection"),
        "review_status": "unreviewed",
        "replacement_of": str(case_id),
        "selection_rationale": f"replacement of rejected {case_id}: existing validated pool",
    })
    selected["cases"] = [
        candidate if str(x.get("case_id")) == str(case_id)
        else x for x in selected.get("cases") or []
    ]
    record = academic_writer.artifact_record(state, "selected_cases")
    academic_writer._save_artifact(
        state, job_dir(job_id), "selected_cases", selected,
        str(record.get("dependency_hash") or ""),
        str(record.get("version") or "case-review-v1"))
    _mark_case_downstream_stale(
        job_id, state, str(candidate.get("case_id")),
        "作者排除案例后从已验证候选池替换，需重组其写作下游",
        actor=actor, action="case_replaced", root_stale=False)
    state.setdefault("case_reviews", {})[str(candidate.get("case_id"))] = {
        "review_status": "unreviewed",
        "case_origin": candidate.get("case_origin"),
        "text_role": dict(candidate.get("text_role") or {}),
        "review_reason": "替换被排除案例；需作者重新终审",
        "reviewed_at": _finalization.now_iso(),
        "updated_at": _finalization.now_iso(),
        "actor": actor,
        "translation_truth_version": int(
            (state.get("translation_truth") or {}).get("version") or 0),
        "content_stale": False,
    }
    state.setdefault("human_actions", []).append({
        "finding_id": f"case:{candidate.get('case_id')}",
        "action": "case_replaced",
        "note": f"从已验证候选池替换 {case_id}",
        "timestamp": _finalization.now_iso(),
        "actor": actor,
    })
    save_job_state(job_id, state)
    return state, True, [str(candidate.get("case_id"))]


def update_synthetic_baseline(job_id, case_id, text, *, status="modified", note="", actor="user"):
    """Modify or reject a synthetic baseline; never mutate translation truth."""
    state = load_job_state(job_id)
    if state is None:
        return None, False, "任务不存在"
    _selected, case = _case_artifact_case(job_id, case_id)
    if case is None:
        return state, False, "找不到案例"
    from transpraxis import case_provenance
    if not case_provenance.is_synthetic(case):
        return state, False, "只有合成对照案例可以修改或拒绝模拟初译"
    if status not in {"modified", "rejected", "approved"}:
        return state, False, "模拟初译状态无效"
    record = dict((state.setdefault("case_review_overrides", {}).get(str(case_id)) or {}))
    if status == "modified":
        text = str(text or "").strip()
        if not text:
            return state, False, "模拟初译不能为空"
        record["synthetic_baseline_text"] = text
    record.update({
        "baseline_status": status,
        "note": str(note or "")[:700],
        "updated_at": _finalization.now_iso(),
        "actor": actor,
    })
    state["case_review_overrides"][str(case_id)] = record
    reason = ("修改模拟初译，需重跑该案例下游"
              if status == "modified" else "拒绝模拟初译，需替换或重新确认该案例")
    _finalization.mark_case_reviews_stale(state, [str(case_id)], reason)
    _mark_case_downstream_stale(
        job_id, state, str(case_id), reason, actor=actor,
        action="synthetic_baseline_modified" if status == "modified"
        else "synthetic_baseline_rejected")
    save_job_state(job_id, state)
    return state, True, "已保存模拟初译决定"


def record_final_qa(job_id, field, status, note="", actor="user"):
    """Persist one of the four independent final-QA facts."""
    if field not in _finalization.QA_FIELDS:
        raise ValueError(f"未知最终 QA 项：{field}")
    state = load_job_state(job_id)
    if state is None:
        return None
    allowed = {"PASS", "FAIL", "NOT_RUN"} if field in {
        "structural_qa", "libreoffice_render"} else {"CONFIRMED", "NOT_CONFIRMED"}
    if status not in allowed:
        raise ValueError(f"{field} 状态无效：{status}")
    qa = _finalization.normalize_final_qa(state.get("final_qa"))
    qa[field] = status
    qa["translation_truth_version"] = int(
        (state.get("translation_truth") or {}).get("version") or 0)
    notes = dict(qa.get("notes") or {})
    if note:
        notes[field] = str(note)[:700]
    qa["notes"] = notes
    qa["updated_at"] = _finalization.now_iso()
    state["final_qa"] = qa
    save_job_state(job_id, state)
    return state


def run_libreoffice_render_qa(job_id, state=None):
    """Render through LibreOffice, then run separate deterministic PDF QA."""
    state = state or load_job_state(job_id)
    if state is None:
        raise ValueError(f"找不到任务 {job_id}")
    docx = report_docx_bytes(job_id, state)
    document_kind = "report"
    if docx is None:
        report_record = (state.get("academic_state") or {}).get(
            "artifacts", {}).get("report") or {}
        if state.get("report_enabled") and state.get("p3_done") and \
                report_record.get("status") in {"stale", "missing", "failed"}:
            raise RuntimeError("当前实践报告 artifact 已 stale，请先按影响范围重建报告")
        try:
            docx = build_delivery_assets(job_id, state).get("translation.docx")
            document_kind = "translation"
        except Exception as exc:
            raise RuntimeError(f"当前 DOCX 不可生成：{str(exc)[:180]}") from exc
    if not docx:
        raise RuntimeError("当前没有可渲染的 DOCX")
    previous_render = load_academic_artifact(job_id, "libreoffice_render") or {}
    current_qa = _finalization.normalize_final_qa(state.get("final_qa"))
    if (previous_render.get("rendered_pdf_hash") or
            previous_render.get("qa_status") == "PASS" or
            current_qa.get("author_visual_review") == "CONFIRMED" or
            current_qa.get("word_final_review") == "CONFIRMED"):
        _reset_final_qa(state, "LibreOffice render rerun; author and Word reviews reset")
        if current_qa.get("structural_qa") in {"PASS", "FAIL"}:
            state["final_qa"]["structural_qa"] = current_qa["structural_qa"]
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        qa = _finalization.normalize_final_qa(state.get("final_qa"))
        qa.update({"libreoffice_render": "NOT_RUN", "rendered_at": None,
                   "page_count": None, "source_docx_hash": _rendered_qa.sha256(docx),
                   "rendered_pdf_hash": None,
                   "updated_at": _finalization.now_iso(),
                   "translation_truth_version": int(
                       (state.get("translation_truth") or {}).get("version") or 0)})
        qa.setdefault("notes", {})["libreoffice_render"] = \
            "LibreOffice not installed; render NOT_RUN"
        state["final_qa"] = qa
        from transpraxis import academic_writer
        academic_writer._save_artifact(
            state, job_dir(job_id), "libreoffice_render", {
                "schema_version": _rendered_qa.VERSION,
                "status": "not_run", "qa_status": "NOT_RUN",
                "render_engine": "libreoffice", "render_engine_version": None,
                "source_docx_hash": _rendered_qa.sha256(docx),
                "rendered_pdf_hash": None, "rendered_at": None, "page_count": None,
                "stale_reason": "LibreOffice unavailable; render was not run",
                "analysis": {"warnings": [], "manual_reviews": [{
                    "type": "libreoffice_unavailable", "severity": "manual_review"}]},
            }, "no-engine", _rendered_qa.VERSION,
            input_artifact_ids=["final_docx_validation"], status="missing")
        generate_report_qa(job_id, state, save_file=True)
        save_job_state(job_id, state)
        return state, qa
    from transpraxis import academic_evidence, academic_writer
    version_result = subprocess.run([soffice, "--version"], capture_output=True,
                                    text=True, timeout=10, check=False)
    engine_version = (version_result.stdout or version_result.stderr or "").strip()
    with tempfile.TemporaryDirectory(prefix=f"transpraxis-lo-{job_id}-") as tmp:
        tmp_path = Path(tmp)
        source_path = tmp_path / "current.docx"
        source_path.write_bytes(docx)
        profile = tmp_path / "profile"
        profile.mkdir()
        command = [soffice, "--headless", "-env:UserInstallation=file://"
                   + str(profile), "--convert-to", "pdf", "--outdir", tmp,
                   str(source_path)]
        result = subprocess.run(command, capture_output=True, text=True,
                                timeout=120, check=False)
        pdf_path = tmp_path / "current.pdf"
        if result.returncode != 0 or not pdf_path.is_file():
            detail = (result.stderr or result.stdout or "未知 LibreOffice 错误").strip()
            qa = _finalization.normalize_final_qa(state.get("final_qa"))
            qa.update({"libreoffice_render": "FAIL",
                       "rendered_at": _finalization.now_iso(),
                       "translation_truth_version": int(
                           (state.get("translation_truth") or {}).get("version") or 0),
                       "updated_at": _finalization.now_iso()})
            qa.setdefault("notes", {})["libreoffice_render"] = detail[:700]
            state["final_qa"] = qa
            render_value = {
                "status": "fail", "qa_status": "FAIL", "document_kind": document_kind,
                "detail": detail[:700],
                "render_engine": "libreoffice", "render_engine_version": engine_version,
                "source_docx_hash": _rendered_qa.sha256(docx),
                "rendered_pdf_hash": None,
                "rendered_at": _finalization.now_iso(), "page_count": None,
                "stale_reason": {"code": "render_failed",
                                 "source_type": "artifact",
                                 "source_id": "final_docx_validation"},
            }
            academic_writer._save_artifact(
                state, job_dir(job_id), "libreoffice_render", render_value,
                academic_evidence.stable_hash({
                    "final_docx_validation": (state.get("academic_state") or {}).get(
                        "artifacts", {}).get("final_docx_validation", {}).get(
                            "content_hash"),
                    "version": _finalization.VERSION,
                }), _finalization.VERSION,
                input_artifact_ids=["final_docx_validation"], status="failed",
                stale_reason={"code": "render_failed", "source_type": "artifact",
                              "source_id": "final_docx_validation"})
            generate_report_qa(job_id, state, save_file=True)
            save_job_state(job_id, state)
            raise RuntimeError(f"LibreOffice 渲染失败：{detail[:180]}")
        pdf_bytes = pdf_path.read_bytes()
    output = job_dir(job_id) / "libreoffice-render.pdf"
    output.write_bytes(pdf_bytes)
    analysis = _rendered_qa.analyze_pdf(pdf_bytes)
    page_count = int(analysis.get("page_count") or 0)
    page_metrics = analysis.get("pages") or []
    render_status = "PASS" if page_count and not analysis.get(
        "definite_failures") else "FAIL"
    qa = _finalization.normalize_final_qa(state.get("final_qa"))
    qa.update({
        "libreoffice_render": render_status,
        "rendered_at": _finalization.now_iso(),
        "page_count": page_count,
        "source_docx_hash": _rendered_qa.sha256(docx),
        "rendered_pdf_hash": _rendered_qa.sha256(pdf_bytes),
        "translation_truth_version": int(
            (state.get("translation_truth") or {}).get("version") or 0),
        "updated_at": _finalization.now_iso(),
    })
    qa.setdefault("notes", {})["document_kind"] = document_kind
    qa["page_metrics"] = page_metrics
    state["final_qa"] = qa
    render_value = _rendered_qa.build_render_record(
        document_kind=document_kind, source_docx=docx, rendered_pdf=pdf_bytes,
        engine="libreoffice", engine_version=engine_version, analysis=analysis)
    render_value["qa_status"] = render_status
    render_value["status"] = "pass" if render_status == "PASS" else "fail"
    academic_writer._save_artifact(
        state, job_dir(job_id), "libreoffice_render", render_value,
        academic_evidence.stable_hash({
            "final_docx_validation": (state.get("academic_state") or {}).get(
                "artifacts", {}).get("final_docx_validation", {}).get("content_hash"),
            "version": _finalization.VERSION,
        }), _finalization.VERSION,
        input_artifact_ids=["final_docx_validation"],
        status="valid" if qa["libreoffice_render"] == "PASS" else "failed")
    generate_report_qa(job_id, state, save_file=True)
    save_job_state(job_id, state)
    return state, qa


def _reconcile_final_delivery_snapshot(job_id, state):
    latest = _snapshots.latest_snapshot(job_dir(job_id))
    if latest and state.get("delivery_status") == "final" \
            and _snapshots.state_identity(state) != latest.get("translation_state_identity"):
        return _invalidate_final_delivery_state(state)
    return state


def new_job_state(filename):
    state = {
        "filename": filename,
        "p1_done": False,
        "p2_done": False,
        "p3_done": False,
        "report_enabled": False,
        "paras": [],
        "pairs": [],
        "auto_terms": {},
        "findings": [],
        "review_stats": {
            "reviewed_segments": 0, "batches_reviewed": 0,
            "blocking": 0, "actionable": 0, "informational": 0, "review_failed": 0,
        },
        "tm_used_count": 0,
        "has_blocking": False,
        "p3_md": "",
        "p3_sections": [],
        "theory": "",
        "warnings": [],
        "annotations": {},
        "annotations_done": False,
        "annotations_done_offset": 0,
    }
    # 术语治理 / 交付门禁新增字段（默认值集中在 state_migration，保持单一来源）
    state.update(_state_migration._default_new_fields())
    return state


def load_job_state(job_id):
    p = job_state_path(job_id)
    if not p.is_file():
        return None
    try:
        raw = json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return None
    from transpraxis import delivery as _delivery
    state = _delivery.normalize_state_findings(_state_migration.migrate_state(raw))
    return _reconcile_final_delivery_snapshot(job_id, state)


def save_job_state(job_id, state):
    """原子写入（先写临时文件再替换），避免中断写坏 state.json。"""
    from transpraxis import delivery as _delivery
    _delivery.normalize_state_findings(state)
    _reconcile_final_delivery_snapshot(job_id, state)
    d = job_dir(job_id)
    d.mkdir(parents=True, exist_ok=True)
    tmp = d / "state.json.tmp"
    tmp.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(d / "state.json")


def _recount_reviewed_segments(state):
    """Keep the segment review count in sync with the canonical pair records."""
    state.setdefault("review_stats", {})["reviewed_segments"] = sum(
        bool(pair.get("reviewed")) for pair in state.get("pairs") or [])


def translation_terms_for_pair(state, pair):
    """Resolve only glossary entries explicitly attached to one pair."""
    ids = {str(item) for item in pair.get("glossary_entry_ids") or []}
    entries = state.get("glossary") or []
    if not entries:
        frozen = state.get("glossary_frozen") or {}
        entries = frozen.get("entries") or [] if isinstance(frozen, dict) else []
    if not entries:
        versions = state.get("glossary_versions") or []
        if versions and isinstance(versions[-1], dict):
            entries = versions[-1].get("entries") or []
    if not ids or not isinstance(entries, list):
        return []
    terms = []
    for entry in entries:
        if not isinstance(entry, dict) or str(entry.get("id") or "") not in ids:
            continue
        source = str(entry.get("source") or "")
        if source:
            terms.append((source, str(entry.get("preferred") or entry.get("target") or "—"),
                          "项目术语"))
    return terms[:8]


def translation_visible_indexes(state, search="", status_filter="全部",
                               filter_terms=False, filter_edited=False,
                               filter_issues=False, filter_tm=False,
                               issue_indexes=None):
    """Return visible pair indexes without mutating the loaded state."""
    pairs = state.get("pairs") or []
    query = str(search or "").strip().casefold()
    issue_indexes = set(issue_indexes or [])
    visible = []
    for index, pair in enumerate(pairs):
        source = str(pair.get("source") or "")
        target = str(pair.get("target") or "")
        if query:
            paragraph_query = query.lstrip("#").strip()
            if paragraph_query.isdigit():
                if int(paragraph_query) != index + 1:
                    continue
            elif query not in f"{source}\n{target}".casefold():
                continue
        if status_filter == "待审" and pair.get("reviewed"):
            continue
        if status_filter == "已审校" and not pair.get("reviewed"):
            continue
        if filter_terms and not translation_terms_for_pair(state, pair):
            continue
        if filter_edited and not pair.get("human_edited"):
            continue
        if filter_tm and not pair.get("from_tm"):
            continue
        if filter_issues and index not in issue_indexes:
            continue
        visible.append(index)
    return visible


def save_translation_edit(job_id, index, target, actor="user"):
    """Save one human translation edit through the state/business layer.

    A manual edit is a new working translation: it preserves the paragraph
    identity, clears segment review, removes TM reuse provenance, and invalidates
    mutable final approval while leaving historical frozen snapshots untouched.
    """
    from transpraxis import delivery as _delivery

    state = load_job_state(job_id)
    if state is None:
        raise ValueError(f"找不到任务 {job_id}")
    pairs = state.get("pairs") or []
    if not (0 <= index < len(pairs)):
        raise IndexError(f"段落索引超出范围：{index}")
    pair = pairs[index]
    new_target = str(target or "").strip()
    if not pair.get("human_edited"):
        pair["_translation_edit_restore"] = {
            "target": pair.get("target") or "",
            "reviewed": bool(pair.get("reviewed")),
            "review_status": pair.get("review_status", "not_reviewed"),
            "target_provenance": pair.get("target_provenance", "generated"),
            "from_tm": bool(pair.get("from_tm")),
        }
    pair["target"] = new_target
    pair["human_edited"] = True
    pair["reviewed"] = False
    pair["review_status"] = "not_reviewed"
    pair["target_provenance"] = "human_edit"
    pair["from_tm"] = False
    _recount_reviewed_segments(state)

    _mark_translation_truth_changed(
        job_id, state, [index], "人工修改 CURRENT_TRANSLATION；相关案例与学术下游需要重建",
        actor=actor, action="translation_edit")
    _recheck_delivery_invariants_for_segments(state, [index], actor=actor)
    save_job_state(job_id, state)
    return state


def _recheck_delivery_invariants_for_segments(state, indexes, *, actor="user"):
    """Refresh deterministic target blockers after a working-text edit.

    A target invariant describes the text that existed when it was recorded.
    When a user edits that target, a resolved old invariant must not continue
    to block delivery; if the same problem remains, the existing finding stays
    open and the current validation report remains authoritative.
    """
    from transpraxis import delivery as _delivery

    indexes = {int(index) for index in indexes}
    report = validate_delivery_translation_state(state)
    active = {
        (issue.get("segment_index"), issue.get("code"))
        for issue in report.get("issues") or []
    }
    for finding in state.setdefault("findings", []):
        if not isinstance(finding, dict) \
                or finding.get("type") != "delivery_invariant" \
                or finding.get("segment_index") not in indexes \
                or finding.get("resolved"):
            continue
        key = (finding.get("segment_index"), finding.get("invariant_code"))
        if key in active:
            continue
        finding["resolved"] = True
        finding["resolution"] = {
            "action": "target_rechecked",
            "note": "CURRENT_TRANSLATION 编辑后重新通过目标文本门禁",
            "timestamp": _finalization.now_iso(),
            "actor": actor,
        }
        _delivery.add_human_action(
            state, _delivery.finding_id(finding), "target_rechecked",
            "CURRENT_TRANSLATION 编辑后重新通过目标文本门禁", actor)
    _record_delivery_validation_findings(state, report)
    state["delivery_status"] = _delivery.compute_delivery_status(state)
    return report


def restore_translation_edit(job_id, index, actor="user"):
    """Restore the translation state that existed before the first edit."""
    from transpraxis import delivery as _delivery

    state = load_job_state(job_id)
    if state is None:
        raise ValueError(f"找不到任务 {job_id}")
    pairs = state.get("pairs") or []
    if not (0 <= index < len(pairs)):
        raise IndexError(f"段落索引超出范围：{index}")
    pair = pairs[index]
    restore = pair.pop("_translation_edit_restore", None)
    if restore is None and not pair.get("human_edited"):
        return state
    restore = restore or {}
    pair["target"] = restore.get("target") or pair.get("initial_target") or ""
    pair["reviewed"] = bool(restore.get("reviewed"))
    pair["review_status"] = restore.get("review_status", "not_reviewed")
    pair["target_provenance"] = restore.get("target_provenance", "generated")
    pair["from_tm"] = bool(restore.get("from_tm"))
    pair.pop("human_edited", None)
    _recount_reviewed_segments(state)

    _mark_translation_truth_changed(
        job_id, state, [index], "恢复前版本也改变了 CURRENT_TRANSLATION；相关下游需要重建",
        actor=actor, action="translation_restore")
    save_job_state(job_id, state)
    return state


def list_jobs():
    jobs = []
    _ensure_output_dir()
    for d in sorted(OUTPUT_DIR.iterdir()):
        sp = d / "state.json"
        if sp.is_file():
            try:
                s = json.loads(sp.read_text(encoding="utf-8"))
            except Exception:
                continue
            from transpraxis import delivery as _delivery
            s = _delivery.normalize_state_findings(_state_migration.migrate_state(s))
            s = _reconcile_final_delivery_snapshot(d.name, s)
            jobs.append({"job_id": d.name, "state": s})
    return jobs


def delete_job(job_id):
    d = job_dir(job_id)
    if d.exists():
        shutil.rmtree(d)


def file_job_id(file_bytes):
    """以文件内容哈希作为任务 ID：同一文件重传可自动续传，不同文件不会串状态。"""
    return hashlib.sha256(file_bytes).hexdigest()[:16]


def save_source(job_id, file_bytes):
    """留存源文件，刷新页面后即使不重新上传也能继续（如重做阶段一）。"""
    d = job_dir(job_id)
    d.mkdir(parents=True, exist_ok=True)
    (d / "source.bin").write_bytes(file_bytes)


def load_source(job_id):
    p = job_dir(job_id) / "source.bin"
    return p.read_bytes() if p.is_file() else None


def save_report_template(job_id, filename, template_bytes):
    """Parse and persist the uploaded DOCX template before academic stages run."""
    from transpraxis import report_template

    raw = _bytes(template_bytes)
    contract = report_template.parse_docx_template(filename, raw)
    directory = job_dir(job_id)
    directory.mkdir(parents=True, exist_ok=True)
    (directory / "report-template.docx").write_bytes(raw)
    (directory / "template-contract.json").write_text(
        json.dumps(contract, ensure_ascii=False, indent=2), encoding="utf-8")
    state = load_job_state(job_id)
    if state is not None:
        identity = contract.get("template_identity") or {}
        state["report_template"] = {
            "filename": identity.get("filename"),
            "template_id": identity.get("template_id"),
            "template_hash": identity.get("sha256"),
            "schema_version": contract.get("schema_version"),
            "status": "parsed",
        }
        state["report_template_contract"] = contract
        settings = dict(state.get("research_settings") or {})
        settings["report_template_contract"] = contract
        state["research_settings"] = settings
        save_job_state(job_id, state)
    return contract


def clear_report_template(job_id):
    """Remove the saved report template and its contract from a task."""
    directory = job_dir(job_id)
    for name in ("report-template.docx", "template-contract.json"):
        path = directory / name
        if path.is_file():
            path.unlink()
    state = load_job_state(job_id)
    if state is not None:
        state["report_template"] = None
        state["report_template_contract"] = None
        settings = dict(state.get("research_settings") or {})
        settings.pop("report_template_contract", None)
        settings.pop("template_contract", None)
        state["research_settings"] = settings
        save_job_state(job_id, state)
    return state


def load_report_template(job_id):
    """Load the persisted DOCX bytes and canonical contract, if configured."""
    from transpraxis import report_template

    directory = job_dir(job_id)
    docx_path = directory / "report-template.docx"
    contract_path = directory / "template-contract.json"
    state = load_job_state(job_id)
    contract = None
    if contract_path.is_file():
        try:
            value = json.loads(contract_path.read_text(encoding="utf-8"))
            contract = value if isinstance(value, dict) else None
        except (OSError, ValueError):
            contract = None
    if contract is None and state:
        contract = state.get("report_template_contract") or \
            (state.get("research_settings") or {}).get("report_template_contract")
    if not contract or not docx_path.is_file():
        return None
    raw = docx_path.read_bytes()
    return {
        "bytes": raw,
        "contract": contract,
        "metadata": state.get("report_template") if state else None,
        "summary": report_template.contract_summary(contract),
    }


def _bytes(value):
    return value.getvalue() if hasattr(value, "getvalue") else bytes(value)


def report_docx_bytes(job_id, state=None, frozen_assets=None):
    """Render the structured report with its template, or use the legacy fallback."""
    if frozen_assets and frozen_assets.get("stage3_report.docx"):
        return frozen_assets["stage3_report.docx"]
    state = state or load_job_state(job_id) or {}
    report = state.get("p3_md")
    if not report:
        return None
    template = load_report_template(job_id)
    if template:
        from transpraxis import academic_evidence, academic_writer, compliance
        from transpraxis import final_docx, report_template
        artifact = load_academic_artifact(job_id, "report")
        report_record = academic_writer.artifact_record(state, "report")
        if isinstance((state.get("academic_state") or {}).get("artifacts", {}).get(
                "report"), dict) and report_record.get("status") != "valid":
            return None
        if not artifact:
            raise report_template.TemplateParseError(
                "模板化报告缺少结构化 report artifact，请重新生成报告。")
        if artifact.get("report_status") not in {
                "generated", "review_required", "literature_required"} or artifact.get(
                "template_compliance") not in {"pass", "pass_with_warnings"}:
            return None
        rendered = _bytes(report_template.render_report_docx(
            artifact, template["bytes"], template["contract"]))
        source_docx_hash = _rendered_qa.sha256(rendered)
        previous_docx = academic_writer.artifact_record(
            state, "final_docx_validation")
        previous_qa = _finalization.normalize_final_qa(state.get("final_qa"))
        previous_docx_hash = str(
            previous_docx.get("source_docx_hash") or
            previous_qa.get("source_docx_hash") or "")
        if (previous_docx_hash and previous_docx_hash != source_docx_hash) or (
                isinstance(previous_docx, dict) and previous_docx and (
                    previous_qa.get("author_visual_review") == "CONFIRMED" or
                    previous_qa.get("word_final_review") == "CONFIRMED")):
            academic_writer.propagate_artifact_staleness(
                state, input_artifact_ids=["final_docx_validation"])
            _reset_final_qa(state, "DOCX bytes changed; render and human reviews reset")
        final_validation = final_docx.validate_final_docx(rendered, artifact)
        final_validation.update({
            "source_docx_hash": source_docx_hash,
            "report_content_hash": (state.get("academic_state") or {}).get(
                "artifacts", {}).get("report", {}).get("content_hash") or
                artifact.get("content_hash"),
            "layout_facts": compliance.inspect_docx_layout(rendered),
        })
        final_validation["content_hash"] = academic_evidence.stable_hash({
            key: value for key, value in final_validation.items()
            if key != "content_hash"
        })
        final_validation_dep = academic_evidence.stable_hash({
            "report": (state.get("academic_state") or {}).get("artifacts", {}).get(
                "report", {}).get("content_hash") or artifact.get("content_hash"),
            "template": template.get("contract", {}).get("template_identity") or
            template.get("contract", {}).get("template_hash"),
            "version": final_docx.SCHEMA_VERSION,
        })
        academic_writer._save_artifact(
            state, job_dir(job_id), "final_docx_validation", final_validation,
            final_validation_dep, final_docx.SCHEMA_VERSION,
            input_artifact_ids=["report"],
            status="valid" if final_validation.get("status") != "fail" else "failed")
        qa = _finalization.normalize_final_qa(state.get("final_qa"))
        qa.update({
            "structural_qa": "FAIL" if final_validation.get("status") == "fail"
            else "PASS",
            "source_docx_hash": source_docx_hash,
            "updated_at": _finalization.now_iso(),
        })
        state["final_qa"] = qa
        save_compliance_record(job_id, state)
        generate_report_qa(job_id, state, save_file=True)
        save_job_state(job_id, state)
        if final_validation.get("status") == "fail":
            return None
        return rendered
    if state.get("report_status") in {
            "incomplete", "failed_template_validation", "review_required"}:
        return None
    from transpraxis import academic_writer
    if isinstance((state.get("academic_state") or {}).get("artifacts", {}).get(
            "report"), dict) and academic_writer.artifact_record(
                state, "report").get("status") != "valid":
        return None
    return _bytes(markdown_to_word(report, state.get("theory") or ""))


def validate_translation_target(source, target, *, segment_index=None,
                                allow_json=False):
    """Public target invariant entry point used by runtime and delivery."""
    return _translation_target.validate_translation_target(
        source, target, segment_index=segment_index, allow_json=allow_json)


def validate_translation_pairs(pairs, glossary=None, target_lang=""):
    """Validate pairs independently of the model response parser."""
    pairs = list(pairs or [])
    target_report = _translation_target.validate_translation_pairs(pairs)
    qa_findings = check_translation_batch(
        [str(pair.get("source") or "") for pair in pairs if isinstance(pair, dict)],
        [str(pair.get("target") or "") for pair in pairs if isinstance(pair, dict)],
        glossary or [], target_lang or "简体中文")
    return {
        **target_report,
        "qa_findings": qa_findings,
        "blocking_findings": [
            finding for finding in qa_findings
            if finding.get("severity") == "blocking"
        ],
    }


def validate_delivery_translation_state(state):
    """Run the final, state-level translation gate before any delivery asset."""
    state = state if isinstance(state, dict) else {}
    pairs = state.get("pairs") or []
    report = validate_translation_pairs(
        pairs, state.get("glossary") or [], state.get("target_lang") or "简体中文")
    issues = list(report.get("issues") or [])
    if state.get("p2_done") and len(pairs) != len(state.get("paras") or []):
        issues.append({
            "code": "pair_count_mismatch",
            "message": "双语 pairs 数量与源文段落数量不一致，不能生成完整交付物",
            "severity": "blocking",
        })
    entity_findings = _entity_registry.EntityRegistry(
        state.get("entity_registry") or []
    ).consistency_findings()
    return {
        **report,
        "issues": issues,
        "entity_findings": entity_findings,
        "blocking": bool(issues or report.get("blocking_findings")),
        "status": "fail" if issues or report.get("blocking_findings") else (
            "review_required" if entity_findings else "pass"),
    }


def _record_delivery_validation_findings(state, report):
    """Persist only new invariant/entity findings for the review queue."""
    state["delivery_validation"] = {
        "status": report.get("status"),
        "blocking": bool(report.get("blocking")),
        "checked_pairs": report.get("checked_pairs", 0),
        "issues": list(report.get("issues") or []),
        "entity_conflicts": [
            {
                "source": item.get("source"),
                "preferred_target": item.get("preferred_target"),
                "observed_targets": item.get("observed_targets"),
            }
            for item in report.get("entity_findings") or []
        ],
    }
    existing = {
        (item.get("type"), item.get("segment_index"), item.get("invariant_code"),
         item.get("reason"))
        for item in state.setdefault("findings", [])
        if isinstance(item, dict) and not item.get("resolved")
    }
    findings = _translation_target.target_invariant_findings(report)
    for item in report.get("entity_findings") or []:
        findings.append({
            **item,
            "segment_index": item.get("segment_index"),
            "segment_id": item.get("segment_id"),
            "summary": item.get("reason"),
            "explanation": item.get("reason"),
            "recommendation": "核对全文同一实体的译名，并在人工作区确认一个一致形式。",
            "confidence": None,
            "diagnostic_version": 1,
        })
    for finding in findings:
        key = (finding.get("type"), finding.get("segment_index"),
               finding.get("invariant_code"), finding.get("reason"))
        if key not in existing:
            state["findings"].append(finding)
            existing.add(key)
    stats = state.setdefault("review_stats", {})
    stats["blocking"] = sum(
        1 for item in state["findings"]
        if item.get("severity") == "blocking" and not item.get("resolved")
    )
    stats["actionable"] = sum(
        1 for item in state["findings"]
        if item.get("severity") == "actionable" and not item.get("resolved")
    )
    stats["informational"] = sum(
        1 for item in state["findings"]
        if item.get("severity") == "informational" and not item.get("resolved")
    )
    state["has_blocking"] = stats["blocking"] > 0
    if report.get("blocking") and state.get("delivery_status") in ("approved", "final"):
        # A state-level mutation (for example a manually injected target) must
        # invalidate the mutable approval even when no snapshot exists yet.
        _invalidate_final_delivery_state(state)
    return state


def _academic_workspace_archive(job_id):
    from transpraxis import academic_writer

    names = ("research_model", "argument_plan", "selected_cases", "outline",
             "case_analysis_plans", "human_evidence_questions")
    output = io.BytesIO()
    written = 0
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as bundle:
        for name in names:
            filename = academic_writer.ARTIFACT_FILES[name]
            path = job_dir(job_id) / filename
            if path.is_file():
                bundle.writestr(filename, path.read_bytes())
                written += 1
    return output.getvalue() if written else None


def _delivery_asset_bundle(job_id, state, target_lang, provider, model):
    """Build the configured delivery set used by both preview and snapshots."""
    from transpraxis import assets as _assets
    from transpraxis import report_evidence as _report_evidence

    validation = validate_delivery_translation_state(state)
    _record_delivery_validation_findings(state, validation)
    if validation["blocking"]:
        save_job_state(job_id, state)
        reasons = "；".join(
            issue.get("message", "目标文本未通过交付检查")
            for issue in validation.get("issues") or []
        )
        raise RuntimeError(
            "交付被 Translation Target Invariant 阻止"
            + (f"：{reasons}" if reasons else "")
        )

    source = load_source(job_id)
    snapshot_state = dict(state)
    if source is not None:
        snapshot_state["_source_bin"] = source
    filename = state.get("filename", "")
    configured = state.get("delivery_config")
    if isinstance(configured, dict) and configured:
        config = normalize_delivery_config(
            configured, enable_report=state.get("report_enabled", False),
            enable_annotate=state.get("enable_annotate", False))
        bundle = {}
        pairs = state.get("pairs") or []
        glossary = state.get("glossary") or []
        if state.get("p2_done") and pairs:
            if config["deliver_plain_docx"]:
                bundle["translation.docx"] = _bytes(translations_to_word(pairs))
            if config["deliver_bilingual_docx"]:
                bundle["bilingual.docx"] = _bytes(pairs_to_word(pairs))
            if config["deliver_pdf"]:
                bundle["translation.pdf"] = translations_to_pdf(pairs)
            if config["enable_annotate"]:
                bundle["annotated_bilingual.docx"] = _bytes(pairs_to_word(
                    pairs, annotations=state.get("annotations"),
                    colors=ANNOTATION_COLORS))
            if config["deliver_terms_xlsx"]:
                bundle["terms.xlsx"] = _bytes(glossary_to_excel(
                    glossary, state.get("auto_terms")))
            if config["deliver_tbx"]:
                bundle["terms.tbx"] = _assets.build_tbx(glossary)
            if config["deliver_tmx"]:
                bundle["memory.tmx"] = _assets.build_tmx(state, job_id=job_id)
            if config["deliver_jsonl"]:
                bundle["bilingual.jsonl"] = _assets.build_jsonl(
                    state, job_id=job_id).encode("utf-8")
            if config["deliver_evidence"]:
                bundle["segment_evidence.jsonl"] = \
                    _report_evidence.export_segment_evidence_jsonl(
                        state, job_id).encode("utf-8")
            if config["deliver_review_report"]:
                bundle["review_report.md"] = findings_report_md(state).encode("utf-8")
        if config["deliver_cases"]:
            selected_cases = load_academic_artifact(job_id, "selected_cases")
            if selected_cases:
                bundle["selected_cases.json"] = (
                    json.dumps(selected_cases, ensure_ascii=False, indent=2) + "\n"
                ).encode("utf-8")
        if config["deliver_academic_workspace"]:
            workspace = _academic_workspace_archive(job_id)
            if workspace:
                bundle["academic_workspace.zip"] = workspace
        if config["enable_report"] and state.get("p3_md"):
            report_docx = report_docx_bytes(job_id, state)
            if report_docx is not None:
                bundle["report.docx"] = report_docx
            bundle["report.md"] = state["p3_md"].encode("utf-8")
        generated_assets = sorted([*bundle, "delivery_manifest.json"])
        manifest = _assets.build_delivery_manifest(
            snapshot_state, job_id, target_lang, provider, model,
            generated_assets=generated_assets, source_filename=filename,
            translator_config=state.get("translator_config"),
            reviewer_config=state.get("reviewer_config"))
        bundle["delivery_manifest.json"] = (
            json.dumps(manifest, ensure_ascii=False, indent=2) + "\n").encode("utf-8")
        return bundle

    # Existing jobs without a saved selection retain their historical bundle.
    bundle = {}
    if state.get("p1_done") and state.get("paras"):
        bundle["stage1_cleaned.docx"] = _bytes(paragraphs_to_word(state["paras"]))
    if state.get("auto_terms"):
        bundle["auto_terms.xlsx"] = _bytes(dict_to_excel(state["auto_terms"]))
    if state.get("p2_done") and state.get("pairs"):
        bundle["stage2_bilingual.docx"] = _bytes(pairs_to_word(
            state["pairs"], annotations=state.get("annotations"), colors=ANNOTATION_COLORS))
    if state.get("p3_md"):
        report_docx = report_docx_bytes(job_id, state)
        if report_docx is not None:
            bundle["stage3_report.docx"] = report_docx
    bundle.update(_assets.export_all(
        snapshot_state, job_id, target_lang, provider, model,
        source_filename=filename, source_bin=source))
    if state.get("p2_done"):
        bundle["segment_evidence.jsonl"] = _report_evidence.export_segment_evidence_jsonl(
            state, job_id).encode("utf-8")
        if state.get("findings"):
            bundle["review_report.md"] = findings_report_md(state).encode("utf-8")
    manifest = _assets.build_delivery_manifest(
        snapshot_state, job_id, target_lang, provider, model,
        generated_assets=sorted(bundle), source_filename=filename,
        translator_config=state.get("translator_config"),
        reviewer_config=state.get("reviewer_config"))
    bundle["delivery_manifest.json"] = (
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n").encode("utf-8")
    return bundle


def build_delivery_assets(job_id, state=None, target_lang="", provider="", model=""):
    state = state or load_job_state(job_id) or {}
    return _delivery_asset_bundle(
        job_id, state, target_lang or state.get("target_lang") or "",
        provider or state.get("provider") or "",
        model or state.get("model") or "")


def create_delivery_snapshot(job_id, state, target_lang="", provider="", model=""):
    validation = validate_delivery_translation_state(state)
    _record_delivery_validation_findings(state, validation)
    if validation["blocking"]:
        save_job_state(job_id, state)
        raise RuntimeError("最终交付未通过 Translation Target Invariant")
    source = load_source(job_id)
    frozen = state.get("glossary_frozen") or {}
    source_hash = hashlib.sha256(source).hexdigest() if source else \
        str(frozen.get("source_hash") or "")
    manifest = _snapshots.create_snapshot(
        job_dir(job_id), job_id, state,
        _delivery_asset_bundle(job_id, state, target_lang, provider, model),
        source_identity={
            "filename": state.get("filename", ""),
            "source_hash": source_hash,
            "source_page_count": state.get("source_page_count"),
        },
        active_terminology_version={
            "version": frozen.get("version"),
            "glossary_hash": frozen.get("glossary_hash"),
            "status": "已冻结" if frozen else "未冻结",
        })
    return manifest


def list_delivery_snapshots(job_id):
    return _snapshots.list_snapshots(job_dir(job_id))


def delivery_snapshot_status(job_id, state=None):
    state = state if state is not None else load_job_state(job_id)
    latest = _snapshots.latest_snapshot(job_dir(job_id))
    if latest is None:
        return {"latest": None, "current": False, "diverged": False, "integrity": False}
    current_identity = _snapshots.state_identity(state or {})
    matches = current_identity == latest.get("translation_state_identity")
    assets = delivery_snapshot_assets(job_id, latest["snapshot_version"])
    integrity = bool(assets) and all(data is not None for data in assets.values())
    return {
        "latest": latest,
        "current": bool(matches and integrity and state
                        and state.get("delivery_status") == "final"),
        "diverged": not matches,
        "integrity": integrity,
    }


def delivery_snapshot_assets(job_id, version):
    manifest = next((item for item in list_delivery_snapshots(job_id)
                     if item.get("snapshot_version") == int(version)), None)
    if manifest is None:
        return {}
    return {
        item["name"]: _snapshots.load_asset(job_dir(job_id), version, item["name"])
        for item in manifest.get("assets") or []
    }


def delivery_snapshot_archive(job_id, version):
    return _snapshots.archive(job_dir(job_id), version)


def progress_label(state):
    academic = state.get("academic_state") or {}
    if academic.get("status") == "failed":
        return "翻译完成 · 学术写作失败"
    if academic.get("status") in ("in_progress", "stale"):
        return "翻译完成 · 学术写作中"
    if academic.get("quality_status") == "review_required":
        return "翻译完成 · 报告待学术复核"
    if academic.get("quality_status") == "fail":
        return "翻译完成 · 报告验证失败"
    if state.get("p1_done") and state.get("p2_done") and \
            (state.get("p3_done") or not state.get("report_enabled", True)):
        return "已完成"
    if state.get("p2_done"):
        return "报告生成中"
    if state.get("p1_done"):
        return "翻译中"
    return "待处理"


def recovery_summary(job_id, state=None):
    """Read-only durable progress summary used by History and the workspace."""
    state = state if state is not None else load_job_state(job_id)
    return _checkpoint.recovery_summary(job_dir(job_id), state or {})


def load_context_artifacts(job_id, state=None):
    """Load persisted context artifacts, tolerating older or partial jobs."""
    state = state if state is not None else load_job_state(job_id)
    state = state or {}

    def read_json(name, default):
        path = job_dir(job_id) / name
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except (OSError, TypeError, ValueError):
            return default

    units = state.get("semantic_units") or read_json("semantic_units.json", [])
    digests = state.get("section_digests") or read_json("section_digests.json", [])
    synopsis = state.get("document_synopsis")
    stored_synopsis = read_json("document_synopsis.json", None)
    if not isinstance(synopsis, dict):
        synopsis = {}
    if not synopsis or (synopsis.get("status") == "pending"
                        and isinstance(stored_synopsis, dict)
                        and stored_synopsis.get("status") != "pending"):
        synopsis = stored_synopsis or {}
    return {
        "semantic_units": units if isinstance(units, list) else [],
        "section_digests": digests if isinstance(digests, list) else [],
        "document_synopsis": synopsis if isinstance(synopsis, dict) else {},
        "warnings": list(state.get("understanding_warnings") or []),
    }


def task_status_label(state, job_id=""):
    """User-facing task status derived from persisted workflow state."""
    from transpraxis import delivery as _delivery

    if job_id:
        snapshot = delivery_snapshot_status(job_id, state)
        latest = snapshot.get("latest") or {}
        version = latest.get("snapshot_version")
        if snapshot.get("current"):
            return (f"已冻结交付 v{version}" if version is not None else
                    "已冻结交付")
        if snapshot.get("diverged"):
            return (f"工作版本已偏离冻结交付 v{version}" if version is not None else
                    "工作版本已偏离冻结交付")

    if state.get("delivery_status") == "final":
        return "已冻结交付"
    if state.get("p2_done") and _delivery.unresolved_blocking(state):
        return "待审校"
    academic = state.get("academic_state") or {}
    if state.get("p2_done") and state.get("p3_done") and (
            academic.get("quality_status") in ("review_required", "fail", "failed")
            or academic.get("status") == "failed"):
        return "待学术复核"
    if state.get("p2_done") and (
            state.get("p3_done") or not state.get("report_enabled", True)):
        if state.get("report_enabled", True):
            final_qa = state.get("final_qa") or {}
            report_status = state.get("report_status") or academic.get("report_status")
            if (report_status != "generated" or
                    final_qa.get("author_visual_review") != "CONFIRMED" or
                    final_qa.get("word_final_review") != "CONFIRMED"):
                return "暂不满足交付条件"
        return "可以冻结交付"
    if state.get("p1_done") and not state.get("p2_done"):
        if (state.get("stage") == "TERMS_PREPARED" and state.get("quality_mode")
                and state.get("glossary") is not None
                and not state.get("glossary_frozen")
                and not state.get("quality_bypass")):
            return "待术语确认"
        return "处理中断"
    if state.get("p2_done") and not state.get("p3_done"):
        return "处理中断"
    return "待处理"


# ================= 术语审核状态（草稿 / 锁定 / 拒绝 / 冻结）=================
def save_glossary_draft(job_id, entries):
    """保存术语审核草稿（不冻结）。刷新/重启后从 TERMS_PREPARED 恢复。"""
    state = load_job_state(job_id)
    if state is None:
        return None
    norm = normalize_glossary(entries)
    state["glossary_draft"] = norm
    state["glossary"] = norm
    state["stage"] = "TERMS_PREPARED"
    save_job_state(job_id, state)
    return state


def _apply_glossary_staleness(state, job_id=None):
    """把受冻结术语表变更影响的段落标记 stale，并清除其 TM 信任。

    权威集合：每次调用都重新计算（先清旧标记/旧 stale finding，
    再标记当前受影响段）。stale 段：
    - stale_due_to_glossary = True，reviewed = False，from_tm = False；
    - 追加 blocking finding（type=glossary_stale），交付回到 review_required；
    - 从 translation_memory.json 中清除，防止后续任务精确命中旧译文。
    """
    from transpraxis import delivery as _delivery
    from transpraxis.terminology import stale_segments_for_glossary

    stale = stale_segments_for_glossary(state)
    pairs = state.get("pairs") or []
    for p in pairs:
        p.pop("stale_due_to_glossary", None)
    state["findings"] = [f for f in state.get("findings") or []
                         if f.get("type") != "glossary_stale"]
    if not stale:
        return state, []

    state["knowledge_candidates"] = _knowledge.discard_candidates_for_segments(
        state.get("knowledge_candidates") or [], stale)
    state["translation_continuity"] = list(state["knowledge_candidates"])

    fg = state.get("glossary_frozen") or {}
    for i in stale:
        p = pairs[i]
        p["stale_due_to_glossary"] = True
        p["reviewed"] = False
        p["from_tm"] = False
        p["review_status"] = "not_reviewed"
        for key in ("accepted_target", "human_accepted", "accepted_by_human",
                    "target_provenance"):
            p.pop(key, None)
        state["findings"].append({
            "segment_id": i, "segment_index": i,
            "severity": "blocking",
            "type": "glossary_stale",
            "entry_id": None,
            "category": "terminology_consistency",
            "summary": "冻结术语表已变更，本段需要重新确认",
            "source_span": None, "target_span": None,
            "explanation": f"冻结术语表已从当前版本变更为 v{fg.get('version')}，本段原有译文可能不再符合最新术语规则。",
            "recommendation": "检查本段涉及的术语；确认译法后标记已处理，或使用最新术语重新翻译。",
            "confidence": None, "detector": "Terminology QA",
            "diagnostic_version": 1,
            "reason": f"冻结术语表已变更（当前 v{fg.get('version')}），本段需复核或重译",
        })

    # 受影响段不得继续作为可信翻译记忆
    tm = load_tm()
    dirty = False
    for i in stale:
        src = pairs[i]["source"]
        if src in tm:
            del tm[src]
            dirty = True
    if dirty:
        save_tm(tm)
    state["delivery_approved_by_human"] = False
    state["delivery_approval"] = None

    stats = state.setdefault("review_stats", {})
    stats["blocking"] = sum(1 for f in state["findings"]
                            if f["severity"] == "blocking")
    stats["actionable"] = sum(1 for f in state["findings"]
                              if f["severity"] == "actionable")
    stats["informational"] = sum(1 for f in state["findings"]
                                 if f["severity"] == "informational")
    state["has_blocking"] = stats["blocking"] > 0
    state["delivery_status"] = _delivery.compute_delivery_status(state)
    if job_id:
        from transpraxis import academic_writer
        segment_ids = [
            _finalization.segment_id(job_id, index, pairs[index])
            for index in stale if 0 <= index < len(pairs)
        ]
        propagated = academic_writer.propagate_artifact_staleness(
            state, input_segment_ids=segment_ids)
        enriched = dict(state)
        enriched["_finalization_artifacts"] = _finalization_artifacts(job_id)
        state["dependency_impact"] = _finalization.build_dependency_impact(
            enriched, job_id, stale, "冻结术语表变化")
        affected_cases = [name.split(":", 1)[1] for name in propagated
                          if str(name).startswith("case:")]
        _finalization.mark_case_reviews_stale(
            state, affected_cases, "冻结术语表变化影响了本案例绑定段落")
        if propagated:
            academic = state.setdefault("academic_state", {})
            if any(name in propagated for name in {"sections", "report", "validation", "review"}):
                state["p3_done"] = False
                state["p3_md"] = ""
                state["p3_sections"] = []
                academic["status"] = "stale"
            _reset_final_qa(state, "冻结术语表变化；相关案例与学术下游需重新检查")
            _invalidate_final_delivery_state(state)
    return state, stale


def set_glossary_entry_status(job_id, entry_ids, status):
    """批量修改术语状态（candidate/provisional/locked/rejected）。"""
    if status not in ("candidate", "provisional", "locked", "rejected"):
        raise ValueError(f"非法状态：{status}")
    state = load_job_state(job_id)
    if state is None:
        return None
    ids = set(entry_ids or [])
    for e in state.get("glossary") or []:
        if e.get("id") in ids:
            e["status"] = status
    save_job_state(job_id, state)
    return state


def review_knowledge_candidate(job_id, candidate_id, decision, actor="user"):
    """Apply an explicit human decision to one persisted knowledge candidate.

    The only project-wide promotion path is the existing per-task glossary
    freeze/version workflow.  There is intentionally no global glossary store.
    """
    allowed = {"project_term", "task_only", "rejected"}
    if decision not in allowed:
        raise ValueError(f"非法知识候选决策：{decision}")
    state = load_job_state(job_id)
    if state is None:
        return None, False, "任务不存在"
    candidate = next((item for item in state.get("knowledge_candidates") or []
                      if _knowledge.candidate_id(item) == str(candidate_id)), None)
    if candidate is None:
        return state, False, "找不到待确认词条"
    if candidate.get("decision"):
        return state, False, "该词条已经处理过"

    context = _knowledge.candidate_context(candidate, state)
    timestamp = datetime.now(timezone.utc).isoformat(timespec="seconds")

    def record(action, note, event):
        candidate["decision"] = decision
        candidate["decision_at"] = timestamp
        candidate["decision_by"] = actor
        candidate["decision_note"] = note
        candidate["status"] = {
            "project_term": "promoted_project_term",
            "task_only": "accepted_task",
            "rejected": "rejected",
        }[decision]
        state["translation_continuity"] = [
            dict(item) for item in state.get("knowledge_candidates") or []
            if isinstance(item, dict)
        ]
        state.setdefault("knowledge_events", []).append({
            "type": "human_candidate_decision",
            "candidate_id": context["candidate_id"],
            "source": context["source"],
            "observed_target": context["proposed_target"],
            "decision": decision,
            "scope": "document" if decision == "project_term" else "task",
            "timestamp": timestamp,
            "actor": actor,
            "note": note,
            **event,
        })
        state.setdefault("human_actions", []).append({
            "action": action,
            "finding_id": f"knowledge:{context['candidate_id']}",
            "note": note,
            "timestamp": timestamp,
            "actor": actor,
        })

    if decision == "project_term":
        if context["conflicts"]:
            return state, False, "与现有项目术语冲突，未覆盖现有术语"
        entries = normalize_glossary(
            state.get("glossary") or (state.get("glossary_frozen") or {}).get("entries") or [])
        matching = [entry for entry in entries
                    if entry.get("source", "").casefold() == context["source"].casefold()]
        target = context["proposed_target"]
        changed = False
        if matching:
            entry = matching[0]
            current = str(entry.get("preferred") or entry.get("target") or "").strip()
            if current.casefold() != target.casefold():
                return state, False, "与现有术语译名冲突，未覆盖现有术语"
            if entry.get("status") != "locked":
                entry["status"] = "locked"
                entry["preferred"] = target
                entry["target"] = target
                changed = True
            entry_id = entry.get("id")
        else:
            entry = _models.normalize_glossary_entry({
                "source": context["source"],
                "proposed_target": target,
                "target": target,
                "preferred": target,
                "behavior": "translate",
                "status": "locked",
                "scope": "document",
                "occurrences": context["occurrences"],
                "note": "由人工从翻译流知识候选提升为项目术语",
                "evidence": [{
                    "evidence_type": "user",
                    "source_name": "TransPraxis 知识候选",
                    "note": f"来自第 {(context['first_observed_segment'] + 1) if context['first_observed_segment'] is not None else '?'} 段的人工确认",
                    "quote": context["source_context"],
                    "url": "",
                }],
            })
            if entry is None:
                return state, False, "词条内容无效，未加入项目术语"
            entries.append(entry)
            entry_id = entry.get("id")
            changed = True
        candidate["promotion_entry_id"] = entry_id
        record("knowledge_project_term", "人工确认并加入项目术语；术语版本将更新", {
            "entry_id": entry_id,
        })
        state["glossary"] = entries
        save_job_state(job_id, state)
        if changed or not state.get("glossary_frozen"):
            state = freeze_glossary(job_id, entries=entries, frozen_by=actor)
        return state, True, "已加入项目术语，并通过术语版本冻结流程保存"

    if decision == "task_only":
        record("knowledge_task_only", "人工确认仅在本任务采用，不加入项目术语", {})
        save_job_state(job_id, state)
        return state, True, "已记录为仅本任务采用"

    record("knowledge_rejected", "人工拒绝该知识候选", {})
    save_job_state(job_id, state)
    return state, True, "已拒绝该知识候选"


def freeze_glossary(job_id, entries=None, frozen_by="user"):
    """冻结术语表：生成新版本 + 确定性 glossary_hash。

    修改后再次冻结 -> 新版本追加到 glossary_versions，不悄悄覆盖旧冻结状态。
    相同 canonical 内容再次冻结 -> 不创建新版本（幂等）。
    冻结新版本后立即对已翻译段落执行术语依赖失效（stale 标记 + TM 清除）。
    """
    state = load_job_state(job_id)
    if state is None:
        return None
    was_final = state.get("delivery_status") == "final"
    norm = normalize_glossary(entries if entries is not None
                              else state.get("glossary") or [])
    versions = state.get("glossary_versions") or []
    if versions and isinstance(versions[-1], dict) \
            and _models.entries_equal(versions[-1].get("entries") or [], norm):
        # 相同内容：不创建新版本（决策 A），保持原 frozen 快照
        state["glossary"] = norm
        state["glossary_frozen"] = versions[-1]
        state["stage"] = "GLOSSARY_FROZEN"
        save_job_state(job_id, state)
        return state
    source_hash = ""
    src = job_dir(job_id) / "source.bin"
    if src.is_file():
        source_hash = hashlib.sha256(src.read_bytes()).hexdigest()
    version = len(versions) + 1
    frozen = {
        "version": version,
        "source_hash": source_hash,
        "entries": norm,
        "frozen_at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "glossary_hash": _models.glossary_hash(norm),
        "frozen_by": frozen_by,
    }
    state["glossary"] = norm
    state["glossary_frozen"] = frozen
    versions.append(frozen)
    state["glossary_versions"] = versions
    state["stage"] = "GLOSSARY_FROZEN"
    if state.get("delivery_status") not in ("approved", "final"):
        state["delivery_status"] = "draft"
    # 术语决策变化 -> 立即失效受影响段落
    state, stale = _apply_glossary_staleness(state, job_id)
    if was_final and not stale:
        _invalidate_final_delivery_state(state)
    save_job_state(job_id, state)
    return state


def unfreeze_glossary(job_id):
    """返回修改：解除冻结（翻译开始前），回到 TERMS_PREPARED；旧冻结版本保留。"""
    state = load_job_state(job_id)
    if state is None:
        return None
    if state.get("p2_done"):
        # 翻译已开始：不允许解除冻结，只能生成新版本（见 freeze_glossary）
        return state
    state["glossary_frozen"] = None
    state["stage"] = "TERMS_PREPARED"
    if state.get("delivery_status") not in ("approved", "final"):
        state["delivery_status"] = "draft"
    save_job_state(job_id, state)
    return state


def save_document_profile(job_id, profile):
    """人工填写/修改文档画像后保存。"""
    state = load_job_state(job_id)
    if state is None:
        return None
    state["document_profile"] = _models.normalize_document_profile(profile)
    state["profile_done"] = True
    save_job_state(job_id, state)
    return state


def set_entity_translation(job_id, source_form, target, *, entity_type="other_proper_noun",
                           locked=True, actor="user", note=""):
    """Persist a human entity choice; it outranks generated continuity hints."""
    state = load_job_state(job_id)
    if state is None:
        return None
    registry = _entity_registry.EntityRegistry(state.get("entity_registry") or [])
    record = (registry.lock if locked else registry.accept)(
        source_form, target, entity_type=entity_type, note=note)
    state["entity_registry"] = registry.to_list()
    state.setdefault("human_actions", []).append({
        "finding_id": f"entity:{record.get('id') or source_form}",
        "action": "entity_locked" if locked else "entity_accepted",
        "note": note or f"人工确认实体译名：{source_form} -> {target}",
        "timestamp": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "actor": actor,
    })
    if (state.get("delivery_status") in ("approved", "final")
            or state.get("delivery_approved_by_human")):
        _invalidate_final_delivery_state(state)
    save_job_state(job_id, state)
    return state


def write_profile_artifacts(job_id, document_profile, style_profile):
    """把 Step 01 的画像产物落盘为版本化 artifact：
    document_profile.json / style_profile.json（含 style_profile_id 哈希）。
    返回写入的 style_profile_id，失败返回 None。
    """
    try:
        from transpraxis.style_profile import style_profile_id
        job_root = job_dir(job_id)
        job_root.mkdir(parents=True, exist_ok=True)
        (job_root / "document_profile.json").write_text(
            json.dumps(_models.normalize_document_profile(document_profile),
                       ensure_ascii=False, indent=2), encoding="utf-8")
        profile_id = style_profile_id(style_profile or {})
        artifact = dict(style_profile or {})
        artifact["style_profile_id"] = profile_id
        (job_root / "style_profile.json").write_text(
            json.dumps(artifact, ensure_ascii=False, indent=2), encoding="utf-8")
        return profile_id
    except Exception:  # noqa: BLE001 - artifact 落盘失败不影响主流程
        return None


def bypass_freeze(job_id, frozen_by="user"):
    """快速模式跳过人工冻结：允许以 provisional 术语直接翻译（记录审计标记）。"""
    state = load_job_state(job_id)
    if state is None:
        return None
    state["quality_bypass"] = True
    state.setdefault("human_actions", []).append({
        "action": "bypass_freeze",
        "note": "快速模式：跳过人工术语冻结，以 provisional 术语直接翻译",
        "timestamp": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "actor": frozen_by,
    })
    save_job_state(job_id, state)
    return state


# ================= 交付状态与人工处理记录 =================
def mark_findings_resolved(job_id, finding_ids, action, note="", actor="user"):
    """标记 findings 已人工处理，并重算交付状态。"""
    from transpraxis import delivery as _delivery
    state = load_job_state(job_id)
    if state is None:
        return None
    state, _marked = _delivery.mark_findings(state, finding_ids, action, note, actor)
    state["delivery_status"] = _delivery.compute_delivery_status(state)
    save_job_state(job_id, state)
    return state


def approve_delivery(job_id, note="", accept_blocking=False, actor="user",
                     target_lang="", provider="", model=""):
    """人工交付确认 -> final + immutable snapshot."""
    from transpraxis import delivery as _delivery
    state = load_job_state(job_id)
    if state is None:
        return None, False, ["任务不存在"]
    if state.get("report_enabled") and state.get("report_status") in {
            "incomplete", "failed_template_validation", "review_required"}:
        return state, False, ["实践报告尚未完成，不能冻结最终交付"]
    academic_records = (state.get("academic_state") or {}).get("artifacts") or {}
    # A pre-v0.4 job may have only p3_md and the old delivery formats.  It can
    # still be read and delivered through its historical path; strict report
    # compliance/QA begins once a structured report artifact is present.
    strict_report_gate = state.get("report_enabled") and (
        isinstance(academic_records.get("report"), dict) or
        any(name in academic_records for name in (
            "compliance", "final_docx_validation", "libreoffice_render",
            "report_qa")))
    case_gate = _finalization.case_review_gate(
        state, load_academic_artifact(job_id, "selected_cases"),
        require_artifact_status=bool(strict_report_gate))
    if case_gate.get("status") == "blocked":
        labels = ", ".join(case_gate.get("blocked_case_ids") or [])
        return state, False, [
            f"案例人工终审未通过：{labels}。作者拒绝或过期案例不能进入最终交付。"
        ]
    if strict_report_gate:
        compliance = compliance_profile_view(job_id, state)
        profile_rules = compliance.get("profile_compliance") or {}
        project = compliance.get("project_constraints") or {}
        if profile_rules.get("status") == "fail":
            return state, False, [
                "Default profile compliance failed: " +
                ", ".join(profile_rules.get("blocking_failures") or [])]
        if project.get("status") == "fail":
            return state, False, [
                "Project compliance constraint failed: " +
                ", ".join(project.get("failures") or [])]
        from transpraxis import academic_writer
        report_record = academic_writer.artifact_record(state, "report")
        docx_record = academic_writer.artifact_record(
            state, "final_docx_validation")
        render_record_state = academic_writer.artifact_record(
            state, "libreoffice_render")
        stale_artifacts = [name for name, record in (
            ("report", report_record),
            ("DOCX", docx_record),
            ("LibreOffice render", render_record_state),
        ) if record.get("status") in {"stale", "missing", "failed"}]
        if stale_artifacts:
            return state, False, [
                "Finalization artifacts are stale or unavailable: " +
                ", ".join(stale_artifacts)]
        final_docx = load_academic_artifact(job_id, "final_docx_validation") or {}
        render_record = load_academic_artifact(job_id, "libreoffice_render") or {}
        qa = _finalization.normalize_final_qa(state.get("final_qa"))
        structural = "PASS" if final_docx.get("status") in {
            "pass", "pass_with_warnings"} else "FAIL" if final_docx.get(
            "status") == "fail" else "NOT_RUN"
        qa_reasons = []
        if structural != "PASS":
            qa_reasons.append(f"Structural QA={structural}")
        if render_record.get("qa_status") != "PASS":
            qa_reasons.append(f"LibreOffice Render={render_record.get('qa_status', 'NOT_RUN')}")
        if qa.get("author_visual_review") != "CONFIRMED":
            qa_reasons.append("Author Visual Review=NOT_CONFIRMED")
        if qa.get("word_final_review") != "CONFIRMED":
            qa_reasons.append("Word Final Review=NOT_CONFIRMED")
        if qa_reasons:
            return state, False, ["Final QA gate blocked: " + "; ".join(qa_reasons)]
    validation = validate_delivery_translation_state(state)
    if validation["blocking"]:
        _record_delivery_validation_findings(state, validation)
        save_job_state(job_id, state)
        reasons = [issue.get("message", "译文未通过交付检查")
                   for issue in validation.get("issues") or []]
        return state, False, reasons or ["译文未通过最终交付检查"]
    state, ok, errors = _delivery.approve_delivery(state, note, actor, accept_blocking)
    if not ok:
        save_job_state(job_id, state)
        return state, ok, errors
    try:
        manifest = create_delivery_snapshot(
            job_id, state,
            target_lang=state.get("target_lang") or target_lang or "",
            provider=state.get("provider") or provider or "",
            model=state.get("model") or model or "")
    except Exception as exc:  # fail closed: final state is not saved without bytes
        persisted = load_job_state(job_id) or state
        return persisted, False, [f"无法冻结最终交付版本：{str(exc)[:200]}"]
    state.setdefault("delivery_snapshots", []).append({
        "version": manifest["snapshot_version"],
        "created_at": manifest["created_at"],
        "approval": dict(manifest.get("approval") or {}),
        "asset_count": len(manifest.get("assets") or []),
        "translation_state_identity": manifest.get("translation_state_identity"),
    })
    state["latest_delivery_snapshot_version"] = manifest["snapshot_version"]
    save_job_state(job_id, state)
    return state, True, []


def retranslate_segments(job_id, indexes, provider, api_key, model, target_lang,
                         style_rules="", glossary=None, on_status=None,
                         on_caption=None, actor="user"):
    """定点重译（抽取自 scripts/fix_segments.py 的能力）。"""
    from transpraxis import delivery as _delivery
    return _delivery.retranslate_segments(
        job_id, indexes, provider, api_key, model, target_lang,
        style_rules, glossary, on_status, on_caption, actor)


def delivery_status_label(state):
    labels = {"draft": "草稿（draft）", "review_required": "待审（review_required）",
              "approved": "已批准（approved）", "final": "最终交付（final）"}
    return labels.get(state.get("delivery_status"), str(state.get("delivery_status")))


def academic_status_label(state):
    status = state.get("report_status") or \
        (state.get("academic_state") or {}).get("report_status") or \
        (state.get("academic_state") or {}).get("quality_status") or \
        (state.get("academic_state") or {}).get("status") or "not_started"
    labels = {
        "not_started": "尚未开始", "stale": "需要更新（按影响范围）",
        "in_progress": "学术写作中", "failed": "学术写作失败",
        "pass": "验证通过", "pass_with_warnings": "通过（有警告）",
        "review_required": "需要人工学术复核", "fail": "验证失败",
        "generated": "报告已生成", "incomplete": "当前证据不足，报告不完整",
        "failed_template_validation": "报告生成未通过模板校验",
    }
    return labels.get(status, status)


def invalidate_academic_report(job_id, scope="all", section_id=None):
    """Invalidate academic artifacts only; translation stages remain intact."""
    from transpraxis import academic_writer
    state = load_job_state(job_id)
    if state is None:
        raise ValueError(f"找不到任务 {job_id}")
    academic_writer.invalidate_academic_state(state, scope, section_id)
    save_job_state(job_id, state)
    return state


def load_academic_artifact(job_id, name):
    """Read a canonical academic JSON artifact for UI/CLI inspection."""
    from transpraxis import academic_writer
    ARTIFACT_FILES = academic_writer.ARTIFACT_FILES
    if name not in ARTIFACT_FILES:
        raise ValueError(f"未知学术 artifact：{name}")
    path = job_dir(job_id) / ARTIFACT_FILES[name]
    if not path.is_file():
        return None
    try:
        return academic_writer._read_artifact(path)
    except Exception:
        return None


def record_human_evidence(job_id, question_id, answer, interface="academic_workspace"):
    """Record a human author answer for an open evidence question.

    The answer is stored verbatim with provenance; the question is marked
    answered; the case-analysis/section staleness is propagated through the
    existing dependency-hash architecture (only affected sections rewrite).
    """
    from transpraxis import academic_writer, human_evidence
    state = load_job_state(job_id)
    if state is None:
        raise ValueError(f"找不到任务 {job_id}")
    questions = academic_writer._read_artifact(
        job_dir(job_id) / academic_writer.ARTIFACT_FILES["human_evidence_questions"])
    if not questions:
        raise ValueError("当前没有待回答的人类证据问题；请先生成学术报告。")
    evidence = academic_writer._read_artifact(
        job_dir(job_id) / academic_writer.ARTIFACT_FILES["evidence"])
    entry, updated_questions = human_evidence.record_human_answer(
        questions, question_id, answer, evidence or {}, interface,
        existing=state.get("human_evidence") or [])
    entries = list(state.get("human_evidence") or [])
    entries = [x for x in entries if x.get("question_id") != question_id]
    entries.append(entry)
    state["human_evidence"] = entries
    academic_writer._write_artifact(
        job_dir(job_id) / academic_writer.ARTIFACT_FILES["human_evidence_questions"],
        updated_questions)
    record = state.get("academic_state", {}).get("artifacts", {}).get(
        "human_evidence_questions")
    if record:
        record["content_hash"] = updated_questions["content_hash"]
        record["updated_at"] = academic_writer._now()
    save_job_state(job_id, state)
    return entry


# ================= 阶段三：证据约束型学术写作 =================
def generate_mti_report(bilingual_pairs, termbase_dict, theory, provider, api_key,
                        model, state, job_id, on_status=None,
                        research_settings=None, literature_sources=None):
    """Compatibility wrapper for the evidence-grounded academic pipeline.

    ``bilingual_pairs`` and ``termbase_dict`` remain in the signature for old
    callers; the canonical inputs are now the saved translation state and the
    durable academic artifacts beside ``state.json``.
    """
    from transpraxis import academic_writer
    report_md = academic_writer.run_academic_pipeline(
        state, job_id, theory, provider, api_key, model,
        artifact_dir=job_dir(job_id), call_llm=call_llm,
        save_state=lambda current: save_job_state(job_id, current),
        research_settings=research_settings, literature_sources=literature_sources,
        on_status=on_status,
    )
    save_compliance_record(job_id, load_job_state(job_id) or state)
    return report_md


# ================= 主流程：单文档完整流水线 =================
def run_job_pipeline(job_id, filename, file_bytes, *, provider, api_key, model,
                     target_lang, auto_term, enable_report, translation_theory,
                     user_glossary=None, style_rules="", enable_review=True,
                     enable_annotate=True, use_tm=True,
                     strict_terminology_governance=False, mode=None,
                     research_settings=None, literature_sources=None,
                     delivery_config=None,
                     enable_understanding=None, reviewer_provider=None,
                     reviewer_api_key=None, reviewer_model=None,
                     reviewer_base_url=None, translator_base_url=None,
                     on_status=None, on_caption=None):
    """执行单个文档的完整流程；每个里程碑实时落盘，刷新/重启后均可继续。

    strict_terminology_governance=True：翻译前建立文档画像，并要求自动候选
    术语完成审核/冻结。导入的锁定术语视为已固定，不会阻塞翻译。
    ``mode`` 仅保留给旧调用方；quality 映射到严格术语治理，quick 映射到关闭。
    """
    _ensure_output_dir()
    base = new_job_state(filename)
    state = load_job_state(job_id) or base
    state = {**base, **state}  # 兼容旧版本状态缺字段
    state = _state_migration.migrate_state(state)
    saved_understanding = (state.get("pipeline_config") or {}).get(
        "enable_understanding")
    if mode is not None:
        strict_terminology_governance = mode == "quality"
    translator_config = _model_roles.normalize_role_config(
        None, fallback_provider=provider, fallback_model=model,
        fallback_api_key=api_key, fallback_base_url=translator_base_url)
    reviewer_config = _model_roles.normalize_role_config(
        {
            "provider": reviewer_provider,
            "model": reviewer_model,
            "api_key": reviewer_api_key,
            "base_url": reviewer_base_url,
        },
        fallback_provider=provider, fallback_model=model,
        fallback_api_key=api_key)
    pipeline_config = {
        "target_lang": target_lang,
        "auto_term": bool(auto_term),
        "enable_report": bool(enable_report),
        "translation_theory": translation_theory,
        "style_rules": style_rules,
        "enable_review": bool(enable_review),
        "enable_annotate": bool(enable_annotate),
        "use_tm": bool(use_tm),
        "strict_terminology_governance": bool(strict_terminology_governance),
        "enable_understanding": enable_understanding,
        "translator": _model_roles.public_role_config(translator_config),
        "reviewer": _model_roles.public_role_config(reviewer_config),
    }
    state["pipeline_config"] = pipeline_config
    state.update(
        target_lang=target_lang, auto_term_enabled=bool(auto_term),
        report_enabled=bool(enable_report), theory=translation_theory,
        style_rules=style_rules, enable_review=bool(enable_review),
        enable_annotate=bool(enable_annotate), use_tm=bool(use_tm),
        provider=provider, model=model,
        translator_config=_model_roles.public_role_config(translator_config),
        reviewer_config=_model_roles.public_role_config(reviewer_config),
    )
    if delivery_config is not None:
        state["delivery_config"] = normalize_delivery_config(
            delivery_config, enable_report=enable_report,
            enable_annotate=enable_annotate)
    elif state.get("delivery_config"):
        state["delivery_config"] = normalize_delivery_config(
            state["delivery_config"], enable_report=enable_report,
            enable_annotate=enable_annotate)
    if enable_understanding is None:
        if isinstance(saved_understanding, bool):
            # A resumed task keeps the decision it was created with.  This is
            # important for old Quick jobs whose legacy state has no new flag.
            enable_understanding = saved_understanding
        elif state.get("p1_done") or state.get("p2_done"):
            # Legacy in-progress jobs had no understanding pass outside strict
            # terminology governance; do not unexpectedly add API cost on resume.
            enable_understanding = bool(
                state.get("profile_done") or state.get("understanding_done")
                or state.get("quality_mode"))
        else:
            # New tasks: Quick skips it; Standard and Academic keep it enabled.
            enable_understanding = mode != "quick"
    pipeline_config["enable_understanding"] = bool(enable_understanding)
    state["pipeline_config"] = pipeline_config
    state["quality_mode"] = bool(strict_terminology_governance)
    warnings = state.setdefault("warnings", [])

    if enable_report:
        from transpraxis import academic_writer
        academic_writer.prepare_academic_inputs(
            state, translation_theory, research_settings, literature_sources)
        academic_writer.sync_versions(state)
    save_job_state(job_id, state)

    # 术语依赖失效：必须在“全部完成”早退之前执行，
    # 否则冻结术语表变更后的旧译文会继续以 reviewed/final/TM 状态存在。
    state, stale_segs = _apply_glossary_staleness(state, job_id)
    if stale_segs:
        save_job_state(job_id, state)

    # 全部完成 -> 直接返回
    if state["p1_done"] and state["p2_done"] and (not enable_report or state["p3_done"]) \
            and (not enable_annotate or state.get("annotations_done")):
        state["stage"] = _state_migration.derive_stage(state)
        return state

    # ---------------- 阶段一：排版清洗 ----------------
    if not state["p1_done"]:
        if on_status:
            on_status("【阶段一】排版解析与段落重建（确定性提取）...")
        if file_bytes is None:
            file_bytes = load_source(job_id)
        if file_bytes is None:
            raise ValueError("缺少源文件，请重新上传后再继续")

        paragraphs = []
        if filename.lower().endswith(".pdf"):
            paragraphs = [clean_xml_chars(p) for p in extract_pdf_paragraphs(file_bytes)]
        elif filename.lower().endswith(".docx"):
            doc_word = Document(io.BytesIO(file_bytes))
            for p in doc_word.paragraphs:
                for sub_p in re.split(r'\n+', clean_xml_chars(p.text)):
                    t = sub_p.strip()
                    if len(t) > 1 and not _ORNAMENT_RE.match(t):
                        paragraphs.append(t)

        if not paragraphs:
            raise ValueError("未提取到有效文本（若为扫描版 PDF，请先做 OCR 生成文本层）")
        state["paras"] = paragraphs
        if filename.lower().endswith(".pdf"):
            with fitz.open(stream=file_bytes, filetype="pdf") as source_pdf:
                state["source_page_count"] = source_pdf.page_count
        state["p1_done"] = True
        save_source(job_id, file_bytes)  # 留存源文件，刷新后无需重新上传
        save_job_state(job_id, state)

    # ---------------- 阶段 1.2：文档画像（长文理解；失败仅警告，不阻断） ----------------
    if enable_understanding and not state.get("profile_done"):
        if on_status:
            on_status("【阶段1.2】文档画像（分布式采样 + 结构化校验）...")
        from transpraxis.document_profile import profile_document
        profile, profile_warnings = profile_document(
            state["paras"], provider, api_key, model, target_lang)
        state["document_profile"] = profile
        state["profile_done"] = True
        for w in profile_warnings:
            if w not in warnings:
                warnings.append(w)
        if on_caption and profile:
            on_caption(f"✅ 文档画像完成：领域「{profile.get('domain') or '未知'}」"
                       f"· 文本类型「{profile.get('genre') or '未知'}」")
        elif on_caption:
            on_caption("⚠️ 文档画像失败，已跳过（可在 UI 中人工填写）。")
        save_job_state(job_id, state)

    # ---------------- 阶段 1.3：全文语义理解 ----------------
    if enable_understanding and not state.get("understanding_done"):
        if on_status:
            on_status("【阶段1.3】全文语义理解（语义单元摘要 + 全文概要）...")
        if on_caption:
            on_caption("🧭 正在建立全文概要与当前单元摘要...")
        understanding_call = call_llm
        understanding_base_url = getattr(_LLM_CTX, "base_url", None)
        if understanding_base_url:
            def understanding_call(provider_name, api_key_value, model_name,
                                   system_prompt, user_prompt, temperature=0.1):
                try:
                    return call_llm(
                        provider_name, api_key_value, model_name, system_prompt,
                        user_prompt, temperature=temperature,
                        base_url=understanding_base_url)
                except TypeError:
                    return call_llm(provider_name, api_key_value, model_name,
                                    system_prompt, user_prompt,
                                    temperature=temperature)
        units, digests, synopsis, understanding_warnings = \
            _context.build_document_understanding(
                state["paras"], state.get("document_profile"), provider, api_key,
                model, target_lang, call_llm=understanding_call,
                checkpoint_dir=job_dir(job_id))
        state["semantic_units"] = units
        state["section_digests"] = digests
        state["document_synopsis"] = synopsis
        state["understanding_warnings"] = list(understanding_warnings)
        state["understanding_done"] = True
        for warning in understanding_warnings:
            if warning not in warnings:
                warnings.append(warning)
        _context.write_understanding_artifacts(
            job_dir(job_id), units, digests, synopsis)
        if on_caption:
            status = synopsis.get("status") or "unavailable"
            on_caption(f"✅ 全文理解完成：{len(digests)} 个语义单元 · 概要状态 {status}")
        save_job_state(job_id, state)

    # ---------------- 阶段 1.5：智能抽取术语 ----------------
    if auto_term and not state["auto_terms"]:
        if on_status:
            on_status("【阶段1.5】正在 AI 智能抽取全文核心术语...")
        if on_caption:
            on_caption("🤖 正在从全文分布式样本中提取专业术语...")
        from transpraxis.terminology import extract_auto_terms_v2
        entries, extract_warnings = extract_auto_terms_v2(
            state["paras"], target_lang, provider, api_key, model,
            document_profile=state.get("document_profile"))
        state["auto_term_entries"] = entries
        state["auto_terms"] = {e["source"]: e["target"] for e in entries}
        if entries:
            if on_caption:
                on_caption(f"✅ 成功提取 {len(entries)} 个候选术语（全部出现位置已记录）")
        else:
            msg = "术语抽取失败（限流或返回格式异常），已跳过该步骤；可稍后点击“继续处理”重试。"
            if msg not in warnings and msg not in extract_warnings:
                warnings.append(msg)
        for w in extract_warnings:
            if w not in warnings:
                warnings.append(w)
        save_job_state(job_id, state)
    legacy_auto = [{"source": k, "target": v, "behavior": "translate",
                    "status": "provisional"}
                   for k, v in (state["auto_terms"] or {}).items()]
    auto_entries = normalize_glossary(state.get("auto_term_entries") or legacy_auto)
    user_entries = normalize_glossary(list(user_glossary or []))
    if not strict_terminology_governance:
        for e in auto_entries:
            if e["status"] == "candidate":
                e["status"] = "provisional"
    working = normalize_glossary(state.get("glossary") or [])
    if not working:
        working = normalize_glossary(user_entries + auto_entries)
    else:
        # 新上传/新抽取的术语若不在已保存审核表中，追加（不覆盖人工审核结果）
        known = {e["source"].casefold() for e in working}
        for e in user_entries + auto_entries:
            if e["source"].casefold() not in known:
                working.append(e)
    if not strict_terminology_governance:
        for e in working:
            if e["status"] == "candidate":
                e["status"] = "provisional"
    state["glossary"] = working
    glossary = working
    final_termbase = glossary_to_terms(glossary)

    # ---------------- 严格术语治理门禁：候选术语需人工审核/冻结后才能翻译 ----------------
    if strict_terminology_governance:
        pending = [e for e in working
                   if (e.get("status") or "").lower() == "candidate"]
        if not state.get("glossary_frozen") and not state.get("quality_bypass") \
                and pending:
            if on_status:
                on_status(f"⏸ 严格术语治理：{len(pending)} 条候选术语等待人工审核与冻结...")
            msg = (f"严格术语治理：术语表尚未冻结（{len(pending)} 条自动抽取的"
                   "候选术语待审核），翻译未开始。请在“术语准备与审核”面板"
                   "完成审核并冻结后继续；导入术语库中的锁定术语已视为固定，"
                   "无需再次审核。")
            if msg not in warnings:
                warnings.append(msg)
            state["stage"] = _state_migration.derive_stage(state)
            save_job_state(job_id, state)
            return state

    # ---------------- 阶段二：双语翻译（批次 + 确定性检查 + 独立审校 + 翻译记忆）----------------
    if not state["p2_done"]:
        if on_status:
            on_status("【阶段二】双语翻译与术语严格注入（批次翻译 + 确定性检查 + 独立审校）...")
        translate_stage(state, job_id, glossary, provider, api_key, model, target_lang,
                        style_rules, enable_review, use_tm=use_tm,
                        document_profile=state.get("document_profile"),
                        translator_config=translator_config,
                        reviewer_config=reviewer_config,
                        on_status=on_status, on_caption=on_caption)
        state["p2_done"] = True
        state["delivery_status"] = "review_required" if state.get("has_blocking") \
            else "draft"
        save_job_state(job_id, state)

    # ---------------- 阶段 2.5：三色自动标注 ----------------
    if enable_annotate and state["p2_done"] and not state.get("annotations_done"):
        if on_status:
            on_status("【阶段 2.5】自动标注学习重点（红=生僻词 / 黄=专业名词 / 青绿=难点句）...")
        annotate_stage(state, job_id, glossary, provider, api_key, model, target_lang,
                       on_caption=on_caption)

    # ---------------- 阶段三：报告生成 ----------------
    if enable_report and not state["p3_done"]:
        if on_status:
            on_status(f"【阶段三】基于《{translation_theory}》生成报告...")
        report_md = generate_mti_report(state["pairs"], final_termbase, translation_theory,
                                        provider, api_key, model, state, job_id,
                                        on_status=on_status,
                                        research_settings=research_settings,
                                        literature_sources=literature_sources)
        if not report_md.strip():
            if state.get("report_status") == "blocked_final_case_policy":
                state["stage"] = _state_migration.derive_stage(state)
                save_job_state(job_id, state)
                return state
            raise RuntimeError("报告内容为空，请点击“继续处理”重试")
        state["p3_md"] = report_md
        state["theory"] = translation_theory
        state["p3_done"] = True
        save_job_state(job_id, state)

    state["stage"] = _state_migration.derive_stage(state)
    save_job_state(job_id, state)
    return state
